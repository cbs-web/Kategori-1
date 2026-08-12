import os
import tempfile

from tkinter import ttk, filedialog, messagebox

from docx import Document

from ekler import (
    ekler_pdf_olustur as ekler_pdf_birlestir,
    ek_taahhutname_mi,
    ekleri_denetle,
)
from on_deger import is_durumu_degistir, normalize_tdth
from taahhutname_islemleri import taahhut_docx_pdfye_cevir
from word_jeoloji_birlestirme import (
    word_jeoloji_bolumunu_ekle,
    word_muhendislik_jeolojisi_cumlelerini_degistir,
)
from word_numaralandirma import word_baslik_numaralarini_normallestir


def benzersiz_cikti_yolu(yol):
    kok, uzanti = os.path.splitext(yol)
    sira = 2
    aday = yol
    while os.path.exists(aday):
        aday = f"{kok}_{sira}{uzanti}"
        sira += 1
    return aday


def word_alanlarini_guncelle(
    docx_yolu,
    jeoloji_word_yolu="",
    muhendislik_jeolojisi_cumlesi="",
    bolgesel_jeoloji_verisi=None,
):
    """TOC, SEQ, REF ve üstbilgi/altbilgi alanlarını Microsoft Word ile güncelle."""
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise RuntimeError(
            "Word alanlarını güncellemek için pywin32 kurulmalıdır. requirements.txt bağımlılıklarını kurun."
        ) from exc

    word = None
    belge = None
    pythoncom.CoInitialize()
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        belge = word.Documents.Open(
            FileName=os.path.abspath(docx_yolu),
            ReadOnly=False,
            AddToRecentFiles=False,
            Visible=False,
        )
        if jeoloji_word_yolu:
            word_jeoloji_bolumunu_ekle(
                word,
                belge,
                jeoloji_word_yolu,
                muhendislik_jeolojisi_cumlesi=muhendislik_jeolojisi_cumlesi,
                bolgesel_jeoloji_verisi=bolgesel_jeoloji_verisi,
            )
        else:
            word_muhendislik_jeolojisi_cumlelerini_degistir(
                belge,
                muhendislik_jeolojisi_cumlesi,
            )

        # Jeoloji Word'ü eklendikten sonra şekil ve tablo sayıları değişebilir.
        # Bütün Caption paragraflarını belge sırasındaki gerçek SEQ alanlarına
        # bağlamadan içindekiler/şekiller listelerini güncellemek eski önbellek
        # değerlerini ve yinelenen numaraları koruyordu.
        word_baslik_numaralarini_normallestir(word, belge)
        belge.Repaginate()
        for index in range(1, belge.TablesOfContents.Count + 1):
            belge.TablesOfContents(index).Update()
        for index in range(1, belge.TablesOfFigures.Count + 1):
            belge.TablesOfFigures(index).Update()
        belge.Fields.Update()
        for story_type in range(1, 18):
            try:
                alan = belge.StoryRanges(story_type)
            except Exception:
                continue
            while alan is not None:
                alan.Fields.Update()
                try:
                    alan = alan.NextStoryRange
                except Exception:
                    alan = None
        belge.Repaginate()
        for index in range(1, belge.TablesOfContents.Count + 1):
            belge.TablesOfContents(index).UpdatePageNumbers()
        for index in range(1, belge.TablesOfFigures.Count + 1):
            belge.TablesOfFigures(index).UpdatePageNumbers()
        belge.Save()
    except Exception as exc:
        raise RuntimeError(f"Microsoft Word alanları güncelleyemedi: {exc}") from exc
    finally:
        if belge is not None:
            try:
                belge.Close(SaveChanges=False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def atomik_docx_kaydet(doc, hedef_yol, son_islem=None):
    hedef_yol = os.path.abspath(hedef_yol)
    hedef_klasor = os.path.dirname(hedef_yol)
    if not os.path.isdir(hedef_klasor):
        raise ValueError(f"Hedef klasör bulunamadı: {hedef_klasor}")
    gecici_yol = ""
    try:
        with tempfile.NamedTemporaryFile(
            prefix=f".{os.path.basename(hedef_yol)}.",
            suffix=".tmp.docx",
            dir=hedef_klasor,
            delete=False,
        ) as f:
            gecici_yol = f.name
        doc.save(gecici_yol)
        if son_islem is not None:
            son_islem(gecici_yol)
        Document(gecici_yol)
        os.replace(gecici_yol, hedef_yol)
        gecici_yol = ""
    finally:
        if gecici_yol:
            try:
                os.remove(gecici_yol)
            except OSError:
                pass


class RaporlamaIslemleri:
    def __init__(self, app):
        object.__setattr__(self, "app", app)

    def __getattr__(self, name):
        return getattr(self.app, name)

    def __setattr__(self, name, value):
        if name == "app":
            object.__setattr__(self, name, value)
        else:
            setattr(self.app, name, value)

    def sekme8_rapor(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="8. Raporlama")
        page = ttk.Frame(frame, padding=16)
        page.pack(fill="both", expand=True)

        ttk.Label(page, text="Raporlama", style="Baslik.TLabel").pack(anchor="w", pady=(0, 10))
        ttk.Separator(page).pack(fill="x", pady=(0, 12))

        akis = ttk.LabelFrame(page, text="Rapor İş Akışı", padding=(12, 10), bootstyle="secondary")
        akis.pack(fill="x", pady=(0, 12))
        ttk.Button(akis, text="1. Şablon Seç", command=self.sablon_sec, bootstyle="secondary").pack(side="left", padx=(0, 8))
        ttk.Button(akis, text="2. Ön Kontrol", command=self.rapor_on_kontrol, bootstyle="warning").pack(side="left", padx=(0, 8))
        ttk.Button(akis, text="3. Şablonu Kontrol Et", command=self.sablon_kontrol_et, bootstyle="info").pack(side="left", padx=(0, 8))
        ttk.Button(akis, text="4. Raporu Oluştur", command=self.rapor_olustur, bootstyle="success").pack(side="left", padx=(0, 8))
        ttk.Button(akis, text="5. Nihai PDF Oluştur", command=self.nihai_rapor_pdf_olustur, bootstyle="success outline").pack(side="left")

        sablon_frame = ttk.LabelFrame(page, text="Word Şablonu", padding=(12, 10), bootstyle="secondary")
        sablon_frame.pack(fill="x")
        sablon_metni = f"Şablon: {self.sablon_yolu}" if getattr(self, "sablon_yolu", "") else "Seçilen Şablon: Yok"
        self.lbl_sablon = ttk.Label(sablon_frame, text=sablon_metni, style="Muted.TLabel", wraplength=980, justify="left")
        self.lbl_sablon.pack(anchor="w")

        jeoloji_frame = ttk.LabelFrame(page, text="Jeoloji Bölümü Şablonu", padding=(12, 10), bootstyle="secondary")
        jeoloji_frame.pack(fill="x", pady=(10, 0))
        ttk.Label(
            jeoloji_frame,
            text=(
                "Ana raporda [JEOLOJI_BOLUMU] etiketi varsa seçili kütüphane kaydının tam Word bölümü; "
                "kütüphane seçimi yoksa elle seçilen veya ilçe/köye göre bulunan Word içeriği eklenir."
            ),
            style="Muted.TLabel",
            wraplength=980,
            justify="left",
        ).pack(anchor="w", pady=(0, 8))
        butonlar = ttk.Frame(jeoloji_frame)
        butonlar.pack(fill="x", pady=(0, 6))
        ttk.Button(butonlar, text="Jeoloji Şablonu Seç", command=self.jeoloji_sablonu_sec, bootstyle="secondary").pack(side="left", padx=(0, 8))
        ttk.Button(butonlar, text="Otomatik Kullan", command=self.jeoloji_sablonu_otomatik_kullan, bootstyle="info").pack(side="left")
        self.lbl_jeoloji_sablon = ttk.Label(
            jeoloji_frame,
            text=self.jeoloji_sablon_etiket_metni(),
            style="Muted.TLabel",
            wraplength=980,
            justify="left",
        )
        self.lbl_jeoloji_sablon.pack(anchor="w")

    def sablon_sec(self):
        self.sablon_yolu = filedialog.askopenfilename(
            initialdir=self.sablon_alt_klasoru("rapor"),
            filetypes=[("Word Dosyaları", "*.docx")],
        )
        if self.sablon_yolu and hasattr(self, "lbl_sablon"):
            self.lbl_sablon.config(text=f"Şablon: {self.sablon_yolu}")

    def jeoloji_sablon_etiket_metni(self):
        if getattr(self, "jeoloji_kutuphanesi_bolumu_aktif", False):
            record_id = getattr(self, "jeoloji_kutuphanesi_kayit_id", None)
            path = getattr(self, "jeoloji_kutuphanesi_bolum_yolu", "")
            prefix = f"Kütüphane kaydı #{record_id}" if record_id else "Kütüphane kaydı"
            return f"{prefix}: Tam Word bölümü kullanılacak · {path or 'dosya yeniden bulunacak'}"
        yol = getattr(self, "jeoloji_sablon_yolu", "")
        if yol:
            return f"Seçilen Jeoloji Şablonu: {yol}"
        return "Jeoloji Şablonu: Otomatik bulunacak"

    def _jeoloji_kutuphanesi_bolum_secimini_temizle(self):
        self.jeoloji_kutuphanesi_bolumu_aktif = False
        self.jeoloji_kutuphanesi_kayit_id = None
        self.jeoloji_kutuphanesi_bolum_yolu = ""
        self.jeoloji_kutuphanesi_bolum_hash = ""
        self.jeoloji_kutuphanesi_uygulanan_genel = ""
        self.jeoloji_kutuphanesi_uygulanan_inceleme = ""

    def jeoloji_sablonu_sec(self):
        yol = filedialog.askopenfilename(
            initialdir=self.sablon_alt_klasoru("jeoloji"),
            filetypes=[("Word Dosyaları", "*.docx")],
        )
        if yol:
            self._jeoloji_kutuphanesi_bolum_secimini_temizle()
            self.jeoloji_sablon_yolu = yol
            if hasattr(self, "lbl_jeoloji_sablon"):
                self.lbl_jeoloji_sablon.config(text=self.jeoloji_sablon_etiket_metni())

    def jeoloji_sablonu_otomatik_kullan(self):
        self._jeoloji_kutuphanesi_bolum_secimini_temizle()
        self.jeoloji_sablon_yolu = ""
        if hasattr(self, "lbl_jeoloji_sablon"):
            self.lbl_jeoloji_sablon.config(text=self.jeoloji_sablon_etiket_metni())
            if hasattr(self, "durum_mesaji_yaz"):
                self.durum_mesaji_yaz("Rapor şablonu seçildi")

    def sablon_kontrol_et(self):
        if not hasattr(self, "sablon_yolu") or not self.sablon_yolu:
            messagebox.showwarning("Uyarı", "Lütfen önce bir Word şablonu seçin.")
            return
        try:
            bulunan = self.sablon_etiketlerini_oku(self.sablon_yolu)
            desteklenen = self.desteklenen_sablon_etiketleri()
            tanimli = sorted(bulunan & desteklenen)
            tanimsiz = sorted(bulunan - desteklenen)

            mesaj = (
                f"Toplam bulunan etiket: {len(bulunan)}\n"
                f"Program tarafından doldurulacak: {len(tanimli)}\n"
                f"Tanımsız kalabilecek: {len(tanimsiz)}\n\n"
                f"Tanımsız etiketler:\n{self.sablon_etiket_listesi_metni(tanimsiz)}\n\n"
                f"Doldurulacak etiketler:\n{self.sablon_etiket_listesi_metni(tanimli)}"
            )
            if tanimsiz:
                messagebox.showwarning("Şablon Kontrolü", mesaj)
            else:
                messagebox.showinfo("Şablon Kontrolü", mesaj)
            if hasattr(self, "durum_mesaji_yaz"):
                self.durum_mesaji_yaz("Şablon kontrolü tamamlandı")
        except ImportError:
            self.hata_kaydet("Şablon kontrolü için python-docx bulunamadı")
            messagebox.showerror("Hata", "Şablon kontrolü için python-docx kütüphanesi gerekli.")
        except Exception as e:
            self.hata_kaydet("Şablon kontrolü yapılamadı", e)
            messagebox.showerror("Hata", f"Şablon kontrolü yapılamadı:\n{e}")

    def sablon_etiket_listesi_metni(self, degerler):
        if not degerler:
            return "Yok"
        satirlar = degerler[:30]
        metin = "\n".join(satirlar)
        if len(degerler) > 30:
            metin += f"\n... ve {len(degerler) - 30} etiket daha"
        return metin

    def rapor_olustur(self):
        if hasattr(self, "degisiklik_izni_kontrol_et") and not self.degisiklik_izni_kontrol_et("Rapor oluşturma"):
            return
        if not hasattr(self, "sablon_yolu") or not self.sablon_yolu:
            messagebox.showwarning("Uyarı", "Lütfen önce bir Word şablonu seçin.")
            return
        if hasattr(self, "rapor_on_kontrol") and not self.rapor_on_kontrol(devam_sor=True):
            return
        kayit_yolu = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Belgesi", "*.docx")],
            confirmoverwrite=False,
        )
        if not kayit_yolu:
            return

        arsiv_yollari = {
            os.path.normcase(os.path.abspath(rev.get("rapor_yolu", "")))
            for rev in self.is_akisi_verisi.get("tamamlanan_revizyonlar", [])
            if rev.get("rapor_yolu")
        }
        if os.path.normcase(os.path.abspath(kayit_yolu)) in arsiv_yollari:
            messagebox.showerror(
                "Bitmiş Revizyon Koruması",
                "Yeni rapor önceki bitmiş revizyonun Word dosyasının üzerine kaydedilemez. Farklı bir dosya adı seçin.",
            )
            return

        if os.path.normcase(os.path.abspath(kayit_yolu)) == os.path.normcase(os.path.abspath(self.sablon_yolu)):
            messagebox.showerror("Hata", "Rapor çıktısı kaynak şablonun üzerine kaydedilemez.")
            return
        if os.path.exists(kayit_yolu):
            karar = messagebox.askyesnocancel(
                "Mevcut Dosya",
                "Seçilen rapor dosyası zaten var.\n\n"
                "Evet: mevcut dosyanın üzerine güvenli biçimde yaz\n"
                "Hayır: numaralı yeni bir dosya oluştur",
            )
            if karar is None:
                return
            if not karar:
                kayit_yolu = benzersiz_cikti_yolu(kayit_yolu)

        try:
            doc = Document(self.sablon_yolu)
            tamamlama_bilgisi = self.rapor_icerigini_doldur(doc) or {}
            jeoloji_word_yolu = str(tamamlama_bilgisi.get("jeoloji_word_yolu") or "")
            muhendislik_jeolojisi_cumlesi = str(
                tamamlama_bilgisi.get("muhendislik_jeolojisi_cumlesi") or ""
            )
            bolgesel_jeoloji_verisi = tamamlama_bilgisi.get("bolgesel_jeoloji_verisi") or {}

            def raporu_tamamla(gecici_yol):
                word_alanlarini_guncelle(
                    gecici_yol,
                    jeoloji_word_yolu=jeoloji_word_yolu,
                    muhendislik_jeolojisi_cumlesi=muhendislik_jeolojisi_cumlesi,
                    bolgesel_jeoloji_verisi=bolgesel_jeoloji_verisi,
                )
                kontrol_doc = Document(gecici_yol)
                self.rapor_uretici().rapor_cikti_belgesini_dogrula(kontrol_doc)

            atomik_docx_kaydet(doc, kayit_yolu, son_islem=raporu_tamamla)
            self.is_akisi_verisi["son_rapor_yolu"] = os.path.abspath(kayit_yolu)
            if self.is_akisi_verisi.get("durum") == "on_deger_verildi":
                self.is_akisi_verisi = is_durumu_degistir(
                    self.is_akisi_verisi,
                    "yazim_asamasinda",
                    "İlk Word raporu oluşturuldu.",
                )
            if hasattr(self, "on_deger_ekranini_guncelle"):
                self.on_deger_ekranini_guncelle()
            if hasattr(self, "durum_mesaji_yaz"):
                self.durum_mesaji_yaz("Rapor oluşturuldu")
            messagebox.showinfo("Başarı", f"Raporunuz {kayit_yolu} konumuna başarıyla kaydedildi!")
        except ImportError as e:
            self.hata_kaydet("Rapor oluşturma için python-docx bulunamadı", e)
            messagebox.showerror("Hata", "python-docx kütüphanesi gerekli.")
        except Exception as e:
            self.hata_kaydet("Rapor oluşturma sırasında hata oluştu", e)
            messagebox.showerror("Hata", f"Kayıt Hatası:\n{str(e)}")

    def nihai_rapor_pdf_olustur(self):
        if hasattr(self, "degisiklik_izni_kontrol_et") and not self.degisiklik_izni_kontrol_et("Nihai PDF oluşturma"):
            return
        word_yolu = self.is_akisi_verisi.get("son_rapor_yolu", "")
        if not word_yolu or not os.path.isfile(word_yolu):
            messagebox.showwarning("Nihai PDF", "Önce güncel Word raporunu oluşturun.")
            return
        tdth = normalize_tdth(getattr(self, "tdth_verisi", {}))
        aktif_tdth = tdth.get("aktif") or {}
        tdth_yolu = aktif_tdth.get("pdf_yolu", "")
        tdth_gecerli = (
            tdth.get("durum") not in {"eksik", "yenilenmeli"}
            and bool(tdth_yolu)
            and os.path.isfile(tdth_yolu)
        )
        if not tdth_gecerli and not messagebox.askyesno(
            "Nihai PDF Ek Uyarısı",
            "Geçerli ve güncel TDTH PDF bulunamadı. TDTH eki olmadan ana rapor ve "
            "mevcut geçerli eklerle devam edilsin mi?",
        ):
            return

        varsayilan = os.path.splitext(os.path.basename(word_yolu))[0] + "_NIHAI.pdf"
        hedef = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf")],
            initialfile=varsayilan,
            title="Nihai rapor PDF dosyasını kaydet",
            confirmoverwrite=False,
        )
        if not hedef:
            return
        arsiv_pdf_yollari = {
            os.path.normcase(os.path.abspath(rev.get("nihai_pdf_yolu", "")))
            for rev in self.is_akisi_verisi.get("tamamlanan_revizyonlar", [])
            if rev.get("nihai_pdf_yolu")
        }
        if os.path.normcase(os.path.abspath(hedef)) in arsiv_pdf_yollari:
            messagebox.showerror(
                "Bitmiş Revizyon Koruması",
                "Yeni nihai PDF önceki bitmiş revizyonun PDF dosyasının üzerine kaydedilemez. Farklı bir dosya adı seçin.",
            )
            return
        if os.path.exists(hedef):
            karar = messagebox.askyesnocancel(
                "Mevcut PDF",
                "Seçilen nihai PDF zaten var.\n\nEvet: güvenli biçimde üzerine yaz\nHayır: numaralı yeni dosya oluştur",
            )
            if karar is None:
                return
            if not karar:
                hedef = benzersiz_cikti_yolu(hedef)

        try:
            ek_denetim = ekleri_denetle(
                getattr(self, "ekler", {}),
                getattr(self, "ek_kategorileri", []),
                derin=True,
            )
            ek_sorunu_var = bool(
                ek_denetim["eksik_dosyalar"]
                or ek_denetim["donusum_gerekenler"]
                or ek_denetim.get("gecersiz_dosyalar")
            )
            if ek_sorunu_var:
                devam = messagebox.askyesno(
                    "Nihai PDF Ek Uyarısı",
                    "Bazı ek dosyaları eksik, dönüştürülemiyor veya okunamıyor. "
                    "Bu ekler atlanarak ana rapor ve geçerli eklerle devam edilsin mi?",
                )
                if not devam:
                    return
            gecerli_ek_yollari = {
                os.path.normcase(os.path.abspath(ek["yol"]))
                for ek in ek_denetim["sirali"]
                if ek["tag"] in {"var", "otomatik_donusum"}
            }
            hedef_klasor = os.path.dirname(os.path.abspath(hedef))
            with tempfile.TemporaryDirectory(prefix="k1_nihai_", dir=hedef_klasor) as gecici:
                ana_pdf = os.path.join(gecici, "00_ana_rapor.pdf")
                basarili, hata = taahhut_docx_pdfye_cevir(word_yolu, ana_pdf)
                if not basarili:
                    raise RuntimeError(f"Word raporu PDF'ye dönüştürülemedi: {hata}")

                paket = {
                    "ANA RAPOR": [{"baslik": "Ana Rapor", "yol": ana_pdf}],
                }
                kategoriler = ["ANA RAPOR"]
                for kategori in self.ek_kategorileri:
                    liste = [
                        dict(ek)
                        for ek in self.ekler.get(kategori, [])
                        if not ek_taahhutname_mi(ek)
                        and os.path.normcase(os.path.abspath(ek.get("yol", "")))
                        in gecerli_ek_yollari
                    ]
                    if kategori == "TDTH" and tdth_gecerli and not any(
                        os.path.normcase(os.path.abspath(ek.get("yol", "")))
                        == os.path.normcase(os.path.abspath(tdth_yolu))
                        for ek in liste if ek.get("yol")
                    ):
                        liste.append({"baslik": "Sismik Tehlike Haritası Detay Raporu", "yol": tdth_yolu})
                    paket[kategori] = liste
                    kategoriler.append(kategori)
                sonuc = ekler_pdf_birlestir(paket, kategoriler, hedef)

            self.is_akisi_verisi["son_nihai_pdf_yolu"] = os.path.abspath(hedef)
            if hasattr(self, "on_deger_ekranini_guncelle"):
                self.on_deger_ekranini_guncelle()
            self.durum_mesaji_yaz("Nihai rapor PDF oluşturuldu", os.path.basename(hedef))
            messagebox.showinfo(
                "Nihai PDF",
                f"Nihai rapor oluşturuldu:\n{sonuc['dosya']}\n\nToplam sayfa: {sonuc['sayfa']}",
            )
        except Exception as exc:
            self.hata_kaydet("Nihai rapor PDF oluşturulamadı", exc)
            messagebox.showerror("Nihai PDF", f"Nihai rapor oluşturulamadı:\n{exc}")
