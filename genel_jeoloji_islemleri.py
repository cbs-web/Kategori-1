"""Genel jeoloji haritası, yaş sırası ve 2.1 metni için K-1 arayüz akışı."""

from __future__ import annotations

import datetime as dt
import os
from pathlib import Path
import threading
import tkinter as tk
from tkinter import messagebox, ttk

from PIL import Image, ImageTk
from docx import Document

from formasyon_metin_kutuphanesi import (
    FormasyonMetinKutuphanesi,
    YAS_SECENEKLERI,
    YAS_SIRA_SOZLUGU,
    birimleri_yasli_gence_sirala,
    jeolojik_yas_tahmin_et,
)
from genel_jeoloji_haritasi import (
    analiz_gorsellerini_kapat,
    birim_serilestir,
    genel_jeoloji_gorselini_olustur,
    genel_jeoloji_verisini_hazirla,
)
from jeoloji_kutuphanesi import jeoloji_anahtari
from jeoloji_pafta_kutuphanesi import JeolojiPaftaHatasi, JeolojiPaftaKutuphanesi
from jeoloji_yapay_zeka import (
    JeolojiYapayZekaAyarlari,
    JeolojiYapayZekaHatasi,
    JeolojiYapayZekaServisi,
    YapayZekaAnahtariEksik,
)
from word_jeoloji_birlestirme import (
    bolgesel_jeoloji_basligi_mi,
    wordde_stratigrafik_kesit_var_mi,
    yapisal_jeoloji_basligi_mi,
)


KAYNAK_MODLARI = {
    "Programın hazırladığı 2.1": "kutuphane",
    "Seçili Word'deki 2.1": "eski_rapor",
}
KAYNAK_MODU_ETIKETI = {value: key for key, value in KAYNAK_MODLARI.items()}
# Eski projelerdeki karma modu yeni sade arayüzde program seçeneği olarak açılır.
KAYNAK_MODU_ETIKETI["karma"] = "Programın hazırladığı 2.1"


class GenelJeolojiIslemleri:
    def __init__(self, app):
        object.__setattr__(self, "app", app)

    def __getattr__(self, name):
        return getattr(self.app, name)

    def __setattr__(self, name, value):
        if name == "app":
            object.__setattr__(self, name, value)
        else:
            setattr(self.app, name, value)

    def _pafta_kutuphanesi(self):
        root = os.path.join(self.kullanici_veri_klasoru_bul(), "jeoloji_pafta_100k")
        return JeolojiPaftaKutuphanesi(root)

    def _metin_kutuphanesi(self):
        return FormasyonMetinKutuphanesi()

    @staticmethod
    def _birim_anahtari(birim):
        return (jeoloji_anahtari(birim.get("kod")), jeoloji_anahtari(birim.get("ad")))

    def genel_jeoloji_durumunu_guncelle(self):
        variable = getattr(self, "genel_jeoloji_durum", None)
        if variable is None:
            return
        state = getattr(self, "genel_jeoloji_verisi", {})
        path = getattr(self, "img_genel_jeoloji", "")
        mode_code = state.get("kaynak_modu", "kutuphane") if isinstance(state, dict) else ""
        ready = bool(
            isinstance(state, dict)
            and state
            and (
                mode_code == "eski_rapor"
                or (state.get("birimler") and path and os.path.isfile(path))
            )
        )
        if ready:
            mode = KAYNAK_MODU_ETIKETI.get(state.get("kaynak_modu"), state.get("kaynak_modu", ""))
            selected_count = sum(
                1 for unit in state["birimler"] if unit.get("kullan", True)
            )
            if mode_code == "eski_rapor":
                variable.set("Genel jeoloji hazır • 2.1, kesit ve 2.1.1 seçili Word'den")
            else:
                variable.set(f"Genel jeoloji hazır • {selected_count} birim • {mode}")
        elif getattr(self, "yuklu_kml_points", None):
            variable.set("Parsel hazır • genel jeoloji haritası ve 2.1 henüz oluşturulmadı")
        else:
            variable.set("Genel jeoloji için parsel KML'si yükleyin")

    def genel_jeoloji_haritasi_hazirla(self):
        if hasattr(self, "degisiklik_izni_kontrol_et") and not self.degisiklik_izni_kontrol_et(
            "Genel jeoloji haritası ve 2.1 hazırlama"
        ):
            return
        points = getattr(self, "yuklu_kml_points", [])
        if len(points) < 3:
            messagebox.showwarning(
                "Genel Jeoloji Haritası",
                "Önce çalışma parselinin KML dosyasını yükleyin.",
            )
            return
        self.root.configure(cursor="watch")
        self.root.update_idletasks()
        try:
            analysis = genel_jeoloji_verisini_hazirla(
                self._pafta_kutuphanesi(),
                points,
                ana_birim=getattr(self, "jeoloji_pafta_sonucu", {}),
            )
        except Exception as exc:
            self.hata_kaydet("Genel jeoloji haritası hazırlanamadı", exc)
            messagebox.showerror("Genel Jeoloji Haritası", str(exc))
            return
        finally:
            self.root.configure(cursor="")
        self._genel_jeoloji_penceresi(analysis)

    def _birimi_kutuphaneyle_tamamla(
        self, detected, index, library=None, saved_units=None, kullan_default=True
    ):
        library = library or self._metin_kutuphanesi()
        if saved_units is None:
            saved_state = getattr(self, "genel_jeoloji_verisi", {})
            saved_units = {
                self._birim_anahtari(unit): unit
                for unit in (
                    saved_state.get("birimler", []) if isinstance(saved_state, dict) else []
                )
            }
        unit = dict(detected)
        stored = library.getir(unit.get("kod"), unit.get("ad")) or {}
        snapshot = saved_units.get(self._birim_anahtari(unit), {})
        age, age_order = jeolojik_yas_tahmin_et(unit.get("kod"), unit.get("ad"))
        unit.setdefault("oran", 0.0)
        unit.setdefault("guven", 0.0)
        # Birim bu çalıştırmadaki yerel analizden geliyorsa eski proje anlık
        # görüntüsündeki AI-only işareti onu yerel aday olmaktan çıkarmamalıdır.
        # Katalogdan yalnız AI sonucu için eklenen satırlar ise açıkça False gelir.
        local_candidate = bool(
            unit.get("yerel_aday", not bool(unit.get("ai_eklenen", False)))
        )
        gemini_confidence = snapshot.get("ai_gemini_guven", 0)
        openai_confidence = snapshot.get("ai_openai_guven", 0)
        unit.update(
            {
                "iid": f"birim_{index}",
                "kullan": bool(snapshot.get("kullan", kullan_default)),
                "jeolojik_yas": snapshot.get("jeolojik_yas") or stored.get("jeolojik_yas") or age,
                "yas_sirasi": snapshot.get("yas_sirasi", stored.get("yas_sirasi", age_order)),
                "lejant_aciklamasi": snapshot.get(
                    "lejant_aciklamasi", stored.get("lejant_aciklamasi", "")
                ),
                "bolgesel_jeoloji_metni": snapshot.get(
                    "bolgesel_jeoloji_metni", stored.get("bolgesel_jeoloji_metni", "")
                ),
                "kutuphane_id": stored.get("id"),
                "kutuphane_revizyon_no": stored.get("revizyon_no"),
                "metin_kaynagi": snapshot.get(
                    "metin_kaynagi", "kalici_kutuphane" if stored else "proje_taslagi"
                ),
                "yerel_aday": local_candidate,
                "ai_gemini_guven": gemini_confidence,
                "ai_openai_guven": openai_confidence,
                "ai_gemini_kanit": snapshot.get(
                    "ai_gemini_kanit",
                    snapshot.get("ai_kanit", "") if gemini_confidence else "",
                ),
                "ai_openai_kanit": snapshot.get(
                    "ai_openai_kanit",
                    snapshot.get("ai_kanit", "") if openai_confidence else "",
                ),
                "ai_gemini_aciklama": snapshot.get(
                    "ai_gemini_aciklama",
                    snapshot.get("ai_aciklama", "") if gemini_confidence else "",
                ),
                "ai_openai_aciklama": snapshot.get(
                    "ai_openai_aciklama",
                    snapshot.get("ai_aciklama", "") if openai_confidence else "",
                ),
                "ai_birlesik_guven": snapshot.get("ai_birlesik_guven", 0),
                "ai_kanit": snapshot.get("ai_kanit", ""),
                "ai_durum": snapshot.get("ai_durum", "Henüz yok"),
                "ai_oneri": snapshot.get("ai_oneri", ""),
                "ai_aciklama": snapshot.get("ai_aciklama", ""),
                "ai_saglayicilar": list(snapshot.get("ai_saglayicilar", [])),
            }
        )
        return unit

    def _birimleri_kutuphaneyle_tamamla(self, analysis):
        library = self._metin_kutuphanesi()
        try:
            word_transfer = library.eski_jeoloji_wordlerinden_tamamla(
                analysis.get("birimler", [])
            )
        except Exception as exc:
            self.hata_kaydet("Eski Word 2.1 metinleri ayrılamadı", exc)
            word_transfer = {"aktarilan": [], "bulunamayan": []}
        self._genel_jeoloji_word_metin_aktarimi = word_transfer
        saved_state = getattr(self, "genel_jeoloji_verisi", {})
        saved_units = {
            self._birim_anahtari(unit): unit
            for unit in (saved_state.get("birimler", []) if isinstance(saved_state, dict) else [])
        }
        result = [
            self._birimi_kutuphaneyle_tamamla(
                detected, index, library=library, saved_units=saved_units
            )
            for index, detected in enumerate(analysis.get("birimler", []))
        ]
        result_keys = {self._birim_anahtari(unit) for unit in result}
        catalog = {
            self._birim_anahtari(unit): unit
            for unit in analysis.get("lejant_birimleri", [])
        }
        # Önceki proje anlık görüntüsünde AI ile bulunup seçilmiş bir birim, yeni
        # yerel örneklemede ilk 12'ye girmese de düzenleyicide görünmeye devam etsin.
        for key, snapshot in saved_units.items():
            if key in result_keys or key not in catalog:
                continue
            detected = dict(catalog[key])
            detected["oran"] = 0.0
            detected["guven"] = 0.0
            detected["yerel_aday"] = False
            detected["ai_eklenen"] = True
            result.append(
                self._birimi_kutuphaneyle_tamamla(
                    detected,
                    len(result),
                    library=library,
                    saved_units=saved_units,
                    kullan_default=True,
                )
            )
            result_keys.add(key)
        return result

    def genel_jeoloji_eksik_metinlerini_tamamla(self):
        """Projedeki boş 2.1 birim metinlerini kalıcı kütüphaneden tamamla.

        Rapor ön kontrolü sırasında çağrılabilir. Kullanıcının dolu birim
        metinlerini ve elle düzenleyebildiği birleşik 2.1 metin alanını
        değiştirmez. Kalıcı kütüphanede metin yoksa önce onaylı eski raporların
        2.1 bölümlerinden güvenle ayrıştırmayı dener.
        """
        state = getattr(self, "genel_jeoloji_verisi", {})
        if not isinstance(state, dict) or state.get("kaynak_modu") == "eski_rapor":
            return {"tamamlanan": [], "bulunamayan": []}

        units = state.get("birimler", [])
        if not isinstance(units, list):
            return {"tamamlanan": [], "bulunamayan": []}
        missing_units = [
            unit
            for unit in units
            if isinstance(unit, dict)
            and unit.get("kullan", True)
            and not str(unit.get("bolgesel_jeoloji_metni") or "").strip()
        ]
        if not missing_units:
            return {"tamamlanan": [], "bulunamayan": []}

        library = self._metin_kutuphanesi()
        try:
            library.eski_jeoloji_wordlerinden_tamamla(missing_units)
        except Exception as exc:
            self.hata_kaydet("Eksik 2.1 metinleri eski Word kayıtlarından tamamlanamadı", exc)

        completed = []
        unresolved = []
        for unit in missing_units:
            stored = library.getir(unit.get("kod"), unit.get("ad")) or {}
            text = str(stored.get("bolgesel_jeoloji_metni") or "").strip()
            if not text:
                unresolved.append(str(unit.get("ad") or unit.get("kod") or "Adsız birim"))
                continue
            unit["bolgesel_jeoloji_metni"] = text
            unit["kutuphane_id"] = stored.get("id")
            unit["kutuphane_revizyon_no"] = stored.get("revizyon_no")
            unit["metin_kaynagi"] = "kalici_kutuphane"
            completed.append(str(unit.get("ad") or unit.get("kod") or "Adsız birim"))

        if completed:
            state["eksik_metinler"] = unresolved
            self._proje_kirli = True

        return {"tamamlanan": completed, "bulunamayan": unresolved}

    def _genel_jeoloji_penceresi(self, analysis):
        previous = getattr(self, "genel_jeoloji_verisi", {})
        units = self._birimleri_kutuphaneyle_tamamla(analysis)
        previous_hash = (
            str(previous.get("geometri_hash") or "")
            if isinstance(previous, dict) else ""
        )
        current_hash = str(analysis.get("geometri_hash") or "")
        same_geometry = bool(previous_hash and current_hash and previous_hash == current_hash)
        if isinstance(previous, dict) and previous.get("birimler") and not same_geometry:
            # Farklı bir parsel/KML için eski model görüşleri geçerli değildir.
            # Yerel satırların kullanıcı metinleri korunur; yalnız eski AI-only
            # satırlar ve sağlayıcı sonuçları temizlenir.
            units = [unit for unit in units if unit.get("yerel_aday", True)]
            for unit in units:
                for field in (
                    "ai_gemini_guven", "ai_openai_guven", "ai_birlesik_guven",
                ):
                    unit[field] = 0
                for field in (
                    "ai_gemini_kanit", "ai_openai_kanit", "ai_gemini_aciklama",
                    "ai_openai_aciklama", "ai_kanit", "ai_aciklama",
                ):
                    unit[field] = ""
                unit["ai_durum"] = "Yerel analiz"
                unit["ai_oneri"] = "Kontrol"
                unit["ai_saglayicilar"] = []
        previous_mode = previous.get("kaynak_modu", "kutuphane") if isinstance(previous, dict) else "kutuphane"
        if previous_mode not in {"kutuphane", "eski_rapor"}:
            previous_mode = "kutuphane"
        previous_providers = {
            provider
            for provider in (
                previous.get("ai_saglayicilar", [])
                if isinstance(previous, dict) and same_geometry else []
            )
            if provider in {"gemini", "openai"}
        }
        previous_provider_summaries = (
            previous.get("ai_saglayici_ozetleri", {})
            if isinstance(previous, dict)
            and same_geometry
            and isinstance(previous.get("ai_saglayici_ozetleri", {}), dict)
            else {}
        )

        window = self.animasyonlu_pencere()
        window.title("1/100.000 Genel Jeoloji Haritası ve 2.1 Bölgesel Jeoloji")
        window.geometry("1500x900")
        window.minsize(1150, 720)
        state = {
            "analysis": analysis,
            "units": units,
            "active_iid": None,
            "photo": None,
            "closing": False,
            "ai_busy": False,
            "ai_provider_runs": set(previous_providers),
            "ai_provider_meta": {
                provider: dict(summary)
                for provider, summary in previous_provider_summaries.items()
                if provider in {"gemini", "openai"} and isinstance(summary, dict)
            },
            "ai_inceleme_tarihi": (
                previous.get("ai_inceleme_tarihi", "")
                if isinstance(previous, dict) else ""
            ),
            "ai_onbellekten": bool(
                previous.get("ai_onbellekten", False)
                if isinstance(previous, dict) else False
            ),
        }

        header = ttk.Frame(window, padding=(12, 10))
        header.pack(fill="x")
        ttk.Label(
            header,
            text="Parsel merkezli 1/100.000 genel jeoloji haritası",
            font=("Segoe UI", 14, "bold"),
        ).pack(anchor="w")
        ttk.Label(
            header,
            text=(
                "Program seçeneğinde kullanılacak birimleri burada denetleyin; bunlar harita ve 2.1'de "
                "yaşlıdan gence sıralanır. Word seçeneğinde 2.1 doğrudan seçili jeoloji Word'ünden gelir."
            ),
            wraplength=1350,
            justify="left",
        ).pack(anchor="w", pady=(3, 8))

        source_row = ttk.Frame(header)
        source_row.pack(fill="x")
        ttk.Label(source_row, text="2.1 içeriği:").pack(side="left")
        mode_var = tk.StringVar(
            master=window,
            value=previous_mode,
        )
        ttk.Radiobutton(
            source_row,
            text="Program oluştursun (önerilen)",
            variable=mode_var,
            value="kutuphane",
            command=lambda: mode_changed(),
        ).pack(side="left", padx=(10, 14))
        ttk.Radiobutton(
            source_row,
            text="Seçili Word'deki 2.1'i kullan",
            variable=mode_var,
            value="eski_rapor",
            command=lambda: mode_changed(),
        ).pack(side="left")
        ttk.Label(
            source_row,
            text="Stratigrafik kesit ve 2.1.1 her zaman seçili jeoloji Word'ünden alınır.",
            foreground="#555555",
        ).pack(side="left", padx=(20, 0))

        ai_row = ttk.Frame(header)
        ai_row.pack(fill="x", pady=(8, 0))
        ai_status = tk.StringVar(
            master=window,
            value=(
                f"Yerel analiz hazır • {sum(1 for unit in units if unit.get('yerel_aday'))} aday"
                + (
                    " • Eski Word'lerden "
                    f"{len(getattr(self, '_genel_jeoloji_word_metin_aktarimi', {}).get('aktarilan', []))} "
                    "birim metni aktarıldı"
                    if getattr(self, "_genel_jeoloji_word_metin_aktarimi", {}).get("aktarilan")
                    else ""
                )
                + (
                    " • Kayıtlı denetim: "
                    + ", ".join(
                        "Gemini" if provider == "gemini" else "Sol"
                        for provider in sorted(previous_providers)
                    )
                    if previous_providers else ""
                )
                + " • Son seçim sizindir."
            ),
        )
        ttk.Label(ai_row, textvariable=ai_status).pack(side="left", fill="x", expand=True)
        ai_settings_button = ttk.Button(
            ai_row,
            text="AI Ayarları",
            command=lambda: open_ai_settings(),
        )
        ai_settings_button.pack(side="right")
        openai_check_button = ttk.Button(
            ai_row,
            text="GPT-5.6 Sol ile Denetle",
            command=lambda: run_ai("openai", False),
            bootstyle="primary",
        )
        openai_check_button.pack(side="right", padx=(0, 8))
        gemini_check_button = ttk.Button(
            ai_row,
            text="Gemini ile Denetle",
            command=lambda: run_ai("gemini", False),
            bootstyle="info",
        )
        gemini_check_button.pack(side="right", padx=(0, 8))

        pane = ttk.PanedWindow(window, orient="horizontal")
        pane.pack(fill="both", expand=True, padx=12, pady=(0, 8))
        left = ttk.Frame(pane)
        right = ttk.Frame(pane, padding=(10, 0, 0, 0))
        pane.add(left, weight=3)
        pane.add(right, weight=4)

        word_mode_panel = ttk.LabelFrame(
            window,
            text="Seçili jeoloji Word'ü kullanılacak",
            padding=(18, 16),
        )
        ttk.Label(
            word_mode_panel,
            text=(
                "Başka bir ayar yapmanız gerekmiyor. 2.1 ve 2.1.1; seçili Word'deki "
                "yazı, tablo ve görselleriyle birlikte rapora alınacak."
            ),
            font=("Segoe UI", 11),
            wraplength=1000,
            justify="left",
        ).pack(anchor="w")
        word_mode_source = tk.StringVar(master=window, value="")
        ttk.Label(
            word_mode_panel,
            textvariable=word_mode_source,
            foreground="#555555",
        ).pack(anchor="w", pady=(8, 0))

        preview_label = ttk.Label(left, anchor="center")
        preview_label.pack(fill="both", expand=True)
        preview_image = analysis["harita"].copy()
        preview_image.thumbnail((700, 520), Image.Resampling.LANCZOS)
        state["photo"] = ImageTk.PhotoImage(preview_image, master=window)
        preview_image.close()
        preview_label.configure(image=state["photo"])
        ttk.Label(
            left,
            text=(
                f"Merkez: {analysis['merkez'][0]:.6f}, {analysis['merkez'][1]:.6f} • "
                f"Harita paneli: {analysis['genislik_km']:g} × {analysis['yukseklik_km']:g} km"
            ),
            justify="center",
        ).pack(fill="x", pady=(5, 0))

        columns = (
            "kullan", "kod", "ad", "yas", "oran", "yerel",
            "gemini", "sol", "sonuc", "metin",
        )
        tree_frame = ttk.Frame(right)
        tree_frame.pack(fill="x")
        tree = ttk.Treeview(
            tree_frame, columns=columns, show="headings", selectmode="browse", height=8
        )
        labels = {
            "kullan": "Kullan",
            "kod": "Kod",
            "ad": "Birim",
            "yas": "Jeolojik yaş",
            "oran": "Örnek payı %",
            "yerel": "Yerel güven",
            "gemini": "Gemini sonucu",
            "sol": "Sol sonucu",
            "sonuc": "Ortak değerlendirme",
            "metin": "2.1 metni",
        }
        widths = {
            "kullan": 55,
            "kod": 60,
            "ad": 180,
            "yas": 120,
            "oran": 90,
            "yerel": 85,
            "gemini": 140,
            "sol": 140,
            "sonuc": 190,
            "metin": 75,
        }
        for column in columns:
            tree.heading(column, text=labels[column])
            tree.column(column, width=widths[column], anchor="center" if column != "ad" else "w")
        tree.grid(row=0, column=0, sticky="nsew")
        tree_scroll = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
        tree_scroll.grid(row=1, column=0, sticky="ew")
        tree.configure(xscrollcommand=tree_scroll.set)
        tree_frame.columnconfigure(0, weight=1)

        editor = ttk.LabelFrame(right, text="Seçili birimin kalıcı bilgileri", padding=(10, 8))
        editor.pack(fill="both", expand=True, pady=(10, 0))
        form = ttk.Frame(editor)
        form.pack(fill="x")
        ttk.Label(form, text="Jeolojik yaş:").grid(row=0, column=0, sticky="w")
        age_var = tk.StringVar(master=window)
        age_combo = ttk.Combobox(
            form,
            textvariable=age_var,
            values=[name for name, _order in YAS_SECENEKLERI],
            state="readonly",
            width=38,
        )
        age_combo.grid(row=0, column=1, sticky="ew", padx=(8, 0))
        form.columnconfigure(1, weight=1)

        ttk.Label(editor, text="Harita lejantındaki kısa açıklama:").pack(anchor="w", pady=(8, 3))
        legend_text = tk.Text(editor, height=3, wrap="word", font=("Segoe UI", 10))
        legend_text.pack(fill="x")
        ttk.Label(editor, text="2.1 Bölgesel Jeoloji uzun metni:").pack(anchor="w", pady=(8, 3))
        regional_text = tk.Text(editor, height=7, wrap="word", font=("Segoe UI", 10))
        regional_text.pack(fill="both", expand=True)
        editor_status = tk.StringVar(master=window, value="")
        ttk.Label(editor, textvariable=editor_status).pack(anchor="w", pady=(5, 0))

        def unit_by_iid(iid):
            return next((unit for unit in units if unit["iid"] == iid), None)

        evidence_labels = {
            "kod_ve_desen": "kod + desen",
            "kod_okundu": "kod okundu",
            "yalniz_desen": "yalnız desen",
        }

        def recalculate_comparison(unit):
            runs = state.get("ai_provider_runs", set())
            local_present = bool(unit.get("yerel_aday", True))
            gemini_confidence = float(unit.get("ai_gemini_guven") or 0)
            openai_confidence = float(unit.get("ai_openai_guven") or 0)
            gemini_present = gemini_confidence > 0
            openai_present = openai_confidence > 0
            ai_confidences = [
                value for value in (gemini_confidence, openai_confidence) if value > 0
            ]
            unit["ai_birlesik_guven"] = (
                round(sum(ai_confidences) / len(ai_confidences), 1)
                if ai_confidences else 0
            )
            evidence_order = {"kod_ve_desen": 3, "kod_okundu": 2, "yalniz_desen": 1}
            evidences = [
                value for value in (
                    unit.get("ai_gemini_kanit"), unit.get("ai_openai_kanit")
                ) if value
            ]
            unit["ai_kanit"] = (
                max(evidences, key=lambda value: evidence_order.get(value, 0))
                if evidences else ""
            )
            descriptions = []
            for value in (
                unit.get("ai_gemini_aciklama"), unit.get("ai_openai_aciklama")
            ):
                value = str(value or "").strip()
                if value and value not in descriptions:
                    descriptions.append(value)
            unit["ai_aciklama"] = " | ".join(descriptions)

            if not runs:
                status, recommendation = "Yerel analiz", "Kontrol"
            elif len(runs) == 1:
                provider = next(iter(runs))
                provider_name = "Gemini" if provider == "gemini" else "Sol"
                provider_present = gemini_present if provider == "gemini" else openai_present
                provider_confidence = (
                    gemini_confidence if provider == "gemini" else openai_confidence
                )
                if local_present and provider_present:
                    status = f"Yerel + {provider_name} ortak"
                    recommendation = "Kullan" if provider_confidence >= 65 else "Kontrol"
                elif provider_present:
                    status, recommendation = f"{provider_name} ekledi", "Kontrol"
                else:
                    status, recommendation = f"{provider_name} görmedi", "Kontrol"
            else:
                present = []
                if local_present:
                    present.append("Yerel")
                if gemini_present:
                    present.append("Gemini")
                if openai_present:
                    present.append("Sol")
                if len(present) == 3:
                    status = "3 kaynak ortak"
                    recommendation = (
                        "Kullan" if min(gemini_confidence, openai_confidence) >= 60
                        else "Kontrol"
                    )
                elif len(present) == 2:
                    missing = next(
                        name for name in ("Yerel", "Gemini", "Sol") if name not in present
                    )
                    status = f"2 kaynak ortak • {missing} görmedi"
                    recommendation = (
                        "Kullan"
                        if not local_present
                        and gemini_present and openai_present
                        and min(gemini_confidence, openai_confidence) >= 65
                        else "Kontrol"
                    )
                elif present:
                    status, recommendation = f"Çelişki • yalnız {present[0]}", "Kontrol"
                else:
                    status, recommendation = "Kaynaklar birimi görmedi", "Kontrol"
            unit["ai_durum"] = status
            unit["ai_oneri"] = recommendation
            unit["ai_saglayicilar"] = sorted(runs)

        def provider_value(unit, provider):
            if provider not in state.get("ai_provider_runs", set()):
                return "—"
            prefix = "ai_gemini" if provider == "gemini" else "ai_openai"
            confidence = float(unit.get(f"{prefix}_guven") or 0)
            if confidence <= 0:
                return "Görmedi"
            evidence = evidence_labels.get(unit.get(f"{prefix}_kanit"), "")
            return f"%{confidence:g}" + (f" • {evidence}" if evidence else "")

        def tree_values(unit):
            recalculate_comparison(unit)
            local_present = bool(unit.get("yerel_aday", True))
            return (
                "✓" if unit.get("kullan") else "—",
                unit.get("kod", ""),
                unit.get("ad", ""),
                unit.get("jeolojik_yas", ""),
                unit.get("oran", 0) if local_present else "—",
                f"%{float(unit.get('guven') or 0):g}" if local_present else "—",
                provider_value(unit, "gemini"),
                provider_value(unit, "openai"),
                unit.get("ai_durum", "Yerel analiz"),
                "Hazır" if unit.get("bolgesel_jeoloji_metni", "").strip() else "Eksik",
            )

        def commit_editor():
            unit = unit_by_iid(state.get("active_iid"))
            if unit is None:
                return
            unit["jeolojik_yas"] = age_var.get().strip()
            unit["yas_sirasi"] = YAS_SIRA_SOZLUGU.get(unit["jeolojik_yas"], 9999)
            unit["lejant_aciklamasi"] = legend_text.get("1.0", "end").strip()
            unit["bolgesel_jeoloji_metni"] = regional_text.get("1.0", "end").strip()
            if unit.get("metin_kaynagi") == "kalici_kutuphane" and not unit.get("kutuphane_id"):
                unit["metin_kaynagi"] = "proje_taslagi"
            if tree.exists(unit["iid"]):
                tree.item(unit["iid"], values=tree_values(unit))

        def refresh_tree(select_iid=None):
            for item in tree.get_children():
                tree.delete(item)
            for unit in birimleri_yasli_gence_sirala(units):
                tree.insert(
                    "",
                    "end",
                    iid=unit["iid"],
                    values=tree_values(unit),
                )
            if select_iid and tree.exists(select_iid):
                tree.selection_set(select_iid)
                tree.focus(select_iid)

        def selection_changed(_event=None):
            selection = tree.selection()
            if not selection:
                return
            iid = selection[0]
            if state.get("active_iid") and state["active_iid"] != iid:
                commit_editor()
            unit = unit_by_iid(iid)
            state["active_iid"] = iid
            age_var.set(unit.get("jeolojik_yas", ""))
            legend_text.delete("1.0", "end")
            legend_text.insert("1.0", unit.get("lejant_aciklamasi", ""))
            regional_text.delete("1.0", "end")
            regional_text.insert("1.0", unit.get("bolgesel_jeoloji_metni", ""))
            source = unit.get("metin_kaynagi", "proje_taslagi")
            ai_detail = ""
            if state.get("ai_provider_runs"):
                gemini_text = provider_value(unit, "gemini")
                openai_text = provider_value(unit, "openai")
                ai_detail = (
                    f" • Gemini: {gemini_text} • Sol: {openai_text}"
                    f" • Sonuç: {unit.get('ai_durum')} / {unit.get('ai_oneri') or 'Kontrol'}"
                )
            editor_status.set(
                f"{unit.get('ad')} ({unit.get('kod')}) • Kaynak: {source} • "
                f"Kütüphane revizyonu: {unit.get('kutuphane_revizyon_no') or 'yok'}"
                f"{ai_detail}"
            )

        def toggle_selected(_event=None):
            selection = tree.selection()
            if not selection:
                return
            unit = unit_by_iid(selection[0])
            unit["kullan"] = not unit.get("kullan", True)
            refresh_tree(unit["iid"])

        tree.bind("<<TreeviewSelect>>", selection_changed)
        tree.bind("<Double-1>", toggle_selected)

        def open_ai_settings():
            manager = JeolojiYapayZekaAyarlari(self.kullanici_veri_klasoru_bul())
            settings = manager.oku()
            dialog = self.animasyonlu_pencere(window)
            dialog.title("Jeoloji Yapay Zekâ Ayarları")
            dialog.geometry("700x650")
            dialog.minsize(650, 590)
            dialog.resizable(True, True)
            dialog.transient(window)
            dialog.grab_set()

            # Düğme şeridi içerikten önce paketlenir. Böylece Windows DPI/yazı
            # ölçeği büyüdüğünde içerik daralsa bile Kaydet düğmesi altta kalır.
            buttons = ttk.Frame(dialog, padding=(16, 8, 16, 16))
            buttons.pack(side="bottom", fill="x")
            body = ttk.Frame(dialog, padding=16)
            body.pack(side="top", fill="both", expand=True)
            ttk.Label(
                body,
                text="Jeoloji paftası AI denetimi",
                font=("Segoe UI", 13, "bold"),
            ).pack(anchor="w")
            ttk.Label(
                body,
                text=(
                    "Gemini 3.6 Flash ve GPT-5.6 Sol birbirinden bağımsızdır. Yalnız hangi "
                    "modelin düğmesine basarsanız o servis çağrılır. Anahtarlar proje dosyasına "
                    "yazılmaz; Windows Kimlik Bilgileri Yöneticisi'nde saklanır."
                ),
                wraplength=610,
                justify="left",
            ).pack(anchor="w", pady=(4, 12))

            provider_frame = ttk.LabelFrame(body, text="API anahtarları", padding=10)
            provider_frame.pack(fill="x")
            key_vars = {
                "gemini": tk.StringVar(master=dialog),
                "openai": tk.StringVar(master=dialog),
            }
            source_vars = {
                "gemini": tk.StringVar(master=dialog),
                "openai": tk.StringVar(master=dialog),
            }

            def refresh_key_sources():
                source_vars["gemini"].set(
                    f"Kayıt durumu: {manager.anahtar_kaynagi('gemini')}"
                )
                source_vars["openai"].set(
                    f"Kayıt durumu: {manager.anahtar_kaynagi('openai')}"
                )

            for row_no, (provider, title) in enumerate(
                (("gemini", "Gemini API anahtarı"), ("openai", "OpenAI API anahtarı"))
            ):
                ttk.Label(provider_frame, text=title).grid(
                    row=row_no * 2, column=0, sticky="w", pady=(3, 0)
                )
                ttk.Entry(
                    provider_frame,
                    textvariable=key_vars[provider],
                    show="●",
                    width=48,
                ).grid(row=row_no * 2, column=1, sticky="ew", padx=(10, 8), pady=(3, 0))

                def delete_key(selected_provider=provider):
                    if manager.anahtar_kaynagi(selected_provider).endswith("ortam değişkeni"):
                        messagebox.showinfo(
                            "API Anahtarı",
                            "Bu anahtar ortam değişkeninden geliyor; program içinden silinemez.",
                            parent=dialog,
                        )
                        return
                    manager.anahtar_sil(selected_provider)
                    refresh_key_sources()

                ttk.Button(
                    provider_frame,
                    text="Kayıtlıyı Sil",
                    command=delete_key,
                ).grid(row=row_no * 2, column=2, pady=(3, 0))
                ttk.Label(
                    provider_frame,
                    textvariable=source_vars[provider],
                    foreground="#555555",
                ).grid(row=row_no * 2 + 1, column=1, sticky="w", padx=(10, 0), pady=(0, 6))
            provider_frame.columnconfigure(1, weight=1)
            refresh_key_sources()

            option_frame = ttk.LabelFrame(body, text="Çalışma biçimi", padding=10)
            option_frame.pack(fill="x", pady=(12, 0))
            cache_var = tk.BooleanVar(
                master=dialog, value=bool(settings.get("onbellek_etkin", True))
            )
            ttk.Checkbutton(
                option_frame,
                text="Her model için aynı harita ve ayarlardaki önbellek sonucunu kullan",
                variable=cache_var,
            ).grid(row=0, column=0, sticky="w")
            ttk.Label(
                option_frame,
                text="Otomatik ikinci görüş yoktur; Gemini ve Sol düğmeleri ayrı çalışır.",
                foreground="#555555",
            ).grid(row=1, column=0, sticky="w", pady=(8, 0))

            ttk.Label(
                body,
                text=(
                    "Not: Harita ve lejant görüntüleri yalnız siz düğmeye bastığınızda seçili "
                    "API sağlayıcısına gönderilir. Sonuç hiçbir zaman doğrudan rapora uygulanmaz."
                ),
                wraplength=610,
                foreground="#555555",
                justify="left",
            ).pack(anchor="w", pady=(12, 0))

            def save_settings():
                try:
                    for provider in ("gemini", "openai"):
                        key = key_vars[provider].get().strip()
                        if key:
                            manager.anahtar_kaydet(provider, key)
                    manager.kaydet(
                        {
                            "onbellek_etkin": cache_var.get(),
                        }
                    )
                except Exception as exc:
                    messagebox.showerror("AI Ayarları", str(exc), parent=dialog)
                    return
                dialog.destroy()
                ai_status.set(
                    "AI ayarları kaydedildi. Denetim yalnız siz başlattığınızda çalışır."
                )

            ttk.Separator(buttons, orient="horizontal").pack(fill="x", pady=(0, 10))
            ttk.Button(buttons, text="İptal", command=dialog.destroy).pack(side="right")
            ttk.Button(
                buttons,
                text="API Anahtarlarını ve Ayarları Kaydet",
                command=save_settings,
                bootstyle="success",
            ).pack(side="right", padx=(0, 8))
            dialog.bind("<Return>", lambda _event: save_settings())
            dialog.bind("<Escape>", lambda _event: dialog.destroy())

        def apply_ai_result(result):
            provider = result.get("saglayici")
            if provider not in {"gemini", "openai"}:
                raise JeolojiYapayZekaHatasi("AI sonucu sağlayıcı bilgisi içermiyor.")
            provider_name = "Gemini" if provider == "gemini" else "GPT-5.6 Sol"
            prefix = "ai_gemini" if provider == "gemini" else "ai_openai"
            state["ai_provider_runs"].add(provider)
            state["ai_provider_meta"][provider] = dict(result)
            state["ai_inceleme_tarihi"] = dt.datetime.now().isoformat(timespec="seconds")
            state["ai_onbellekten"] = bool(
                state.get("ai_onbellekten") or result.get("onbellekten")
            )

            # Aynı sağlayıcı yeniden çalıştırıldığında eski sonucu temizlenir;
            # diğer sağlayıcının ve yerel analizin sonuçları korunur.
            for unit in units:
                unit[f"{prefix}_guven"] = 0
                unit[f"{prefix}_kanit"] = ""
                unit[f"{prefix}_aciklama"] = ""
            by_code = {
                jeoloji_anahtari(unit.get("kod")): unit
                for unit in units
                if jeoloji_anahtari(unit.get("kod"))
            }
            catalog_by_code = {
                jeoloji_anahtari(unit.get("kod")): unit
                for unit in analysis.get("lejant_birimleri", [])
                if jeoloji_anahtari(unit.get("kod"))
            }
            added = 0
            for ai_unit in result.get("birimler", []):
                key = jeoloji_anahtari(ai_unit.get("kod"))
                unit = by_code.get(key)
                confidence = float(ai_unit.get("guven") or 0)
                if unit is None and confidence > 0:
                    catalog_unit = catalog_by_code.get(key)
                    if catalog_unit is None:
                        continue
                    detected = dict(catalog_unit)
                    preview = catalog_unit.get("onizleme")
                    detected["onizleme"] = preview.copy() if preview is not None else None
                    detected["oran"] = 0.0
                    detected["guven"] = 0.0
                    detected["yerel_aday"] = False
                    detected["ai_eklenen"] = True
                    analysis.setdefault("birimler", []).append(detected)
                    try:
                        self._metin_kutuphanesi().eski_jeoloji_wordlerinden_tamamla(
                            [detected]
                        )
                    except Exception as exc:
                        self.hata_kaydet(
                            "AI birimi için eski Word 2.1 metni ayrılamadı", exc
                        )
                    unit = self._birimi_kutuphaneyle_tamamla(
                        detected,
                        len(units),
                        kullan_default=False,
                    )
                    units.append(unit)
                    by_code[key] = unit
                    added += 1
                if unit is None:
                    continue
                unit[f"{prefix}_guven"] = confidence
                unit[f"{prefix}_kanit"] = ai_unit.get("kanit", "")
                unit[f"{prefix}_aciklama"] = ai_unit.get("konum_aciklamasi", "")

            for unit in units:
                recalculate_comparison(unit)
            recommended = sum(1 for unit in units if unit.get("ai_oneri") == "Kullan")
            common = sum(1 for unit in units if "ortak" in str(unit.get("ai_durum", "")))
            refresh_tree(state.get("active_iid"))
            if state.get("active_iid") and tree.exists(state["active_iid"]):
                selection_changed()
            cache_note = " • önbellekten" if result.get("onbellekten") else ""
            completed = ", ".join(
                "Gemini" if value == "gemini" else "Sol"
                for value in sorted(state["ai_provider_runs"])
            )
            ai_status.set(
                f"{provider_name} denetimi hazır{cache_note} • {len(result.get('birimler') or [])} birim "
                f"• {added} yeni lejant birimi • Çalıştırılan: {completed} • Son seçim sizindir."
            )
            details = [
                f"{provider_name} tarafından görülen birim: {len(result.get('birimler') or [])}",
                f"Kaynakların ortak bulduğu satır: {common}",
                f"Kullan önerisi: {recommended}",
                f"Yeni eklenen lejant birimi: {added}",
                "Kullan/Kullanma seçimleri otomatik değiştirilmedi.",
            ]
            messagebox.showinfo(
                f"{provider_name} Denetimi",
                f"{provider_name} denetimi tamamlandı.\n\n" + "\n".join(details)
                + "\n\nDiğer modeli çalıştırabilir veya son seçimleri doğrudan denetleyebilirsiniz.",
                parent=window,
            )

        def run_ai(provider, force_refresh=False):
            if state.get("ai_busy") or state.get("closing"):
                return
            if provider not in {"gemini", "openai"}:
                return
            provider_name = "Gemini" if provider == "gemini" else "GPT-5.6 Sol"
            service = JeolojiYapayZekaServisi(self.kullanici_veri_klasoru_bul())
            if not service.ayarlar.anahtar_al(provider):
                messagebox.showwarning(
                    f"{provider_name} Anahtarı Eksik",
                    f"Önce AI Ayarları bölümünden {provider_name} API anahtarını kaydedin.",
                    parent=window,
                )
                open_ai_settings()
                return
            ai_status.set("Harita parçaları ve tüm pafta lejantları hazırlanıyor…")
            window.update_idletasks()
            try:
                package = service.paket_hazirla(analysis)
            except Exception as exc:
                messagebox.showerror(f"{provider_name} Denetimi", str(exc), parent=window)
                ai_status.set(f"{provider_name} denetimi başlatılamadı; yerel analiz korundu.")
                return
            state["ai_busy"] = True
            gemini_check_button.configure(state="disabled")
            openai_check_button.configure(state="disabled")
            ai_settings_button.configure(state="disabled")
            ai_status.set(f"{provider_name} haritayı ve lejantları denetliyor…")

            def finish_error(error):
                if state.get("closing"):
                    return
                state["ai_busy"] = False
                gemini_check_button.configure(state="normal")
                openai_check_button.configure(state="normal")
                ai_settings_button.configure(state="normal")
                ai_status.set(f"{provider_name} denetimi tamamlanamadı; yerel analiz korundu.")
                title = (
                    f"{provider_name} Anahtarı Eksik"
                    if isinstance(error, YapayZekaAnahtariEksik)
                    else f"{provider_name} Denetimi"
                )
                messagebox.showerror(title, str(error), parent=window)

            def finish_success(result):
                if state.get("closing"):
                    return
                state["ai_busy"] = False
                gemini_check_button.configure(state="normal")
                openai_check_button.configure(state="normal")
                ai_settings_button.configure(state="normal")
                apply_ai_result(result)

            def worker():
                try:
                    result = service.saglayici_ile_analiz_et(
                        package,
                        saglayici=provider,
                        zorla_yenile=force_refresh,
                    )
                except Exception as exc:
                    try:
                        window.after(0, lambda error=exc: finish_error(error))
                    except tk.TclError:
                        pass
                    return
                try:
                    window.after(0, lambda value=result: finish_success(value))
                except tk.TclError:
                    pass

            threading.Thread(
                target=worker,
                name=f"jeoloji-ai-{provider}",
                daemon=True,
            ).start()

        action_row = ttk.Frame(editor)
        action_row.pack(fill="x", pady=(8, 0))

        def save_unit_to_library():
            commit_editor()
            unit = unit_by_iid(state.get("active_iid"))
            if unit is None:
                return
            try:
                saved = self._metin_kutuphanesi().kaydet(
                    kod=unit.get("kod"),
                    ad=unit.get("ad"),
                    jeolojik_yas=unit.get("jeolojik_yas"),
                    yas_sirasi=unit.get("yas_sirasi"),
                    lejant_aciklamasi=unit.get("lejant_aciklamasi"),
                    bolgesel_jeoloji_metni=unit.get("bolgesel_jeoloji_metni"),
                    kaynak_notu="K-1 genel jeoloji düzenleyicisi",
                )
            except Exception as exc:
                messagebox.showerror("Birim Kütüphanesi", str(exc), parent=window)
                return
            unit["kutuphane_id"] = saved.get("id")
            unit["kutuphane_revizyon_no"] = saved.get("revizyon_no")
            unit["metin_kaynagi"] = "kalici_kutuphane"
            editor_status.set(
                f"Kalıcı kütüphaneye kaydedildi • Revizyon {saved.get('revizyon_no')}"
            )
            refresh_tree(unit["iid"])

        ttk.Button(
            action_row,
            text="Bu Birimi Kalıcı Kütüphaneye Kaydet",
            command=save_unit_to_library,
            bootstyle="success",
        ).pack(side="left")
        ttk.Button(
            action_row,
            text="Kullan / Kullanma",
            command=toggle_selected,
        ).pack(side="left", padx=(8, 0))

        footer = ttk.Frame(window, padding=(12, 0, 12, 12))
        # Alt işlem düğmeleri, içerik panellerinin talep ettiği yükseklikten bağımsız
        # olarak her zaman pencerenin görünür alt kenarında kalır.
        footer.pack(fill="x", side="bottom", before=pane)
        footer_status = tk.StringVar(
            master=window,
            value="Çift tıklayarak birimi lejant ve 2.1 kapsamına alabilir/çıkarabilirsiniz.",
        )
        ttk.Label(footer, textvariable=footer_status).pack(side="left", fill="x", expand=True)

        def combined_text(selected, mode):
            if mode == "eski_rapor":
                return str(
                    getattr(self, "jeoloji_kutuphanesi_uygulanan_genel", "") or ""
                ).strip()
            paragraphs = []
            for unit in birimleri_yasli_gence_sirala(selected):
                heading = f"{unit.get('ad')} ({unit.get('kod')})" if unit.get("kod") else unit.get("ad", "")
                text = str(unit.get("bolgesel_jeoloji_metni") or "").strip()
                paragraphs.append(f"{heading}\n{text}".strip())
            return "\n\n".join(paragraphs).strip()

        def save_project():
            commit_editor()
            mode = mode_var.get() if mode_var.get() in {"kutuphane", "eski_rapor"} else "kutuphane"
            source_word = self.rapor_uretici().rapor_jeoloji_sablon_yolu_bul()
            if not source_word:
                source_word = self.jeoloji_kutuphanesi_islemleri().jeoloji_2_1_1_kaynagini_bagla(
                    parent=window,
                    bolgesel_de_kullanilacak=(mode == "eski_rapor"),
                )
            if not source_word:
                return
            try:
                source_document = Document(source_word)
                has_structural_geology = any(
                    yapisal_jeoloji_basligi_mi(paragraph.text)
                    for paragraph in source_document.paragraphs
                )
                has_regional_geology = any(
                    bolgesel_jeoloji_basligi_mi(paragraph.text)
                    for paragraph in source_document.paragraphs
                )
            except Exception as exc:
                messagebox.showerror("2.1.1 Kaynağı", str(exc), parent=window)
                return
            if not has_structural_geology:
                messagebox.showwarning(
                    "2.1.1 Kaynağı",
                    "Seçili Word'de '2.1.1 Yapısal Jeoloji ve Aktif Tektonik' başlığı yok. "
                    "Başka bir kütüphane kaydı uygulayın.",
                    parent=window,
                )
                return
            if not wordde_stratigrafik_kesit_var_mi(source_word):
                messagebox.showwarning(
                    "Stratigrafik Kesit Kaynağı",
                    "Seçili Word'de görseliyle birlikte aktarılabilir stratigrafik "
                    "kesit bulunamadı. Başka bir kütüphane kaydı uygulayın.",
                    parent=window,
                )
                return
            if mode == "eski_rapor" and not has_regional_geology:
                messagebox.showwarning(
                    "2.1 Kaynağı",
                    "Seçili Word'de '2.1 Bölgesel Jeoloji' başlığı yok. "
                    "Başka bir kütüphane kaydı uygulayın.",
                    parent=window,
                )
                return
            selected = [unit for unit in units if unit.get("kullan")]
            if not selected and mode == "kutuphane":
                messagebox.showwarning("Genel Jeoloji", "En az bir formasyon kullanılmalıdır.", parent=window)
                return
            selected = birimleri_yasli_gence_sirala(selected)
            missing = [unit.get("ad", "") for unit in selected if not unit.get("bolgesel_jeoloji_metni", "").strip()]
            if missing and mode != "eski_rapor":
                proceed = messagebox.askyesno(
                    "Eksik 2.1 Metinleri",
                    "Şu birimlerin uzun 2.1 açıklaması henüz boş:\n- "
                    + "\n- ".join(missing)
                    + "\n\nHaritayı ve proje taslağını yine de kaydetmek ister misiniz?",
                    parent=window,
                )
                if not proceed:
                    return
            project_path = getattr(self, "guncel_dosya_yolu", "")
            if project_path:
                target_dir = Path(project_path).resolve().parent / "Haritalar"
            else:
                target_dir = Path(self.kullanici_veri_klasoru_bul()) / "Genel_Jeoloji"
            output_path = ""
            if mode == "kutuphane":
                target_path = target_dir / "Genel_Jeoloji_Haritasi.jpg"
                try:
                    output_path = genel_jeoloji_gorselini_olustur(
                        analysis["harita"], selected, target_path
                    )
                except Exception as exc:
                    self.hata_kaydet("Genel jeoloji görseli kaydedilemedi", exc)
                    messagebox.showerror("Genel Jeoloji Haritası", str(exc), parent=window)
                    return
            text = combined_text(selected, mode)
            # Rapor yalnız `kullan=True` satırlarını tüketir. Buna karşılık tüm
            # satırları saklamak, yeniden açıldığında iki modelin "gördü/görmedi"
            # karşılaştırmasının kaybolmasını önler.
            serial_units = []
            for unit in birimleri_yasli_gence_sirala(units):
                snapshot = birim_serilestir(unit)
                snapshot["kullan"] = bool(unit.get("kullan", False))
                serial_units.append(snapshot)
            if output_path:
                self.img_genel_jeoloji = output_path
            provider_summaries = {
                provider: {
                    "model": result.get("model", ""),
                    "genel_guven": result.get("genel_guven", 0),
                    "ana_parsel_kodu": result.get("ana_parsel_kodu", ""),
                    "notlar": result.get("notlar", ""),
                    "onbellekten": bool(result.get("onbellekten", False)),
                }
                for provider, result in state.get("ai_provider_meta", {}).items()
            }
            self.genel_jeoloji_verisi = {
                "surum": 4,
                "kaynak_modu": mode,
                "eski_rapor_kayit_id": None,
                "eski_rapor_revizyon_no": None,
                "merkez": list(analysis.get("merkez", [])),
                "sinir": dict(analysis.get("sinir", {})),
                "genislik_km": analysis.get("genislik_km"),
                "yukseklik_km": analysis.get("yukseklik_km"),
                "geometri_hash": analysis.get("geometri_hash", ""),
                "pafta_idleri": list(analysis.get("pafta_idleri", [])),
                "pafta_adlari": list(analysis.get("pafta_adlari", [])),
                "birimler": serial_units,
                "bolgesel_jeoloji_metni": text,
                "eksik_metinler": missing if mode == "kutuphane" else [],
                "ai_inceleme_tarihi": state.get("ai_inceleme_tarihi", ""),
                "ai_saglayicilar": sorted(state.get("ai_provider_runs", set())),
                "ai_saglayici_ozetleri": provider_summaries,
                "ai_onbellekten": bool(state.get("ai_onbellekten", False)),
                "ai_uyarilar": [],
                "olusturma_tarihi": dt.datetime.now().isoformat(timespec="seconds"),
            }
            if hasattr(self, "txt_formasyon_rapor"):
                self.txt_formasyon_rapor.delete("1.0", "end")
                self.txt_formasyon_rapor.insert("1.0", text)
            self.genel_jeoloji_durumunu_guncelle()
            if hasattr(self, "proje_durum_seridi_guncelle"):
                self.proje_durum_seridi_guncelle()
            close()
            messagebox.showinfo(
                "Genel Jeoloji Hazır",
                (
                    "Programın hazırladığı harita ve 2.1 kaydedildi; stratigrafik kesit "
                    "seçili Word'den alınacak."
                    if mode == "kutuphane"
                    else "Seçili Word'ün 2.1 bölümü, stratigrafik kesiti ve 2.1.1'i kullanılacak."
                )
                + (f"\n\n{output_path}" if output_path else "")
                + (
                    f"\n\nEksik uzun metin: {len(missing)}"
                    if missing and mode == "kutuphane" else ""
                ),
            )

        ttk.Button(
            footer,
            text="2.1 Seçimini Projeye Uygula",
            command=save_project,
            bootstyle="primary",
        ).pack(side="right", padx=(8, 0))
        ttk.Button(footer, text="İptal", command=lambda: close()).pack(side="right")

        def mode_changed(_event=None):
            mode = mode_var.get()
            if mode == "eski_rapor":
                ai_row.pack_forget()
                pane.pack_forget()
                if not word_mode_panel.winfo_manager():
                    word_mode_panel.pack(
                        fill="x",
                        padx=12,
                        pady=(0, 8),
                        after=footer,
                    )
                try:
                    source_word = self.rapor_uretici().rapor_jeoloji_sablon_yolu_bul()
                except Exception:
                    source_word = ""
                word_mode_source.set(
                    f"Seçili dosya: {Path(source_word).name}"
                    if source_word
                    else "Henüz bir jeoloji Word'ü seçilmedi."
                )
            else:
                word_mode_panel.pack_forget()
                if not ai_row.winfo_manager():
                    ai_row.pack(fill="x", pady=(8, 0), after=source_row)
                if not pane.winfo_manager():
                    pane.pack(
                        fill="both",
                        expand=True,
                        padx=12,
                        pady=(0, 8),
                        after=footer,
                    )
            footer_status.set(
                "2.1, stratigrafik kesit ve 2.1.1 seçili jeoloji Word'ünden biçimleriyle alınır."
                if mode == "eski_rapor"
                else (
                    "2.1 program tarafından hazırlanır; stratigrafik kesit ve 2.1.1 "
                    "seçili jeoloji Word'ünden biçimiyle otomatik alınır."
                )
            )

        mode_changed()

        def close():
            if state["closing"]:
                return
            state["closing"] = True
            analiz_gorsellerini_kapat(analysis)
            state["photo"] = None
            try:
                window.destroy()
            except tk.TclError:
                pass

        window.protocol("WM_DELETE_WINDOW", close)
        refresh_tree()
        children = tree.get_children()
        if children:
            tree.selection_set(children[0])
            tree.focus(children[0])
            selection_changed()


__all__ = ["GenelJeolojiIslemleri"]
