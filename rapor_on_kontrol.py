import os
from tkinter import messagebox

from docx import Document
from PIL import Image

from ekler import ekleri_denetle
from harita_dosyalari import proje_klasorundeki_rapor_haritalarini_bul
from genel_jeoloji_haritasi import parsel_geometri_hashi
from on_deger import normalize_tdth
from word_jeoloji_birlestirme import (
    bolgesel_jeoloji_basligi_mi,
    wordde_stratigrafik_kesit_var_mi,
    yapisal_jeoloji_basligi_mi,
)


ZORUNLU_PROJE_ALANLARI = [
    ("PROJE_ADI", "Proje Sahibi"),
    ("IL", "İl"),
    ("ILCE", "İlçe"),
    ("PAFTA", "Pafta"),
    ("ADA", "Ada"),
    ("PARSEL", "Parsel"),
]

JEOFIZIK_VERI_ETIKETLERI = {"[JEOFIZIK_PARAMETRE]", "[MASW]", "[VP]"}
GORSEL_ETIKETLERI = {
    "[MJH]": "img_mjh",
    "[JEOFIZIK_LOKASYON]": "img_jeofizik_lok",
    "[JEOLOJI_LOKASYON]": "img_jeoloji_lok",
    "[YERBULDURU]": "img_yerbulduru",
    "[PARSEL_HARITASI]": "img_parsel_haritasi",
}
TASIMA_GIRDI_ETIKETLERI = {
    "c": "Kohezyon",
    "phi": "İçsel sürtünme açısı",
    "gn": "Doğal birim hacim ağırlık",
    "gsat": "Doygun birim hacim ağırlık",
    "yass": "Yeraltı suyu derinliği",
    "B": "Temel genişliği",
    "L": "Temel uzunluğu",
    "Df": "Temel derinliği",
    "RvGk": "Dayanım katsayısı",
}
KAYA_TASIMA_GIRDI_ETIKETLERI = {
    "qt": "Karakteristik kaya taşıma gücü",
    "Gk": "Kaya güvenlik katsayısı",
    "ks_carpani": "Yatak katsayısı çarpanı",
}
TASIMA_RAPOR_ETIKETLERI = {"[TASIMA_GUCU]", "[QK]", "[QT]", "[KS]", "[DF]"}
RAPOR_TAAHHUT_BILGI_ALANLARI = (
    "JEOFIZIK_MUH_AD",
    "JEOFIZIK_MUH_SICIL",
    "JEOFIZIK_MUH_ADRES",
    "JEOFIZIK_MUH_TELEFON",
    "JEOLOJI_MUH_AD",
    "JEOLOJI_MUH_SICIL",
    "JEOLOJI_MUH_ADRES",
    "JEOLOJI_MUH_TELEFON",
)


class RaporOnKontrol:
    def __init__(self, app):
        object.__setattr__(self, "app", app)

    def __getattr__(self, name):
        return getattr(self.app, name)

    def deger_bos_mu(self, deger):
        return str(deger or "").strip() in ("", "-")

    def anlamli_satir_var_mi(self, satirlar):
        for satir in satirlar or []:
            for deger in satir:
                if not self.deger_bos_mu(deger):
                    return True
        return False

    def sonuclari_hazirla(self):
        hatalar = []
        uyarilar = []
        etiketler = set()

        sablon_yolu = getattr(self, "sablon_yolu", "")
        if not sablon_yolu:
            hatalar.append("Word rapor şablonu seçilmemiş.")
        elif not os.path.isfile(sablon_yolu):
            hatalar.append(f"Word rapor şablonu bulunamadı: {sablon_yolu}")
        else:
            try:
                etiketler = self.sablon_etiketlerini_oku(sablon_yolu)
                tanimsiz = sorted(etiketler - self.desteklenen_sablon_etiketleri())
                if tanimsiz:
                    hatalar.append("Word şablonunda tanımsız etiketler var: " + ", ".join(tanimsiz))
                if "[JEOLOJI_BOLUMU]" in etiketler:
                    jeoloji_sablonu = self.rapor_uretici().rapor_jeoloji_sablon_yolu_bul()
                    if not jeoloji_sablonu:
                        if getattr(self, "jeoloji_kutuphanesi_bolumu_aktif", False):
                            record_id = getattr(self, "jeoloji_kutuphanesi_kayit_id", None)
                            ek = f" #{record_id}" if record_id else ""
                            hatalar.append(
                                f"Seçili kütüphane kaydının{ek} JEOLOJİ Word paketi bulunamadı; "
                                "kayıt yeniden uygulanmalı veya başka bir jeoloji şablonu seçilmelidir."
                            )
                        else:
                            hatalar.append("Ana şablonda [JEOLOJI_BOLUMU] etiketi var; ancak ilçe/köy için jeoloji Word şablonu bulunamadı.")
                    else:
                        try:
                            jeoloji_doc = Document(jeoloji_sablonu)
                            if not list(self.rapor_uretici().rapor_docx_govde_elemanlari(jeoloji_doc)):
                                hatalar.append(f"Jeoloji Word şablonu boş: {jeoloji_sablonu}")
                            genel_state = getattr(self, "genel_jeoloji_verisi", {})
                            genel_mode = (
                                str(genel_state.get("kaynak_modu") or "kutuphane")
                                if isinstance(genel_state, dict) else ""
                            )
                            if (
                                isinstance(genel_state, dict)
                                and (
                                    genel_state.get("birimler")
                                    or genel_state.get("kaynak_modu") == "eski_rapor"
                                )
                                and not any(
                                    yapisal_jeoloji_basligi_mi(paragraph.text)
                                    for paragraph in jeoloji_doc.paragraphs
                                )
                            ):
                                hatalar.append(
                                    "Seçili jeoloji Word'ünde '2.1.1 Yapısal Jeoloji ve Aktif "
                                    "Tektonik' başlığı bulunamadı; başka bir kütüphane Word'ü seçin."
                                )
                            if (
                                isinstance(genel_state, dict)
                                and (
                                    genel_state.get("birimler")
                                    or genel_state.get("kaynak_modu") == "eski_rapor"
                                )
                                and not wordde_stratigrafik_kesit_var_mi(jeoloji_sablonu)
                            ):
                                hatalar.append(
                                    "Seçili jeoloji Word'ünde görseliyle birlikte aktarılabilir "
                                    "stratigrafik kesit bulunamadı; başka bir kütüphane Word'ü seçin."
                                )
                            if (
                                genel_mode == "eski_rapor"
                                and not any(
                                    bolgesel_jeoloji_basligi_mi(paragraph.text)
                                    for paragraph in jeoloji_doc.paragraphs
                                )
                            ):
                                hatalar.append(
                                    "Seçili jeoloji Word'ünde '2.1 Bölgesel Jeoloji' başlığı "
                                    "bulunamadı; başka bir kütüphane Word'ü seçin."
                                )
                        except Exception as e:
                            hatalar.append(f"Jeoloji Word şablonu okunamadı: {e}")

                    if (
                        getattr(self, "jeoloji_kutuphanesi_bolumu_aktif", False)
                        and not getattr(self, "genel_jeoloji_verisi", {}).get("birimler")
                    ):
                        current_general = (
                            self.txt_formasyon_rapor.get("1.0", "end").strip()
                            if hasattr(self, "txt_formasyon_rapor") else ""
                        )
                        applied_general = str(
                            getattr(self, "jeoloji_kutuphanesi_uygulanan_genel", "") or ""
                        ).strip()
                        if current_general != applied_general:
                            uyarilar.append(
                                "Genel jeoloji metni kütüphane kaydı uygulandıktan sonra değiştirilmiş; "
                                "çıktıda resim ve tablolarla birlikte kayıtlı tam Word bölümü kullanılacak."
                            )
            except Exception as e:
                self.hata_kaydet("Rapor Word şablonu ön kontrolü yapılamadı", e)
                hatalar.append(f"Word rapor şablonu okunamadı: {e}")

        if "[JEOLOJI_BOLUMU]" in etiketler:
            genel = getattr(self, "genel_jeoloji_verisi", {})
            if (
                isinstance(genel, dict)
                and genel.get("kaynak_modu") != "eski_rapor"
                and hasattr(self, "genel_jeoloji_eksik_metinlerini_tamamla")
            ):
                try:
                    tamamlama = self.genel_jeoloji_eksik_metinlerini_tamamla()
                    tamamlananlar = tamamlama.get("tamamlanan", []) if isinstance(tamamlama, dict) else []
                    if tamamlananlar:
                        uyarilar.append(
                            "2.1 Bölgesel Jeoloji metni kalıcı birim metin kütüphanesinden "
                            "otomatik tamamlanan birimler: " + ", ".join(tamamlananlar[:12])
                        )
                except Exception as exc:
                    self.hata_kaydet("2.1 Bölgesel Jeoloji metinleri otomatik tamamlanamadı", exc)
            if not isinstance(genel, dict) or not genel:
                hatalar.append(
                    "2.1 Bölgesel Jeoloji için parsel merkezli genel jeoloji haritası ve birim listesi hazırlanmamış."
                )
            else:
                mode = str(genel.get("kaynak_modu") or "kutuphane")
                if mode != "eski_rapor":
                    if not genel.get("birimler"):
                        hatalar.append(
                            "2.1 Bölgesel Jeoloji için kullanılacak birim listesi hazırlanmamış."
                        )
                    genel_yol = getattr(self, "img_genel_jeoloji", "")
                    if not genel_yol or not os.path.isfile(genel_yol):
                        hatalar.append(
                            "2.1 Bölgesel Jeoloji için oluşturulmuş "
                            "Genel_Jeoloji_Haritasi.jpg bulunamadı."
                        )
                    missing = [
                        str(unit.get("ad") or unit.get("kod") or "Adsız birim")
                        for unit in genel.get("birimler", [])
                        if unit.get("kullan", True)
                        if self.deger_bos_mu(unit.get("bolgesel_jeoloji_metni"))
                    ]
                    if missing:
                        hatalar.append(
                            "2.1 Bölgesel Jeoloji açıklaması eksik birimler: "
                            + ", ".join(missing[:12])
                            + ". Haritalar sekmesinde 'Genel Jeoloji Haritası ve 2.1 Hazırla' "
                            "ekranını açıp ilgili birimin 2.1 metnini girin veya bu birimi "
                            "harita/2.1 kapsamından çıkarın."
                        )
                stored_hash = str(genel.get("geometri_hash") or "")
                current_points = getattr(self, "yuklu_kml_points", [])
                current_hash = parsel_geometri_hashi(current_points) if len(current_points) >= 3 else ""
                if stored_hash and current_hash and stored_hash != current_hash:
                    hatalar.append(
                        "Genel jeoloji haritası yüklü parsel KML'sinden daha eski; haritayı ve 2.1'i yeniden hazırlayın."
                    )

        for kod, etiket in ZORUNLU_PROJE_ALANLARI:
            if self.deger_bos_mu(self.proje_deger(kod, "")):
                hatalar.append(f"{etiket} [{kod}] alanı boş.")

        taahhut_bilgileri = getattr(self, "taahhut_bilgileri", {})
        if not isinstance(taahhut_bilgileri, dict):
            taahhut_bilgileri = {}
        for kod in RAPOR_TAAHHUT_BILGI_ALANLARI:
            tag = f"[{kod}]"
            if tag in etiketler and self.deger_bos_mu(taahhut_bilgileri.get(kod, "")):
                hatalar.append(f"Word şablonundaki {tag} etiketi için mühendis bilgisi boş.")

        ac_yn_kayitlari = self.ac_yn_sekme_kayitlari() if hasattr(self, "ac_yn_sekme_kayitlari") else []
        if not ac_yn_kayitlari:
            uyarilar.append("AÇ/YN sekmesinde kayıt bulunmuyor.")
        else:
            bos_kayitlar = [
                kayit.get("isim", "İsimsiz")
                for kayit in ac_yn_kayitlari
                if not self.anlamli_satir_var_mi(self.ac_yn_satirlari(kayit))
            ]
            if bos_kayitlar:
                uyarilar.append("AÇ/YN satırı boş görünen kayıtlar: " + ", ".join(bos_kayitlar[:5]))

            rapor_uretici = self.rapor_uretici()
            for kayit in ac_yn_kayitlari:
                for satir_no, satir in enumerate(self.ac_yn_satirlari(kayit), start=1):
                    derinlik = satir[0] if satir else ""
                    if self.deger_bos_mu(derinlik):
                        continue
                    try:
                        rapor_uretici.rapor_derinlik_araligi_coz(derinlik)
                    except ValueError as e:
                        hatalar.append(
                            f"{kayit.get('isim', 'İsimsiz')} kaydının {satir_no}. satırında geçersiz derinlik aralığı: {e}"
                        )

        lab_ac = self.lab_ac_satirlari_al() if hasattr(self, "lab_ac_satirlari_al") else []
        lab_yn = self.lab_yn_satirlari_al() if hasattr(self, "lab_yn_satirlari_al") else []
        if not self.anlamli_satir_var_mi(lab_ac) and not self.anlamli_satir_var_mi(lab_yn):
            uyarilar.append("Laboratuvar verisi bulunmuyor.")

        jeofizik_yolu = self.jeofizik_excel_yolu_al() if hasattr(self, "jeofizik_excel_yolu_al") else ""
        if etiketler & JEOFIZIK_VERI_ETIKETLERI:
            if self.deger_bos_mu(jeofizik_yolu):
                hatalar.append("Şablon jeofizik tabloları istiyor; parametre Excel/CSV dosyası seçilmemiş.")
            elif not os.path.isfile(jeofizik_yolu):
                hatalar.append(f"Jeofizik parametre dosyası bulunamadı: {jeofizik_yolu}")
            else:
                try:
                    self.rapor_uretici().rapor_jeofizik_parametrelerini_oku()
                except Exception as e:
                    hatalar.append(f"Jeofizik parametre dosyası geçersiz: {e}")

        bulunan_rapor_haritalari = proje_klasorundeki_rapor_haritalarini_bul(
            getattr(self, "guncel_dosya_yolu", "")
        )
        pga_caption_var = False
        pga_caption_key = "canakkale_bolgesi_deprem_tehlike_haritasi"
        if sablon_yolu and os.path.isfile(sablon_yolu):
            try:
                sablon_document = Document(sablon_yolu)
                rapor_uretici = self.rapor_uretici()
                pga_caption_var = any(
                    pga_caption_key in rapor_uretici.rapor_metin_normalize(paragraph.text)
                    for paragraph in rapor_uretici.docx_paragraflarini_dolas(sablon_document)
                )
            except Exception:
                pga_caption_var = False
        if pga_caption_var:
            pga_yolu = getattr(self, "img_pga_haritasi", "") or ""
            if not os.path.isfile(pga_yolu):
                pga_yolu = bulunan_rapor_haritalari.get("img_pga_haritasi", "")
                if pga_yolu:
                    setattr(self.app, "img_pga_haritasi", pga_yolu)
            if not pga_yolu or not os.path.isfile(pga_yolu):
                uyarilar.append(
                    "Çanakkale bölgesi deprem tehlike haritası için PGA görseli bulunamadı; "
                    "rapor üretimi durmayacak ve bu konum uyarıyla bırakılacak."
                )
            else:
                try:
                    with Image.open(pga_yolu) as img:
                        img.verify()
                except Exception as exc:
                    uyarilar.append(f"PGA haritası görseli okunamadı; rapor üretimi durmayacak: {exc}")
        for tag, alan in GORSEL_ETIKETLERI.items():
            if tag not in etiketler:
                continue
            yol = getattr(self, alan, "")
            if (not yol or not os.path.isfile(yol)) and alan != "img_parsel_haritasi":
                yol = bulunan_rapor_haritalari.get(alan, "")
                if yol:
                    setattr(self.app, alan, yol)
            if not yol or not os.path.isfile(yol):
                hatalar.append(f"{tag} için hazırlanmış görsel bulunamadı.")
                continue
            try:
                with Image.open(yol) as img:
                    img.verify()
            except Exception as e:
                hatalar.append(f"{tag} görseli okunamadı: {e}")
                continue
            if tag == "[PARSEL_HARITASI]":
                stored_hash = str(getattr(self, "parsel_haritasi_geometri_hash", "") or "")
                try:
                    current_hash = self.cizim_uretici().parsel_geometri_hashi()
                except Exception:
                    current_hash = ""
                if stored_hash and current_hash and stored_hash != current_hash:
                    hatalar.append(
                        "[PARSEL_HARITASI] yüklü KML'den daha eski; parsel haritasını yeniden hazırlayın."
                    )
                stored_ada = str(getattr(self, "parsel_haritasi_ada", "") or "").strip()
                stored_parsel = str(getattr(self, "parsel_haritasi_parsel", "") or "").strip()
                current_ada = str(self.proje_deger("ADA", "") or "").strip()
                current_parsel = str(self.proje_deger("PARSEL", "") or "").strip()
                if (stored_ada, stored_parsel) != (current_ada, current_parsel):
                    hatalar.append(
                        "[PARSEL_HARITASI] başka ada/parsel bilgisiyle hazırlanmış; görüntüyü yeniden oluşturun."
                    )

        if "[JEOFIZIK_KOORDINAT]" in etiketler:
            koordinatlar = self.jeofizik_koordinatlari_al() if hasattr(self, "jeofizik_koordinatlari_al") else []
            if not koordinatlar:
                hatalar.append("Şablon jeofizik koordinat tablosu istiyor; koordinat kaydı yok.")
        if "[JEOLOJI_KOORDINAT]" in etiketler:
            jeoloji_noktalari = [
                d for d in getattr(self, "harita_isaretleri", {}).values()
                if d.get("tip") in ("AÇ", "YN")
            ]
            if not jeoloji_noktalari:
                hatalar.append("Şablon jeoloji koordinat tablosu istiyor; AÇ/YN harita noktası yok.")

        if etiketler & {"[KESIT]", "[AC_LOGLARI]"} and not ac_yn_kayitlari:
            hatalar.append("Şablon kesit/AÇ logu istiyor; AÇ/YN kaydı bulunmuyor.")
        if "[LAB]" in etiketler and not (
            self.anlamli_satir_var_mi(lab_ac) or self.anlamli_satir_var_mi(lab_yn)
        ):
            hatalar.append("Şablon laboratuvar tablosu istiyor; laboratuvar verisi bulunmuyor.")

        if etiketler & TASIMA_RAPOR_ETIKETLERI:
            if hasattr(self, "tasima_raporu_guncel_mi"):
                try:
                    if not self.tasima_raporu_guncel_mi():
                        hatalar.append("Taşıma gücü girdileri değişmiş; taşıma rapor metnini yeniden oluşturun.")
                except Exception as e:
                    hatalar.append(f"Taşıma gücü rapor güncelliği doğrulanamadı: {e}")
            tg_girdiler = getattr(self, "tg_girdiler", {})
            tasima_turu = ""
            if getattr(self, "zemin_kaya_var", None) is not None:
                tasima_turu = self.zemin_kaya_var.get()
            gerekli_girdiler = (
                dict(KAYA_TASIMA_GIRDI_ETIKETLERI)
                if tasima_turu == "kaya"
                else dict(TASIMA_GIRDI_ETIKETLERI)
            )
            if tasima_turu != "kaya":
                yass_variable = getattr(self, "tasima_yass_var", None)
                if yass_variable is not None and not bool(yass_variable.get()):
                    gerekli_girdiler.pop("yass", None)
            tasima_girdi_hatasi = False
            for anahtar, etiket in gerekli_girdiler.items():
                entry = tg_girdiler.get(anahtar)
                ham = str(entry.get()).strip() if entry is not None else ""
                if not ham:
                    hatalar.append(f"Taşıma gücü için {etiket} [{anahtar}] alanı boş.")
                    tasima_girdi_hatasi = True
                    continue
                try:
                    float(ham.replace(",", "."))
                except ValueError:
                    hatalar.append(f"Taşıma gücü için {etiket} [{anahtar}] sayısal değil.")
                    tasima_girdi_hatasi = True
            if not tasima_girdi_hatasi and tasima_turu != "kaya":
                try:
                    self.rapor_uretici().rapor_tasima_gucu_tablosu_olustur(Document())
                except Exception as e:
                    hatalar.append(f"Taşıma gücü girdileri geçersiz: {e}")
            if "[QK]" in etiketler:
                try:
                    qk = float(str(getattr(self, "son_qk", "")).replace(",", "."))
                    if not qk > 0:
                        raise ValueError
                except (TypeError, ValueError):
                    hatalar.append("Şablondaki [QK] için doğrulanmış pozitif değer bulunamadı.")
            for tag, alan in (("[QT]", "entry_qt_nihai"), ("[KS]", "entry_ks_nihai")):
                if tag not in etiketler:
                    continue
                entry = getattr(self, alan, None)
                try:
                    deger = float(str(entry.get()).replace(",", ".")) if entry is not None else 0.0
                    if not deger > 0:
                        raise ValueError
                except (TypeError, ValueError):
                    hatalar.append(f"Şablondaki {tag} için doğrulanmış pozitif değer bulunamadı.")
            if "[DF]" in etiketler:
                entry = tg_girdiler.get("Df")
                try:
                    df = float(str(entry.get()).replace(",", ".")) if entry is not None else -1.0
                    if df < 0:
                        raise ValueError
                except (TypeError, ValueError):
                    hatalar.append("Şablondaki [DF] için doğrulanmış negatif olmayan değer bulunamadı.")
            if "[TASIMA_GUCU]" in etiketler:
                tasima_metni = ""
                if hasattr(self, "txt_tasima_rapor"):
                    tasima_metni = self.txt_tasima_rapor.get("1.0", "end").strip()
                if not tasima_metni:
                    hatalar.append("Taşıma gücü rapor metni oluşturulmamış.")

        denetim = ekleri_denetle(
            getattr(self, "ekler", {}),
            getattr(self, "ek_kategorileri", []),
            derin=True,
        )
        if denetim["bos_kategoriler"]:
            uyarilar.append("Boş EKLER kategorileri: " + ", ".join(denetim["bos_kategoriler"]))
        if denetim["eksik_dosyalar"]:
            uyarilar.append(f"EKLER içinde bulunamayan dosya sayısı: {len(denetim['eksik_dosyalar'])}")
        if denetim["donusum_gerekenler"]:
            uyarilar.append(f"EKLER içinde PDF'e çevrilmesi gereken dosya sayısı: {len(denetim['donusum_gerekenler'])}")
        if denetim.get("gecersiz_dosyalar"):
            uyarilar.append(f"EKLER içinde bozuk/okunamayan dosya sayısı: {len(denetim['gecersiz_dosyalar'])}")

        tdth = normalize_tdth(getattr(self, "tdth_verisi", {}))
        if tdth.get("durum") == "eksik":
            uyarilar.append("TDTH PDF seçilmemiş; ek dosya olmadan da rapor üretilebilir.")
        elif tdth.get("durum") == "yenilenmeli":
            uyarilar.append(
                "Zemin sınıfı değiştiği için TDTH PDF yenilenmeli; mevcut ek olmadan da rapor üretilebilir."
            )
        elif tdth.get("uyarilar"):
            uyarilar.extend(f"TDTH: {uyari}" for uyari in tdth["uyarilar"])

        return {
            "hatalar": hatalar,
            "uyarilar": uyarilar,
            "gecerli": not hatalar,
        }

    def mesaj_hazirla(self, sonuc):
        satirlar = []
        if sonuc["hatalar"]:
            satirlar.append("Rapor üretimini durduran eksikler:")
            satirlar.extend(f"- {hata}" for hata in sonuc["hatalar"])
        if sonuc["uyarilar"]:
            if satirlar:
                satirlar.append("")
            satirlar.append("Kontrol edilmesi önerilen uyarılar:")
            satirlar.extend(f"- {uyari}" for uyari in sonuc["uyarilar"])
        if not satirlar:
            satirlar.append("Rapor üretimi için temel kontroller uygun görünüyor.")
        return "\n".join(satirlar)

    def calistir(self, devam_sor=True):
        sonuc = self.sonuclari_hazirla()
        mesaj = self.mesaj_hazirla(sonuc)

        if sonuc["hatalar"]:
            messagebox.showerror("Rapor Ön Kontrolü", mesaj)
            if hasattr(self, "durum_mesaji_yaz"):
                self.durum_mesaji_yaz("Rapor ön kontrolünde kritik eksik var")
            return False

        if sonuc["uyarilar"]:
            if hasattr(self, "durum_mesaji_yaz"):
                self.durum_mesaji_yaz("Rapor ön kontrolünde uyarılar var")
            if not devam_sor:
                messagebox.showwarning("Rapor Ön Kontrolü", mesaj)
                return False
            return messagebox.askyesno(
                "Rapor Ön Kontrolü",
                mesaj + "\n\nBu uyarılara rağmen rapor üretimine devam edilsin mi?",
            )

        messagebox.showinfo("Rapor Ön Kontrolü", mesaj)
        if hasattr(self, "durum_mesaji_yaz"):
            self.durum_mesaji_yaz("Rapor ön kontrolü tamamlandı")
        return True
