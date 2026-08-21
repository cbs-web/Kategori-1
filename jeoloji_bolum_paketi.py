"""JEOLOJİ bölümünü kayıpsız mini DOCX olarak ayırma ve yerleştirme.

Bu modül metni yeniden kurmaz. Kaynak belgedeki ``w:p`` ve ``w:tbl``
bloklarını, DrawingML nesnelerini ve bunların ilişkilerini koruyarak taşır.
Bu sayede inline/anchored görseller, tablolar ve doğrudan biçimlendirme
düz metin tabanlı bir aktarımda olduğu gibi kaybolmaz.
"""

from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO
import inspect
import os
from pathlib import Path
import re
import tempfile
import unicodedata

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.image import Image as _DocxImage
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shape import CT_Inline
from docx.shape import InlineShape
from docx.shared import Cm
from lxml import etree

try:  # Tercih edilen birleştirici; kurulu değilse aşağıdaki OOXML yolu kullanılır.
    from docxcompose.composer import Composer as _Composer
except ImportError:  # pragma: no cover - geliştirme ortamında docxcompose kurulu değil.
    _Composer = None


VARSAYILAN_YER_TUTUCU = "[JEOLOJI_BOLUMU]"

_W_VAL = qn("w:val")
_W_STYLE_ID = qn("w:styleId")
_R_ILISKI_NITELIKLERI = (qn("r:id"), qn("r:embed"), qn("r:link"))
_BOLUM_ICERIK_ILISKI_TURLERI = {
    RT.A_F_CHUNK,
    RT.CHART,
    RT.CHARTSHEET,
    RT.CHART_USER_SHAPES,
    RT.HYPERLINK,
    RT.IMAGE,
    RT.OLE_OBJECT,
    RT.PACKAGE,
}
_STIL_REFERANS_ETIKETLERI = {
    qn("w:pStyle"),
    qn("w:rStyle"),
    qn("w:tblStyle"),
}
_STIL_BAGIMLILIK_ETIKETLERI = {
    qn("w:basedOn"),
    qn("w:next"),
    qn("w:link"),
}

_JEOLOJI_ALT_BASLIKLARI = (
    "bolgesel jeoloji",
    "stratigraf",
    "yapisal jeoloji",
    "aktif tektonik",
    "inceleme alani jeolojisi",
    "calisma alani jeolojisi",
    "formasyonu",
    "volkaniti",
    "ignimbiriti",
    "melanji",
    "uyesi",
)

_BILINEN_ANA_BASLIKLAR = (
    "hidrojeoloji",
    "arazi calismalari",
    "laboratuvar calismalari",
    "sondaj calismalari",
    "jeofizik calismalar",
    "jeoteknik",
    "zemin ve temel",
    "dogal afet",
    "yerlesime uygunluk",
    "sonuc ve oneriler",
    "sonuclar ve oneriler",
)

_BASLIK_OLMAYAN_ON_EKLER = (
    "sekil",
    "tablo",
    "cizelge",
    "fotograf",
    "resim",
    "harita",
    "levha",
)

_BASLIK_KUCUK_KELIMELERI = {
    "ve",
    "veya",
    "ile",
    "icin",
    "ait",
    "dair",
    "hakkinda",
}


class JeolojiBolumPaketiHatasi(ValueError):
    """Bölüm ayırma veya yerleştirme kuralları sağlanmadığında oluşur."""


@dataclass(frozen=True)
class JeolojiBolumAyirmaSonucu:
    kaynak_docx: str
    paket_docx: str
    baslik_bulundu: bool
    blok_sayisi: int
    baslangic_blok_no: int
    bitis_blok_no: int


@dataclass(frozen=True)
class JeolojiBolumYerlestirmeSonucu:
    paket_docx: str
    hedef_docx: str
    cikti_docx: str
    blok_sayisi: int
    motor: str


@dataclass(frozen=True)
class StratigrafikKesitAyirmaSonucu:
    kaynak_docx: str
    paket_docx: str
    blok_sayisi: int
    baslangic_blok_no: int
    bitis_blok_no: int
    sekil_no: str


@dataclass(frozen=True)
class _StratigrafikKesitSecimi:
    baslangic_blok_no: int
    bitis_blok_no: int
    sekil_no: str
    bolgesel_bitis_blok_no: int
    caption_blok_no: int
    gorsel_blok_no: int
    gorsel_blob: bytes
    gorsel_uzantisi: str
    gorsel_turu: str
    gorsel_adayi: object


@dataclass(frozen=True)
class _StratigrafikGorselAdayi:
    """Kesit adayı olan tek bir DrawingML raster ilişkisi."""

    blok_no: int
    tasiyici: str
    part: object
    drawing: object
    blip: object
    behind_doc: str
    relative_from: str
    position_offset: int | None


# ``word_numaralandirma`` sıfır yer tutucusunu SEQ alanıyla değiştirir; nokta
# alan sonucunun arkasında kalacağı için nihai metin ``Şekil X. ...`` olur.
STRATIGRAFIK_KESIT_CAPTION = (
    "Şekil 0. Çalışma alanı ve yakın çevresinin stratigrafik kesiti"
)

_INLINE_GORSEL_UZANTILARI = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/jpg": ".jpg",
    "image/gif": ".gif",
    "image/bmp": ".bmp",
    "image/tiff": ".tif",
}
_INLINE_GORSEL_DOSYA_UZANTILARI = set(_INLINE_GORSEL_UZANTILARI.values()) | {
    ".jpeg",
    ".tiff",
}


def _anahtar(value: object) -> str:
    text = unicodedata.normalize("NFKD", str(value or "").strip().casefold())
    text = "".join(char for char in text if not unicodedata.combining(char))
    text = text.replace("ı", "i")
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return " ".join(text.split())


def _paragraf_metni(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def _ana_jeoloji_basligi_mi(text: str) -> bool:
    # "Hidrojeoloji" ve "Bölgesel Jeoloji" bu kalıba girmez; başlık yalnızca
    # numara + JEOLOJİ biçiminde olmalıdır.
    return bool(re.fullmatch(r"(?:\d+\s+)*jeoloji", _anahtar(text)))


def _numarali_baslik_seviyesi(text: str) -> int | None:
    # Bölüm numaraları kısa olur. ``100000 ölçekli ...`` gibi jeoloji
    # açıklamaları bölüm numarası kabul edilmemelidir.
    match = re.match(r"^\s*(\d{1,3}(?:\.\d{1,3})*)[.)]?\s+", text or "")
    if not match:
        return None
    return match.group(1).count(".") + 1


def _baslik_metni_gibi_mi(text: str) -> bool:
    """Yanlış başlık stili verilmiş uzun gövde paragraflarını ayıkla."""
    value = " ".join(str(text or "").split())
    if not value or len(value) > 140 or len(value.split()) > 14:
        return False
    # Rapor ana başlıkları normalde tam cümle sonlandırıcısıyla bitmez.
    return re.search(r"[.!?;]\s*$", value) is None


def _baslik_yazimi_gibi_mi(text: str) -> bool:
    words = re.findall(r"[A-Za-zÇĞİÖŞÜçğıöşü]+", str(text or ""))
    significant = [word for word in words if _anahtar(word) not in _BASLIK_KUCUK_KELIMELERI]
    if not significant:
        return False
    initial_upper = sum(word[0].isupper() for word in significant)
    return initial_upper / len(significant) >= 0.65


def _stil_baslik_seviyesi(document, paragraph_element) -> int | None:
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is None:
        return None

    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is not None:
        try:
            return int(outline.get(_W_VAL)) + 1
        except (TypeError, ValueError):
            pass

    p_style = p_pr.find(qn("w:pStyle"))
    style_id = p_style.get(_W_VAL) if p_style is not None else ""
    visited: set[str] = set()
    while style_id and style_id not in visited:
        visited.add(style_id)
        match = re.search(r"(?:heading|baslik)\s*(\d+)", _anahtar(style_id))
        if match:
            return int(match.group(1))
        style = _stil_elementi(document, style_id)
        if style is None:
            break
        name = style.find(qn("w:name"))
        match = re.search(
            r"(?:heading|baslik)\s*(\d+)",
            _anahtar(name.get(_W_VAL) if name is not None else ""),
        )
        if match:
            return int(match.group(1))
        style_outline = style.find("./" + qn("w:pPr") + "/" + qn("w:outlineLvl"))
        if style_outline is not None:
            try:
                return int(style_outline.get(_W_VAL)) + 1
            except (TypeError, ValueError):
                pass
        based_on = style.find(qn("w:basedOn"))
        style_id = based_on.get(_W_VAL) if based_on is not None else ""
    return None


def _stil_elementi(document, style_id: str):
    for style in document.styles.element.findall(qn("w:style")):
        if style.get(_W_STYLE_ID) == style_id:
            return style
    return None


def _baslik_seviyesi(document, paragraph_element) -> int | None:
    return (
        _stil_baslik_seviyesi(document, paragraph_element)
        or _numarali_baslik_seviyesi(_paragraf_metni(paragraph_element))
    )


def _toc_paragrafi_mi(document, paragraph_element) -> bool:
    p_pr = paragraph_element.find(qn("w:pPr"))
    p_style = p_pr.find(qn("w:pStyle")) if p_pr is not None else None
    style_id = p_style.get(_W_VAL) if p_style is not None else ""
    if _anahtar(style_id).startswith(("toc", "icindekiler")):
        return True
    return any("TOC" in (node.text or "").upper() for node in paragraph_element.iter(qn("w:instrText")))


def _jeoloji_alt_basligi_mi(text: str) -> bool:
    key = _anahtar(text)
    return any(marker in key for marker in _JEOLOJI_ALT_BASLIKLARI)


def _dogrudan_baslik_bicimi_mi(paragraph_element) -> bool:
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is not None and (
        p_pr.find(qn("w:keepNext")) is not None
        or p_pr.find(qn("w:pageBreakBefore")) is not None
    ):
        return True
    text_runs = [
        run
        for run in paragraph_element.findall(qn("w:r"))
        if any((node.text or "").strip() for node in run.iter(qn("w:t")))
    ]
    if not text_runs:
        return False
    bold_runs = sum(
        1
        for run in text_runs
        if run.find("./" + qn("w:rPr") + "/" + qn("w:b")) is not None
    )
    return bold_runs == len(text_runs)


def jeoloji_sinir_basligi_mi(
    text: str,
    ana_seviye: int,
    *,
    style_level: int | None = None,
    dogrudan_baslik_bicimi: bool = False,
    jeoloji_alt_basligi: bool = False,
) -> bool:
    """Bir paragrafın ana JEOLOJİ bölümünü gerçekten bitirip bitirmediğini sınar."""
    if not text or jeoloji_alt_basligi or _jeoloji_alt_basligi_mi(text):
        return False

    key = _anahtar(text)
    plain = re.sub(r"^(?:\d+\s+)+", "", key)
    if any(
        plain == prefix or plain.startswith(f"{prefix} ")
        for prefix in _BASLIK_OLMAYAN_ON_EKLER
    ):
        return False
    if not _baslik_metni_gibi_mi(text):
        return False

    numbered_level = _numarali_baslik_seviyesi(text)
    letters = [char for char in text if char.isalpha()]
    upper_ratio = (
        sum(char.isupper() for char in letters) / len(letters)
        if letters
        else 0.0
    )
    strong_appearance = upper_ratio >= 0.72 or _baslik_yazimi_gibi_mi(text)
    known_major = any(plain.startswith(marker) for marker in _BILINEN_ANA_BASLIKLAR)
    if known_major and (
        numbered_level is not None
        or style_level is not None
        or strong_appearance
        or dogrudan_baslik_bicimi
    ):
        return True

    if numbered_level is not None and numbered_level <= ana_seviye:
        return True

    return bool(
        style_level is not None
        and style_level <= ana_seviye
        and strong_appearance
    )


def _sonraki_ana_baslik_mi(document, paragraph_element, ana_seviye: int) -> bool:
    text = _paragraf_metni(paragraph_element)
    return bool(
        jeoloji_sinir_basligi_mi(
            text,
            ana_seviye,
            style_level=_stil_baslik_seviyesi(document, paragraph_element),
            dogrudan_baslik_bicimi=_dogrudan_baslik_bicimi_mi(paragraph_element),
        )
    )


def _govde_bloklari(document) -> list:
    return [child for child in document.element.body.iterchildren() if child.tag != qn("w:sectPr")]


def _blok_gorsel_iceriyor(block) -> bool:
    gorsel_etiketleri = {qn("w:drawing"), qn("w:pict"), qn("w:object")}
    return any(node.tag in gorsel_etiketleri for node in block.iter())


def _blok_inline_gorsel_iceriyor(block) -> bool:
    """Body bloğunda Word'ün akış içi görseli var mı?"""
    return any(node.tag == qn("wp:inline") for node in block.iter())


def _blok_desteklenen_gorsel_tasiyicisi_var_mi(block) -> bool:
    """Body bloğunda işleyebildiğimiz inline veya floating görsel var mı?"""
    return any(
        node.tag in {qn("wp:inline"), qn("wp:anchor")}
        for node in block.iter()
    )


def _cizim_gorsel_adaylarini_bul(document, drawing, tasiyici: str, blok_no: int):
    unsupported_tags = {
        qn("c:chart"),
        qn("w:oleObject"),
        qn("w:pict"),
        qn("w:object"),
    }
    if any(node.tag in unsupported_tags for node in drawing.iter()):
        raise JeolojiBolumPaketiHatasi(
            "Stratigrafik kesit çiziminde chart/OLE veya desteklenmeyen Word "
            "nesnesi bulundu; güvenli görsel seçilemedi."
        )

    blips = list(drawing.iter(qn("a:blip")))
    if not blips:
        raise JeolojiBolumPaketiHatasi(
            "Stratigrafik kesit çiziminde görsel blob'u bulunamadı; kaynak Word "
            "güvenle dönüştürülemedi."
        )

    candidates = []
    for blip in blips:
        relation_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        relationship = (
            document.part.rels.get(relation_id)
            if relation_id
            else None
        )
        relation = (
            document.part.related_parts.get(relation_id)
            if relation_id
            else None
        )
        if (
            relationship is None
            or relationship.reltype != RT.IMAGE
            or relation is None
        ):
            raise JeolojiBolumPaketiHatasi(
                "Stratigrafik kesit çiziminin özgün görsel ilişkisi okunamadı; "
                "kaynak Word güvenle dönüştürülemedi."
            )

        behind_doc = ""
        relative_from = ""
        position_offset = None
        if tasiyici == "anchor":
            behind_doc = str(
                drawing.get("behindDoc")
                or drawing.get(qn("wp:behindDoc"))
                or ""
            ).strip().casefold()
            position_v = drawing.find(qn("wp:positionV"))
            if position_v is not None:
                relative_from = str(
                    position_v.get("relativeFrom")
                    or position_v.get(qn("wp:relativeFrom"))
                    or ""
                ).strip().casefold()
                position_node = position_v.find(qn("wp:posOffset"))
                if position_node is not None:
                    try:
                        position_offset = int((position_node.text or "").strip())
                    except (TypeError, ValueError):
                        position_offset = None
        candidates.append(
            _StratigrafikGorselAdayi(
                blok_no=blok_no,
                tasiyici=tasiyici,
                part=relation,
                drawing=drawing,
                blip=blip,
                behind_doc=behind_doc,
                relative_from=relative_from,
                position_offset=position_offset,
            )
        )
    return candidates


def _blok_gorsel_adaylarini_bul(
    document,
    block,
    blok_no: int = -1,
    *,
    tasiyicilar: tuple[str, ...] = ("inline", "anchor"),
    strict: bool = False,
):
    """Body bloğundaki inline/anchor raster ilişkilerini metadata ile döndür."""
    candidates = []
    if "inline" in tasiyicilar:
        for inline in block.iter(qn("wp:inline")):
            candidates.extend(
                _cizim_gorsel_adaylarini_bul(document, inline, "inline", blok_no)
            )
    if "anchor" in tasiyicilar:
        for anchor in block.iter(qn("wp:anchor")):
            candidates.extend(
                _cizim_gorsel_adaylarini_bul(document, anchor, "anchor", blok_no)
            )

    if strict and not candidates and any(
        node.tag in {qn("w:pict"), qn("w:object")} for node in block.iter()
    ):
        raise JeolojiBolumPaketiHatasi(
            "Stratigrafik kesit paragrafında desteklenmeyen VML/OLE Word nesnesi "
            "bulundu; güvenli raster görsel seçilemedi."
        )
    return candidates


def _blok_gorsel_parcalarini_bul(document, block) -> list:
    """Bir body bloğundaki doğrudan inline image ilişkilerini döndür.

    Bu yardımcı yalnız inline şekilleri döndürür. Floating şekillerin seçim
    önceliği, caption bağlamındaki anchor metadata'sını da değerlendiren
    ``_blok_gorsel_adaylarini_bul`` katmanında yapılır.
    """
    return [
        candidate.part
        for candidate in _blok_gorsel_adaylarini_bul(
            document,
            block,
            tasiyicilar=("inline",),
        )
    ]


def _gorsel_uzantisi_ve_turu(part) -> tuple[str, str]:
    content_type = str(getattr(part, "content_type", "") or "").casefold()
    extension = _INLINE_GORSEL_UZANTILARI.get(content_type)
    if extension is None:
        extension = Path(str(getattr(part, "partname", ""))).suffix.casefold()
        if extension not in _INLINE_GORSEL_DOSYA_UZANTILARI:
            raise JeolojiBolumPaketiHatasi(
                "Stratigrafik kesit görselinin formatı Word InlineShape için "
                f"desteklenmiyor: {content_type or extension or 'bilinmiyor'}."
            )
    return extension, content_type or extension


def _stratigrafik_anchor_onceligi(candidate: _StratigrafikGorselAdayi) -> int:
    """Anchor'ın foreground/paragraph-relative güvenilirlik düzeyini döndür."""
    if candidate.tasiyici == "inline":
        return 3
    if candidate.behind_doc in {"1", "true", "yes"}:
        # Sayfa arkasında kalan eski harita gibi floating nesneler kesit adayı
        # değildir. Caption paragrafına bağlı olsalar bile gerçek kesit seçimini
        # gölgelememelidir.
        return 0
    if candidate.relative_from == "paragraph":
        return 3
    if candidate.relative_from:
        return 2
    return 1


def _stratigrafik_blok_gorselini_sec(candidates):
    """Tek bir paragraftaki image adayını belirsizlikte durarak seç."""
    if not candidates:
        return None
    scored = [
        ( _stratigrafik_anchor_onceligi(candidate), candidate)
        for candidate in candidates
    ]
    scored = [item for item in scored if item[0] > 0]
    if not scored:
        return None
    best_score = max(item[0] for item in scored)
    best = [item[1] for item in scored if item[0] == best_score]
    if len(best) > 1:
        raise JeolojiBolumPaketiHatasi(
            "Stratigrafik kesit adayı paragrafında birden fazla görsel eşdeğer "
            "bulundu; yanlış görsel seçilmemesi için aktarım durduruldu."
        )
    return best[0]


def _stratigrafik_sonraki_ayri_gorsel_adayini_bul(
    document,
    blocks,
    caption_index: int,
    regional_end: int,
):
    """Caption sonrasındaki ilk ayrı image-bearing paragrafı değerlendir."""
    for index in range(caption_index + 1, regional_end):
        block = blocks[index]
        if not _blok_gorsel_iceriyor(block):
            continue
        candidates = _blok_gorsel_adaylarini_bul(
            document,
            block,
            blok_no=index,
            strict=True,
        )
        selected = _stratigrafik_blok_gorselini_sec(candidates)
        if selected is not None:
            return selected
        # Yalnız behindDoc=1 floating nesne varsa bu eski/sarkan görüntüdür;
        # bir sonraki gerçek raster paragrafına kadar aramaya devam et.
    return None


def _stratigrafik_caption_gorsel_adayini_bul(document, block, blok_no: int):
    candidates = _blok_gorsel_adaylarini_bul(
        document,
        block,
        blok_no=blok_no,
        strict=True,
    )
    return _stratigrafik_blok_gorselini_sec(candidates)


def _stratigrafik_onceki_referans_mesafesi(
    blocks,
    regional_start: int,
    caption_index: int,
    figure_number: str,
):
    if not figure_number:
        return None
    pattern = re.compile(
        rf"\b(?:sekil|resim|levha)\s+{re.escape(figure_number)}\b"
    )
    for index in range(caption_index - 1, regional_start - 1, -1):
        key = _anahtar(_paragraf_metni(blocks[index]))
        if "stratigraf" in key and pattern.search(key):
            return caption_index - index
    return None


def _stratigrafik_onceki_referans_numarasi(
    blocks,
    regional_start: int,
    caption_index: int,
):
    pattern = re.compile(r"\b(?:sekil|resim|levha)\s+(\d+)\b")
    for index in range(caption_index - 1, regional_start - 1, -1):
        key = _anahtar(_paragraf_metni(blocks[index]))
        if "stratigraf" not in key:
            continue
        match = pattern.search(key)
        if match:
            return match.group(1)
    return ""


def _stratigrafik_referans_cumlesi_mi(key: str) -> bool:
    """Şekil numarasıyla başlayan düz cümleyi tam caption'dan ayır."""
    return bool(
        re.search(
            r"\b(?:ver|goster|sunul|belirt|bak|yer\s+isareti)\w*\b",
            key,
        )
    )


def _normal_stil_id(document) -> str:
    """Belgedeki Normal paragraf stilinin gerçek XML kimliğini döndür."""
    for style in document.styles.element.findall(qn("w:style")):
        style_id = style.get(_W_STYLE_ID) or ""
        name = style.find(qn("w:name"))
        style_name = name.get(_W_VAL) if name is not None else ""
        if style_id.casefold() == "normal" or _anahtar(style_name) == "normal":
            return style_id or "Normal"
    return "Normal"


def _paragrafi_normal_stile_getir(document, paragraph_element) -> bool:
    """Heading 3'ten taşınan doğrudan paragraf stilini Normal yap.

    ``w:numPr`` korunur; böylece kaynakta numaralandırılmış bir şekil yazısı
    veya başlık varsa numarası kaybolmaz. Yalnızca istemeden taşınan Heading 3
    seviyesi ve buna ait outline bilgisi temizlenir.
    """
    if _stil_baslik_seviyesi(document, paragraph_element) != 3:
        return False
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        paragraph_element.insert(0, p_pr)
    p_style = p_pr.find(qn("w:pStyle"))
    normal_id = _normal_stil_id(document)
    if p_style is None:
        p_style = OxmlElement("w:pStyle")
        p_pr.insert(0, p_style)
    changed = p_style.get(_W_VAL) != normal_id
    p_style.set(_W_VAL, normal_id)
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is not None:
        p_pr.remove(outline)
        changed = True
    return changed


def _stratigrafik_kesit_stillerini_normalize_et(document, start, end) -> int:
    blocks = _govde_bloklari(document)
    changed = 0
    for block in blocks[start:end]:
        if block.tag == qn("w:p") and _paragrafi_normal_stile_getir(document, block):
            changed += 1
    return changed


def _bolgesel_jeoloji_araligi(document) -> tuple[int, int]:
    """2.1 gövdesinin body blok aralığını, bozuk eski başlıklara da toleranslı bul."""
    blocks = _govde_bloklari(document)
    start = 0
    for index, block in enumerate(blocks):
        if block.tag != qn("w:p"):
            continue
        key = _anahtar(_paragraf_metni(block))
        if "bolgesel jeoloji" in key and len(key.split()) <= 8:
            start = index + 1
            break

    end = len(blocks)
    for index in range(start, len(blocks)):
        block = blocks[index]
        if block.tag != qn("w:p"):
            continue
        key = _anahtar(_paragraf_metni(block))
        if "yapisal jeoloji" in key or re.match(r"^2 1 [1-9]\d*\b", key):
            end = index
            break
        if re.match(r"^2 [2-9]\d*\b", key):
            end = index
            break
    return start, end


def _stratigrafik_sekil_numaralari(blocks: list) -> set[str]:
    numbers: set[str] = set()
    for block in blocks:
        key = _anahtar(_paragraf_metni(block))
        if "stratigraf" not in key:
            continue
        patterns = (
            r"stratigraf(?:ik)?(?: [a-z0-9]+){0,10} sekil (\d+)",
            r"sekil (\d+)(?: [a-z0-9]+){0,14} stratigraf",
        )
        for pattern in patterns:
            numbers.update(re.findall(pattern, key))
    return numbers


def _stratigrafik_kesit_secim_detaylari(document):
    """Kesit bloğunun gövde aralığını ve güvenli görsel adayını bul."""
    blocks = _govde_bloklari(document)
    regional_start, regional_end = _bolgesel_jeoloji_araligi(document)
    regional_blocks = blocks[regional_start:regional_end]
    figure_numbers = _stratigrafik_sekil_numaralari(regional_blocks)

    candidates: list[tuple[int, int, str]] = []
    for index in range(regional_start, regional_end):
        block = blocks[index]
        key = _anahtar(_paragraf_metni(block))
        figure_match = re.match(r"^(?:sekil|resim|levha) (\d+)\b", key)
        figure_number = figure_match.group(1) if figure_match else ""
        has_stratigraf = "stratigraf" in key
        reference_sentence = _stratigrafik_referans_cumlesi_mi(key)
        heading_level = _stil_baslik_seviyesi(document, block)
        heading_caption = bool(
            key.startswith("stratigraf")
            and (
                heading_level == 3
                or (len(key.split()) <= 4 and "sekil" not in key)
            )
        )
        full_caption = bool(
            figure_match
            and has_stratigraf
            and not reference_sentence
        )
        short_caption = bool(
            figure_match
            and figure_number in figure_numbers
            and not reference_sentence
        )
        direct_caption = full_caption or short_caption or heading_caption
        referenced_caption = bool(figure_match and figure_number in figure_numbers)
        if not direct_caption and not referenced_caption:
            continue

        if not figure_number:
            figure_number = _stratigrafik_onceki_referans_numarasi(
                blocks,
                regional_start,
                index,
            )

        score = 0
        if full_caption:
            # Tam şekil yazısı, öncesindeki düz referans cümlesinden açıkça
            # daha güçlü adaydır.
            score += 500
        elif short_caption:
            score += 420
        elif heading_caption:
            score += 280
        if has_stratigraf:
            score += 50
        if figure_match:
            score += 30
        if _blok_desteklenen_gorsel_tasiyicisi_var_mi(block):
            score += 25
        reference_distance = (
            _stratigrafik_onceki_referans_mesafesi(
                blocks,
                regional_start,
                index,
                figure_number,
            )
            if figure_number
            else None
        )
        if reference_distance is not None:
            score += max(1, 25 - reference_distance)
        candidates.append((score, index, figure_number))

    if not candidates:
        return None
    best_score = max(item[0] for item in candidates)
    best_candidates = [item for item in candidates if item[0] == best_score]
    if len(best_candidates) > 1:
        raise JeolojiBolumPaketiHatasi(
            "Stratigrafik kesit için birden fazla olası açıklama bulundu; "
            "kaynak Word güvenli biçimde sadeleştirilemedi."
        )
    _score, caption_index, figure_number = best_candidates[0]

    # Önce caption'dan sonraki ilk ayrı image-bearing paragrafını değerlendir.
    # Caption içine eski bir floating harita ilişmiş olsa bile gerçek kesit bu
    # paragraftaki inline/anchor raster olabilir.
    visual_candidate = _stratigrafik_sonraki_ayri_gorsel_adayini_bul(
        document,
        blocks,
        caption_index,
        regional_end,
    )
    if visual_candidate is None:
        # Ayrı görsel yoksa caption bloğundaki foreground/paragraph-relative
        # anchor veya inline adayını değerlendir. behindDoc=1/page anchor burada
        # özellikle dışarıda kalır.
        visual_candidate = _stratigrafik_caption_gorsel_adayini_bul(
            document,
            blocks[caption_index],
            caption_index,
        )

    if visual_candidate is None:
        # Eski kütüphane belgelerinde görsel caption'dan önce ayrı paragrafta
        # bulunabiliyor. Yeni önceliklerin hiçbir adayı yoksa bu sınırlı geri
        # dönüşü koru; eşit/çoklu adayda yine güvenli hata üret.
        for index in range(caption_index - 1, max(regional_start, caption_index - 4) - 1, -1):
            if not _blok_gorsel_iceriyor(blocks[index]):
                continue
            candidates_before = _blok_gorsel_adaylarini_bul(
                document,
                blocks[index],
                blok_no=index,
                strict=True,
            )
            visual_candidate = _stratigrafik_blok_gorselini_sec(candidates_before)
            if visual_candidate is not None:
                break

    if visual_candidate is None:
        return None

    start = min(caption_index, visual_candidate.blok_no)
    end = max(caption_index, visual_candidate.blok_no) + 1
    # Bazı eski Word'lerde başlık, numaralı şekil yazısından ayrı bir Heading 3
    # paragrafıdır. Kesit paketine ait olduğu açıkça görülen bu komşu başlığı da
    # al; görsel ve şekil yazısı ile birlikte Normal stile dönüştürülebilsin.
    while start > regional_start:
        previous = blocks[start - 1]
        if previous.tag != qn("w:p"):
            break
        previous_key = _anahtar(_paragraf_metni(previous))
        if (
            "stratigraf" not in previous_key
            or _stil_baslik_seviyesi(document, previous) != 3
        ):
            break
        start -= 1
    return start, end, figure_number, regional_end, caption_index, visual_candidate


def _stratigrafik_kesit_araligi(document) -> tuple[int, int, str] | None:
    """Kesit görseli ile eski açıklamasının 2.1 içindeki aralığını bul."""
    details = _stratigrafik_kesit_secim_detaylari(document)
    if details is None:
        return None
    start, end, figure_number, _regional_end, _caption_index, _visual_candidate = details
    return start, end, figure_number


def _stratigrafik_kesit_secimini_hazirla(document) -> _StratigrafikKesitSecimi:
    details = _stratigrafik_kesit_secim_detaylari(document)
    if details is None:
        raise JeolojiBolumPaketiHatasi(
            "Seçili jeoloji Word'ünde görseliyle birlikte aktarılabilir "
            "stratigrafik kesit bulunamadı."
        )
    (
        start,
        end,
        figure_number,
        regional_end,
        caption_index,
        visual_candidate,
    ) = details
    part = visual_candidate.part
    extension, content_type = _gorsel_uzantisi_ve_turu(part)
    try:
        blob = bytes(part.blob)
    except Exception as exc:
        raise JeolojiBolumPaketiHatasi(
            "Stratigrafik kesit görselinin özgün DOCX blob'u okunamadı."
        ) from exc
    if not blob:
        raise JeolojiBolumPaketiHatasi(
            "Stratigrafik kesit görselinin özgün DOCX blob'u boş."
        )
    return _StratigrafikKesitSecimi(
        baslangic_blok_no=start,
        bitis_blok_no=end,
        sekil_no=figure_number,
        bolgesel_bitis_blok_no=regional_end,
        caption_blok_no=caption_index,
        gorsel_blok_no=visual_candidate.blok_no,
        gorsel_blob=blob,
        gorsel_uzantisi=extension,
        gorsel_turu=content_type,
        gorsel_adayi=visual_candidate,
    )


def _jeoloji_araligi(document) -> tuple[int, int, bool]:
    blocks = _govde_bloklari(document)
    candidates: list[tuple[int, int, int, int]] = []
    for index, block in enumerate(blocks):
        if block.tag != qn("w:p"):
            continue
        text = _paragraf_metni(block)
        if not _ana_jeoloji_basligi_mi(text) or _toc_paragrafi_mi(document, block):
            continue
        level = _baslik_seviyesi(document, block) or 1
        end = len(blocks)
        for next_index in range(index + 1, len(blocks)):
            following = blocks[next_index]
            if following.tag == qn("w:p") and _sonraki_ana_baslik_mi(document, following, level):
                end = next_index
                break
        section_blocks = blocks[index + 1 : end]
        text_size = sum(len(_paragraf_metni(item)) for item in section_blocks)
        nested = sum(
            1
            for item in section_blocks
            if item.tag == qn("w:p") and _jeoloji_alt_basligi_mi(_paragraf_metni(item))
        )
        explicit_bonus = 500 if _stil_baslik_seviyesi(document, block) == 1 else 0
        level_bonus = 250 if level == 1 else 0
        score = explicit_bonus + level_bonus + (nested * 100) + min(text_size, 5000)
        candidates.append((score, index + 1, end, level))

    if not candidates:
        return 0, len(blocks), False
    _score, start, end, _level = max(candidates, key=lambda item: (item[0], -item[1]))
    return start, end, True


def _docx_yolu(path: str | os.PathLike[str], *, mevcut_olmali: bool) -> Path:
    result = Path(path)
    if result.suffix.lower() != ".docx":
        raise JeolojiBolumPaketiHatasi(f"DOCX dosyası bekleniyor: {result}")
    if mevcut_olmali and not result.is_file():
        raise JeolojiBolumPaketiHatasi(f"Dosya bulunamadı: {result}")
    return result


def _atomik_kaydet(saver, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    handle = tempfile.NamedTemporaryFile(
        prefix=f".{output_path.stem}_",
        suffix=".docx",
        dir=output_path.parent,
        delete=False,
    )
    temp_path = Path(handle.name)
    handle.close()
    try:
        saver.save(str(temp_path))
        os.replace(temp_path, output_path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def _kullanilan_iliski_kimlikleri(elements) -> set[str]:
    result: set[str] = set()
    for element in elements:
        for node in element.iter():
            for attribute in _R_ILISKI_NITELIKLERI:
                rid = node.get(attribute)
                if rid:
                    result.add(rid)
    return result


def _header_footer_referanslarini_kaldir(document) -> None:
    for sect_pr in document.element.iter(qn("w:sectPr")):
        for tag in (qn("w:headerReference"), qn("w:footerReference")):
            for reference in list(sect_pr.findall(tag)):
                sect_pr.remove(reference)


def _kullanilmayan_bolum_iliskilerini_kaldir(document) -> None:
    _header_footer_referanslarini_kaldir(document)
    used = _kullanilan_iliski_kimlikleri([document.element.body])
    for rid, relationship in list(document.part.rels.items()):
        remove = relationship.reltype in (RT.HEADER, RT.FOOTER)
        remove = remove or (
            relationship.reltype in _BOLUM_ICERIK_ILISKI_TURLERI and rid not in used
        )
        if remove:
            document.part.drop_rel(rid)


def _cekirdek_ozellikleri_temizle(document) -> None:
    properties = document.core_properties
    properties.author = ""
    properties.last_modified_by = ""
    properties.title = "Jeoloji Bölüm Paketi"
    properties.subject = ""
    properties.keywords = ""
    properties.comments = ""
    properties.category = ""
    properties.content_status = ""
    properties.identifier = ""
    properties.language = ""
    properties.version = ""
    properties.revision = 1


def jeoloji_bolumunu_ayir(
    kaynak_docx: str | os.PathLike[str],
    paket_docx: str | os.PathLike[str],
) -> JeolojiBolumAyirmaSonucu:
    """Ana JEOLOJİ başlığının altını bağımsız mini DOCX'e ayırır.

    Ana başlığın kendisi pakete alınmaz. Bir sonraki aynı veya daha üst
    düzey başlık sınırdır. Kaynak zaten elle ayrılmış bir parçaysa ve ana
    JEOLOJİ başlığı yoksa belgenin tüm body blokları paket kabul edilir.
    """

    source = _docx_yolu(kaynak_docx, mevcut_olmali=True)
    output = _docx_yolu(paket_docx, mevcut_olmali=False)
    document = Document(str(source))
    blocks = _govde_bloklari(document)
    start, end, heading_found = _jeoloji_araligi(document)
    selected = set(blocks[start:end])
    if not selected:
        raise JeolojiBolumPaketiHatasi("JEOLOJİ bölümünde aktarılacak blok bulunamadı.")

    for child in list(document.element.body.iterchildren()):
        if child.tag != qn("w:sectPr") and child not in selected:
            document.element.body.remove(child)

    _kullanilmayan_bolum_iliskilerini_kaldir(document)
    _cekirdek_ozellikleri_temizle(document)
    _atomik_kaydet(document, output)
    return JeolojiBolumAyirmaSonucu(
        kaynak_docx=str(source.resolve()),
        paket_docx=str(output.resolve()),
        baslik_bulundu=heading_found,
        blok_sayisi=len(selected),
        baslangic_blok_no=start,
        bitis_blok_no=end,
    )


def stratigrafik_kesit_var_mi(kaynak_docx: str | os.PathLike[str]) -> bool:
    """Word içinde görseliyle birlikte aktarılabilir bir stratigrafik kesit var mı?"""
    try:
        source = _docx_yolu(kaynak_docx, mevcut_olmali=True)
        document = Document(str(source))
        _stratigrafik_kesit_secimini_hazirla(document)
    except Exception:
        return False
    return True


def stratigrafik_kesit_gorselini_cikar(
    kaynak_docx: str | os.PathLike[str],
) -> tuple[bytes, str]:
    """Kesit görselinin özgün blob'unu ve güvenli dosya uzantısını döndür."""
    source = _docx_yolu(kaynak_docx, mevcut_olmali=True)
    selection = _stratigrafik_kesit_secimini_hazirla(Document(str(source)))
    return selection.gorsel_blob, selection.gorsel_uzantisi


def _stratigrafik_kesit_gorselini_paragrafa_ekle(
    document,
    paragraph,
    gorsel_blob: bytes,
):
    """Blob'u özgün bytes olarak wp:inline'a ekle ve oranını koru.

    Bazı gerçek Word JPEG'lerinde DPI metadata'sı ``0`` olur. Python-docx'un
    standart ``add_picture`` yolu bu durumda genişliği DPI ile hesaplarken
    sıfıra bölünebilir. Fallback yalnız XML extent'ini piksel oranından kurar;
    image part'a yazılan blob değişmeden kalır.
    """
    run = paragraph.add_run()
    try:
        return run.add_picture(BytesIO(gorsel_blob), width=Cm(15.0))
    except ZeroDivisionError:
        image = _DocxImage.from_blob(gorsel_blob)
        if image.px_width <= 0 or image.px_height <= 0:
            raise JeolojiBolumPaketiHatasi(
                "Stratigrafik kesit görselinin piksel boyutu okunamadı; "
                "kaynak Word güvenle dönüştürülemedi."
            )
        relation_id, image_part = document.part.get_or_add_image(
            BytesIO(gorsel_blob)
        )
        width = int(Cm(15.0))
        height = max(1, round(width * image.px_height / image.px_width))
        inline = CT_Inline.new_pic_inline(
            document.part.next_id,
            relation_id,
            image_part.filename,
            width,
            height,
        )
        run._r.add_drawing(inline)
        return InlineShape(inline)


def _stratigrafik_kesit_paragraflarini_ekle(document, hedef_blok, gorsel_blob: bytes):
    """Hedef bloğun önüne yalnız standart Caption ve InlineShape ekle."""
    caption = document.add_paragraph(STRATIGRAFIK_KESIT_CAPTION, style="Caption")
    caption_p = caption._p
    caption_p.getparent().remove(caption_p)
    hedef_blok.addprevious(caption_p)

    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _stratigrafik_kesit_gorselini_paragrafa_ekle(
        document,
        image_paragraph,
        gorsel_blob,
    )
    image_p = image_paragraph._p
    image_p.getparent().remove(image_p)
    caption_p.addnext(image_p)
    return caption, image_paragraph


def _stratigrafik_kesit_paket_belgesini_olustur(gorsel_blob: bytes):
    document = Document()
    caption = document.add_paragraph(STRATIGRAFIK_KESIT_CAPTION, style="Caption")
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _stratigrafik_kesit_gorselini_paragrafa_ekle(
        document,
        image_paragraph,
        gorsel_blob,
    )
    return document


def _stratigrafik_secili_cizimi_kaldir(block, gorsel_adayi):
    """Yalnız seçilen blip'i kaldır; aynı drawing'deki diğerlerini koru."""
    blip = gorsel_adayi.blip
    if blip.getparent() is not None:
        blip.getparent().remove(blip)
    drawing = gorsel_adayi.drawing
    if drawing.getparent() is not None and not any(drawing.iter(qn("a:blip"))):
        drawing.getparent().remove(drawing)


def _stratigrafik_eski_caption_metinini_kaldir(block):
    """Caption paragrafındaki Word metnini temizle, drawing'leri bırak."""
    text_tags = {
        qn("w:t"),
        qn("w:delText"),
        qn("w:instrText"),
        qn("w:tab"),
        qn("w:br"),
        qn("w:cr"),
        qn("w:fldChar"),
        qn("w:noBreakHyphen"),
        qn("w:softHyphen"),
    }
    for node in list(block.iter()):
        if node.tag not in text_tags:
            continue
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)


def _stratigrafik_blokta_anlamli_icerik_var_mi(block) -> bool:
    if _paragraf_metni(block):
        return True
    return any(
        node.tag in {qn("w:drawing"), qn("w:pict"), qn("w:object")}
        for node in block.iter()
    )


def _stratigrafik_eski_icerigi_temizle(document, selection):
    """Eski caption/kesiti temizlemeden ilişkisiz drawing'leri koru."""
    blocks = _govde_bloklari(document)
    caption_index = selection.caption_blok_no
    visual_index = selection.gorsel_blok_no
    selected_indexes = range(
        selection.baslangic_blok_no,
        selection.bitis_blok_no,
    )
    for index in selected_indexes:
        block = blocks[index]
        if block.getparent() is None:
            continue
        if block.tag != qn("w:p"):
            continue

        key = _anahtar(_paragraf_metni(block))
        old_heading = (
            "stratigraf" in key
            and _stil_baslik_seviyesi(document, block) == 3
        )
        if index == caption_index or old_heading:
            _stratigrafik_eski_caption_metinini_kaldir(block)
        if index == visual_index:
            _stratigrafik_secili_cizimi_kaldir(
                block,
                selection.gorsel_adayi,
            )

        if not _stratigrafik_blokta_anlamli_icerik_var_mi(block):
            block.getparent().remove(block)


def stratigrafik_kesit_bolumunu_ayir(
    kaynak_docx: str | os.PathLike[str],
    paket_docx: str | os.PathLike[str],
) -> StratigrafikKesitAyirmaSonucu:
    """Kesiti yalnız standart Caption + InlineShape mini DOCX'i olarak ayır."""
    source = _docx_yolu(kaynak_docx, mevcut_olmali=True)
    output = _docx_yolu(paket_docx, mevcut_olmali=False)
    document = Document(str(source))
    selection = _stratigrafik_kesit_secimini_hazirla(document)
    package = _stratigrafik_kesit_paket_belgesini_olustur(selection.gorsel_blob)
    _cekirdek_ozellikleri_temizle(package)
    package.core_properties.title = "Stratigrafik Kesit Paketi"
    _atomik_kaydet(package, output)
    return StratigrafikKesitAyirmaSonucu(
        kaynak_docx=str(source.resolve()),
        paket_docx=str(output.resolve()),
        blok_sayisi=selection.bitis_blok_no - selection.baslangic_blok_no,
        baslangic_blok_no=selection.baslangic_blok_no,
        bitis_blok_no=selection.bitis_blok_no,
        sekil_no=selection.sekil_no,
    )


def stratigrafik_kesit_bolumunu_gorsel_olarak_yenile(
    kaynak_docx: str | os.PathLike[str],
    cikti_docx: str | os.PathLike[str] | None = None,
) -> StratigrafikKesitAyirmaSonucu:
    """Kaynak 2.1 içindeki eski kesit bloğunu standart görselle değiştir."""
    source = _docx_yolu(kaynak_docx, mevcut_olmali=True)
    output = _docx_yolu(cikti_docx or source, mevcut_olmali=False)
    document = Document(str(source))
    selection = _stratigrafik_kesit_secimini_hazirla(document)
    blocks = _govde_bloklari(document)
    if selection.bolgesel_bitis_blok_no >= len(blocks):
        raise JeolojiBolumPaketiHatasi(
            "Stratigrafik kesit için 2.1.1 Yapısal Jeoloji başlığı öncesi "
            "güvenli ekleme konumu bulunamadı."
        )
    hedef_blok = blocks[selection.bolgesel_bitis_blok_no]
    if hedef_blok.tag != qn("w:p") or "yapisal jeoloji" not in _anahtar(
        _paragraf_metni(hedef_blok)
    ):
        raise JeolojiBolumPaketiHatasi(
            "Stratigrafik kesit için 2.1.1 Yapısal Jeoloji başlığı öncesi "
            "güvenli ekleme konumu doğrulanamadı."
        )

    _stratigrafik_eski_icerigi_temizle(document, selection)
    _stratigrafik_kesit_paragraflarini_ekle(
        document,
        hedef_blok,
        selection.gorsel_blob,
    )
    _atomik_kaydet(document, output)
    return StratigrafikKesitAyirmaSonucu(
        kaynak_docx=str(source.resolve()),
        paket_docx=str(output.resolve()),
        blok_sayisi=selection.bitis_blok_no - selection.baslangic_blok_no,
        baslangic_blok_no=selection.baslangic_blok_no,
        bitis_blok_no=selection.bitis_blok_no,
        sekil_no=selection.sekil_no,
    )


def stratigrafik_kesit_stilini_normalize_et(
    kaynak_docx: str | os.PathLike[str],
    cikti_docx: str | os.PathLike[str] | None = None,
) -> int:
    """Kaynak Word'deki aktarılabilir kesitte Heading 3 kalıntısını temizle.

    Belgenin diğer paragraflarına ve kesit görselinin XML'ine dokunulmaz.
    ``cikti_docx`` verilmezse kaynak belge güvenli biçimde kendi üzerine yazılır.
    """
    source = _docx_yolu(kaynak_docx, mevcut_olmali=True)
    output = _docx_yolu(cikti_docx or source, mevcut_olmali=False)
    document = Document(str(source))
    result = _stratigrafik_kesit_araligi(document)
    if result is None:
        raise JeolojiBolumPaketiHatasi(
            "Seçili jeoloji Word'ünde görseliyle birlikte aktarılabilir "
            "stratigrafik kesit bulunamadı."
        )
    start, end, _figure_number = result
    changed = _stratigrafik_kesit_stillerini_normalize_et(document, start, end)
    if changed or output != source:
        _atomik_kaydet(document, output)
    return changed


def _yer_tutucu_paragrafi(document, marker: str):
    matches = [
        child
        for child in _govde_bloklari(document)
        if child.tag == qn("w:p") and _paragraf_metni(child) == marker
    ]
    if not matches:
        raise JeolojiBolumPaketiHatasi(
            f"Hedef Word'de tek başına duran {marker} paragrafı bulunamadı."
        )
    if len(matches) > 1:
        raise JeolojiBolumPaketiHatasi(
            f"Hedef Word'de birden fazla {marker} paragrafı bulundu."
        )
    return matches[0]


def _stil_referanslari(elements: list) -> set[str]:
    result: set[str] = set()
    for element in elements:
        for node in element.iter():
            if node.tag in _STIL_REFERANS_ETIKETLERI:
                value = node.get(_W_VAL)
                if value:
                    result.add(value)
    return result


def _numaralandirma_kimlikleri(elements: list) -> set[int]:
    result: set[int] = set()
    for element in elements:
        for node in element.iter(qn("w:numId")):
            try:
                result.add(int(node.get(_W_VAL)))
            except (TypeError, ValueError):
                continue
    return result


def _numbering_part(document):
    try:
        return document.part.numbering_part
    except (KeyError, NotImplementedError):
        return None


def _numaralandirma_stilleri(document, num_ids: set[int]) -> set[str]:
    part = _numbering_part(document)
    if part is None or not num_ids:
        return set()
    root = part.element
    abstract_ids: set[int] = set()
    for num in root.findall(qn("w:num")):
        try:
            num_id = int(num.get(qn("w:numId")))
        except (TypeError, ValueError):
            continue
        if num_id not in num_ids:
            continue
        ref = num.find(qn("w:abstractNumId"))
        if ref is not None:
            try:
                abstract_ids.add(int(ref.get(_W_VAL)))
            except (TypeError, ValueError):
                pass
    result: set[str] = set()
    for abstract in root.findall(qn("w:abstractNum")):
        try:
            abstract_id = int(abstract.get(qn("w:abstractNumId")))
        except (TypeError, ValueError):
            continue
        if abstract_id not in abstract_ids:
            continue
        for node in abstract.iter(qn("w:pStyle")):
            if node.get(_W_VAL):
                result.add(node.get(_W_VAL))
    return result


def _stillerdeki_numaralandirma_kimlikleri(document, style_ids: set[str]) -> set[int]:
    result: set[int] = set()
    for style_id in _stil_bagimlilik_kapanimi(document, style_ids):
        style = _stil_elementi(document, style_id)
        if style is None:
            continue
        result.update(_numaralandirma_kimlikleri([style]))
    return result


def _stil_xml_ayni(left, right) -> bool:
    return etree.tostring(left, with_tail=False) == etree.tostring(right, with_tail=False)


def _benzersiz_stil_id(base: str, used: set[str]) -> str:
    stem = ("JEO_" + re.sub(r"[^A-Za-z0-9_]", "_", base))[:220]
    candidate = stem
    number = 2
    while candidate in used:
        candidate = f"{stem}_{number}"
        number += 1
    used.add(candidate)
    return candidate


def _stil_bagimlilik_kapanimi(document, style_ids: set[str]) -> set[str]:
    result = set(style_ids)
    pending = list(style_ids)
    while pending:
        style_id = pending.pop()
        style = _stil_elementi(document, style_id)
        if style is None:
            continue
        for node in style:
            if node.tag not in _STIL_BAGIMLILIK_ETIKETLERI:
                continue
            dependency = node.get(_W_VAL)
            if dependency and dependency not in result:
                result.add(dependency)
                pending.append(dependency)
    return result


def _stil_referanslarini_guncelle(element, mapping: dict[str, str]) -> None:
    for node in element.iter():
        if node.tag in _STIL_REFERANS_ETIKETLERI or node.tag in _STIL_BAGIMLILIK_ETIKETLERI:
            current = node.get(_W_VAL)
            if current in mapping:
                node.set(_W_VAL, mapping[current])


def _stilleri_kopyala(
    source_document,
    target_document,
    style_ids: set[str],
) -> tuple[dict[str, str], list]:
    source_root = source_document.styles.element
    target_root = target_document.styles.element
    required = _stil_bagimlilik_kapanimi(source_document, style_ids)
    source_by_id = {
        style.get(_W_STYLE_ID): style
        for style in source_root.findall(qn("w:style"))
        if style.get(_W_STYLE_ID)
    }
    target_by_id = {
        style.get(_W_STYLE_ID): style
        for style in target_root.findall(qn("w:style"))
        if style.get(_W_STYLE_ID)
    }
    used = set(target_by_id)
    mapping: dict[str, str] = {}
    imported: list = []
    for style_id in required:
        source_style = source_by_id.get(style_id)
        target_style = target_by_id.get(style_id)
        if source_style is None or target_style is None:
            mapping[style_id] = style_id
        elif _stil_xml_ayni(source_style, target_style):
            mapping[style_id] = style_id
        else:
            mapping[style_id] = _benzersiz_stil_id(style_id, used)

    for style_id in required:
        source_style = source_by_id.get(style_id)
        if source_style is None:
            continue
        mapped_id = mapping[style_id]
        if mapped_id in target_by_id and mapped_id == style_id:
            continue
        clone = deepcopy(source_style)
        clone.set(_W_STYLE_ID, mapped_id)
        # Hedef belgenin varsayılan stilleri değişmesin.
        clone.attrib.pop(qn("w:default"), None)
        _stil_referanslarini_guncelle(clone, mapping)
        if mapped_id != style_id:
            name = clone.find(qn("w:name"))
            if name is not None:
                name.set(_W_VAL, f"{name.get(_W_VAL, style_id)} (Jeoloji)")
        target_root.append(clone)
        target_by_id[mapped_id] = clone
        imported.append(clone)
    return mapping, imported


def _sonraki_sayi(elements: list, attribute_qname: str) -> int:
    values: list[int] = []
    for element in elements:
        try:
            values.append(int(element.get(attribute_qname)))
        except (TypeError, ValueError):
            continue
    return max(values, default=-1) + 1


def _once_ekle(parent, child, before_tags: tuple[str, ...]) -> None:
    for existing in parent:
        if existing.tag in before_tags:
            existing.addprevious(child)
            return
    parent.append(child)


def _parca_adi_sablonu(partname: object) -> str:
    value = str(partname)
    match = re.match(r"^(.*?)(\d+)(\.[^./]+)$", value)
    if match:
        return f"{match.group(1)}%d{match.group(3)}"
    dot = value.rfind(".")
    if dot > value.rfind("/"):
        return f"{value[:dot]}%d{value[dot:]}"
    return value + "%d"


def _parcayi_kopyala(source_part, target_package, cache: dict[int, Part]):
    cache_key = id(source_part)
    if cache_key in cache:
        return cache[cache_key]
    if source_part.content_type.startswith("image/"):
        image_part = target_package.get_or_add_image_part(BytesIO(source_part.blob))
        cache[cache_key] = image_part
        return image_part

    partname = target_package.next_partname(_parca_adi_sablonu(source_part.partname))
    clone = Part(partname, source_part.content_type, source_part.blob, target_package)
    cache[cache_key] = clone
    for relationship in source_part.rels.values():
        if relationship.is_external:
            target = relationship.target_ref
        else:
            target = _parcayi_kopyala(relationship.target_part, target_package, cache)
        clone.rels.add_relationship(
            relationship.reltype,
            target,
            relationship.rId,
            relationship.is_external,
        )
    return clone


def _iliskiyi_kopyala(source_owner, target_owner, old_rid: str, cache: dict[int, Part]) -> str:
    relationship = source_owner.rels.get(old_rid)
    if relationship is None:
        raise JeolojiBolumPaketiHatasi(
            f"Kaynak Word ilişkisi bulunamadı: {old_rid}"
        )
    if relationship.is_external:
        return target_owner.relate_to(
            relationship.target_ref,
            relationship.reltype,
            is_external=True,
        )
    target_part = _parcayi_kopyala(relationship.target_part, target_owner.package, cache)
    return target_owner.relate_to(target_part, relationship.reltype)


def _element_iliskilerini_kopyala(
    element,
    source_owner,
    target_owner,
    cache: dict[int, Part],
) -> None:
    rid_mapping: dict[str, str] = {}
    for node in element.iter():
        for attribute in _R_ILISKI_NITELIKLERI:
            old_rid = node.get(attribute)
            if not old_rid:
                continue
            if old_rid not in rid_mapping:
                rid_mapping[old_rid] = _iliskiyi_kopyala(
                    source_owner,
                    target_owner,
                    old_rid,
                    cache,
                )
            node.set(attribute, rid_mapping[old_rid])


def _numaralandirmayi_kopyala(
    source_document,
    target_document,
    num_ids: set[int],
    style_mapping: dict[str, str],
    part_cache: dict[int, Part],
) -> dict[int, int]:
    if not num_ids:
        return {}
    source_part = _numbering_part(source_document)
    target_part = _numbering_part(target_document)
    if source_part is None:
        return {}
    if target_part is None:
        raise JeolojiBolumPaketiHatasi(
            "Hedef Word numaralandırma parçası içermediği için liste aktarılamadı."
        )

    source_root = source_part.element
    target_root = target_part.element
    source_nums = {
        int(node.get(qn("w:numId"))): node
        for node in source_root.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    }
    source_abstracts = {
        int(node.get(qn("w:abstractNumId"))): node
        for node in source_root.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    }
    next_num = _sonraki_sayi(target_root.findall(qn("w:num")), qn("w:numId"))
    next_abstract = _sonraki_sayi(
        target_root.findall(qn("w:abstractNum")), qn("w:abstractNumId")
    )
    next_picture = _sonraki_sayi(
        target_root.findall(qn("w:numPicBullet")), qn("w:numPicBulletId")
    )
    num_mapping: dict[int, int] = {}
    abstract_mapping: dict[int, int] = {}
    picture_mapping: dict[int, int] = {}

    for old_num_id in sorted(num_ids):
        source_num = source_nums.get(old_num_id)
        if source_num is None:
            continue
        abstract_ref = source_num.find(qn("w:abstractNumId"))
        if abstract_ref is None:
            continue
        old_abstract_id = int(abstract_ref.get(_W_VAL))
        if old_abstract_id not in abstract_mapping:
            source_abstract = source_abstracts.get(old_abstract_id)
            if source_abstract is None:
                continue
            abstract_clone = deepcopy(source_abstract)
            new_abstract_id = next_abstract
            next_abstract += 1
            abstract_mapping[old_abstract_id] = new_abstract_id
            abstract_clone.set(qn("w:abstractNumId"), str(new_abstract_id))
            _stil_referanslarini_guncelle(abstract_clone, style_mapping)

            for picture_ref in abstract_clone.iter(qn("w:lvlPicBulletId")):
                old_picture_id = int(picture_ref.get(_W_VAL))
                if old_picture_id not in picture_mapping:
                    source_picture = next(
                        (
                            item
                            for item in source_root.findall(qn("w:numPicBullet"))
                            if int(item.get(qn("w:numPicBulletId"))) == old_picture_id
                        ),
                        None,
                    )
                    if source_picture is not None:
                        picture_clone = deepcopy(source_picture)
                        new_picture_id = next_picture
                        next_picture += 1
                        picture_mapping[old_picture_id] = new_picture_id
                        picture_clone.set(qn("w:numPicBulletId"), str(new_picture_id))
                        _element_iliskilerini_kopyala(
                            picture_clone,
                            source_part,
                            target_part,
                            part_cache,
                        )
                        _once_ekle(
                            target_root,
                            picture_clone,
                            (qn("w:abstractNum"), qn("w:num")),
                        )
                if old_picture_id in picture_mapping:
                    picture_ref.set(_W_VAL, str(picture_mapping[old_picture_id]))

            _once_ekle(target_root, abstract_clone, (qn("w:num"),))

        num_clone = deepcopy(source_num)
        new_num_id = next_num
        next_num += 1
        num_mapping[old_num_id] = new_num_id
        num_clone.set(qn("w:numId"), str(new_num_id))
        num_clone.find(qn("w:abstractNumId")).set(
            _W_VAL,
            str(abstract_mapping[old_abstract_id]),
        )
        target_root.append(num_clone)
    return num_mapping


def _numaralandirma_referanslarini_guncelle(element, mapping: dict[int, int]) -> None:
    for node in element.iter(qn("w:numId")):
        try:
            old_id = int(node.get(_W_VAL))
        except (TypeError, ValueError):
            continue
        if old_id in mapping:
            node.set(_W_VAL, str(mapping[old_id]))


def _cizim_kimliklerini_yenile(elements: list, target_document) -> None:
    drawing_tags = (qn("wp:docPr"), qn("pic:cNvPr"))
    existing_ids: list[int] = []
    for node in target_document.element.iter():
        if node.tag not in drawing_tags:
            continue
        try:
            existing_ids.append(int(node.get("id")))
        except (TypeError, ValueError):
            continue
    next_id = max(existing_ids, default=0) + 1
    for element in elements:
        for node in element.iter():
            if node.tag not in drawing_tags:
                continue
            node.set("id", str(next_id))
            next_id += 1


def _ooxml_ile_yerlestir(source_document, target_document, placeholder) -> int:
    source_blocks = _govde_bloklari(source_document)
    if not source_blocks:
        raise JeolojiBolumPaketiHatasi("Bölüm paketinde aktarılacak blok bulunamadı.")
    copies = [deepcopy(block) for block in source_blocks]
    _cizim_kimliklerini_yenile(copies, target_document)
    num_ids = _numaralandirma_kimlikleri(copies)
    style_ids = _stil_referanslari(copies)
    # Bazı liste tanımları bir stile, bazı stiller de bir numId'ye bağlıdır.
    # Kapanım sabitlenene kadar iki yönlü bağımlılıkları genişlet.
    while True:
        previous = (set(style_ids), set(num_ids))
        style_ids.update(_numaralandirma_stilleri(source_document, num_ids))
        num_ids.update(_stillerdeki_numaralandirma_kimlikleri(source_document, style_ids))
        if previous == (style_ids, num_ids):
            break
    style_mapping, imported_styles = _stilleri_kopyala(
        source_document,
        target_document,
        style_ids,
    )
    for block in copies:
        _stil_referanslarini_guncelle(block, style_mapping)

    part_cache: dict[int, Part] = {}
    num_mapping = _numaralandirmayi_kopyala(
        source_document,
        target_document,
        num_ids,
        style_mapping,
        part_cache,
    )
    for style in imported_styles:
        _numaralandirma_referanslarini_guncelle(style, num_mapping)
    for block in copies:
        _numaralandirma_referanslarini_guncelle(block, num_mapping)
        _element_iliskilerini_kopyala(
            block,
            source_document.part,
            target_document.part,
            part_cache,
        )
        placeholder.addprevious(block)
    placeholder.getparent().remove(placeholder)
    return len(copies)


def _docxcompose_ile_yerlestir(source_document, target_document, placeholder) -> tuple[int, object]:
    source_blocks = _govde_bloklari(source_document)
    if not source_blocks:
        raise JeolojiBolumPaketiHatasi("Bölüm paketinde aktarılacak blok bulunamadı.")
    _cizim_kimliklerini_yenile(source_blocks, target_document)
    body = target_document.element.body
    index = body.index(placeholder)
    body.remove(placeholder)
    composer = _Composer(target_document)
    insert_signature = inspect.signature(composer.insert)
    if "remove_property_fields" in insert_signature.parameters:
        composer.insert(index, source_document, remove_property_fields=False)
    else:  # pragma: no cover - eski docxcompose sürümleri
        composer.insert(index, source_document)
    return len(source_blocks), composer


def jeoloji_bolumunu_yerlestir(
    paket_docx: str | os.PathLike[str],
    hedef_docx: str | os.PathLike[str],
    cikti_docx: str | os.PathLike[str],
    *,
    yer_tutucu: str = VARSAYILAN_YER_TUTUCU,
) -> JeolojiBolumYerlestirmeSonucu:
    """Mini DOCX'i hedefteki tek başına duran yer tutucunun yerine koyar.

    Hedef belgenin bölüm başlıkları ve ``sectPr`` ayarları aynen kalır;
    paket belgenin son ``sectPr`` elemanı hedefe taşınmaz.
    """

    source_path = _docx_yolu(paket_docx, mevcut_olmali=True)
    target_path = _docx_yolu(hedef_docx, mevcut_olmali=True)
    output_path = _docx_yolu(cikti_docx, mevcut_olmali=False)
    marker = str(yer_tutucu or "").strip()
    if not marker:
        raise JeolojiBolumPaketiHatasi("Yer tutucu boş olamaz.")

    source_document = Document(str(source_path))
    target_document = Document(str(target_path))
    placeholder = _yer_tutucu_paragrafi(target_document, marker)

    if _Composer is not None and hasattr(_Composer, "insert"):
        block_count, saver = _docxcompose_ile_yerlestir(
            source_document,
            target_document,
            placeholder,
        )
        engine = "docxcompose"
    else:
        block_count = _ooxml_ile_yerlestir(
            source_document,
            target_document,
            placeholder,
        )
        saver = target_document
        engine = "ooxml"

    _atomik_kaydet(saver, output_path)
    return JeolojiBolumYerlestirmeSonucu(
        paket_docx=str(source_path.resolve()),
        hedef_docx=str(target_path.resolve()),
        cikti_docx=str(output_path.resolve()),
        blok_sayisi=block_count,
        motor=engine,
    )


__all__ = [
    "JeolojiBolumAyirmaSonucu",
    "JeolojiBolumPaketiHatasi",
    "JeolojiBolumYerlestirmeSonucu",
    "StratigrafikKesitAyirmaSonucu",
    "STRATIGRAFIK_KESIT_CAPTION",
    "VARSAYILAN_YER_TUTUCU",
    "jeoloji_bolumunu_ayir",
    "jeoloji_bolumunu_yerlestir",
    "stratigrafik_kesit_bolumunu_ayir",
    "stratigrafik_kesit_bolumunu_gorsel_olarak_yenile",
    "stratigrafik_kesit_gorselini_cikar",
    "stratigrafik_kesit_var_mi",
]
