import os
import time

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT


TAAHHUT_BILGI_ALANLARI = (
    "JEOFIZIK_MUH_AD",
    "JEOFIZIK_MUH_SICIL",
    "JEOFIZIK_MUH_ADRES",
    "JEOFIZIK_MUH_TELEFON",
    "JEOLOJI_MUH_AD",
    "JEOLOJI_MUH_SICIL",
    "JEOLOJI_MUH_ADRES",
    "JEOLOJI_MUH_TELEFON",
)


class TaahhutnameUretici:
    def __init__(self, app):
        self.app = app

    @property
    def taahhut_word_sablon_yolu(self):
        return getattr(self.app, "taahhut_word_sablon_yolu", "")

    def proje_deger(self, kod, varsayilan=""):
        if hasattr(self.app, "proje_deger"):
            return self.app.proje_deger(kod, varsayilan)
        entry = getattr(self.app, "veri_alanlari", {}).get(kod)
        deger = entry.get().strip() if entry else ""
        return deger or varsayilan

    def bina_deger(self, etiket, varsayilan=""):
        if hasattr(self.app, "bina_deger"):
            return self.app.bina_deger(etiket, varsayilan)
        entry = getattr(self.app, "bina_alanlari", {}).get(etiket)
        deger = entry.get().strip() if entry else ""
        return deger or varsayilan

    def docx_xml_metin_dugumlerini_dolas(self, doc):
        from docx.oxml.ns import qn

        for part in doc.part.package.parts:
            partname = str(getattr(part, "partname", ""))
            if not partname.startswith("/word/") or not hasattr(part, "element"):
                continue
            if not any(parca in partname for parca in ("document", "header", "footer", "footnotes", "endnotes")):
                continue
            for node in part.element.iter(qn("w:t")):
                yield node

    def taahhut_deger(self, kod, varsayilan=""):
        yapilandirma = getattr(self.app, "taahhut_bilgileri", {})
        if isinstance(yapilandirma, dict) and kod in yapilandirma:
            return str(yapilandirma.get(kod, "")).strip()
        eski_yapilandirma = getattr(self.app, "taahhut_varsayilanlari", {})
        if isinstance(eski_yapilandirma, dict) and kod in eski_yapilandirma:
            return str(eski_yapilandirma.get(kod, "")).strip()
        return str(varsayilan or "").strip()

    def taahhut_bilgilerini_dogrula(self):
        eksikler = [kod for kod in TAAHHUT_BILGI_ALANLARI if not self.taahhut_deger(kod, "").strip()]
        if not self.taahhut_ilgili_idare().strip():
            eksikler.append("ILGILI_IDARE")
        if eksikler:
            raise ValueError(
                "Taahhütname mühendis bilgileri eksik: " + ", ".join(eksikler)
            )

    def taahhut_turkce_kucuk(self, metin):
        return str(metin).replace("I", "ı").replace("İ", "i").lower()

    def taahhut_turkce_buyuk(self, metin):
        return str(metin).replace("i", "İ").replace("ı", "I").upper()

    def taahhut_kelime_yazim_duzeni(self, kelime, tireli_buyuk_koru=False):
        import string

        if not kelime or not any(ch.isalpha() for ch in kelime):
            return kelime

        bas = 0
        son = len(kelime)
        kirpilacak = string.punctuation.replace("-", "").replace("'", "")
        while bas < son and kelime[bas] in kirpilacak:
            bas += 1
        while son > bas and kelime[son - 1] in kirpilacak:
            son -= 1

        on = kelime[:bas]
        govde = kelime[bas:son]
        arka = kelime[son:]
        if not govde:
            return kelime

        kisaltmalar = {
            "AŞ": "A.Ş.",
            "LTD": "Ltd.",
            "ŞTİ": "Şti.",
            "STİ": "Şti.",
            "TİC": "Tic.",
            "TIC": "Tic.",
            "SAN": "San.",
            "İNŞ": "İnş.",
            "INS": "İnş.",
            "TUR": "Tur.",
            "GID": "Gıd.",
            "GIDA": "Gıd.",
            "MOB": "Mob.",
            "MAH": "Mah.",
            "MH": "Mh.",
            "CAD": "Cad.",
            "CD": "Cd.",
            "SOK": "Sok.",
            "SK": "Sk.",
            "NO": "No",
            "DA": "Da",
        }
        anahtar = self.taahhut_turkce_buyuk(govde.replace(".", ""))
        if anahtar in kisaltmalar:
            deger = kisaltmalar[anahtar]
            if arka.startswith(".") and deger.endswith("."):
                arka = arka[1:]
            return f"{on}{deger}{arka}"

        if tireli_buyuk_koru and "-" in govde and govde == self.taahhut_turkce_buyuk(govde):
            return kelime

        if "-" in govde:
            parcalar = [
                self.taahhut_kelime_yazim_duzeni(parca, tireli_buyuk_koru)
                for parca in govde.split("-")
            ]
            return f"{on}{'-'.join(parcalar)}{arka}"

        if "'" in govde:
            ilk, kalan = govde.split("'", 1)
            ilk = self.taahhut_kelime_yazim_duzeni(ilk, tireli_buyuk_koru)
            kalan = self.taahhut_turkce_kucuk(kalan)
            return f"{on}{ilk}'{kalan}{arka}"

        kucuk = self.taahhut_turkce_kucuk(govde)
        return f"{on}{self.taahhut_turkce_buyuk(kucuk[:1])}{kucuk[1:]}{arka}"

    def taahhut_yazim_duzeni(self, metin, tireli_buyuk_koru=False):
        return " ".join(
            self.taahhut_kelime_yazim_duzeni(kelime, tireli_buyuk_koru)
            for kelime in str(metin).split()
        )

    def taahhut_proje_adi_yazim_duzeni(self):
        return self.taahhut_yazim_duzeni(self.proje_deger("PROJE_ADI", ""), tireli_buyuk_koru=True)

    def taahhut_verisini_topla(self):
        return {kod: self.taahhut_deger(kod, "") for kod in TAAHHUT_BILGI_ALANLARI}

    def taahhut_verisini_yerlestir(self, veriler):
        if not isinstance(veriler, dict):
            return
        mevcut = dict(getattr(self.app, "taahhut_bilgileri", {}))
        for kod in TAAHHUT_BILGI_ALANLARI:
            if kod in veriler:
                mevcut[kod] = str(veriler[kod]).strip()
        self.app.taahhut_bilgileri = mevcut

    def taahhut_dosya_adi_temizle(self, metin):
        izinli = []
        for ch in str(metin):
            if ch.isalnum() or ch in (" ", "_", "-", "."):
                izinli.append(ch)
        temiz = "".join(izinli).strip().replace(" ", "_")
        return temiz[:60] or "K1"

    def taahhut_ilgili_idare(self):
        explicit = str(self.proje_deger("ILGILI_IDARE", "")).strip()
        if explicit:
            return explicit
        # İlgili idare alanı proje formunda ayrıca sorulmuyor. K-1 köy projelerinde
        # idare, proje ilinden güvenli ve tekrarlanabilir biçimde türetilir.
        il = str(self.proje_deger("IL", "")).strip()
        if not il:
            return ""
        return f"{self.taahhut_yazim_duzeni(il)} İl Özel İdaresi"

    def taahhut_yapi_adresi(self):
        il = self.taahhut_yazim_duzeni(self.proje_deger("IL", ""))
        ilce = self.taahhut_yazim_duzeni(self.proje_deger("ILCE", ""))
        koy = self.proje_deger("KOY", "")
        koy_adresi = self.taahhut_yazim_duzeni(koy)
        koy_kucuk = koy_adresi.lower()
        if koy_adresi and not any(ifade in koy_kucuk for ifade in ("mah", "mh", "köy", "koy")):
            koy_adresi = f"{koy_adresi} Köyü"
        return " ".join(parca for parca in [koy_adresi, ilce, il] if parca).strip()

    def taahhut_pafta_ada_parsel_metni(self):
        pafta = self.proje_deger("PAFTA", "-")
        ada = self.proje_deger("ADA", "-")
        parsel = self.proje_deger("PARSEL", "-")
        parcalar = [str(deger).strip() or "-" for deger in (pafta, ada, parsel)]
        return " / ".join(parcalar)

    def taahhut_tarihi(self):
        t = time.localtime()
        return f"{t.tm_mday}.{t.tm_mon:02d}.{t.tm_year}"

    def taahhut_word_hucre_metni_ayarla(self, hucre, metin, kalin=None, hizalama=None):
        paragraf = hucre.paragraphs[0] if hucre.paragraphs else hucre.add_paragraph()
        if hizalama is not None:
            paragraf.alignment = hizalama
        paragraf.paragraph_format.space_before = Pt(0)
        paragraf.paragraph_format.space_after = Pt(0)
        if paragraf.runs:
            paragraf.runs[0].text = str(metin)
            run = paragraf.runs[0]
            for diger_run in paragraf.runs[1:]:
                diger_run.text = ""
        else:
            run = paragraf.add_run(str(metin))
        if kalin is not None:
            run.bold = kalin
        run.font.name = "Arial"
        run.font.size = Pt(10)
        for ek_paragraf in hucre.paragraphs[1:]:
            hucre._tc.remove(ek_paragraf._p)

    def taahhut_word_tablo_bosluklarini_temizle(self, doc):
        for tablo in doc.tables:
            for satir in tablo.rows:
                for hucre in satir.cells:
                    hucre.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraf in hucre.paragraphs:
                        paragraf.paragraph_format.space_before = Pt(0)
                        paragraf.paragraph_format.space_after = Pt(0)
                        for run in paragraf.runs:
                            run.font.name = "Arial"

    def taahhut_word_hucre_genisligi_ayarla(self, hucre, genislik):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = hucre._tc.get_or_add_tcPr()
        tc_w = tc_pr.tcW
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(genislik))
        tc_w.set(qn("w:type"), "dxa")

    def taahhut_word_hucre_kenarliklari_ayarla(self, hucre, kenarlar):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = hucre._tc.get_or_add_tcPr()
        for mevcut in tc_pr.findall(qn("w:tcBorders")):
            tc_pr.remove(mevcut)

        if not kenarlar:
            return

        borders = OxmlElement("w:tcBorders")
        for kenar in ("top", "left", "bottom", "right"):
            if kenar not in kenarlar:
                continue
            border = OxmlElement(f"w:{kenar}")
            if kenarlar[kenar]:
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "8")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "000000")
            else:
                border.set(qn("w:val"), "nil")
            borders.append(border)
        tc_pr.append(borders)

    def taahhut_word_hucre_ic_boslugu_ayarla(self, hucre, ust=0, sol=0, alt=0, sag=0):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = hucre._tc.get_or_add_tcPr()
        tc_mar = tc_pr.find(qn("w:tcMar"))
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for kenar, deger in (("top", ust), ("left", sol), ("bottom", alt), ("right", sag)):
            eleman = tc_mar.find(qn(f"w:{kenar}"))
            if eleman is None:
                eleman = OxmlElement(f"w:{kenar}")
                tc_mar.append(eleman)
            eleman.set(qn("w:w"), str(deger))
            eleman.set(qn("w:type"), "dxa")

    def taahhut_word_proje_tablo_kenarliklarini_duzenle(self, tablo):
        for satir_no, satir in enumerate(tablo.rows):
            hucreler = satir.cells
            if satir_no == 0:
                self.taahhut_word_hucre_kenarliklari_ayarla(
                    hucreler[0],
                    {"top": True, "left": True, "bottom": True, "right": True},
                )
            elif satir_no == 1:
                self.taahhut_word_hucre_kenarliklari_ayarla(hucreler[0], {"bottom": True})
            elif satir_no == len(tablo.rows) - 1:
                for i, hucre in enumerate(hucreler):
                    kenarlar = {"bottom": True}
                    if i == 0:
                        kenarlar["left"] = True
                    if i == len(hucreler) - 1:
                        kenarlar["right"] = True
                    self.taahhut_word_hucre_kenarliklari_ayarla(hucre, kenarlar)
            else:
                for i, hucre in enumerate(hucreler):
                    kenarlar = {}
                    if i == 0:
                        kenarlar["left"] = True
                    if i == len(hucreler) - 1:
                        kenarlar["right"] = True
                    self.taahhut_word_hucre_kenarliklari_ayarla(hucre, kenarlar)

    def taahhut_word_tablo_genisligi_ayarla(self, tablo, genislik):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tbl_pr = tablo._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tablo._tbl.insert(0, tbl_pr)

        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.insert(0, tbl_w)
        tbl_w.set(qn("w:w"), str(genislik))
        tbl_w.set(qn("w:type"), "dxa")

        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:w"), "137")
        tbl_ind.set(qn("w:type"), "dxa")

        jc = tbl_pr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            tbl_pr.append(jc)
        jc.set(qn("w:val"), "left")

    def taahhut_word_proje_satirlari(self):
        il = self.taahhut_yazim_duzeni(self.proje_deger("IL", ""))
        ilce = self.taahhut_yazim_duzeni(self.proje_deger("ILCE", ""))
        pafta = self.proje_deger("PAFTA", "-")
        ada = self.proje_deger("ADA", "")
        parsel = self.proje_deger("PARSEL", "")
        proje_adi = self.taahhut_proje_adi_yazim_duzeni()
        yapi_adresi = self.taahhut_yapi_adresi()
        ilgili_idare = self.taahhut_ilgili_idare()
        return [
            ("İl / İlçe", " ".join(p for p in (il, ilce) if p), ""),
            ("İlgili İdare", ilgili_idare, ""),
            ("Pafta/Ada/Parsel No", self.taahhut_pafta_ada_parsel_metni(), ""),
            ("Yapı Adresi", yapi_adresi, ""),
            ("Yapı Sahibi", proje_adi, ""),
            ("Yapı Sahibinin Adresi", yapi_adresi, ""),
            ("Projenin Türü", "ZEMİN ETÜDÜ RAPORU", ""),
        ]

    def taahhut_word_proje_tablo_elemanlari_olustur(self, doc, govde_genislikleri=None):
        govde_genislikleri = govde_genislikleri or [2070, 586, 4149, 2779]
        elemanlar = []
        baslik_genisligi = sum(govde_genislikleri)

        baslik_tablo = doc.add_table(rows=1, cols=1)
        baslik_tablo.alignment = WD_TABLE_ALIGNMENT.LEFT
        baslik_tablo.autofit = False
        self.taahhut_word_tablo_genisligi_ayarla(baslik_tablo, baslik_genisligi)
        baslik_hucre = baslik_tablo.cell(0, 0)
        self.taahhut_word_hucre_genisligi_ayarla(baslik_hucre, baslik_genisligi)
        self.taahhut_word_hucre_ic_boslugu_ayarla(baslik_hucre)
        self.taahhut_word_hucre_kenarliklari_ayarla(
            baslik_hucre,
            {"top": True, "left": True, "bottom": True, "right": True},
        )
        self.taahhut_word_hucre_metni_ayarla(
            baslik_hucre,
            "Müellifliği Üstlenilen Proje",
            True,
            WD_ALIGN_PARAGRAPH.CENTER,
        )
        elemanlar.append(baslik_tablo._tbl)

        bosluk = doc.add_paragraph()
        bosluk.paragraph_format.space_before = Pt(0)
        bosluk.paragraph_format.space_after = Pt(6)
        elemanlar.append(bosluk._p)

        govde_tablo = doc.add_table(rows=1, cols=1)
        govde_tablo.alignment = WD_TABLE_ALIGNMENT.LEFT
        govde_tablo.autofit = False
        self.taahhut_word_tablo_genisligi_ayarla(govde_tablo, baslik_genisligi)
        govde_hucre = govde_tablo.cell(0, 0)
        self.taahhut_word_hucre_genisligi_ayarla(govde_hucre, baslik_genisligi)
        self.taahhut_word_hucre_ic_boslugu_ayarla(govde_hucre, ust=90, alt=90)
        self.taahhut_word_hucre_kenarliklari_ayarla(
            govde_hucre,
            {"top": True, "left": True, "bottom": True, "right": True},
        )
        govde_hucre.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for paragraf in list(govde_hucre.paragraphs):
            govde_hucre._tc.remove(paragraf._p)

        ic_tablo = govde_hucre.add_table(rows=7, cols=4)
        ic_tablo.autofit = False
        self.taahhut_word_tablo_genisligi_ayarla(ic_tablo, baslik_genisligi)
        tbl_grid = ic_tablo._tbl.tblGrid
        if tbl_grid is not None:
            from docx.oxml.ns import qn
            for grid_col, genislik in zip(tbl_grid.gridCol_lst, govde_genislikleri):
                grid_col.set(qn("w:w"), str(genislik))

        for satir_no, (etiket, deger, ek_deger) in enumerate(self.taahhut_word_proje_satirlari()):
            satir = ic_tablo.rows[satir_no]
            for hucre, genislik in zip(satir.cells, govde_genislikleri):
                self.taahhut_word_hucre_genisligi_ayarla(hucre, genislik)
                self.taahhut_word_hucre_ic_boslugu_ayarla(hucre)
                self.taahhut_word_hucre_kenarliklari_ayarla(hucre, {})
                hucre.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self.taahhut_word_hucre_metni_ayarla(satir.cells[0], etiket, True)
            self.taahhut_word_hucre_metni_ayarla(satir.cells[1], ":", False, WD_ALIGN_PARAGRAPH.CENTER)
            self.taahhut_word_hucre_metni_ayarla(satir.cells[2], deger, False)
            self.taahhut_word_hucre_metni_ayarla(satir.cells[3], ek_deger, False)
            for hucre, girinti in zip(satir.cells, [35, 149, 396, 0]):
                for paragraf in hucre.paragraphs:
                    paragraf.paragraph_format.left_indent = Pt(girinti / 20)
                    paragraf.paragraph_format.first_line_indent = Pt(0)
                    paragraf.paragraph_format.space_before = Pt(0)
                    paragraf.paragraph_format.space_after = Pt(0)

        for paragraf in govde_hucre.paragraphs:
            paragraf.paragraph_format.space_before = Pt(0)
            paragraf.paragraph_format.space_after = Pt(0)
            for run in paragraf.runs:
                run.font.size = Pt(1)

        elemanlar.append(govde_tablo._tbl)
        return elemanlar

    def taahhut_word_proje_tablolarini_yeniden_ciz(self, doc):
        proje_tablolari = [
            tablo
            for tablo in list(doc.tables)
            if tablo.rows and "Müellifliği" in tablo.rows[0].cells[0].text
        ]
        if not proje_tablolari:
            # Eski şablonlarda başlık bozuk kodlanmış olabilir. Yalnızca
            # beklenen proje tablosu geometrisine uyan ilk iki tabloyu yenileriz.
            proje_tablolari = [
                tablo
                for tablo in list(doc.tables)
                if len(tablo.rows) >= 7 and len(tablo.columns) >= 3
            ][:2]
        for index, tablo in enumerate(proje_tablolari):
            eski_tablo = tablo._tbl
            govde = eski_tablo.getparent()
            ekleme_index = govde.index(eski_tablo)
            govde_genislikleri = [2070, 586, 4149, 2779] if index == 0 else [2070, 667, 4050, 2633]
            yeni_elemanlar = self.taahhut_word_proje_tablo_elemanlari_olustur(doc, govde_genislikleri)
            for eleman in yeni_elemanlar:
                eleman.getparent().remove(eleman)
            for offset, eleman in enumerate(yeni_elemanlar):
                govde.insert(ekleme_index + offset, eleman)
            govde.remove(eski_tablo)

    def taahhut_word_proje_tablolarini_hizala(self, doc):
        from copy import deepcopy
        from docx.oxml.ns import qn

        proje_tablolari = []
        for tablo in doc.tables:
            if tablo.rows and "Müellifliği" in tablo.rows[0].cells[0].text:
                proje_tablolari.append(tablo)
        if not proje_tablolari:
            return

        kaynak = proje_tablolari[0]
        kaynak_satir_sayisi = len(kaynak.rows)
        govde_genislikleri = [2070, 586, 4149, 2779]
        baslik_genisligi = sum(govde_genislikleri)

        for tablo in proje_tablolari[1:]:
            if kaynak_satir_sayisi > len(tablo.rows) and len(tablo.rows) > 1:
                bos_satir = deepcopy(kaynak.rows[1]._tr)
                tablo.rows[0]._tr.addnext(bos_satir)

        for tablo in proje_tablolari:
            tbl_grid = tablo._tbl.tblGrid
            if tbl_grid is not None:
                for grid_col, genislik in zip(tbl_grid.gridCol_lst, govde_genislikleri):
                    grid_col.set(qn("w:w"), str(genislik))
            for satir_no, satir in enumerate(tablo.rows):
                genislikler = [baslik_genisligi] * len(satir.cells) if satir_no in (0, 1) else govde_genislikleri
                for hucre, genislik in zip(satir.cells, genislikler):
                    self.taahhut_word_hucre_genisligi_ayarla(hucre, genislik)
                if satir_no >= 2:
                    for hucre, girinti in zip(satir.cells, [35, 149, 396, 0]):
                        for paragraf in hucre.paragraphs:
                            paragraf.paragraph_format.left_indent = Pt(girinti / 20)
                            paragraf.paragraph_format.first_line_indent = Pt(0)
            self.taahhut_word_proje_tablo_kenarliklarini_duzenle(tablo)

    def taahhut_word_tablolari_doldur(self, doc):
        il = self.taahhut_yazim_duzeni(self.proje_deger("IL", ""))
        ilce = self.taahhut_yazim_duzeni(self.proje_deger("ILCE", ""))
        proje_adi = self.taahhut_proje_adi_yazim_duzeni()
        yapi_adresi = self.taahhut_yapi_adresi()
        ilgili_idare = self.taahhut_ilgili_idare()
        pafta_ada_parsel = self.taahhut_pafta_ada_parsel_metni()

        for tablo in doc.tables:
            for satir in tablo.rows:
                hucreler = satir.cells
                if len(hucreler) < 3:
                    continue
                etiket = " ".join(hucreler[0].text.split()).lower().replace("i̇", "i").replace("ı", "i")
                if not etiket:
                    continue
                if "il / il" in etiket or "il/il" in etiket:
                    self.taahhut_word_hucre_metni_ayarla(hucreler[0], "İl / İlçe", True)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[1], ":", False, WD_ALIGN_PARAGRAPH.CENTER)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[2], " ".join(p for p in (il, ilce) if p), False)
                    if len(hucreler) > 3:
                        self.taahhut_word_hucre_metni_ayarla(hucreler[3], "", False)
                elif "ilgili" in etiket and "idare" in etiket:
                    self.taahhut_word_hucre_metni_ayarla(hucreler[0], "İlgili İdare", True)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[1], ":", False, WD_ALIGN_PARAGRAPH.CENTER)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[2], ilgili_idare, False)
                    if len(hucreler) > 3:
                        self.taahhut_word_hucre_metni_ayarla(hucreler[3], "", False)
                elif "pafta" in etiket and "parsel" in etiket:
                    self.taahhut_word_hucre_metni_ayarla(hucreler[0], "Pafta/Ada/Parsel No", True)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[1], ":", False, WD_ALIGN_PARAGRAPH.CENTER)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[2], pafta_ada_parsel, False)
                    if len(hucreler) > 3:
                        self.taahhut_word_hucre_metni_ayarla(hucreler[3], "", False)
                elif "yapı adresi" in etiket or "yapi adresi" in etiket:
                    self.taahhut_word_hucre_metni_ayarla(hucreler[0], "Yapı Adresi", True)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[1], ":", False, WD_ALIGN_PARAGRAPH.CENTER)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[2], yapi_adresi, False)
                    if len(hucreler) > 3:
                        self.taahhut_word_hucre_metni_ayarla(hucreler[3], "", False)
                elif "sahibinin" in etiket and ("adresi" in etiket or "adres" in etiket):
                    self.taahhut_word_hucre_metni_ayarla(hucreler[0], "Yapı Sahibinin Adresi", True)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[1], ":", False, WD_ALIGN_PARAGRAPH.CENTER)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[2], yapi_adresi, False)
                    if len(hucreler) > 3:
                        self.taahhut_word_hucre_metni_ayarla(hucreler[3], "", False)
                elif "yapı sahibi" in etiket or "yapi sahibi" in etiket:
                    self.taahhut_word_hucre_metni_ayarla(hucreler[0], "Yapı Sahibi", True)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[1], ":", False, WD_ALIGN_PARAGRAPH.CENTER)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[2], proje_adi, False)
                    if len(hucreler) > 3:
                        self.taahhut_word_hucre_metni_ayarla(hucreler[3], "", False)
                elif "projenin" in etiket and ("türü" in etiket or "turu" in etiket):
                    self.taahhut_word_hucre_metni_ayarla(hucreler[0], "Projenin Türü", True)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[1], ":", False, WD_ALIGN_PARAGRAPH.CENTER)
                    self.taahhut_word_hucre_metni_ayarla(hucreler[2], "ZEMİN ETÜDÜ RAPORU", False)
                    if len(hucreler) > 3:
                        self.taahhut_word_hucre_metni_ayarla(hucreler[3], "", False)

    def taahhut_word_xml_metnini_guncelle(self, doc):
        import re
        from docx.oxml.ns import qn

        yer_tutucular = {
            "[JEOFIZIK_MUH_AD]": self.taahhut_deger("JEOFIZIK_MUH_AD", ""),
            "[JEOFIZIK_MUH_SICIL]": self.taahhut_deger("JEOFIZIK_MUH_SICIL", ""),
            "[JEOFIZIK_MUH_ADRES]": self.taahhut_deger("JEOFIZIK_MUH_ADRES", ""),
            "[JEOFIZIK_MUH_TELEFON]": self.taahhut_deger("JEOFIZIK_MUH_TELEFON", ""),
            "[JEOLOJI_MUH_AD]": self.taahhut_deger("JEOLOJI_MUH_AD", ""),
            "[JEOLOJI_MUH_SICIL]": self.taahhut_deger("JEOLOJI_MUH_SICIL", ""),
            "[JEOLOJI_MUH_ADRES]": self.taahhut_deger("JEOLOJI_MUH_ADRES", ""),
            "[JEOLOJI_MUH_TELEFON]": self.taahhut_deger("JEOLOJI_MUH_TELEFON", ""),
            "[TAAHHUT_TARIHI]": self.taahhut_tarihi(),
        }
        for eski, deger in yer_tutucular.items():
            self.taahhut_word_xml_etiketini_degistir(doc, eski, deger)

        # Etiketsiz eski şablonlarda imza adı ve tarihini, gerçek bir eski
        # kimliği sabit olarak aramadan, meslek unvanının bulunduğu imza
        # bloğuna göre yerleştir. Böylece belgedeki başka tarih ve sayılar
        # değiştirilmez.
        dugumler = list(doc._element.iter(qn("w:t")))
        tarih_deseni = re.compile(r"^\d{2}\.\d{2}\.\d{4}$")
        for index, dugum in enumerate(dugumler):
            metin = dugum.text or ""
            kucuk = metin.casefold()
            rol = next((aday for aday in ("jeofizik", "jeoloji") if aday in kucuk), None)
            if rol is None:
                continue
            sonraki = " ".join(
                (aday.text or "").strip().casefold()
                for aday in dugumler[index:index + 4]
                if (aday.text or "").strip()
            )
            if "hendisi" not in sonraki:
                continue

            bilgi_kodu = "JEOFIZIK_MUH_AD" if rol == "jeofizik" else "JEOLOJI_MUH_AD"
            ad = self.taahhut_deger(bilgi_kodu, "")
            rol_baslangici = kucuk.index(rol)
            ad_dugum_index = index
            if metin[:rol_baslangici].strip():
                sol = metin[:len(metin) - len(metin.lstrip())]
                dugum.text = f"{sol}{ad} {metin[rol_baslangici:].lstrip()}"
            else:
                onceki = index - 1
                while onceki >= 0 and not (dugumler[onceki].text or "").strip():
                    onceki -= 1
                if onceki >= 0:
                    eski_ad = dugumler[onceki].text or ""
                    sol = eski_ad[:len(eski_ad) - len(eski_ad.lstrip())]
                    sag = eski_ad[len(eski_ad.rstrip()):]
                    dugumler[onceki].text = f"{sol}{ad}{sag}"
                    ad_dugum_index = onceki

            kalan = 4
            tarih_index = ad_dugum_index - 1
            while tarih_index >= 0 and kalan:
                aday = (dugumler[tarih_index].text or "").strip()
                if aday:
                    kalan -= 1
                    if tarih_deseni.fullmatch(aday):
                        eski_tarih = dugumler[tarih_index].text or ""
                        sol = eski_tarih[:len(eski_tarih) - len(eski_tarih.lstrip())]
                        sag = eski_tarih[len(eski_tarih.rstrip()):]
                        dugumler[tarih_index].text = f"{sol}{self.taahhut_tarihi()}{sag}"
                        break
                tarih_index -= 1

    def taahhut_word_xml_etiketini_degistir(self, doc, etiket, deger):
        yeni_deger = str(deger)
        for part in doc.part.package.parts:
            partname = str(getattr(part, "partname", ""))
            if not partname.startswith("/word/") or not hasattr(part, "element"):
                continue
            try:
                dugumler = list(part.element.xpath(".//w:t"))
            except Exception:
                continue
            arama_baslangici = 0
            while dugumler:
                parcalar = [dugum.text or "" for dugum in dugumler]
                birlesik = "".join(parcalar)
                baslangic = birlesik.find(etiket, arama_baslangici)
                if baslangic < 0:
                    break
                bitis = baslangic + len(etiket)
                konum = 0
                ilk = son = None
                ilk_ofset = son_ofset = 0
                for index, parca in enumerate(parcalar):
                    sonraki = konum + len(parca)
                    if ilk is None and baslangic < sonraki:
                        ilk = index
                        ilk_ofset = baslangic - konum
                    if ilk is not None and bitis <= sonraki:
                        son = index
                        son_ofset = bitis - konum
                        break
                    konum = sonraki
                if ilk is None or son is None:
                    break
                if ilk == son:
                    dugumler[ilk].text = parcalar[ilk][:ilk_ofset] + yeni_deger + parcalar[ilk][son_ofset:]
                else:
                    dugumler[ilk].text = parcalar[ilk][:ilk_ofset] + yeni_deger
                    for index in range(ilk + 1, son):
                        dugumler[index].text = ""
                    dugumler[son].text = parcalar[son][son_ofset:]
                arama_baslangici = baslangic + len(yeni_deger)

    def taahhut_word_bilgilerini_doldur(self, doc):
        self.taahhut_word_tablolari_doldur(doc)
        self.taahhut_word_proje_tablolarini_hizala(doc)
        self.taahhut_word_tablo_bosluklarini_temizle(doc)
        self.taahhut_word_xml_metnini_guncelle(doc)

    def taahhut_word_metin_dizisini_degistir(self, kok, eski_degerler, yeni_degerler):
        from docx.oxml.ns import qn

        if not eski_degerler or not yeni_degerler:
            return

        dugumler = [
            (dugum, (dugum.text or "").strip())
            for dugum in kok.iter(qn("w:t"))
            if (dugum.text or "").strip()
        ]
        if len(eski_degerler) > len(dugumler):
            return

        i = 0
        while i <= len(dugumler) - len(eski_degerler):
            aday = [deger for _, deger in dugumler[i:i + len(eski_degerler)]]
            if aday == eski_degerler:
                for j in range(len(eski_degerler)):
                    dugum, _ = dugumler[i + j]
                    eski_metin = dugum.text or ""
                    sol_bosluk = eski_metin[:len(eski_metin) - len(eski_metin.lstrip())]
                    sag_bosluk = eski_metin[len(eski_metin.rstrip()):]
                    yeni = yeni_degerler[j] if j < len(yeni_degerler) else ""
                    if j == len(eski_degerler) - 1 and len(yeni_degerler) > len(eski_degerler):
                        yeni = " ".join([yeni, *yeni_degerler[len(eski_degerler):]]).strip()
                    dugum.text = f"{sol_bosluk}{yeni}{sag_bosluk}"
                i += len(eski_degerler)
            else:
                i += 1

    def taahhut_word_ust_bilgi_degerlerini_guncelle(self, kok, tur):
        from docx.oxml.ns import qn

        muhendis = self.taahhut_word_muhendis_bilgileri(tur)
        unvan = "JEOFİZİK" if tur == "jeofizik" else "JEOLOJİ"
        degisimler = {
            "[MUH_SICIL]": muhendis["sicil"],
            "[MUH_UNVAN]": muhendis["unvan"],
            "[MUH_ADRES]": muhendis["adres"],
            "[MUH_TELEFON]": muhendis["telefon"],
        }
        for dugum in kok.iter(qn("w:t")):
            metin = dugum.text or ""
            for eski, yeni in degisimler.items():
                metin = metin.replace(eski, yeni)
            if "JEOLOJİ" in metin or "JEOFİZİK" in metin:
                metin = metin.replace("JEOLOJİ", unvan).replace("JEOFİZİK", unvan)
            dugum.text = metin

    def taahhut_word_ust_bilgi_tablosu_olustur(self, doc, tur):
        muhendis = self.taahhut_word_muhendis_bilgileri(tur)
        tablo = doc.add_table(rows=5, cols=3)
        self.taahhut_word_tablo_stili_uygula(tablo)
        baslik = tablo.cell(0, 0).merge(tablo.cell(0, 2))
        self.taahhut_word_govde_hucre_yaz(
            baslik,
            "TAAHHÜTNAME",
            True,
            WD_ALIGN_PARAGRAPH.CENTER,
        )
        satirlar = (
            ("Oda Sicil No", muhendis["sicil"]),
            ("Unvanı", muhendis["unvan"]),
            ("Adresi", muhendis["adres"]),
            ("Telefonu", muhendis["telefon"]),
        )
        for satir_no, (etiket, deger) in enumerate(satirlar, start=1):
            self.taahhut_word_govde_hucre_yaz(tablo.cell(satir_no, 0), etiket, True)
            self.taahhut_word_govde_hucre_yaz(
                tablo.cell(satir_no, 1),
                ":",
                False,
                WD_ALIGN_PARAGRAPH.CENTER,
            )
            self.taahhut_word_govde_hucre_yaz(tablo.cell(satir_no, 2), deger)
        return tablo

    def taahhut_word_ust_alt_bilgileri_govdeye_tasi(self, doc):
        from docx.oxml.ns import qn

        body = doc._body._element
        aktif_header_var = any(
            child.tag == qn("w:headerReference")
            for sectPr in body.xpath(".//w:sectPr")
            for child in list(sectPr)
        )
        if not aktif_header_var:
            return

        tablolar = [
            tablo
            for tablo in list(doc.tables)
            if len(tablo.rows) >= 7 and len(tablo.columns) >= 3
        ][:2]
        for index, hedef_tablo in enumerate(tablolar):
            tur = "jeofizik" if index == 0 else "jeoloji"
            ust_bilgi = self.taahhut_word_ust_bilgi_tablosu_olustur(doc, tur)
            ust_bilgi_xml = ust_bilgi._tbl
            ust_bilgi_xml.getparent().remove(ust_bilgi_xml)
            hedef_xml = hedef_tablo._tbl
            hedef_xml.getparent().insert(hedef_xml.getparent().index(hedef_xml), ust_bilgi_xml)

        for section in doc.sections:
            for header in (section.header, section.first_page_header, section.even_page_header):
                for child in list(header._element):
                    header._element.remove(child)

        for sectPr in body.xpath(".//w:sectPr"):
            for child in list(sectPr):
                if child.tag == qn("w:headerReference"):
                    sectPr.remove(child)

    def taahhut_word_sablondan_belge_olustur(self):
        doc = Document(self.taahhut_word_sablon_yolu)
        self.taahhut_word_bilgilerini_doldur(doc)
        self.taahhut_word_ust_alt_bilgileri_govdeye_tasi(doc)
        self.taahhut_word_proje_tablolarini_yeniden_ciz(doc)
        self.taahhut_word_belgesini_dogrula(doc)
        return doc

    def taahhut_word_belgesini_dogrula(self, doc):
        import re
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        metin = "\n".join(node.text or "" for node in self.docx_xml_metin_dugumlerini_dolas(doc))
        bitisik_metin = "".join(node.text or "" for node in self.docx_xml_metin_dugumlerini_dolas(doc))
        kalan = sorted(set(re.findall(r"\[[A-Z0-9_]+\]", metin + "\n" + bitisik_metin)))
        if kalan:
            raise ValueError("Taahhütname şablonunda çözülemeyen etiketler var: " + ", ".join(kalan))
        settings = doc.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    def taahhut_word_paragraf_ekle(self, doc, metin="", hizalama=None, punto=10, kalin=False, bosluk_sonra=0):
        paragraf = doc.add_paragraph()
        if hizalama is not None:
            paragraf.alignment = hizalama
        paragraf.paragraph_format.space_after = Pt(bosluk_sonra)
        run = paragraf.add_run(str(metin))
        run.font.name = "Arial"
        run.font.size = Pt(punto)
        run.bold = kalin
        return paragraf

    def taahhut_word_tablo_stili_uygula(self, tablo):
        try:
            tablo.style = "Table Grid"
        except Exception:
            pass
        tablo.alignment = WD_TABLE_ALIGNMENT.CENTER
        for satir in tablo.rows:
            for hucre in satir.cells:
                hucre.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraf in hucre.paragraphs:
                    paragraf.paragraph_format.space_after = Pt(0)
                    for run in paragraf.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(10)

    def taahhut_word_govde_hucre_yaz(self, hucre, metin, kalin=False, hizalama=WD_ALIGN_PARAGRAPH.LEFT):
        paragraf = hucre.paragraphs[0] if hucre.paragraphs else hucre.add_paragraph()
        paragraf.alignment = hizalama
        paragraf.paragraph_format.space_after = Pt(0)
        if paragraf.runs:
            paragraf.runs[0].text = str(metin)
            run = paragraf.runs[0]
            for diger in paragraf.runs[1:]:
                diger.text = ""
        else:
            run = paragraf.add_run(str(metin))
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.bold = kalin

    def taahhut_resmi_metinleri(self):
        ana_metin_1 = (
            "Yukarıdaki bilgilere sahip projenin müellifliğini üstlenmemde 6235 sayılı Türk Mühendis ve Mimar "
            "Odaları Birliği Kanunu, 3194 sayılı İmar Kanunu ve ilgili mevzuat kapsamında süreli veya süresiz "
            "olarak mesleki faaliyet haklarımda herhangi bir kısıtlılık bulunmadığını ve odama üyeliğimin devam "
            "ettiğini taahhüt ederim."
        )
        ana_metin_2 = (
            "Yukarıdaki bilgilere sahip yapıya ilişkin hazırlanacak tüm projelerde, 3194 sayılı Kanun ve deprem, "
            "yangın,enerji verimliliği, asansör gibi ilgili tüm mevzuat hükümlerini eksiksiz uygulayacağımı "
            "taahhüt ederim"
        )
        uyarı_metin = (
            "Gerçeğe aykırı beyanda bulunduğu tespit edilenlerin işlemleri iptal edilecek ve bu kişiler hakkında "
            "5237 sayılı Türk Ceza Kanununun ilgili hükümleri gereği Cumhuriyet Savcılığına suç duyurusunda "
            "bulunulacak, ayrıca 6235 sayılı Türk Mühendis ve Mimar Odaları Birliği Kanunu ve ilgili mevzuatı "
            "uyarınca işlem yapılmak üzere ilgili Meslek Odasına bilgi verilecektir."
        )
        return ana_metin_1, ana_metin_2, uyarı_metin

    def taahhut_word_sayfa_duzenini_uygula(self, doc):
        for section in doc.sections:
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(1.2)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(1.5)
            section.header_distance = Cm(0)
            section.footer_distance = Cm(0)

    def taahhut_word_muhendis_bilgileri(self, tur):
        if tur == "jeofizik":
            return {
                "sicil": self.taahhut_deger("JEOFIZIK_MUH_SICIL", ""),
                "unvan": "JEOFİZİK MÜHENDİSİ",
                "adres": self.taahhut_deger("JEOFIZIK_MUH_ADRES", ""),
                "telefon": self.taahhut_deger("JEOFIZIK_MUH_TELEFON", ""),
                "ad": self.taahhut_deger("JEOFIZIK_MUH_AD", ""),
                "imza_unvan": "Jeofizik Mühendisi",
            }
        return {
            "sicil": self.taahhut_deger("JEOLOJI_MUH_SICIL", ""),
            "unvan": "JEOLOJİ MÜHENDİSİ",
            "adres": self.taahhut_deger("JEOLOJI_MUH_ADRES", ""),
            "telefon": self.taahhut_deger("JEOLOJI_MUH_TELEFON", ""),
            "ad": self.taahhut_deger("JEOLOJI_MUH_AD", ""),
            "imza_unvan": "Jeoloji Mühendisi",
        }

    def taahhut_word_sayfasi_ekle(self, doc, tur):
        muhendis = self.taahhut_word_muhendis_bilgileri(tur)
        ana_metin_1, ana_metin_2, uyarı_metin = self.taahhut_resmi_metinleri()

        self.taahhut_word_paragraf_ekle(doc, "TAAHHÜTNAME", WD_ALIGN_PARAGRAPH.CENTER, 12, True, 6)

        bilgi_tablo = doc.add_table(rows=4, cols=3)
        self.taahhut_word_tablo_stili_uygula(bilgi_tablo)
        bilgi_satirlari = [
            ("Oda Sicil No", muhendis["sicil"]),
            ("Unvanı", muhendis["unvan"]),
            ("Adresi", muhendis["adres"]),
            ("Telefonu", muhendis["telefon"]),
        ]
        for i, (etiket, deger) in enumerate(bilgi_satirlari):
            self.taahhut_word_govde_hucre_yaz(bilgi_tablo.cell(i, 0), etiket, True)
            self.taahhut_word_govde_hucre_yaz(bilgi_tablo.cell(i, 1), ":", False, WD_ALIGN_PARAGRAPH.CENTER)
            self.taahhut_word_govde_hucre_yaz(bilgi_tablo.cell(i, 2), deger)

        self.taahhut_word_paragraf_ekle(doc, "", bosluk_sonra=4)

        proje_tablo = doc.add_table(rows=8, cols=4)
        self.taahhut_word_tablo_stili_uygula(proje_tablo)
        baslik_hucre = proje_tablo.cell(0, 0).merge(proje_tablo.cell(0, 3))
        self.taahhut_word_govde_hucre_yaz(baslik_hucre, "Müellifliği Üstlenilen Proje", True, WD_ALIGN_PARAGRAPH.CENTER)

        il = self.taahhut_yazim_duzeni(self.proje_deger("IL", ""))
        ilce = self.taahhut_yazim_duzeni(self.proje_deger("ILCE", ""))
        proje_adi = self.taahhut_proje_adi_yazim_duzeni()
        yapi_adresi = self.taahhut_yapi_adresi()
        ilgili_idare = self.taahhut_ilgili_idare()
        pafta_ada_parsel = self.taahhut_pafta_ada_parsel_metni()
        proje_satirlari = [
            ("İl / İlçe", " ".join(p for p in (il, ilce) if p), ""),
            ("İlgili İdare", ilgili_idare, ""),
            ("Pafta/Ada/Parsel No", pafta_ada_parsel, ""),
            ("Yapı Adresi", yapi_adresi, ""),
            ("Yapı Sahibi", proje_adi, ""),
            ("Yapı Sahibinin Adresi", yapi_adresi, ""),
            ("Projenin Türü", "ZEMİN ETÜDÜ RAPORU", ""),
        ]
        for satir_no, (etiket, deger, ek_deger) in enumerate(proje_satirlari, start=1):
            self.taahhut_word_govde_hucre_yaz(proje_tablo.cell(satir_no, 0), etiket, True)
            self.taahhut_word_govde_hucre_yaz(proje_tablo.cell(satir_no, 1), ":", False, WD_ALIGN_PARAGRAPH.CENTER)
            self.taahhut_word_govde_hucre_yaz(proje_tablo.cell(satir_no, 2), deger)
            self.taahhut_word_govde_hucre_yaz(proje_tablo.cell(satir_no, 3), ek_deger)

        self.taahhut_word_paragraf_ekle(doc, "", bosluk_sonra=4)
        p1 = self.taahhut_word_paragraf_ekle(doc, ana_metin_1, WD_ALIGN_PARAGRAPH.JUSTIFY, 10, False, 4)
        p1.paragraph_format.first_line_indent = Cm(0.5)
        p2 = self.taahhut_word_paragraf_ekle(doc, ana_metin_2, WD_ALIGN_PARAGRAPH.JUSTIFY, 10, False, 10)
        p2.paragraph_format.first_line_indent = Cm(0.5)

        self.taahhut_word_paragraf_ekle(doc, self.taahhut_tarihi(), WD_ALIGN_PARAGRAPH.RIGHT, 10, False, 0)
        self.taahhut_word_paragraf_ekle(doc, muhendis["ad"], WD_ALIGN_PARAGRAPH.RIGHT, 10, False, 0)
        self.taahhut_word_paragraf_ekle(doc, muhendis["imza_unvan"], WD_ALIGN_PARAGRAPH.RIGHT, 10, False, 10)

        uyari = self.taahhut_word_paragraf_ekle(doc, uyarı_metin, WD_ALIGN_PARAGRAPH.JUSTIFY, 8, False, 0)
        uyari.paragraph_format.keep_together = True

    def taahhut_word_belgesi_olustur(self):
        self.taahhut_bilgilerini_dogrula()
        if self.taahhut_word_sablon_yolu and not os.path.isfile(self.taahhut_word_sablon_yolu):
            raise ValueError(f"Seçili taahhütname Word şablonu bulunamadı: {self.taahhut_word_sablon_yolu}")
        if self.taahhut_word_sablon_yolu:
            return self.taahhut_word_sablondan_belge_olustur()
        doc = Document()
        self.taahhut_word_sayfa_duzenini_uygula(doc)
        self.taahhut_word_sayfasi_ekle(doc, "jeofizik")
        doc.add_page_break()
        self.taahhut_word_sayfasi_ekle(doc, "jeoloji")
        self.taahhut_word_belgesini_dogrula(doc)
        return doc
