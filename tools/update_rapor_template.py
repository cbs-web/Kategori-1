from __future__ import annotations

import os
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "ornek_sablonlar" / "rapor" / "TASLAK.docx"


def paragraph_with(document, marker: str):
    for paragraph in document.paragraphs:
        if marker in paragraph.text:
            return paragraph
    raise ValueError(f"Şablonda paragraf bulunamadı: {marker}")


def paragraph_with_any(document, *markers: str):
    for marker in markers:
        for paragraph in document.paragraphs:
            if marker in paragraph.text:
                return paragraph
    raise ValueError(f"Şablonda paragraf bulunamadı: {' / '.join(markers)}")


def caption_with(document, marker: str, excluded: tuple[str, ...] = ()):
    for paragraph in document.paragraphs:
        style_name = str(getattr(paragraph.style, "name", "") or "").casefold()
        if "caption" not in style_name and "resim yaz" not in style_name:
            continue
        if marker in paragraph.text and not any(value in paragraph.text for value in excluded):
            return paragraph
    raise ValueError(f"Şablonda resim yazısı bulunamadı: {marker}")


def set_text(paragraph, text: str):
    paragraph.clear()
    paragraph.add_run(text)


def next_bookmark_id(document) -> int:
    values = []
    for part in document.part.package.parts:
        if not hasattr(part, "element"):
            continue
        for node in part.element.xpath(".//w:bookmarkStart"):
            value = node.get(qn("w:id"))
            if str(value or "").isdigit():
                values.append(int(value))
    return max(values, default=0) + 1


def add_bookmarked_seq_field(
    paragraph,
    sequence: str,
    cached_value: str,
    name: str,
    bookmark_id: int,
):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    paragraph._p.append(start)

    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    code_run = paragraph.add_run()
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = f" SEQ {sequence} \\* ARABIC "
    code_run._r.append(code)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    paragraph.add_run(cached_value)

    end_run = paragraph.add_run()
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(field_end)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)


def add_ref_field(paragraph, bookmark: str, cached_value: str):
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    code_run = paragraph.add_run()
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = f" REF {bookmark} \\h "
    code_run._r.append(code)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    paragraph.add_run(cached_value)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def set_caption(
    paragraph,
    prefix: str,
    sequence: str,
    cached_value: str,
    suffix: str,
    bookmark: str,
    bookmark_id: int,
):
    paragraph.clear()
    paragraph.style = "Caption"
    paragraph.add_run(prefix)
    add_bookmarked_seq_field(
        paragraph,
        sequence,
        cached_value,
        bookmark,
        bookmark_id,
    )
    paragraph.add_run(suffix)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True


def set_paragraph_with_ref(paragraph, prefix: str, bookmark: str, cached_value: str, suffix: str):
    paragraph.clear()
    paragraph.add_run(prefix)
    add_ref_field(paragraph, bookmark, cached_value)
    paragraph.add_run(suffix)


def update_headers(document):
    first_header = document.sections[0].first_page_header
    if first_header.tables and len(first_header.tables[0].rows[0].cells) >= 2:
        cell = first_header.tables[0].rows[0].cells[1]
        cell.text = (
            "Proje Adı: [PROJE_ADI]\n"
            "İmar Bilgileri: [IL] İli, [ILCE] İlçesi, [KOY] Köyü, "
            "[PAFTA] Pafta, [ADA] Ada, [PARSEL] Parsel"
        )
    for section in document.sections[1:]:
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
        section.first_page_header.is_linked_to_previous = True
        section.first_page_footer.is_linked_to_previous = True
        section.even_page_header.is_linked_to_previous = True
        section.even_page_footer.is_linked_to_previous = True


def update_template():
    document = Document(TEMPLATE)
    bookmark_id = next_bookmark_id(document)

    geology_anchor = paragraph_with(document, "[JEOLOJI_BOLUMU]")
    if not any(p.text.strip() == "2. JEOLOJİ" for p in document.paragraphs):
        heading = geology_anchor.insert_paragraph_before("2. JEOLOJİ", style="Heading 1")
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.keep_together = True

    set_text(
        paragraph_with_any(
            document,
            "Çalışma alanı [BIRIM_TANIMI] üzerindedir.",
            "Çalışma alanı [BIRIM_TANIMI] birimleri üzerinde yer almaktadır.",
        ),
        "Çalışma alanı [BIRIM_TANIMI] birimleri üzerinde yer almaktadır. "
        "[MUHENDISLIK_JEOLOJISI_METNI] [ARAZI_CALISMA_OZETI]",
    )

    numune_caption = caption_with(
        document,
        "Numune Lokasyonları",
        excluded=("Koordinatları",),
    )
    set_caption(
        numune_caption,
        "Şekil ",
        "Şekil",
        "7",
        " Numune Lokasyonları",
        "K1_Sekil_Numune_Lokasyon",
        bookmark_id,
    )
    bookmark_id += 1

    numune_tablo_caption = caption_with(document, "Numune Lokasyonları Koordinatları")
    set_caption(
        numune_tablo_caption,
        "Tablo ",
        "Tablo",
        "3",
        " Numune Lokasyonları Koordinatları",
        "K1_Tablo_Numune_Koordinatlari",
        bookmark_id,
    )
    bookmark_id += 1

    set_paragraph_with_ref(
        paragraph_with(document, "[ARAZI_CALISMA_ACIKLAMA]"),
        "[ARAZI_CALISMA_ACIKLAMA] Araştırma çukuru lokasyonları Şekil ",
        "K1_Sekil_Numune_Lokasyon",
        "7",
        " ile gösterilmiştir.",
    )

    mjh_caption = caption_with(document, "Mühendislik Jeolojisi Haritası")
    set_caption(
        mjh_caption,
        "Şekil ",
        "Şekil",
        "9",
        " Mühendislik Jeolojisi Haritası",
        "K1_Sekil_Muhendislik_Jeolojisi",
        bookmark_id,
    )
    bookmark_id += 1
    muhendislik = paragraph_with_any(
        document,
        "1/100.000 lik jeoloji haritalarında",
        "1/100.000 ölçekli jeoloji haritasında",
    )
    set_paragraph_with_ref(
        muhendislik,
        "İnceleme alanı, 1/100.000 ölçekli jeoloji haritasında "
        "“[FORMASYON_KISA]” simgesiyle gösterilen [BIRIM_TANIMI] birimleri üzerinde yer almaktadır. "
        "[MUHENDISLIK_JEOLOJISI_METNI] Çalışma alanına ait mühendislik jeolojisi haritası Şekil ",
        "K1_Sekil_Muhendislik_Jeolojisi",
        "9",
        " ile verilmiştir.",
    )

    zemin_caption = caption_with(document, "Yerel zemin sınıfları")
    set_caption(
        zemin_caption,
        "Tablo ",
        "Tablo",
        "10",
        ": Yerel zemin sınıfları.",
        "K1_Tablo_Yerel_Zemin_Sinifi",
        bookmark_id,
    )
    set_paragraph_with_ref(
        paragraph_with_any(
            document,
            "Çalışma alanı zemini genel olarak",
            "[YEREL_ZEMIN_SINIFI]",
        ),
        "Çalışma alanı zemini [BIRIM_TANIMI] birimleri üzerinde yer almaktadır. "
        "[VS30_CUMLESI] Türkiye Bina Deprem Yönetmeliği'ne göre Yerel Zemin Sınıfı "
        "“[YEREL_ZEMIN_SINIFI]” olarak belirlenmiştir (Tablo ",
        "K1_Tablo_Yerel_Zemin_Sinifi",
        "10",
        ").",
    )

    set_text(
        paragraph_with_any(
            document,
            "[PROJE_ADI]’ ya ait taşınmazda",
            "statik tasarımına temel oluşturacak zemin parametrelerinin belirlenmesi",
        ),
        "[IL] ili, [ILCE] ilçesi, [KOY] Köyü [ADA] ada, [PARSEL] numaralı parselde, "
        "proje sahibi [PROJE_ADI] adına yapılması planlanan yapının statik tasarımına temel "
        "oluşturacak zemin parametrelerinin belirlenmesi amacıyla bu çalışma yapılmıştır.",
    )
    set_text(
        paragraph_with_any(
            document,
            "[PROJE_ADI]’ ya ait taşınmazda yapılması planlanan yapının projelerine",
            "nihai tasarım parametreleri aşağıda özetlenmiştir",
        ),
        "[IL] ili, [ILCE] ilçesi, [KOY] Köyü [ADA] ada, [PARSEL] numaralı parselde, "
        "proje sahibi [PROJE_ADI] adına yapılması planlanan yapı için belirlenen zemin koşulları "
        "ve nihai tasarım parametreleri aşağıda özetlenmiştir.",
    )
    set_text(
        paragraph_with_any(
            document,
            "İnceleme alanı literatürde",
            "İnceleme alanı [BIRIM_TANIMI] birimleri üzerinde yer almaktadır.",
        ),
        "İnceleme alanı [BIRIM_TANIMI] birimleri üzerinde yer almaktadır. "
        "[MUHENDISLIK_JEOLOJISI_METNI]",
    )
    set_text(
        paragraph_with_any(
            document,
            "Çalışma alanında yeraltı suyuna rastlanmamıştır.",
            "[YASS_SONUC_METNI]",
        ),
        "[YASS_SONUC_METNI] Yüzey ve atık sularının yapı temeline ulaşmasını önlemek için "
        "uygun yalıtım ve drenaj önlemleri alınmalıdır.",
    )
    set_text(
        paragraph_with(document, "Yol, altyapı ve komşu parsel güvenliği"),
        "Yol, altyapı ve komşu parsel güvenliği sağlanmadan kazı yapılmamalı; gerekli kazı "
        "destekleri uygulama projesine göre oluşturulmalıdır.",
    )

    update_headers(document)

    for paragraph in document.paragraphs:
        style_name = str(getattr(paragraph.style, "name", "") or "").casefold()
        if "heading" in style_name or "caption" in style_name:
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True

    with tempfile.NamedTemporaryFile(
        prefix="taslak_guncel_", suffix=".docx", dir=TEMPLATE.parent, delete=False
    ) as temp_file:
        temp_path = Path(temp_file.name)
    try:
        document.save(temp_path)
        Document(temp_path)
        os.replace(temp_path, TEMPLATE)
    finally:
        if temp_path.exists():
            temp_path.unlink()


if __name__ == "__main__":
    update_template()
