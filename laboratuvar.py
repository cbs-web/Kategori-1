import csv
import io
import os
import re
import unicodedata

LAB_AC_KOLONLARI = (
    "No",
    "Derinlik",
    "Çakıl",
    "Kum",
    "Silt+Kil",
    "LL",
    "PL",
    "PI",
    "Wn",
    "DBHA",
    "KBHA",
    "Sınıflama",
    "Kohezyon",
    "Φ",
)

LAB_YN_KOLONLARI = ("No", "Derinlik", "BHA", "IS50")
LAB_AZAMI_DOSYA_BOYUTU = 50 * 1024 * 1024
LAB_AZAMI_SATIR_SAYISI = 10_000
LAB_DESTEKLENEN_UZANTILAR = {".xlsx", ".xls", ".csv"}
LAB_PANO_AZAMI_KARAKTER = 5 * 1024 * 1024
LAB_1_KOLON_SAYISI = 29


class EksikLaboratuvarSutunu(ValueError):
    pass


def sayi_formatla(val, hane=2):
    v_str = str(val).strip()
    if v_str.lower() in ["-", "", "nan", "none"]:
        return "-"
    try:
        f_val = float(v_str.replace(",", "."))
        if f_val.is_integer():
            return str(int(f_val))
        return str(round(f_val, hane))
    except ValueError:
        return v_str


def _hucre_bos_mu(value):
    return str(value if value is not None else "").strip().casefold() in {
        "",
        "-",
        "nan",
        "none",
    }


def _hucre_metni(value, default="-"):
    if _hucre_bos_mu(value):
        return default
    return str(value).strip()


def _sayisal_deger(value):
    if _hucre_bos_mu(value):
        return None
    text = str(value).strip().replace(" ", "").replace(",", ".")
    try:
        return float(text)
    except (TypeError, ValueError):
        return None


def _silt_kil_toplami(silt, kil):
    values = [value for value in (_sayisal_deger(silt), _sayisal_deger(kil)) if value is not None]
    if not values:
        return "-"
    return sayi_formatla(sum(values), 2)


def _lab_no_normalize(value):
    text = str(value or "").strip().casefold().replace("ı", "i")
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return re.sub(r"[^a-z0-9]+", "", text).upper()


def laboratuvar_numune_anahtari(value):
    """AÇ1, AÇ-1, AÇ 1 ve AC1 için tek teknik numune anahtarı üret."""
    return _lab_no_normalize(value)


def laboratuvar_numune_etiketlerini_uyarla(satirlar, bilinen_etiketler):
    """Gelen satırlardaki numune adlarını projedeki görünen etiketlere uyarla.

    Örneğin projede ``AÇ1`` bulunurken laboratuvar dosyası ``AÇ-1``
    getirirse teknik anahtarları aynı olduğundan satırda projenin ``AÇ1``
    yazımı korunur.
    """
    etiket_by_anahtar = {}
    for etiket in bilinen_etiketler or []:
        anahtar = laboratuvar_numune_anahtari(etiket)
        if anahtar and anahtar not in etiket_by_anahtar:
            etiket_by_anahtar[anahtar] = etiket

    sonuc = []
    for satir in satirlar or []:
        degerler = list(satir)
        if degerler:
            anahtar = laboratuvar_numune_anahtari(degerler[0])
            if anahtar in etiket_by_anahtar:
                degerler[0] = etiket_by_anahtar[anahtar]
        sonuc.append(tuple(degerler))
    return sonuc


def _lab_kayit_tipi(no):
    normalized = _lab_no_normalize(no)
    if normalized.startswith("AC"):
        return "ac"
    if normalized.startswith("YN"):
        return "yn"
    return ""


def _lab1_satiri_donustur(hucreler):
    """LAB_1 A:AC satırını K-1 AÇ veya YN satırına dönüştürür."""
    cells = list(hucreler or [])
    if len(cells) < LAB_1_KOLON_SAYISI:
        cells.extend([""] * (LAB_1_KOLON_SAYISI - len(cells)))

    no = _hucre_metni(cells[0], default="") or _hucre_metni(cells[1], default="")
    derinlik = _hucre_metni(cells[2], default="")
    kayit_tipi = _lab_kayit_tipi(no)
    if not kayit_tipi or not derinlik:
        return None

    if kayit_tipi == "yn":
        return "yn", (
            no,
            derinlik,
            sayi_formatla(cells[11], 3),
            sayi_formatla(cells[28], 2),
        )

    return "ac", (
        no,
        derinlik,
        sayi_formatla(cells[3], 2),
        sayi_formatla(cells[4], 2),
        _silt_kil_toplami(cells[5], cells[6]),
        sayi_formatla(cells[7], 2),
        sayi_formatla(cells[8], 2),
        sayi_formatla(cells[9], 2),
        sayi_formatla(cells[10], 2),
        sayi_formatla(cells[11], 3),
        sayi_formatla(cells[12], 3),
        _hucre_metni(cells[14]),
        sayi_formatla(cells[20], 2),
        sayi_formatla(cells[21], 2),
    )


def _lab1_satirlari_donustur(satirlar):
    ac_satirlari = []
    yn_satirlari = []
    atlanan = 0
    for cells in satirlar:
        if not cells or all(_hucre_bos_mu(cell) for cell in cells):
            continue
        converted = _lab1_satiri_donustur(cells)
        if converted is None:
            atlanan += 1
            continue
        target, row = converted
        if target == "yn":
            yn_satirlari.append(row)
        else:
            ac_satirlari.append(row)
    return {
        "ac_satirlari": ac_satirlari,
        "yn_satirlari": yn_satirlari,
        "eklenen_ac": len(ac_satirlari),
        "eklenen_yn": len(yn_satirlari),
        "atlanan": atlanan,
    }


def _lab1_duzeni_mi(satirlar):
    for cells in satirlar:
        if len(cells) < LAB_1_KOLON_SAYISI:
            continue
        first = _lab_no_normalize(cells[0] if cells else "")
        second = _lab_no_normalize(cells[1] if len(cells) > 1 else "")
        third = _lab_no_normalize(cells[2] if len(cells) > 2 else "")
        if (
            first.startswith(("AC", "YN", "SONDAJNO", "BORINGNO"))
            or second.startswith(("AC", "YN", "NUMUNENO", "SAMPLENO"))
            or third.startswith(("DERINLIK", "DEPTH"))
        ):
            return True
    return False


def laboratuvar_pano_metnini_ayristir(metin):
    """Excel panosunu boş hücreleri koruyan satır/sütun matrisine dönüştürür."""
    text = str(metin or "")
    if len(text) > LAB_PANO_AZAMI_KARAKTER:
        raise ValueError("Panodaki laboratuvar verisi 5 MB güvenlik sınırını aşıyor.")
    try:
        rows = list(csv.reader(io.StringIO(text), delimiter="\t", quotechar='"'))
    except csv.Error as exc:
        raise ValueError(f"Excel pano verisi okunamadı: {exc}") from exc
    rows = [[str(cell).strip() for cell in row] for row in rows]
    rows = [row for row in rows if row and not all(_hucre_bos_mu(cell) for cell in row)]
    if len(rows) > LAB_AZAMI_SATIR_SAYISI:
        raise ValueError(f"Panodan en fazla {LAB_AZAMI_SATIR_SAYISI} veri satırı aktarılabilir.")
    return rows


def laboratuvar_pano_verisini_donustur(metin, hedef):
    """LAB_1 veya eski konumsal pano düzenini K-1 satırlarına dönüştürür."""
    rows = laboratuvar_pano_metnini_ayristir(metin)
    if not rows:
        return {
            "format": "bos",
            "ac_satirlari": [],
            "yn_satirlari": [],
            "standart_satirlar": [],
            "atlanan": 0,
        }
    if _lab1_duzeni_mi(rows):
        result = _lab1_satirlari_donustur(rows)
        result.update({"format": "lab1", "standart_satirlar": []})
        return result

    if hedef not in {"ac", "yn"}:
        raise ValueError("Laboratuvar pano hedefi AÇ veya YN olmalıdır.")
    width = len(LAB_AC_KOLONLARI) if hedef == "ac" else len(LAB_YN_KOLONLARI)
    normalized_rows = []
    for row in rows:
        cells = list(row[:width])
        if len(cells) < width:
            cells.extend(["-"] * (width - len(cells)))
        normalized_rows.append(tuple(cells))
    return {
        "format": f"standart_{hedef}",
        "ac_satirlari": [],
        "yn_satirlari": [],
        "standart_satirlar": normalized_rows,
        "atlanan": 0,
    }


def _lab_derinlik_anahtari(value):
    text = str(value or "").strip().replace(",", ".")
    matches = re.fullmatch(
        r"\s*(-?\d+(?:\.\d+)?)\s*[-–—]\s*(-?\d+(?:\.\d+)?)\s*",
        text,
    )
    if matches:
        return f"{float(matches.group(1)):.8f}:{float(matches.group(2)):.8f}"
    return re.sub(r"\s+", "", text).casefold()


def laboratuvar_satirlarini_birlestir(mevcut_satirlar, gelen_satirlar):
    """No+Derinlik anahtarıyla satırları sırasını koruyarak ekler veya günceller."""
    result = [tuple(row) for row in (mevcut_satirlar or [])]
    index_by_key = {}
    for index, row in enumerate(result):
        if len(row) >= 2:
            key = (laboratuvar_numune_anahtari(row[0]), _lab_derinlik_anahtari(row[1]))
            if key[0] and key[1] and key not in index_by_key:
                index_by_key[key] = index

    added = 0
    updated = 0
    repeated = 0
    incoming_seen = set()
    for incoming in gelen_satirlar or []:
        row = tuple(incoming)
        if len(row) < 2:
            continue
        key = (laboratuvar_numune_anahtari(row[0]), _lab_derinlik_anahtari(row[1]))
        if key in incoming_seen:
            repeated += 1
        incoming_seen.add(key)
        if key[0] and key[1] and key in index_by_key:
            existing = list(result[index_by_key[key]])
            incoming_values = list(row)
            if existing and str(existing[0] or "").strip():
                # Laboratuvar Excel'i AÇ-1 getirse de kullanıcıdaki görünen
                # etiketi (ör. AÇ1) mümkün olduğunca koru.
                incoming_values[0] = existing[0]
            result[index_by_key[key]] = tuple(incoming_values)
            updated += 1
        else:
            result.append(row)
            if key[0] and key[1]:
                index_by_key[key] = len(result) - 1
            added += 1
    return {
        "satirlar": result,
        "eklenen": added,
        "guncellenen": updated,
        "pano_tekrari": repeated,
    }


def _baslik_normalize(deger):
    metin = str(deger or "").strip().upper()
    metin = metin.translate(str.maketrans({"İ": "I", "I": "I", "Ş": "S", "Ğ": "G", "Ü": "U", "Ö": "O", "Ç": "C"}))
    metin = unicodedata.normalize("NFKD", metin)
    metin = "".join(ch for ch in metin if not unicodedata.combining(ch))
    return " ".join(re.findall(r"[A-Z0-9ΦΓ]+", metin))


def _aday_eslesme_puani(aday, takma_adlar):
    aday = _baslik_normalize(aday)
    if not aday:
        return 0
    aday_tokens = set(aday.split())
    en_iyi = 0
    for takma_ad in takma_adlar:
        norm = _baslik_normalize(takma_ad)
        if not norm:
            continue
        if aday == norm:
            en_iyi = max(en_iyi, 100)
            continue
        alias_tokens = norm.split()
        if len(alias_tokens) > 1 and set(alias_tokens).issubset(aday_tokens):
            en_iyi = max(en_iyi, 80)
        elif len(norm) >= 4 and norm in aday_tokens:
            en_iyi = max(en_iyi, 70)
    return en_iyi


def _benzersiz_sutun_bul(df, takma_adlar, kullanilan_indeksler):
    """Başlık öncelikli, benzersiz ve kısa-kodlarda tam eşleşmeli kolon seçer."""
    en_iyi = None
    for indeks, kolon in enumerate(df.columns):
        if indeks in kullanilan_indeksler:
            continue
        adaylar = [(str(kolon), 1000)]
        for satir_no in range(min(4, len(df))):
            hucre = df.iloc[satir_no, indeks]
            if hucre is None or str(hucre).strip().casefold() == "nan":
                continue
            hucre_metni = str(hucre).strip()
            if hucre_metni and len(hucre_metni) <= 80:
                adaylar.append((hucre_metni, 100))
        for aday, kaynak_onceligi in adaylar:
            puan = _aday_eslesme_puani(aday, takma_adlar)
            if not puan:
                continue
            toplam = kaynak_onceligi + puan
            secim = (toplam, -indeks, indeks, kolon)
            if en_iyi is None or secim[:2] > en_iyi[:2]:
                en_iyi = secim
    if en_iyi is None:
        return None
    kullanilan_indeksler.add(en_iyi[2])
    return en_iyi[3]


def laboratuvar_dosyasi_oku(dosya_yolu):
    # pandas yalnız Excel/CSV dosyası gerçekten okunduğunda yüklenir.
    import pandas as pd

    uzanti = os.path.splitext(str(dosya_yolu))[1].lower()
    if uzanti not in LAB_DESTEKLENEN_UZANTILAR:
        raise ValueError("Laboratuvar dosyası .xlsx, .xls veya .csv biçiminde olmalıdır.")
    dosya_boyutu = os.path.getsize(dosya_yolu)
    if dosya_boyutu > LAB_AZAMI_DOSYA_BOYUTU:
        raise ValueError("Laboratuvar dosyası 50 MB sınırını aşıyor.")
    if uzanti == ".csv":
        son_hata = None
        for kodlama in ("utf-8-sig", "cp1254", "latin-1"):
            try:
                df = pd.read_csv(
                    dosya_yolu,
                    sep=None,
                    engine="python",
                    encoding=kodlama,
                    nrows=LAB_AZAMI_SATIR_SAYISI + 1,
                )
                break
            except UnicodeDecodeError as exc:
                son_hata = exc
        else:
            raise ValueError("CSV dosyasının metin kodlaması okunamadı.") from son_hata
    else:
        df = pd.read_excel(dosya_yolu, nrows=LAB_AZAMI_SATIR_SAYISI + 1)
    if len(df) > LAB_AZAMI_SATIR_SAYISI:
        raise ValueError(f"Laboratuvar dosyası en fazla {LAB_AZAMI_SATIR_SAYISI} veri satırı içerebilir.")
    return laboratuvar_dataframe_satirlari(df)


def laboratuvar_dataframe_satirlari(df):
    df = df.fillna("-")
    lab1_rows = [list(df.columns)] + df.values.tolist()
    if _lab1_duzeni_mi(lab1_rows):
        # Sütun konumları LAB_1 şablonunda anlamlıdır; çok katlı başlıklar
        # genel başlık eşleştiricisine bırakılırsa boş "Numune No" seçilebilir.
        return _lab1_satirlari_donustur(df.values.tolist())

    kullanilan = set()

    def bul(*takma_adlar):
        return _benzersiz_sutun_bul(df, takma_adlar, kullanilan)

    col_no = bul("NO", "NUMUNE", "NUMUNE NO", "SAMPLE", "SAMPLE NO", "KUYU", "KUYU NO", "CUKUR", "CUKUR NO", "ID")
    col_derinlik = bul("DERINLIK", "DERINLIK M", "DEPTH", "DEPTH M", "METRE")
    col_cakil = bul("CAKIL", "GRAVEL", "CAKIL YUZDESI")
    col_kum = bul("KUM", "SAND", "KUM YUZDESI")
    col_ince = bul("INCE", "FINE", "SILT KIL", "SILT+KIL")
    col_silt = bul("SILT", "SILT YUZDESI")
    col_kil = bul("KIL", "CLAY", "KIL YUZDESI")
    col_ll = bul("LL", "LIKIT LIMIT", "LIQUID LIMIT")
    col_pl = bul("PL", "PLASTIK LIMIT", "PLASTIC LIMIT")
    col_pi = bul("PI", "PLASTISITE INDEKSI", "PLASTICITY INDEX")
    col_wn = bul("WN", "SU ICERIGI", "DOGAL SU ICERIGI", "SU MUHTEVASI", "NEM", "WATER CONTENT", "MOISTURE")
    col_sinif = bul("USCS", "SINIFLAMA", "SINIF", "CLASSIFICATION", "CLASS")
    col_dbha = bul("DBHA", "DOGAL BHA", "DOGAL BIRIM HACIM AGIRLIK", "BHA", "ΓN")
    col_kbha = bul("KBHA", "KURU BHA", "KURU BIRIM HACIM AGIRLIK", "ΓK")
    col_kohezyon = bul("C", "KOH", "KOHEZYON", "COHESION")
    col_phi = bul("Φ", "PHI", "FI", "SURTUNME ACISI")
    col_is50 = bul("IS50", "IS 50", "IS(50)", "NOKTA YUK INDEKSI", "POINT LOAD INDEX")

    if col_no is None or col_derinlik is None:
        raise EksikLaboratuvarSutunu(
            "Yüklenen dosyada 'Numune/Kuyu No' veya 'Derinlik' sütunu bulunamadı. "
            "Lütfen başlıklarınızı kontrol edin."
        )

    ac_satirlari = []
    yn_satirlari = []

    for _, row in df.iterrows():
        val_no = str(row[col_no]).strip()
        val_der = str(row[col_derinlik]).strip()

        if val_no == "-" or val_no.lower() == "nan" or "numune" in val_no.lower() or val_no.lower() == "no":
            continue

        val_ince = "-"
        if col_ince is not None and str(row[col_ince]).strip().lower() not in ["", "nan", "none", "-"]:
            val_ince = sayi_formatla(row[col_ince], 2)
        else:
            s_val, k_val, has_val = 0.0, 0.0, False
            if col_silt is not None and str(row[col_silt]).strip().lower() not in ["", "nan", "none", "-"]:
                try:
                    s_val = float(str(row[col_silt]).replace(",", "."))
                    has_val = True
                except ValueError:
                    pass
            if col_kil is not None and str(row[col_kil]).strip().lower() not in ["", "nan", "none", "-"]:
                try:
                    k_val = float(str(row[col_kil]).replace(",", "."))
                    has_val = True
                except ValueError:
                    pass
            if has_val:
                val_ince = sayi_formatla(s_val + k_val, 2)

        ac_satiri = [
            val_no,
            val_der,
            sayi_formatla(row[col_cakil], 2) if col_cakil is not None else "-",
            sayi_formatla(row[col_kum], 2) if col_kum is not None else "-",
            val_ince,
            sayi_formatla(row[col_ll], 2) if col_ll is not None else "-",
            sayi_formatla(row[col_pl], 2) if col_pl is not None else "-",
            sayi_formatla(row[col_pi], 2) if col_pi is not None else "-",
            sayi_formatla(row[col_wn], 2) if col_wn is not None else "-",
            sayi_formatla(row[col_dbha], 3) if col_dbha is not None else "-",
            sayi_formatla(row[col_kbha], 3) if col_kbha is not None else "-",
            str(row[col_sinif]).strip() if col_sinif is not None and str(row[col_sinif]).lower() != "nan" else "-",
            sayi_formatla(row[col_kohezyon], 2) if col_kohezyon is not None else "-",
            sayi_formatla(row[col_phi], 2) if col_phi is not None else "-",
        ]

        if _baslik_normalize(val_no).replace(" ", "").startswith("YN"):
            bha = sayi_formatla(row[col_dbha], 3) if col_dbha is not None else "-"
            is50 = sayi_formatla(row[col_is50], 2) if col_is50 is not None else "-"
            yn_satirlari.append((val_no, val_der, bha, is50))
        else:
            ac_satirlari.append(ac_satiri)

    return {
        "ac_satirlari": ac_satirlari,
        "yn_satirlari": yn_satirlari,
        "eklenen_ac": len(ac_satirlari),
        "eklenen_yn": len(yn_satirlari),
    }


def rapor_lab_tablosu_olustur(doc, tree, kolon_map, tabloyu_ortala, tablo_stili_uygula=None):
    from docx.shared import Pt

    satirlar = tree.get_children()
    if not satirlar:
        return None

    kolonlar = tree["columns"]
    tablo = doc.add_table(rows=1, cols=len(kolonlar))
    tablo.style = "Table Grid"
    for c_idx, col in enumerate(kolonlar):
        cell = tablo.rows[0].cells[c_idx]
        cell.text = kolon_map.get(col, col)
        for pr in cell.paragraphs:
            for run in pr.runs:
                run.font.size = Pt(10)

    for row_id in satirlar:
        degerler = tree.item(row_id)["values"]
        r = tablo.add_row()
        for c_idx, val in enumerate(degerler):
            cell = r.cells[c_idx]
            cell.text = str(val) if val else "-"
            for pr in cell.paragraphs:
                for run in pr.runs:
                    run.font.size = Pt(10)

    if tablo_stili_uygula:
        tablo_stili_uygula(tablo, header_rows=1)
    else:
        tabloyu_ortala(tablo)
    return tablo


def rapor_lab_tablolarini_ekle(app, doc):
    lab_p = app.rapor_paragraf_bul(doc, "[LAB]")
    if not lab_p:
        return

    app.rapor_metin_degistir(doc, "[LAB]", "")
    anchor = lab_p._p

    if hasattr(app, "tree_lab_ac"):
        kolon_map = {
            "DBHA": "γn (gr/cm^3)",
            "KBHA": "γk (gr/cm^3)",
            "Kohezyon": "c (kPa)",
            "Derinlik": "Derinlik (m)",
        }
        tablo = rapor_lab_tablosu_olustur(
            doc,
            app.tree_lab_ac,
            kolon_map,
            app.rapor_tabloyu_ortala,
            getattr(app, "rapor_tablo_stili_uygula", None),
        )
        if tablo is not None:
            anchor = app.rapor_xml_sonrasina_ekle(anchor, tablo._tbl)
            bosluk_p = doc.add_paragraph("")
            anchor = app.rapor_xml_sonrasina_ekle(anchor, bosluk_p._p)

    if hasattr(app, "tree_lab_yn"):
        kolon_map = {"Derinlik": "Derinlik (m)"}
        tablo = rapor_lab_tablosu_olustur(
            doc,
            app.tree_lab_yn,
            kolon_map,
            app.rapor_tabloyu_ortala,
            getattr(app, "rapor_tablo_stili_uygula", None),
        )
        if tablo is not None:
            anchor = app.rapor_xml_sonrasina_ekle(anchor, tablo._tbl)
            bosluk_p = doc.add_paragraph("")
            app.rapor_xml_sonrasina_ekle(anchor, bosluk_p._p)
