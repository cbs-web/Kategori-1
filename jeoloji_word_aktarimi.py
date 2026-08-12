"""Eski Word raporlarından künye ve jeoloji içeriği çıkarma araçları."""

from __future__ import annotations

from dataclasses import asdict, dataclass
import os
from pathlib import Path
import re
import stat
import unicodedata
from zipfile import BadZipFile, ZipFile, is_zipfile

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from jeoloji_bolum_paketi import jeoloji_sinir_basligi_mi


# KATEGORI_1'in 2. sekmesindeki katalogla aynÄ± sÄ±ra ve yazÄ±m kullanÄ±lÄ±r.
FORMASYON_SECILMEDI = "Seçiniz..."
FORMASYON_ADLARI = (
    FORMASYON_SECILMEDI,
    "Ayvacık Volkaniti (Tmay)",
    "Ezine Volkaniti (Tme)",
    "Hüseyinfakı Volkaniti (Tmhü)",
    "Arıklı İgnimbiriti (Tmar)",
    "İlyasbaşı Formasyonu (Tmi)",
    "Çamkabalak İgnimbiriti (Tmç)",
    "Çetmi Melanjı (Kç)",
    "Üst Oligosen-Alt Miyosen Granitoyitleri (Tg)",
    "Şahinli Formasyonu (Teşa)",
    "Bayramiç Formasyonu (Tplb)",
    "Çanakkale Formasyonu (Tmçk)",
    "Kirazlı Üyesi (Tmki)",
    "Çamrakdere Üyesi (Tmçd)",
    "Alçıtepe Üyesi (Tmal)",
    "Tüf Üyesi (Tmçt)",
)


_BOS_DEGERLER = {
    "",
    "-",
    "—",
    "il",
    "ilce",
    "ilçesi",
    "mahalle",
    "mahallesi",
    "koy",
    "köy",
    "ada",
    "parsel",
}

_KUNYE_ETIKETLERI = {
    "il": ("il", "ili"),
    "ilce": ("ilce", "ilcesi", "ilce adi"),
    "yerlesim": (
        "mah",
        "mahalle",
        "mahallesi",
        "mahalle adi",
        "koy",
        "koyu",
        "koy adi",
        "koy mahalle",
        "koyu mahallesi",
        "koy mahalle adi",
        "yerlesim",
        "yerlesim yeri",
    ),
    "ada": ("ada", "ada no", "ada nosu", "ada numarasi"),
    "parsel": ("parsel", "parsel no", "parsel nosu", "parsel numarasi"),
}

_BIRLESIK_ADA_PARSEL_ETIKETLERI = {
    "ada parsel",
    "ada parsel no",
    "ada parsel numarasi",
}


def _anahtar(value):
    text = unicodedata.normalize("NFKD", str(value or "").strip().casefold())
    text = "".join(char for char in text if not unicodedata.combining(char))
    text = text.replace("ı", "i")
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return " ".join(text.split())


def _temiz_metin(value):
    text = str(value or "").replace("\xa0", " ").strip()
    return re.sub(r"[ \t]+", " ", text)


def _gercek_deger(value):
    text = _temiz_metin(value).strip(" :;|\t")
    key = _anahtar(text)
    if not text or key in _BOS_DEGERLER:
        return ""
    if re.fullmatch(r"\[[^\]]+\]|\{[^}]+\}|<[^>]+>", text):
        return ""
    if key in {"il", "ilce", "mahalle", "koy", "ada", "parsel", "formasyon"}:
        return ""
    return text


def _kunye_degeri(field, value):
    text = _gercek_deger(value)
    if not text or _etiket_alani(text) or _anahtar(text) in _BIRLESIK_ADA_PARSEL_ETIKETLERI:
        return ""
    if field in ("ada", "parsel"):
        compact = re.sub(r"\s+", "", text)
        if not re.fullmatch(r"[0-9]+[A-Za-z]?(?:[/,\-][0-9]+[A-Za-z]?)*", compact):
            return ""
        return compact
    return text


def _etiket_alani(value):
    key = _anahtar(value)
    for field, labels in _KUNYE_ETIKETLERI.items():
        if key in labels:
            return field
    return ""


def _birlesik_ada_parsel(value):
    text = _temiz_metin(value)
    match = re.match(r"^\s*([^/]+?)\s*/\s*(.+?)\s*$", text)
    if not match:
        match = re.match(
            r"^\s*([0-9]+[A-Za-z]?)\s+ada\s*[,;\-]?\s*"
            r"([0-9]+[A-Za-z]?(?:[/,-][0-9A-Za-z]+)*)\s+parsel",
            text,
            re.IGNORECASE,
        )
    if not match:
        return "", ""
    return _kunye_degeri("ada", match.group(1)), _kunye_degeri("parsel", match.group(2))


def _stil_baslik_seviyesi(paragraph):
    style_name = _anahtar(getattr(getattr(paragraph, "style", None), "name", ""))
    match = re.search(r"(?:heading|baslik)\s*(\d+)", style_name)
    if match:
        return int(match.group(1))
    return None


def _tahmini_baslik_seviyesi(text):
    clean = _temiz_metin(text)
    if not clean or len(clean) > 180:
        return None
    numbered = re.match(r"^(\d{1,3}(?:\.\d{1,3})*)[.)]?\s+", clean)
    if numbered:
        return numbered.group(1).count(".") + 1
    letters = [char for char in clean if char.isalpha()]
    if letters and len(clean.split()) <= 14:
        upper_ratio = sum(char.isupper() for char in letters) / len(letters)
        if upper_ratio >= 0.82:
            return 1
    return None


def _baslik_seviyesi(paragraph):
    return _stil_baslik_seviyesi(paragraph) or _tahmini_baslik_seviyesi(paragraph.text)


def _baslik_mi(paragraph):
    return _baslik_seviyesi(paragraph) is not None


def _atlanacak_paragraf(paragraph):
    text = _temiz_metin(paragraph.text)
    key = _anahtar(text)
    style = _anahtar(getattr(getattr(paragraph, "style", None), "name", ""))
    if not text or style.startswith("toc") or style.startswith("icindekiler"):
        return True
    if re.fullmatch(r"\[(?:sekil|resim|harita|figure)[^\]]*\]", text, re.IGNORECASE):
        return True
    return False


def _yer_tutucu_orani(text):
    clean = _temiz_metin(text)
    if not clean:
        return 1.0
    placeholders = re.findall(r"\[[^\]]+\]|\{[^}]+\}|<[^>]+>", clean)
    return sum(len(item) for item in placeholders) / max(1, len(clean))


def _belge_bloklari(document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _tablo_satirlari(table):
    result = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            text = _temiz_metin(cell.text)
            if text and (not cells or cells[-1] != text):
                cells.append(text)
        if cells:
            result.append(" | ".join(cells))
    return result


def _tum_belge_metni(document, *, ana_basligi_atla=False):
    """Ayrılmış jeoloji belgelerinin paragraf ve tablo metnini sırasıyla döndürür."""
    content = []
    main_heading_skipped = False
    for block in _belge_bloklari(document):
        if isinstance(block, Table):
            content.extend(_tablo_satirlari(block))
            continue
        if _atlanacak_paragraf(block):
            continue
        text = _temiz_metin(block.text)
        if ana_basligi_atla and not main_heading_skipped and _ana_jeoloji_basligi(_anahtar(text)):
            main_heading_skipped = True
            continue
        content.append(text)
    return "\n\n".join(item for item in content if item).strip()


def _ayrilmis_jeoloji_belgesi(document):
    """Belgenin tamamının bir köy jeoloji bölümü gibi görünüp görünmediğini sınar."""
    paragraphs = []
    for block in _belge_bloklari(document):
        if not isinstance(block, Paragraph) or _atlanacak_paragraf(block):
            continue
        paragraphs.append(_anahtar(block.text))
        if len(paragraphs) >= 4:
            break
    if not paragraphs:
        return False
    first = paragraphs[0]
    return bool(
        _ana_jeoloji_basligi(first)
        or any(
            marker in first
            for marker in (
                "bolgesel jeoloji",
                "stratigraf",
                "yapisal jeoloji",
                "aktif tektonik",
            )
        )
    )


def _dosya_adindan_kunye(path):
    """``İlçe_Yerleşim.docx`` biçimindeki güvenli dosya adı yedeğini çözer."""
    stem = _temiz_metin(Path(path).stem)
    if "_" not in stem:
        return {"ilce": "", "yerlesim": ""}
    ilce, yerlesim = stem.split("_", 1)
    ilce = _temiz_metin(ilce.replace("_", " ").replace("-", " "))
    yerlesim = _temiz_metin(yerlesim.replace("_", " "))

    def yer_adi_mi(value):
        letters = [char for char in value if char.isalpha()]
        return (
            len(letters) >= 2
            and letters[0].isupper()
            and not any(char.isdigit() for char in value)
        )

    if not yer_adi_mi(ilce) or not yer_adi_mi(yerlesim):
        return {"ilce": "", "yerlesim": ""}
    return {"ilce": ilce, "yerlesim": yerlesim}


def _bolum_adaylari(document, heading_test, nested_heading_test=None):
    blocks = list(_belge_bloklari(document))
    candidates = []
    for index, paragraph in enumerate(blocks):
        if not isinstance(paragraph, Paragraph):
            continue
        heading = _temiz_metin(paragraph.text)
        style = _anahtar(getattr(getattr(paragraph, "style", None), "name", ""))
        if style.startswith("toc") or style.startswith("icindekiler"):
            continue
        if not heading or not heading_test(_anahtar(heading)):
            continue
        explicit_level = _stil_baslik_seviyesi(paragraph)
        level = explicit_level or _tahmini_baslik_seviyesi(heading) or 1
        content = []
        for following in blocks[index + 1 :]:
            if isinstance(following, Table):
                content.extend(_tablo_satirlari(following))
                continue
            following_level = _stil_baslik_seviyesi(following)
            following_key = _anahtar(following.text)
            is_nested = bool(nested_heading_test and nested_heading_test(following_key))
            if jeoloji_sinir_basligi_mi(
                _temiz_metin(following.text),
                level,
                style_level=following_level,
                jeoloji_alt_basligi=is_nested,
            ):
                break
            if _atlanacak_paragraf(following):
                continue
            content.append(_temiz_metin(following.text))
        text = "\n\n".join(item for item in content if item).strip()
        score = len(text) - int(4000 * _yer_tutucu_orani(text))
        if re.search(r"\[(?:FORMASYON|GENEL_JEOLOJI|INCELEME_ALANI_JEOLOJISI)\]", text, re.IGNORECASE):
            score -= 1000
        if len(text.split()) < 4:
            score -= 500
        candidates.append((score, index, text))
    return candidates


def _en_iyi_bolum(document, heading_test, nested_heading_test=None):
    candidates = _bolum_adaylari(document, heading_test, nested_heading_test=nested_heading_test)
    if not candidates:
        return ""
    best = max(candidates, key=lambda item: (item[0], -item[1]))
    return best[2] if best[0] > 0 else ""


def _metadata_metinleri(document):
    lines = [_temiz_metin(paragraph.text) for paragraph in document.paragraphs]
    table_rows = []
    for table in document.tables:
        for row in table.rows:
            cells = [_temiz_metin(cell.text) for cell in row.cells]
            table_rows.append(cells)
            lines.extend(cells)
    for section in document.sections:
        for container in (section.header, section.footer):
            lines.extend(_temiz_metin(paragraph.text) for paragraph in container.paragraphs)
            for table in container.tables:
                for row in table.rows:
                    lines.extend(_temiz_metin(cell.text) for cell in row.cells)
    return [line for line in lines if line], table_rows


def _kunye_bilgileri(document):
    result = {field: "" for field in _KUNYE_ETIKETLERI}
    lines, table_rows = _metadata_metinleri(document)

    for cells in table_rows:
        for index, cell in enumerate(cells):
            if _anahtar(cell.rstrip(" :")) in _BIRLESIK_ADA_PARSEL_ETIKETLERI:
                for candidate in cells[index + 1 :]:
                    ada, parsel = _birlesik_ada_parsel(candidate)
                    if ada and parsel:
                        result["ada"] = result["ada"] or ada
                        result["parsel"] = result["parsel"] or parsel
                        break
            field = _etiket_alani(cell.rstrip(" :"))
            if field and not result[field]:
                for candidate in cells[index + 1 :]:
                    if (
                        _etiket_alani(candidate.rstrip(" :"))
                        or _anahtar(candidate.rstrip(" :")) in _BIRLESIK_ADA_PARSEL_ETIKETLERI
                    ):
                        break
                    value = _kunye_degeri(field, candidate)
                    if value:
                        result[field] = value
                        break

    for line in lines:
        match = re.match(r"^\s*(.+?)\s*[:\-–—]\s*(.+)$", line)
        if not match:
            continue
        label = match.group(1).rstrip(" .")
        if _anahtar(label) in _BIRLESIK_ADA_PARSEL_ETIKETLERI:
            ada, parsel = _birlesik_ada_parsel(match.group(2))
            result["ada"] = result["ada"] or ada
            result["parsel"] = result["parsel"] or parsel
            continue
        field = _etiket_alani(label)
        if field and not result[field]:
            result[field] = _kunye_degeri(field, match.group(2))

    all_text = "\n".join(lines)
    fallbacks = {
        "il": (
            r"\b([A-ZÇĞİÖŞÜ][\wÇĞİÖŞÜçğıöşü-]+)\s+ili\b",
        ),
        "ilce": (
            r"\b([A-ZÇĞİÖŞÜ][\wÇĞİÖŞÜçğıöşü-]+)\s+ilçesi\b",
        ),
        "yerlesim": (
            r"\b([A-ZÇĞİÖŞÜ][\wÇĞİÖŞÜçğıöşü-]+(?:[ \t]+[A-ZÇĞİÖŞÜ][\wÇĞİÖŞÜçğıöşü-]+){0,3})[ \t]+(?:Mahallesi|Köyü)\b",
        ),
        "ada": (
            r"\b([0-9]+[A-Za-z]?)\s+ada\b",
        ),
        "parsel": (
            r"\b([0-9]+[A-Za-z]?(?:[/,-][0-9A-Za-z]+)*)\s+(?:numaralı\s+)?parsel\b",
        ),
    }
    for field, patterns in fallbacks.items():
        if result[field]:
            continue
        for pattern in patterns:
            match = re.search(pattern, all_text, re.IGNORECASE)
            if match:
                result[field] = _kunye_degeri(field, match.group(1))
                break
    return result


def _formasyon_katalog_satirlari():
    for item in FORMASYON_ADLARI:
        if item == FORMASYON_SECILMEDI:
            continue
        match = re.match(r"^(.*?)\s*\(([^()]*)\)\s*$", item)
        if match:
            yield item, match.group(1).strip(), match.group(2).strip()
        else:
            yield item, item, ""


def _formasyon_adi_anahtarlari(name):
    """Katalog adının raporlarda görülen çekimli/çoğul yazımlarını üretir."""
    key = _anahtar(name)
    aliases = {key}
    suffixes = {
        "formasyonu": (
            "formasyon", "formasyonu", "formasyonuna", "formasyonunun",
            "formasyonunda", "formasyonundan",
        ),
        "volkaniti": (
            "volkanit", "volkaniti", "volkanitleri", "volkanitine",
            "volkanitinin", "volkanitlerinde",
        ),
        "ignimbiriti": (
            "ignimbirit", "ignimbiriti", "ignimbiritleri", "ignimbiritine",
            "ignimbiritinin", "ignimbiritlerinde",
        ),
        "melanji": ("melanj", "melanji", "melanjina", "melanjinin", "melanjinda"),
        "uyesi": ("uye", "uyesi", "uyesine", "uyesinin", "uyesinde"),
    }
    for suffix, variants in suffixes.items():
        marker = f" {suffix}"
        if key.endswith(marker):
            root = key[: -len(marker)]
            aliases.update(f"{root} {variant}" for variant in variants)

    # Bu birim eski raporlarda "üst" sözcüğü olmadan ve tekil/çoğul
    # granitoyit yazımlarıyla da geçebiliyor.
    if key == "ust oligosen alt miyosen granitoyitleri":
        aliases.update(
            {
                "oligosen alt miyosen granitoyitleri",
                "oligosen miyosen granitoyitleri",
                "ust oligosen alt miyosen granitoyit",
                "oligosen alt miyosen granitoyit",
                "oligosen miyosen granitoyit",
            }
        )
    return aliases


def _anahtar_ifadesi_var(haystack, needle):
    if not haystack or not needle:
        return False
    return re.search(rf"(?:^| ){re.escape(needle)}(?:$| )", haystack) is not None


def _katalog_formasyonunu_eslestir(text):
    haystack = _anahtar(text)
    if not haystack:
        return ""
    tokens = set(haystack.split())
    for item, name, code in _formasyon_katalog_satirlari():
        if any(
            _anahtar_ifadesi_var(haystack, alias)
            for alias in _formasyon_adi_anahtarlari(name)
        ):
            return item
        code_key = _anahtar(code)
        if code_key and code_key in tokens:
            return item
    return ""


_SERBEST_FORMASYON_DESENI = re.compile(
    r"\b("
    r"[A-ZÇĞİÖŞÜ][\wÇĞİÖŞÜçğıöşü-]+"
    r"(?:[ \t]+[A-ZÇĞİÖŞÜ][\wÇĞİÖŞÜçğıöşü-]+){0,3}[ \t]+"
    r"(?i:Formasyon(?:u|una|unun|unda|undan)?|"
    r"Volkanit(?:i|leri|ine|inin|lerinde)?|"
    r"İgnimbirit(?:i|leri|ine|inin|lerinde)?|"
    r"Melanj(?:ı|ına|ının|ında)?|"
    r"Üye(?:si|sine|sinin|sinde)?)"
    r")\b"
)


def _serbest_formasyon_adi(value):
    text = _temiz_metin(value).strip(" \t\n\r.,;:()[]{}\"'“”‘’")
    text = re.sub(
        r"^(?:İnceleme|Çalışma)\s+Alan(?:ı|ında|ının)\s+",
        "",
        text,
        flags=re.IGNORECASE,
    )
    suffixes = (
        (r"Formasyon(?:u|una|unun|unda|undan)?", "Formasyonu"),
        (r"Volkanit(?:i|leri|ine|inin|lerinde)?", "Volkaniti"),
        (r"İgnimbirit(?:i|leri|ine|inin|lerinde)?", "İgnimbiriti"),
        (r"Melanj(?:ı|ına|ının|ında)?", "Melanjı"),
        (r"Üye(?:si|sine|sinin|sinde)?", "Üyesi"),
    )
    for pattern, replacement in suffixes:
        if re.search(rf"\b{pattern}\b$", text, re.IGNORECASE):
            text = re.sub(rf"\b{pattern}\b$", replacement, text, flags=re.IGNORECASE)
            break
    return text


def _formasyon_adaylarini_bul(text):
    """Metindeki katalog birimleri ve güvenli özel formasyon adlarını döndürür."""
    if not _temiz_metin(text):
        return ()
    candidates = []
    seen = set()

    def add(value):
        value = _temiz_metin(value)
        if not value:
            return
        catalog_value = _katalog_formasyonunu_eslestir(value)
        value = catalog_value or value
        key = _anahtar(re.sub(r"\s*\([^)]*\)\s*$", "", value))
        if key and key not in seen:
            seen.add(key)
            candidates.append(value)

    haystack = _anahtar(text)
    tokens = set(haystack.split())
    for item, name, code in _formasyon_katalog_satirlari():
        aliases = _formasyon_adi_anahtarlari(name)
        if any(_anahtar_ifadesi_var(haystack, alias) for alias in aliases):
            add(item)
            continue
        code_key = _anahtar(code)
        if code_key and code_key in tokens:
            add(item)

    for match in _SERBEST_FORMASYON_DESENI.finditer(str(text or "")):
        add(_serbest_formasyon_adi(match.group(1)))
    return tuple(candidates)


def _formasyon_bul(text):
    candidates = _formasyon_adaylarini_bul(text)
    return candidates[0] if len(candidates) == 1 else ""


def _ana_jeoloji_basligi(key):
    """Numarası ne olursa olsun yalnızca ana ``JEOLOJİ`` başlığını tanır."""
    return bool(re.fullmatch(r"(?:\d+\s+)*jeoloji", key))


def _jeoloji_alt_basligi(key):
    return any(
        marker in key
        for marker in (
            "bolgesel jeoloji",
            "stratigraf",
            "yapisal jeoloji",
            "aktif tektonik",
            "inceleme alani",
            "calisma alani",
            "formasyonu",
            "volkaniti",
            "ignimbiriti",
            "melanji",
            "uyesi",
        )
    )


def _inceleme_alani_jeolojisi_basligi(key):
    if "jeoloji" not in key or any(marker in key for marker in ("genel", "bolgesel", "harita")):
        return False
    return any(
        marker in key
        for marker in ("inceleme alani", "calisma alani", "muhendislik jeolojisi")
    )


@dataclass
class JeolojiWordSonucu:
    dosya_yolu: str
    il: str = ""
    ilce: str = ""
    yerlesim: str = ""
    ada: str = ""
    parsel: str = ""
    formasyon: str = ""
    formasyon_adaylari: tuple[str, ...] = ()
    genel_jeoloji_metni: str = ""
    inceleme_alani_jeolojisi: str = ""
    koy_geneli: bool = False
    guven: int = 0
    uyarilar: tuple[str, ...] = ()
    hata: str = ""

    @property
    def uygun(self):
        return bool(
            not self.hata
            and self.ilce
            and self.yerlesim
            and self.genel_jeoloji_metni
            and ((self.ada and self.parsel) or self.koy_geneli)
        )

    def kutuphane_kaydi(self):
        notes = "Word raporundan otomatik çıkarıldı; onay öncesi kontrol edilmelidir."
        if self.uyarilar:
            notes += " Otomatik uyarılar: " + "; ".join(self.uyarilar)
        return {
            "il": self.il or "Çanakkale",
            "ilce": self.ilce,
            "yerlesim": self.yerlesim,
            "ada": self.ada,
            "parsel": self.parsel,
            "formasyon": self.formasyon,
            "genel_jeoloji_metni": self.genel_jeoloji_metni,
            "inceleme_alani_jeolojisi": self.inceleme_alani_jeolojisi,
            "bolum_docx_path": "",
            "bolum_hash": "",
            "harita_path": "",
            "harita_aciklamasi": "",
            "harita_kaynagi": "",
            "harita_olcegi": "",
            "kaynak_rapor_path": self.dosya_yolu,
            "notlar": notes,
            "onay_durumu": "taslak",
        }

    def sozluk(self):
        result = asdict(self)
        result["uygun"] = self.uygun
        return result


def word_raporunu_oku(path):
    source = Path(path)
    if source.suffix.lower() != ".docx":
        return JeolojiWordSonucu(str(source), hata="Yalnızca .docx Word raporları destekleniyor.")
    if not source.is_file():
        return JeolojiWordSonucu(str(source), hata="Dosya bulunamadı.")
    try:
        document = Document(str(source))
    except Exception as exc:
        return JeolojiWordSonucu(str(source), hata=f"Word dosyası açılamadı: {exc}")

    metadata = _kunye_bilgileri(document)
    filename_metadata = _dosya_adindan_kunye(source)
    filename_fields = []
    for field in ("ilce", "yerlesim"):
        if not metadata[field] and filename_metadata[field]:
            metadata[field] = filename_metadata[field]
            filename_fields.append(field)
    general = _en_iyi_bolum(
        document,
        _ana_jeoloji_basligi,
        nested_heading_test=_jeoloji_alt_basligi,
    )
    isolated_geology = _ayrilmis_jeoloji_belgesi(document)
    if not general and isolated_geology:
        # Bazı eski arşivler ana JEOLOJİ başlığı olmadan, doğrudan 2.1 Bölgesel
        # Jeoloji ile başlayan ve elle ayrılmış bölüm belgeleridir.
        general = _tum_belge_metni(document, ana_basligi_atla=True)
    site = _en_iyi_bolum(document, _inceleme_alani_jeolojisi_basligi)
    conclusions = _en_iyi_bolum(
        document,
        lambda key: "sonuc" in key and "oneri" in key,
    )
    # Bölgesel JEOLOJİ bölümü çok sayıda birim içerir; bu nedenle otomatik seçimde
    # kullanılmaz. Önce çalışma alanına özel bölüm, gerekirse Sonuç ve Öneriler
    # bölümü kullanılır. Birden çok aday varsa yalnız iki bölümün ortak tek adayı
    # otomatik seçilebilir.
    site_candidates = _formasyon_adaylarini_bul(site)
    conclusion_candidates = _formasyon_adaylarini_bul(conclusions)
    document_candidates = ()
    formation = ""
    formation_source = ""
    if len(site_candidates) == 1:
        formation = site_candidates[0]
        formation_source = "inceleme"
    elif len(site_candidates) > 1:
        conclusion_keys = {_anahtar(item) for item in conclusion_candidates}
        common = tuple(item for item in site_candidates if _anahtar(item) in conclusion_keys)
        if len(common) == 1:
            formation = common[0]
            formation_source = "dogrulandi"
    elif len(conclusion_candidates) == 1:
        formation = conclusion_candidates[0]
        formation_source = "sonuc"

    if not formation and not site_candidates and not conclusion_candidates:
        # Başlıkları standart olmayan eski raporlarda son güvenli yedek: bütün
        # belgede yalnız bir birim geçiyorsa onu kullan; birden fazlaysa seçme.
        document_candidates = _formasyon_adaylarini_bul(_tum_belge_metni(document))
        if len(document_candidates) == 1:
            formation = document_candidates[0]
            formation_source = "belge"

    formation_candidates = tuple(
        dict.fromkeys((*site_candidates, *conclusion_candidates, *document_candidates))
    )
    warnings = []
    filename_labels = {"ilce": "İlçe", "yerlesim": "Köy/mahalle"}
    for field in filename_fields:
        warnings.append(
            f"{filename_labels[field]} dosya adından alındı: {metadata[field]}"
        )
    field_labels = {
        "ilce": "ilçe",
        "yerlesim": "köy/mahalle",
        "ada": "ada",
        "parsel": "parsel",
    }
    for field, label in field_labels.items():
        if not metadata[field]:
            warnings.append(f"{label.capitalize()} algılanamadı")
    if not general:
        warnings.append("Ana JEOLOJİ bölümü algılanamadı")
    if not site:
        warnings.append("İnceleme Alanı Jeolojisi bölümü algılanamadı")
    if not formation:
        if formation_candidates:
            warnings.append(
                "Birden fazla formasyon adayı bulundu: " + ", ".join(formation_candidates)
            )
        else:
            warnings.append("Formasyon adı algılanamadı")
    elif formation_source == "sonuc":
        warnings.append(f"Formasyon Sonuç ve Öneriler bölümünden algılandı: {formation}")
    elif formation_source == "belge":
        warnings.append(f"Formasyon belgenin tek birim adayından algılandı: {formation}")
    if re.search(r"hata!\s*yer\s+işareti\s+başvurusu\s+geçersiz", general, re.IGNORECASE):
        warnings.append("Word bölümünde bozuk çapraz başvuru metni bulundu; onaydan önce düzeltin")

    village_level = bool(
        not metadata["ada"]
        and not metadata["parsel"]
        and metadata["ilce"]
        and metadata["yerlesim"]
        and general
        and isolated_geology
    )
    if village_level:
        warnings.append("Ada/parsel bulunmadı; köy-geneli kayıt olarak içe aktarılacak")

    confidence = 0
    confidence += 12 if metadata["ilce"] else 0
    confidence += 8 if metadata["yerlesim"] else 0
    confidence += 7 if metadata["ada"] else 0
    confidence += 7 if metadata["parsel"] else 0
    confidence += 28 if general else 0
    confidence += 28 if site else 0
    confidence += 10 if formation else 0
    return JeolojiWordSonucu(
        dosya_yolu=str(source.resolve()),
        il=metadata["il"] or "Çanakkale",
        ilce=metadata["ilce"],
        yerlesim=metadata["yerlesim"],
        ada=metadata["ada"],
        parsel=metadata["parsel"],
        formasyon=formation,
        formasyon_adaylari=formation_candidates,
        genel_jeoloji_metni=general,
        inceleme_alani_jeolojisi=site,
        koy_geneli=village_level,
        guven=min(100, confidence),
        uyarilar=tuple(warnings),
    )


def word_dosyalari_bul(paths):
    """Dosya ve klasörlerden geçerli Word raporlarını güvenle ve tekrarsız bulur."""

    def gizli_mi(candidate):
        if any(part.startswith(".") for part in candidate.parts if part not in (".", "..")):
            return True
        try:
            attributes = getattr(candidate.stat(), "st_file_attributes", 0)
            return bool(attributes & getattr(stat, "FILE_ATTRIBUTE_HIDDEN", 2))
        except OSError:
            return True

    def gecerli_docx(candidate):
        if candidate.suffix.lower() != ".docx" or candidate.name.startswith("~$"):
            return False
        try:
            if gizli_mi(candidate) or not candidate.is_file() or not is_zipfile(candidate):
                return False
            with ZipFile(candidate) as archive:
                names = set(archive.namelist())
                return "[Content_Types].xml" in names and "word/document.xml" in names
        except (OSError, BadZipFile, ValueError):
            return False

    candidates = []
    for raw_path in paths or ():
        path = Path(raw_path)
        try:
            if path.is_dir():
                for root, directories, filenames in os.walk(
                    path, topdown=True, onerror=lambda _error: None, followlinks=False
                ):
                    root_path = Path(root)
                    directories[:] = sorted(
                        (
                            name
                            for name in directories
                            if not gizli_mi(root_path / name)
                        ),
                        key=str.casefold,
                    )
                    for filename in sorted(filenames, key=str.casefold):
                        candidates.append(root_path / filename)
            else:
                candidates.append(path)
        except (OSError, ValueError):
            continue

    result = []
    seen = set()
    for candidate in candidates:
        if not gecerli_docx(candidate):
            continue
        try:
            resolved = candidate.resolve(strict=True)
        except (OSError, ValueError):
            continue
        key = os.path.normcase(str(resolved)).casefold()
        if key in seen:
            continue
        seen.add(key)
        result.append(str(resolved))
    return sorted(result, key=lambda item: os.path.normcase(item).casefold())


def word_raporlarini_oku(paths, task_context=None):
    files = word_dosyalari_bul(paths)
    results = []
    total = len(files)
    for index, path in enumerate(files, start=1):
        if task_context is not None:
            task_context.check_cancelled()
            task_context.report(index - 1, total, f"Word raporu okunuyor: {Path(path).name}")
        results.append(word_raporunu_oku(path))
    if task_context is not None:
        task_context.report(total, total, f"{total} Word raporu incelendi")
    return results


def word_rapor_adaylarini_oku(paths, task_context=None):
    """Klasördeki Word dosyalarını ana jeoloji raporu olma puanıyla sıralar."""
    sonuclar = word_raporlarini_oku(paths, task_context=task_context)
    adaylar = []
    for sonuc in sonuclar:
        if sonuc.hata:
            continue
        puan = int(sonuc.guven or 0)
        if sonuc.genel_jeoloji_metni:
            puan += 35
        if sonuc.inceleme_alani_jeolojisi:
            puan += 35
        if sonuc.ilce:
            puan += 8
        if sonuc.ada and sonuc.parsel:
            puan += 12
        dosya_anahtari = _anahtar(Path(sonuc.dosya_yolu).stem)
        if any(kelime in dosya_anahtari.split() for kelime in ("rapor", "jeolojik", "jeoteknik")):
            puan += 12
        if any(kelime in dosya_anahtari.split() for kelime in ("taahhut", "dilekce", "log", "koordinat")):
            puan -= 20
        adaylar.append({"sonuc": sonuc, "puan": puan})
    return sorted(
        adaylar,
        key=lambda item: (-item["puan"], os.path.normcase(item["sonuc"].dosya_yolu).casefold()),
    )


__all__ = [
    "JeolojiWordSonucu",
    "word_dosyalari_bul",
    "word_rapor_adaylarini_oku",
    "word_raporlarini_oku",
    "word_raporunu_oku",
]
