import os
import re
import shutil
import tempfile
import unicodedata

from docx import Document
from docx.oxml.ns import qn

from jeoloji_bolum_paketi import (
    JeolojiBolumPaketiHatasi,
    STRATIGRAFIK_KESIT_CAPTION,
    stratigrafik_kesit_bolumunu_gorsel_olarak_yenile,
    stratigrafik_kesit_gorselini_cikar,
    stratigrafik_kesit_var_mi,
)


JEOLOJI_WORD_EKLEME_ISARETI = "K1_JEOLOJI_WORD_EKLEME_NOKTASI_7F3A9C"
BOLGESEL_JEOLOJI_BASLIGI = "2.1 Bölgesel Jeoloji"
WD_COLLAPSE_START = 1
WD_COLLAPSE_END = 0
WD_STYLE_NORMAL = -1
WD_STYLE_HEADING_2 = -3
WD_STYLE_CAPTION = -35
WD_ALIGN_PARAGRAPH_CENTER = 1
GENEL_JEOLOJI_GORSEL_ISARETI = "K1_GENEL_JEOLOJI_GORSEL_91B62A"

MUHENDISLIK_JEOLOJISI_CUMLESI_KALIBI = re.compile(
    r"(?:Çalışma|İnceleme)\s+alanında\s+"
    r"(?:birim(?:ler)?|zemin|kayaçlar?)\b"
    r"[^.!?\r\n]{0,600}?"
    r"(?:gözlenmektedir|izlenmektedir|gözlenmiştir|izlenmiştir)\s*[.!?]?",
    re.IGNORECASE,
)


def muhendislik_jeolojisi_cumlelerini_degistir_metin(metin, yeni_cumle):
    yeni_cumle = str(yeni_cumle or "").strip()
    if not yeni_cumle:
        return str(metin or ""), 0
    return MUHENDISLIK_JEOLOJISI_CUMLESI_KALIBI.subn(yeni_cumle, str(metin or ""))


def word_muhendislik_jeolojisi_cumlelerini_degistir(belge, yeni_cumle):
    yeni_cumle = str(yeni_cumle or "").strip()
    if not yeni_cumle:
        return 0
    degisen = 0
    # Document.Content.Text içindeki Python karakter indisleri, Word'ün tablo,
    # alan ve hücre sonu işaretlerini saydığı Range konumlarıyla aynı değildir.
    # Bu nedenle belge-geneli ofset kullanmak metni ilgisiz paragraf/hücrelere
    # kaydırabilir. Her paragrafı kendi Word aralığında, gerçek metinle bul.
    for paragraf_no in range(belge.Paragraphs.Count, 0, -1):
        paragraf = belge.Paragraphs(paragraf_no)
        paragraf_metni = str(paragraf.Range.Text or "")
        eslesmeler = list(MUHENDISLIK_JEOLOJISI_CUMLESI_KALIBI.finditer(paragraf_metni))
        for eslesme in reversed(eslesmeler):
            eski_cumle = eslesme.group(0)
            if eski_cumle.strip().casefold() == yeni_cumle.casefold():
                continue
            aralik = paragraf.Range.Duplicate
            bul = aralik.Find
            bul.ClearFormatting()
            bul.Text = eski_cumle
            bul.Forward = True
            bul.Wrap = 0
            bul.Format = False
            bul.MatchCase = False
            bul.MatchWholeWord = False
            bul.MatchWildcards = False
            if bul.Execute():
                aralik.Text = yeni_cumle
                degisen += 1
    return degisen


def _baslik_normalize(metin):
    metin = unicodedata.normalize("NFKD", str(metin or ""))
    metin = "".join(ch for ch in metin if not unicodedata.combining(ch))
    metin = metin.casefold().replace("ı", "i")
    return re.sub(r"[^a-z0-9]+", "_", metin).strip("_")


def _word_body_paragraf_metni(block):
    return "".join(node.text or "" for node in block.iter(qn("w:t"))).strip()


def _word_body_baslik_seviyesi(document, block):
    text = _word_body_paragraf_metni(block)
    numbered = re.match(r"^\s*(\d{1,3}(?:\.\d{1,3})*)[.)]?\s*", text)
    if numbered:
        return numbered.group(1).count(".") + 1
    p_pr = block.find(qn("w:pPr"))
    if p_pr is not None:
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is not None:
            try:
                return int(outline.get(qn("w:val"))) + 1
            except (TypeError, ValueError):
                pass
        p_style = p_pr.find(qn("w:pStyle"))
        style_id = p_style.get(qn("w:val")) if p_style is not None else ""
        match = re.search(r"(?:heading|baslik)_?(\d+)", _baslik_normalize(style_id))
        if match:
            return int(match.group(1))
    return None


def _word_body_dogrudan_baslik_mi(block):
    p_pr = block.find(qn("w:pPr"))
    if p_pr is not None and (
        p_pr.find(qn("w:keepNext")) is not None
        or p_pr.find(qn("w:pageBreakBefore")) is not None
    ):
        return True
    metin_dugumleri = list(block.iter(qn("w:t")))
    if not metin_dugumleri:
        return False
    anlamli_runlar = []
    for run in block.findall(qn("w:r")):
        if any((node.text or "").strip() for node in run.iter(qn("w:t"))):
            anlamli_runlar.append(run)
    return bool(anlamli_runlar) and all(
        run.find("./" + qn("w:rPr") + "/" + qn("w:b")) is not None
        for run in anlamli_runlar
    )


def _word_body_toc_paragrafi_mi(block):
    p_pr = block.find(qn("w:pPr"))
    p_style = p_pr.find(qn("w:pStyle")) if p_pr is not None else None
    style_id = p_style.get(qn("w:val")) if p_style is not None else ""
    if _baslik_normalize(style_id).startswith(("toc", "icindekiler")):
        return True
    return any(
        "TOC" in (node.text or "").upper()
        for node in block.iter(qn("w:instrText"))
    )


def _eski_rapor_muhendislik_jeolojisi_bolumunu_cikar(kaynak_yolu):
    """Eski rapordan 1.3.2 bölümünü bir sonraki üst/eş başlığa kadar çıkar."""
    document = Document(os.path.abspath(kaynak_yolu))
    blocks = [
        child
        for child in document.element.body.iterchildren()
        if child.tag != qn("w:sectPr")
    ]
    hedef_index = None
    for index, block in enumerate(blocks):
        if block.tag != qn("w:p"):
            continue
        if _word_body_toc_paragrafi_mi(block):
            continue
        key = _baslik_normalize(_word_body_paragraf_metni(block))
        if re.match(
            r"^1_3_2(?:_|$).*inceleme_alani_muhendislik_jeolojisi(?:_|$)",
            key,
        ):
            hedef_index = index
            break
    if hedef_index is None:
        return False

    hedef_seviye = _word_body_baslik_seviyesi(document, blocks[hedef_index]) or 3
    bitis = len(blocks)
    for index in range(hedef_index + 1, len(blocks)):
        block = blocks[index]
        if block.tag != qn("w:p") or not _word_body_paragraf_metni(block):
            continue
        seviye = _word_body_baslik_seviyesi(document, block)
        if (
            seviye is None
            and _word_body_dogrudan_baslik_mi(block)
            and len(_word_body_paragraf_metni(block)) <= 140
        ):
            seviye = hedef_seviye
        if seviye is not None and seviye <= hedef_seviye:
            bitis = index
            break
    for block in blocks[hedef_index:bitis]:
        document.element.body.remove(block)
    document.save(os.path.abspath(kaynak_yolu))
    return True


def _ana_jeoloji_basligi_mi(metin):
    return bool(re.fullmatch(r"(?:\d+_)*jeoloji", _baslik_normalize(metin)))


def _bolgesel_jeoloji_basligi_mi(metin):
    return bool(
        re.fullmatch(
            r"(?:\d+_)*bolgesel_jeoloji",
            _baslik_normalize(metin),
        )
    )


def bolgesel_jeoloji_basligi_mi(metin):
    """Dış ön kontroller için 2.1 Bölgesel Jeoloji başlık denetimi."""
    return _bolgesel_jeoloji_basligi_mi(metin)


def yapisal_jeoloji_basligi_mi(metin):
    """2.1.1 Yapısal Jeoloji ve Aktif Tektonik başlığını ayırt et."""
    return bool(
        re.fullmatch(
            r"(?:\d+_)*yapisal_jeoloji(?:_ve_aktif_tektonik)?",
            _baslik_normalize(metin),
        )
    )


def _ilk_anlamli_paragraftaki_ana_basligi_sil(word, kaynak_yolu):
    """Kaynağı incele; gerekirse geçici kopyadan mükerrer ana başlığı sil."""
    kaynak = word.Documents.Open(
        FileName=os.path.abspath(kaynak_yolu),
        ReadOnly=True,
        AddToRecentFiles=False,
        Visible=False,
    )
    ilk_paragraf_no = None
    ilk_anlamli_paragraflar = []
    try:
        for index in range(1, kaynak.Paragraphs.Count + 1):
            metin = str(kaynak.Paragraphs(index).Range.Text or "").strip("\r\n\x07 ")
            if metin:
                ilk_anlamli_paragraflar.append((index, metin))
                if len(ilk_anlamli_paragraflar) == 1 and _ana_jeoloji_basligi_mi(metin):
                    ilk_paragraf_no = index
                if len(ilk_anlamli_paragraflar) >= 2:
                    break
    finally:
        kaynak.Close(SaveChanges=False)

    bolgesel_baslik_var = False
    if ilk_anlamli_paragraflar:
        ilk_metin = ilk_anlamli_paragraflar[0][1]
        if _bolgesel_jeoloji_basligi_mi(ilk_metin):
            bolgesel_baslik_var = True
        elif _ana_jeoloji_basligi_mi(ilk_metin) and len(ilk_anlamli_paragraflar) > 1:
            bolgesel_baslik_var = _bolgesel_jeoloji_basligi_mi(
                ilk_anlamli_paragraflar[1][1]
            )

    if ilk_paragraf_no is None:
        return os.path.abspath(kaynak_yolu), "", bolgesel_baslik_var

    with tempfile.NamedTemporaryFile(prefix="k1_jeoloji_", suffix=".docx", delete=False) as f:
        gecici_yol = f.name
    shutil.copy2(kaynak_yolu, gecici_yol)
    gecici = word.Documents.Open(
        FileName=os.path.abspath(gecici_yol),
        ReadOnly=False,
        AddToRecentFiles=False,
        Visible=False,
    )
    try:
        gecici.Paragraphs(ilk_paragraf_no).Range.Delete()
        gecici.Save()
    finally:
        gecici.Close(SaveChanges=False)
    return gecici_yol, gecici_yol, bolgesel_baslik_var


def _bolgesel_jeoloji_basligini_ekle(belge, konum):
    """Eksik 2.1 başlığını hedef raporun Heading 2 stiliyle ekle."""
    baslik_metni = BOLGESEL_JEOLOJI_BASLIGI
    baslik_araligi = belge.Range(Start=konum, End=konum)
    baslik_araligi.Text = baslik_metni + "\r"
    baslik_paragrafi = belge.Range(
        Start=konum,
        End=konum + len(baslik_metni) + 1,
    ).Paragraphs(1)
    baslik_paragrafi.Range.Style = WD_STYLE_HEADING_2
    return baslik_paragrafi.Range.End


def _sonraki_ana_alt_baslik_mi(metin):
    """2.1 gövdesini bitirecek 2.2, 2.3 ... başlıklarını ayırt et."""
    temiz = str(metin or "").strip("\r\n\x07 ")
    eslesme = re.match(r"^2\s*[.\-]\s*(\d+)\b", temiz)
    return bool(eslesme and int(eslesme.group(1)) >= 2)


def _bolgesel_jeoloji_bitis_basligi_mi(metin):
    """Yeni 2.1 içeriğinin aşmaması gereken ilk alt/ardıl başlığı bul."""
    temiz = str(metin or "").strip("\r\n\x07 ")
    if re.match(r"^2\s*[.\-]\s*1\s*[.\-]\s*\d+\b", temiz):
        return True
    return _sonraki_ana_alt_baslik_mi(temiz)


def _wordde_yapisal_jeoloji_var_mi(word, kaynak_yolu):
    """2.1.1 başlığını önce doğrudan DOCX yapısından, gerekirse Word ile ara."""
    try:
        kaynak = Document(os.path.abspath(kaynak_yolu))
        return any(yapisal_jeoloji_basligi_mi(p.text) for p in kaynak.paragraphs)
    except Exception:
        # Eski .doc gibi python-docx'in okuyamadığı kaynaklarda geriye dönük destek.
        kaynak = word.Documents.Open(
            FileName=os.path.abspath(kaynak_yolu),
            ReadOnly=True,
            AddToRecentFiles=False,
            Visible=False,
        )
        try:
            return any(
                yapisal_jeoloji_basligi_mi(kaynak.Paragraphs(index).Range.Text)
                for index in range(1, kaynak.Paragraphs.Count + 1)
            )
        finally:
            kaynak.Close(SaveChanges=False)


def _wordde_bolgesel_jeoloji_var_mi(kaynak_yolu):
    try:
        kaynak = Document(os.path.abspath(kaynak_yolu))
    except Exception:
        return False
    return any(bolgesel_jeoloji_basligi_mi(p.text) for p in kaynak.paragraphs)


def wordde_stratigrafik_kesit_var_mi(kaynak_yolu):
    """Arayüz ve ön kontrol için seçili Word'de aktarılabilir kesit denetimi."""
    return stratigrafik_kesit_var_mi(kaynak_yolu)


def _stratigrafik_kesiti_konuma_ekle(word, belge, kaynak_yolu, konum):
    """Kesit blob'unu Caption üstte, doğrudan InlineShape altta ekle."""
    with tempfile.NamedTemporaryFile(
        prefix="k1_stratigrafik_kesit_",
        suffix=".png",
        delete=False,
    ) as handle:
        gorsel_yolu = handle.name
    try:
        try:
            gorsel_blob, gorsel_uzantisi = stratigrafik_kesit_gorselini_cikar(
                kaynak_yolu
            )
            if gorsel_uzantisi != ".png":
                uzantili_yol = f"{os.path.splitext(gorsel_yolu)[0]}{gorsel_uzantisi}"
                os.replace(gorsel_yolu, uzantili_yol)
                gorsel_yolu = uzantili_yol
            with open(gorsel_yolu, "wb") as image_file:
                image_file.write(gorsel_blob)
        except JeolojiBolumPaketiHatasi as exc:
            raise ValueError(
                f"Seçili jeoloji Word'ünden stratigrafik kesit alınamadı: {exc}"
            ) from exc

        hedef_araligi = belge.Range(Start=konum, End=konum)
        hedef_araligi.Text = STRATIGRAFIK_KESIT_CAPTION + "\r\r"
        caption_end = konum + len(STRATIGRAFIK_KESIT_CAPTION) + 1
        caption = belge.Range(
            Start=konum,
            End=caption_end,
        ).Paragraphs(1)
        caption.Range.Style = WD_STYLE_CAPTION

        image_range = belge.Range(Start=caption_end, End=caption_end)
        shape = image_range.InlineShapes.AddPicture(
            FileName=os.path.abspath(gorsel_yolu),
            LinkToFile=False,
            SaveWithDocument=True,
            Range=image_range,
        )
        try:
            shape.LockAspectRatio = True
            shape.Width = word.CentimetersToPoints(15.0)
        except Exception:
            pass
        try:
            image_paragraph = shape.Range.Paragraphs(1)
            image_paragraph.Range.Style = WD_STYLE_NORMAL
        except Exception:
            pass
        try:
            shape.Range.Paragraphs(1).Alignment = WD_ALIGN_PARAGRAPH_CENTER
        except Exception:
            pass
    finally:
        try:
            os.remove(gorsel_yolu)
        except OSError:
            pass


def _bolgesel_jeoloji_kaynagini_hazirla(word, kaynak_yolu, veri):
    """2.1 kaynağını uygula; 2.1.1 ve devamını kaynak Word'den koru."""
    if not isinstance(veri, dict):
        return os.path.abspath(kaynak_yolu), ""
    if not _wordde_yapisal_jeoloji_var_mi(word, kaynak_yolu):
        raise ValueError(
            "Seçili jeoloji Word'ünde '2.1.1 Yapısal Jeoloji ve Aktif Tektonik' "
            "başlığı bulunamadı. Başka bir kütüphane Word'ü seçin."
        )
    try:
        stratigrafik_kesit_gorselini_cikar(kaynak_yolu)
    except JeolojiBolumPaketiHatasi as exc:
        raise ValueError(
            "Seçili jeoloji Word'ünden stratigrafik kesit alınamadı: "
            f"{exc}"
        ) from exc
    mode = str(veri.get("kaynak_modu") or "kutuphane")
    if mode == "eski_rapor":
        if not _wordde_bolgesel_jeoloji_var_mi(kaynak_yolu):
            raise ValueError(
                "Seçili jeoloji Word'ünde '2.1 Bölgesel Jeoloji' başlığı bulunamadı. "
                "Başka bir kütüphane Word'ü seçin."
            )
        # Eski raporun 1.3.2 mühendislik jeolojisi metni, raporun program
        # tarafından üretilen 3.6 bölümünün yerine geçmemelidir. Kaynak kopyada
        # yalnız bu bölümün bir sonraki eş/üst başlığa kadar olan gövdesini
        # çıkar; kesit bloğunu da 2.1.1 öncesinde standart görsel olarak yenile.
        with tempfile.NamedTemporaryFile(
            prefix="k1_eski_jeoloji_",
            suffix=".docx",
            delete=False,
        ) as handle:
            temporary = handle.name
        try:
            shutil.copy2(kaynak_yolu, temporary)
            _eski_rapor_muhendislik_jeolojisi_bolumunu_cikar(temporary)
            stratigrafik_kesit_bolumunu_gorsel_olarak_yenile(temporary, temporary)
        except Exception:
            try:
                os.remove(temporary)
            except OSError:
                pass
            raise
        return temporary, temporary
    if not veri.get("birimler"):
        return os.path.abspath(kaynak_yolu), ""
    gorsel_yolu = str(veri.get("gorsel_yolu") or "").strip()
    if not gorsel_yolu or not os.path.isfile(gorsel_yolu):
        raise ValueError("2.1 Bölgesel Jeoloji için oluşturulmuş genel jeoloji haritası bulunamadı.")

    with tempfile.NamedTemporaryFile(prefix="k1_bolgesel_jeoloji_", suffix=".docx", delete=False) as f:
        temporary = f.name
    shutil.copy2(kaynak_yolu, temporary)
    # 2.1'i program üretse bile kaynağın başka bir yerindeki eski
    # "1.3.2 İnceleme Alanı Mühendislik Jeolojisi" bölümü taşınmamalıdır.
    _eski_rapor_muhendislik_jeolojisi_bolumunu_cikar(temporary)
    document = word.Documents.Open(
        FileName=os.path.abspath(temporary),
        ReadOnly=False,
        AddToRecentFiles=False,
        Visible=False,
    )
    try:
        heading_index = None
        for index in range(1, document.Paragraphs.Count + 1):
            if _bolgesel_jeoloji_basligi_mi(document.Paragraphs(index).Range.Text):
                heading_index = index
                break
        if heading_index is None:
            start = document.Content.Start
            start = _bolgesel_jeoloji_basligini_ekle(document, start)
            heading_index = 1
            body_start = start
        else:
            body_start = document.Paragraphs(heading_index).Range.End

        body_end = max(body_start, document.Content.End - 1)
        for index in range(heading_index + 1, document.Paragraphs.Count + 1):
            paragraph = document.Paragraphs(index)
            if _bolgesel_jeoloji_bitis_basligi_mi(paragraph.Range.Text):
                body_end = paragraph.Range.Start
                break
        body_range = document.Range(Start=body_start, End=body_end)
        body_range.Delete()

        intro = (
            "Çalışma alanı merkez alınarak hazırlanan 1/100.000 ölçekli genel jeoloji "
            "haritasında çalışma alanı ve çevresindeki jeolojik birimler gösterilmiştir."
        )
        genel_jeoloji_caption = "Şekil Genel jeoloji haritası (Ölçek 1/100.000)"
        parts = [intro, GENEL_JEOLOJI_GORSEL_ISARETI, genel_jeoloji_caption]
        heading_texts = []
        for unit in veri.get("birimler", []):
            if not unit.get("kullan", True):
                continue
            name = str(unit.get("ad") or "").strip()
            code = str(unit.get("kod") or "").strip()
            heading = f"{name} ({code})" if code else name
            text = str(unit.get("bolgesel_jeoloji_metni") or "").strip()
            if heading:
                heading_texts.append(heading)
                parts.append(heading)
            if text:
                parts.append(text)
        insertion = document.Range(Start=body_start, End=body_start)
        insertion.Text = "\r".join(parts) + "\r"
        inserted_end = insertion.End

        marker_range = document.Range(Start=body_start, End=inserted_end)
        finder = marker_range.Find
        finder.ClearFormatting()
        finder.Text = GENEL_JEOLOJI_GORSEL_ISARETI
        if not finder.Execute():
            raise ValueError("Genel jeoloji haritası Word ekleme işareti bulunamadı.")
        marker_range.Text = ""
        shape = marker_range.InlineShapes.AddPicture(
            FileName=os.path.abspath(gorsel_yolu),
            LinkToFile=False,
            SaveWithDocument=True,
            Range=marker_range,
        )
        try:
            shape.LockAspectRatio = True
            shape.Width = word.CentimetersToPoints(15.0)
        except Exception:
            pass
        try:
            shape.Range.Paragraphs(1).Alignment = WD_ALIGN_PARAGRAPH_CENTER
        except Exception:
            pass

        # Heading 2'nin sonuna yazılan yeni paragraflar Word tarafından aynı stili
        # miras alabilir. 2.2 başlığına kadar gövdeyi önce Normal stile döndür.
        for index in range(heading_index + 1, document.Paragraphs.Count + 1):
            paragraph = document.Paragraphs(index)
            if _bolgesel_jeoloji_bitis_basligi_mi(paragraph.Range.Text):
                break
            paragraph.Range.Style = WD_STYLE_NORMAL
            paragraph.Range.Bold = False
            paragraph.Range.Italic = False

        for index in range(heading_index + 1, document.Paragraphs.Count + 1):
            paragraph = document.Paragraphs(index)
            text = str(paragraph.Range.Text or "").strip("\r\n\x07 ")
            if _bolgesel_jeoloji_bitis_basligi_mi(text):
                break
            if text in heading_texts:
                paragraph.Range.Bold = True
                paragraph.Format.KeepWithNext = True
                paragraph.Format.SpaceBefore = 6
            elif text == genel_jeoloji_caption:
                paragraph.Range.Style = WD_STYLE_CAPTION
                paragraph.Alignment = WD_ALIGN_PARAGRAPH_CENTER
                paragraph.Range.Italic = True
                paragraph.Format.KeepWithNext = False

        kesit_konumu = max(body_start, document.Content.End - 1)
        for index in range(heading_index + 1, document.Paragraphs.Count + 1):
            paragraph = document.Paragraphs(index)
            if _bolgesel_jeoloji_bitis_basligi_mi(paragraph.Range.Text):
                kesit_konumu = paragraph.Range.Start
                break
        _stratigrafik_kesiti_konuma_ekle(
            word,
            document,
            kaynak_yolu,
            kesit_konumu,
        )
        document.Save()
    except Exception:
        try:
            document.Close(SaveChanges=False)
        finally:
            try:
                os.remove(temporary)
            except OSError:
                pass
        raise
    else:
        document.Close(SaveChanges=False)
    return temporary, temporary


def word_jeoloji_bolumunu_ekle(
    word,
    belge,
    kaynak_yolu,
    muhendislik_jeolojisi_cumlesi="",
    bolgesel_jeoloji_verisi=None,
):
    """Kaynak Word gövdesini ekleme işaretine Word'ün kendi motoruyla yerleştir."""
    if not kaynak_yolu or not os.path.isfile(kaynak_yolu):
        raise ValueError(f"Jeoloji Word kaynağı bulunamadı: {kaynak_yolu}")

    ekleme_araligi = belge.Content.Duplicate
    bul = ekleme_araligi.Find
    bul.ClearFormatting()
    bul.Text = JEOLOJI_WORD_EKLEME_ISARETI
    if not bul.Execute():
        raise ValueError("Raporda jeoloji Word ekleme noktası bulunamadı.")

    silinecek_yollar = []
    try:
        (
            eklenecek_yol,
            silinecek_yol,
            bolgesel_baslik_var,
        ) = _ilk_anlamli_paragraftaki_ana_basligi_sil(word, kaynak_yolu)
        if silinecek_yol:
            silinecek_yollar.append(silinecek_yol)
        if isinstance(bolgesel_jeoloji_verisi, dict) and (
            bolgesel_jeoloji_verisi.get("birimler")
            or bolgesel_jeoloji_verisi.get("kaynak_modu") == "eski_rapor"
        ):
            bolgesel_veri = dict(bolgesel_jeoloji_verisi)
            bolgesel_veri["gorsel_yolu"] = str(bolgesel_veri.get("gorsel_yolu") or "")
            eklenecek_yol, bolgesel_gecici_yol = _bolgesel_jeoloji_kaynagini_hazirla(
                word,
                eklenecek_yol,
                bolgesel_veri,
            )
            if bolgesel_gecici_yol:
                silinecek_yollar.append(bolgesel_gecici_yol)
            bolgesel_baslik_var = True
        ekleme_araligi.Text = ""
        ekleme_araligi.Collapse(WD_COLLAPSE_START)
        ekleme_konumu = ekleme_araligi.Start
        if not bolgesel_baslik_var:
            ekleme_konumu = _bolgesel_jeoloji_basligini_ekle(
                belge,
                ekleme_konumu,
            )
        ekleme_araligi = belge.Range(Start=ekleme_konumu, End=ekleme_konumu)
        ekleme_araligi.InsertFile(os.path.abspath(eklenecek_yol))

        # Kaynaktan yeni bölüm geldiyse ana raporun üst/alt bilgisini ve devam eden
        # sayfa numarasını kullan. Kaynak gövdesindeki sayfa kırılımları korunur.
        for section_index in range(2, belge.Sections.Count + 1):
            section = belge.Sections(section_index)
            for collection_name in ("Headers", "Footers"):
                collection = getattr(section, collection_name)
                for item_index in range(1, collection.Count + 1):
                    try:
                        collection(item_index).LinkToPrevious = True
                    except Exception:
                        pass
            try:
                for item_index in range(1, section.Footers.Count + 1):
                    page_numbers = section.Footers(item_index).PageNumbers
                    if page_numbers.Count:
                        page_numbers.RestartNumberingAtSection = False
            except Exception:
                pass
        word_muhendislik_jeolojisi_cumlelerini_degistir(
            belge,
            muhendislik_jeolojisi_cumlesi,
        )
    finally:
        for silinecek_yol in dict.fromkeys(silinecek_yollar):
            try:
                os.remove(silinecek_yol)
            except OSError:
                pass

    if JEOLOJI_WORD_EKLEME_ISARETI in str(belge.Content.Text or ""):
        raise ValueError("Jeoloji Word ekleme işareti çıktıdan kaldırılamadı.")

    bolgesel_baslik_sayisi = sum(
        1
        for index in range(1, belge.Paragraphs.Count + 1)
        if _bolgesel_jeoloji_basligi_mi(belge.Paragraphs(index).Range.Text)
    )
    if bolgesel_baslik_sayisi != 1:
        raise ValueError(
            "Raporda 2.1 Bölgesel Jeoloji başlığı bir kez bulunmalıdır; "
            f"bulunan: {bolgesel_baslik_sayisi}."
        )
