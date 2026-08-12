from types import SimpleNamespace

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

from rapor import RaporUretici


def _png(path, color):
    Image.new("RGB", (640, 360), color).save(path)


def _gorsel_sayisi(paragraph):
    return len(list(paragraph._p.iter(qn("w:drawing"))))


def test_pga_sablondaki_sabit_gorseli_yerinde_degistirir(tmp_path):
    eski = tmp_path / "eski.png"
    pga = tmp_path / "pga.png"
    _png(eski, "red")
    _png(pga, "blue")

    doc = Document()
    caption = doc.add_paragraph("Şekil 3: Çanakkale bölgesi deprem tehlike haritası")
    eski_paragraf = doc.add_paragraph()
    eski_paragraf.add_run().add_picture(str(eski))
    doc.add_paragraph("(Afet ve Acil Durum Yönetimi Başkanlığı)")

    app = SimpleNamespace(img_pga_haritasi=str(pga))
    uretici = RaporUretici(app)
    uretici.rapor_pga_haritasini_yerlestir(doc, caption, str(pga))

    assert doc.paragraphs[0].text.startswith("Şekil 3")
    assert _gorsel_sayisi(doc.paragraphs[1]) == 1
    assert sum(_gorsel_sayisi(p) for p in doc.paragraphs) == 1
    assert doc.paragraphs[2].text.startswith("(Afet")


def test_pga_sabit_gorsel_yoksa_basliktan_sonra_eklenir(tmp_path):
    pga = tmp_path / "pga.png"
    _png(pga, "blue")
    doc = Document()
    caption = doc.add_paragraph("Şekil 3: Çanakkale bölgesi deprem tehlike haritası")
    doc.add_paragraph("(Kaynak)")

    uretici = RaporUretici(SimpleNamespace(img_pga_haritasi=str(pga)))
    uretici.rapor_pga_haritasini_yerlestir(doc, caption, str(pga))

    assert _gorsel_sayisi(doc.paragraphs[1]) == 1
    assert doc.paragraphs[2].text == "(Kaynak)"
