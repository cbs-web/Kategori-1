from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import pytest
from PIL import Image
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor

from rapor import RaporUretici
from raporlama_islemleri import atomik_docx_kaydet
from taahhutname import TaahhutnameUretici
from jeoloji_bolum_paketi import (
    stratigrafik_kesit_bolumunu_ayir,
    stratigrafik_kesit_var_mi,
)
from word_jeoloji_birlestirme import (
    JEOLOJI_WORD_EKLEME_ISARETI,
    _eski_rapor_muhendislik_jeolojisi_bolumunu_cikar,
    muhendislik_jeolojisi_cumlelerini_degistir_metin,
)


class Entry:
    def __init__(self, value):
        self.value = str(value)

    def get(self):
        return self.value


class TextValue:
    def __init__(self, value):
        self.value = value

    def get(self, *_args):
        return self.value


class App:
    def __init__(self):
        self.errors = []

    def hata_kaydet(self, message, error=None):
        self.errors.append((message, error))


def _png(path, color):
    Image.new("RGB", (24, 16), color).save(path)


def test_docx_birlestirme_medya_stil_ve_numaralandirmayi_esler(tmp_path):
    hedef_resim = tmp_path / "hedef.png"
    kaynak_resim = tmp_path / "kaynak.png"
    _png(hedef_resim, "red")
    _png(kaynak_resim, "blue")

    hedef = Document()
    hedef_stil = hedef.styles.add_style("CakisanStil", WD_STYLE_TYPE.PARAGRAPH)
    hedef_stil.font.color.rgb = RGBColor(255, 0, 0)
    hedef.add_picture(str(hedef_resim))
    anchor = hedef.add_paragraph("[JEOLOJI_BOLUMU]")

    kaynak = Document()
    kaynak_stil = kaynak.styles.add_style("CakisanStil", WD_STYLE_TYPE.PARAGRAPH)
    kaynak_stil.font.color.rgb = RGBColor(0, 0, 255)
    kaynak.add_paragraph("Kaynak özel stil", style=kaynak_stil)
    kaynak.add_paragraph("Numaralı kaynak satırı", style="List Number")
    kaynak.add_picture(str(kaynak_resim))

    uretici = RaporUretici(App())
    assert uretici.rapor_docx_icerigi_ekle(anchor, kaynak) > 0
    hedef_yol = tmp_path / "birlesik.docx"
    hedef.save(hedef_yol)

    yeniden = Document(hedef_yol)
    assert len(yeniden.inline_shapes) == 2
    kaynak_paragraf = next(p for p in yeniden.paragraphs if p.text == "Kaynak özel stil")
    assert kaynak_paragraf.style.style_id != "CakisanStil"
    assert any(p.text == "Numaralı kaynak satırı" and p.style is not None for p in yeniden.paragraphs)

    with ZipFile(hedef_yol) as paket:
        document_xml = paket.read("word/document.xml").decode("utf-8")
        rels_xml = paket.read("word/_rels/document.xml.rels").decode("utf-8")
        numbering_xml = paket.read("word/numbering.xml").decode("utf-8")
    import re

    embed_ids = set(re.findall(r'r:embed="(rId\d+)"', document_xml))
    image_ids = set(
        re.findall(
            r'<Relationship[^>]+Id="(rId\d+)"[^>]+Type="[^"]+/image"',
            rels_xml,
        )
    )
    kullanilan_num = set(re.findall(r'<w:numId w:val="(\d+)"', document_xml))
    tanimli_num = set(re.findall(r'<w:num w:numId="(\d+)"', numbering_xml))
    assert embed_ids <= image_ids
    assert kullanilan_num <= tanimli_num


@pytest.mark.parametrize(
    "kesit_basligi",
    (
        "Şekil 5 Çalışma alanı ve yakın çevresinin stratigrafik kesiti",
        "Şekil 5 Çalışma alanı ve yakın çevresinin",
    ),
)
def test_secili_wordden_stratigrafik_kesit_gorseli_ve_basligi_ayrilir(
    tmp_path,
    kesit_basligi,
):
    genel_resim = tmp_path / "genel.png"
    kesit_resim = tmp_path / "kesit.png"
    fay_resmi = tmp_path / "fay.png"
    _png(genel_resim, "blue")
    _png(kesit_resim, "green")
    _png(fay_resmi, "red")

    source = Document()
    source.add_paragraph("2.1 Bölgesel Jeoloji", style="Heading 2")
    source.add_paragraph(
        "Bölgenin genel jeoloji haritası Şekil 4'te, stratigrafik kesit "
        "Şekil 5'te verilmiştir."
    )
    source.add_paragraph("Şekil 4 Genel jeoloji haritası")
    source.add_picture(str(genel_resim))
    source.add_paragraph("Bayramiç Formasyonu açıklaması.")
    kesit_paragrafi = source.add_paragraph()
    kesit_paragrafi.add_run().add_picture(str(kesit_resim))
    kesit_paragrafi.add_run(kesit_basligi)
    source.add_paragraph(
        "2.1.1 Yapısal Jeoloji ve Aktif Tektonik",
        style="Heading 3",
    )
    source.add_picture(str(fay_resmi))
    source_path = tmp_path / "secili_jeoloji.docx"
    source.save(source_path)

    output_path = tmp_path / "stratigrafik_kesit.docx"
    assert stratigrafik_kesit_var_mi(source_path)
    result = stratigrafik_kesit_bolumunu_ayir(source_path, output_path)
    snippet = Document(output_path)
    snippet_text = "\n".join(paragraph.text for paragraph in snippet.paragraphs)

    assert result.sekil_no == "5"
    assert result.blok_sayisi == 1
    assert len(snippet.inline_shapes) == 1
    assert kesit_basligi in snippet_text
    assert "Şekil 4" not in snippet_text
    assert "2.1.1" not in snippet_text


def test_stratigrafik_kesit_heading_3_ise_normal_stille_aktarilir(tmp_path):
    kesit_resim = tmp_path / "kesit_heading3.png"
    _png(kesit_resim, "green")
    source = Document()
    source.add_paragraph("2.1 Bölgesel Jeoloji", style="Heading 2")
    source.add_paragraph("Stratigrafik kesit Şekil 5'te verilmiştir.")
    source.add_paragraph("Stratigrafik Kesit", style="Heading 3")
    source.add_picture(str(kesit_resim))
    source.add_paragraph("Şekil 5 Çalışma alanı stratigrafik kesiti")
    source.add_paragraph("2.1.1 Yapısal Jeoloji", style="Heading 3")
    source_path = tmp_path / "heading3_kesit.docx"
    source.save(source_path)

    output_path = tmp_path / "heading3_kesit_paketi.docx"
    stratigrafik_kesit_bolumunu_ayir(source_path, output_path)
    snippet = Document(output_path)

    baslik = next(p for p in snippet.paragraphs if "Stratigrafik Kesit" in p.text)
    assert baslik.style.name == "Normal"
    assert len(snippet.inline_shapes) == 1


def test_eski_wordde_1_3_2_muhendislik_jeolojisi_bolumu_cikarilir(tmp_path):
    source = Document()
    source.add_paragraph("1.3.1. Önceki Bölüm", style="Heading 3")
    source.add_paragraph("Korunacak önceki metin.")
    source.add_paragraph("1.3.2.İnceleme Alanı Mühendislik Jeolojisi", style="Heading 3")
    source.add_paragraph("Seröiler ve çevresinde gözlenen eski mühendislik jeolojisi metni.")
    source.add_paragraph("1.3.3. Sonraki Bölüm", style="Heading 3")
    source.add_paragraph("Korunacak sonraki metin.")
    source_path = tmp_path / "eski_rapor.docx"
    source.save(source_path)

    assert _eski_rapor_muhendislik_jeolojisi_bolumunu_cikar(source_path)
    text = "\n".join(p.text for p in Document(source_path).paragraphs)
    assert "1.3.2" not in text
    assert "Seröiler" not in text
    assert "Korunacak önceki metin" in text
    assert "1.3.3. Sonraki Bölüm" in text
    assert "Korunacak sonraki metin" in text


def test_kutuphane_word_bolumu_resim_ve_tabloyla_rapora_eklenir(tmp_path):
    source_image = tmp_path / "jeoloji.png"
    _png(source_image, "green")
    source = Document()
    source.add_paragraph("2. JEOLOJİ")
    source.add_paragraph("Kütüphaneden seçilen jeoloji açıklaması")
    table = source.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Şekil"
    table.cell(0, 1).text = "Açıklama"
    source.add_picture(str(source_image))
    source_path = tmp_path / "kutuphane_bolumu.docx"
    source.save(source_path)

    target = Document()
    target.add_paragraph("2. JEOLOJİ", style="Heading 1")
    target.add_paragraph("[JEOLOJI_BOLUMU]")
    app = App()
    app.jeoloji_kutuphanesi_bolumu_aktif = True
    app.jeoloji_kutuphanesi_kayit_id = 27
    app.jeoloji_kutuphanesi_bolum_yolu = str(source_path)
    app.jeoloji_sablon_yolu = ""

    kaynak_yolu = RaporUretici(app).rapor_jeoloji_bolumu_ekle(target)
    output = tmp_path / "rapor.docx"
    target.save(output)
    reopened = Document(output)

    assert kaynak_yolu == str(source_path)
    assert sum(paragraph.text == "2. JEOLOJİ" for paragraph in reopened.paragraphs) == 1
    assert "[JEOLOJI_BOLUMU]" not in "\n".join(p.text for p in reopened.paragraphs)
    assert JEOLOJI_WORD_EKLEME_ISARETI in "\n".join(p.text for p in reopened.paragraphs)
    assert len(reopened.tables) == 0
    assert len(reopened.inline_shapes) == 0


def test_jeoloji_ana_basligi_turkce_word_stil_kimliginden_bagimsiz_dogrulanir():
    doc = Document()
    turkce_baslik = doc.styles.add_style("Başlık 1 Türkçe", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("2. JEOLOJİ", style=turkce_baslik)

    RaporUretici(App()).rapor_cikti_belgesini_dogrula(doc)


def test_parsel_haritasi_etiketi_word_gorseline_donusur(tmp_path):
    image_path = tmp_path / "parsel.png"
    _png(image_path, "green")
    document = Document()
    document.add_paragraph("[PARSEL_HARITASI]")
    app = App()
    app.img_parsel_haritasi = str(image_path)

    RaporUretici(app).rapor_resimleri_ekle(document)
    output = tmp_path / "parsel_raporu.docx"
    document.save(output)
    reopened = Document(output)

    assert "[PARSEL_HARITASI]" not in "\n".join(p.text for p in reopened.paragraphs)
    assert len(reopened.inline_shapes) == 1


def test_csv_cp1254_ve_alintili_ayiraci_dogru_okur(tmp_path):
    csv_yolu = tmp_path / "jeofizik.csv"
    csv_metni = (
        '"Sismik Ölçü ve Hesaplarının Sahibi";"Serim; SS1";;\n'
        '"VP =";;100;200\n'
        '"VS =";;50;100\n'
        '"Tabaka Kalınlığı";;2;\n'
        '"Tabaka Yoğunluğu";;1,8;2,0\n'
        '"Poisson Oranı";;0,3;0,25\n'
        '"Elastisite";;10;20\n'
        '"Kayma Modülü";;4;8\n'
        '"Bulk";;7;14\n'
        '"Vs30 =";;150;\n'
    )
    csv_yolu.write_bytes(csv_metni.encode("cp1254"))
    app = App()
    app.jeofizik_excel_yolu_al = lambda: str(csv_yolu)

    sonuc = RaporUretici(app).rapor_jeofizik_parametrelerini_oku()

    assert len(sonuc) == 1
    assert sonuc[0]["ad"] == "Serim; SS1"
    assert [layer["vp"] for layer in sonuc[0]["layers"]] == ["100", "200"]


def test_cozulemeyen_etiket_bloklar_ve_update_fields_ayarlanir():
    uretici = RaporUretici(App())
    doc = Document()
    doc.add_paragraph("[BILINMEYEN]")
    with pytest.raises(ValueError, match="BILINMEYEN"):
        uretici.rapor_cozulemeyen_etiketleri_dogrula(doc)

    doc.paragraphs[0].text = "Tamamlandı"
    uretici.rapor_update_fields_ayarla(doc)
    cikti = BytesIO()
    doc.save(cikti)
    with ZipFile(BytesIO(cikti.getvalue())) as paket:
        settings = paket.read("word/settings.xml").decode("utf-8")
    assert "updateFields" in settings
    assert 'w:val="true"' in settings


def test_rapor_muhendis_etiketleri_proje_taahhut_bilgilerinden_gelir():
    app = App()
    app.veri_alanlari = {"ILGILI_IDARE": Entry("Test Belediyesi")}
    app.taahhut_bilgileri = {
        "JEOFIZIK_MUH_AD": "Jeofizik Kişi",
        "JEOLOJI_MUH_AD": "Jeoloji Kişi",
    }
    app.ac_yn_sekme_kayitlari = lambda: []
    app.formasyon_bilgilerini_hazirla = lambda: {
        "adi": "",
        "kisa": "",
        "birim_tanimi": "",
        "muhendislik_metni": "",
    }
    app.tg_girdiler = {}

    uretici = RaporUretici(app)
    etiketler = uretici.desteklenen_sablon_etiketleri()
    degerler, *_ = uretici.rapor_etiket_verilerini_hazirla()

    assert "[JEOFIZIK_MUH_AD]" in etiketler
    assert "[JEOLOJI_MUH_TELEFON]" in etiketler
    assert degerler["[JEOFIZIK_MUH_AD]"] == "Jeofizik Kişi"
    assert degerler["[JEOLOJI_MUH_AD]"] == "Jeoloji Kişi"
    assert degerler["[JEOLOJI_MUH_TELEFON]"] == ""


def test_muhendislik_jeolojisi_cumlesi_programdaki_litolojiden_olusturulur():
    uretici = RaporUretici(App())
    uzun = (
        "Bölgesel jeolojiye ilişkin uzun açıklama. "
        "Çalışma alanında birim kahverengi kumlu kil olarak izlenmektedir. "
        "Sonraki uzun jeoloji açıklaması."
    )

    sonuc = uretici.rapor_muhendislik_jeolojisi_metnini_hazirla(
        uzun,
        ["kahverengi kumlu kil"],
    )

    assert sonuc == "Çalışma alanında birimler, kahverengi kumlu kil olarak gözlenmiştir."


def test_kutuphane_wordundeki_eski_litoloji_cumlesi_guncel_cumleyle_degistirilir():
    eski = (
        "Önceki metin. Çalışma alanında kayaçlar; orta derece altere olmuş "
        "granodiyorit olarak gözlenmektedir. Sonraki metin."
    )
    yeni = "Çalışma alanında birimler, kahverengi kumlu kil olarak gözlenmiştir."

    sonuc, adet = muhendislik_jeolojisi_cumlelerini_degistir_metin(eski, yeni)

    assert adet == 1
    assert yeni in sonuc
    assert "granodiyorit olarak gözlenmektedir" not in sonuc


def test_tasima_tablosunun_ustundeki_eski_katsayi_satiri_temizlenir():
    class TestRaporUretici(RaporUretici):
        def rapor_tasima_gucu_tablosu_olustur(self, doc):
            return doc.add_table(rows=1, cols=1)

    app = App()
    uretici = TestRaporUretici(app)
    metin = (
        "Taşıma gücü açıklaması.\n\n"
        "N_c = 9.285; N_q = 2.974; N_γ = 0.839; s_c = 1.214; "
        "s_q = 1.142; d_c = 1.040; d_q = 1.027.\n\n"
        "[TABLO_BURADA]\n\n"
    )

    satirlar = uretici.rapor_tasima_metin_satirlarini_temizle(metin)

    assert satirlar == ["Taşıma gücü açıklaması.", "[TABLO_BURADA]"]

    app.txt_tasima_rapor = TextValue(metin)
    belge = Document()
    belge.add_paragraph("[TASIMA_GUCU]")

    uretici.rapor_tasima_gucu_ekle(belge)

    assert [paragraf.text for paragraf in belge.paragraphs] == [
        "Taşıma gücü açıklaması."
    ]
    assert belge.paragraphs[0].paragraph_format.space_after.pt == pytest.approx(8)
    assert len(belge.tables) == 1


def test_tasima_tablosu_kritik_girdide_fallback_kullanmaz():
    app = App()
    app.tg_girdiler = {
        "c": Entry(10),
        "phi": Entry(30),
        "gn": Entry(1.8),
        "B": Entry(1.5),
        "L": Entry(2.0),
        "Df": Entry(1.0),
        "RvGk": Entry(1.4),
        "ks_carpani": Entry(56),
    }
    with pytest.raises(ValueError, match="gsat"):
        RaporUretici(app).rapor_tasima_gucu_tablosu_olustur(Document())

    app.tg_girdiler.update({"gsat": Entry(2.0), "yass": Entry(999)})
    app.entry_qt_nihai = Entry(90)
    app.entry_ks_nihai = Entry(5040)
    tablo = RaporUretici(app).rapor_tasima_gucu_tablosu_olustur(Document())
    tablo_metni = " ".join(cell.text for row in tablo.rows for cell in row.cells)
    assert len(tablo.rows) == 9
    assert "Kohezyon (c) [t/m2]" in tablo_metni
    assert "Temel Genişliği (B) [metre]" in tablo_metni
    assert "Temel Uzunluğu (L) [metre]" in tablo_metni
    assert "Yükleme Eğikliği Açısı [derece]" in tablo_metni
    assert "Temel Zemini Eğim Açısı [derece]" in tablo_metni
    assert "Temel Taban Eğim Açısı [derece]" in tablo_metni
    assert "Raporda Kullanılan Nihai" not in tablo_metni
    assert "Yatak Katsayısı" not in tablo_metni

    app.entry_qt_nihai = Entry(95)
    app.entry_ks_nihai = Entry(5320)
    with pytest.raises(ValueError, match="büyük olamaz"):
        RaporUretici(app).rapor_tasima_gucu_tablosu_olustur(Document())

    app.entry_qt_nihai = Entry(90)
    app.entry_ks_nihai = Entry(5000)
    with pytest.raises(ValueError, match="uyuşmuyor"):
        RaporUretici(app).rapor_tasima_gucu_tablosu_olustur(Document())


def test_ters_derinlik_araligi_sessizce_duzeltilmez():
    with pytest.raises(ValueError, match="ters girilmiş"):
        RaporUretici(App()).rapor_derinlik_araligi_coz("3,50 - 1,25 m")


def test_eski_cikti_atomik_docx_hatasinda_korunur(tmp_path):
    hedef = tmp_path / "rapor.docx"
    hedef.write_bytes(b"eski-icerik")

    class BozukBelge:
        def save(self, _path):
            raise RuntimeError("yazma hatası")

    with pytest.raises(RuntimeError):
        atomik_docx_kaydet(BozukBelge(), hedef)
    assert hedef.read_bytes() == b"eski-icerik"
    assert not list(tmp_path.glob(".*.tmp.docx"))


def test_taahhut_yer_tutucusu_degisir_ilgisiz_tarih_ve_sayi_korunur():
    app = App()
    app.taahhut_bilgileri = {
        "JEOFIZIK_MUH_AD": "Yeni Jeofizik",
        "JEOFIZIK_MUH_SICIL": "1111",
        "JEOFIZIK_MUH_ADRES": "Adres 1",
        "JEOFIZIK_MUH_TELEFON": "0555 111 11 11",
        "JEOLOJI_MUH_AD": "Yeni Jeoloji",
        "JEOLOJI_MUH_SICIL": "2222",
        "JEOLOJI_MUH_ADRES": "Adres 2",
        "JEOLOJI_MUH_TELEFON": "0555 222 22 22",
    }
    doc = Document()
    doc.add_paragraph("[JEOFIZIK_MUH_AD]")
    doc.add_paragraph("Revizyon: 01.01.2020 / Proje no: 19820")

    uretici = TaahhutnameUretici(app)
    uretici.taahhut_word_xml_metnini_guncelle(doc)

    assert doc.paragraphs[0].text == "Yeni Jeofizik"
    assert doc.paragraphs[1].text == "Revizyon: 01.01.2020 / Proje no: 19820"
    assert uretici.taahhut_verisini_topla()["JEOLOJI_MUH_AD"] == "Yeni Jeoloji"


def test_taahhut_etiketsiz_imza_blogunu_baglamla_gunceller():
    app = App()
    app.taahhut_bilgileri = {
        "JEOFIZIK_MUH_AD": "Yeni Jeofizik",
        "JEOFIZIK_MUH_SICIL": "1111",
        "JEOFIZIK_MUH_ADRES": "Adres 1",
        "JEOFIZIK_MUH_TELEFON": "0555 111 11 11",
        "JEOLOJI_MUH_AD": "Yeni Jeoloji",
        "JEOLOJI_MUH_SICIL": "2222",
        "JEOLOJI_MUH_ADRES": "Adres 2",
        "JEOLOJI_MUH_TELEFON": "0555 222 22 22",
    }
    doc = Document()
    doc.add_paragraph("Revizyon tarihi 01.01.2020")
    doc.add_paragraph("02.02.2020")
    doc.add_paragraph("Eski Kimlik")
    doc.add_paragraph("Jeofizik Mühendisi")

    TaahhutnameUretici(app).taahhut_word_xml_metnini_guncelle(doc)

    assert doc.paragraphs[0].text == "Revizyon tarihi 01.01.2020"
    assert doc.paragraphs[1].text != "02.02.2020"
    assert doc.paragraphs[2].text == "Yeni Jeofizik"


def test_taahhut_ilgili_idareyi_proje_alanindan_ister():
    app = App()
    app.taahhut_bilgileri = {
        "JEOFIZIK_MUH_AD": "Jeofizik Kişi",
        "JEOFIZIK_MUH_SICIL": "1111",
        "JEOFIZIK_MUH_ADRES": "Adres 1",
        "JEOFIZIK_MUH_TELEFON": "0500 111 11 11",
        "JEOLOJI_MUH_AD": "Jeoloji Kişi",
        "JEOLOJI_MUH_SICIL": "2222",
        "JEOLOJI_MUH_ADRES": "Adres 2",
        "JEOLOJI_MUH_TELEFON": "0500 222 22 22",
    }
    app.proje_deger = lambda kod, varsayilan="": varsayilan
    uretici = TaahhutnameUretici(app)

    with pytest.raises(ValueError, match="ILGILI_IDARE"):
        uretici.taahhut_bilgilerini_dogrula()

    app.proje_deger = lambda kod, varsayilan="": "Test Belediyesi" if kod == "ILGILI_IDARE" else varsayilan
    uretici.taahhut_bilgilerini_dogrula()
    assert uretici.taahhut_ilgili_idare() == "Test Belediyesi"


def test_taahhut_ilgili_idareyi_proje_ilinden_turetir():
    app = App()
    app.proje_deger = lambda kod, varsayilan="": "Çanakkale" if kod == "IL" else varsayilan

    assert TaahhutnameUretici(app).taahhut_ilgili_idare() == "Çanakkale İl Özel İdaresi"
