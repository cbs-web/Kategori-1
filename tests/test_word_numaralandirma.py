from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from word_numaralandirma import (
    _docx_paragraf_alan_kodlari,
    docx_baslik_numaralandirma_hatalari,
)


def _seq_basligi_ekle(doc, etiket, sayi, baslik):
    paragraf = doc.add_paragraph(style="Caption")
    paragraf.add_run(f"{etiket} ")
    alan = OxmlElement("w:fldSimple")
    alan.set(qn("w:instr"), f" SEQ {etiket} \\* ARABIC ")
    sonuc = OxmlElement("w:r")
    metin = OxmlElement("w:t")
    metin.text = str(sayi)
    sonuc.append(metin)
    alan.append(sonuc)
    paragraf._p.append(alan)
    paragraf.add_run(f" {baslik}")
    return paragraf


def test_seq_basliklari_kesintisizse_dogrulama_gecer():
    doc = Document()
    _seq_basligi_ekle(doc, "Şekil", 1, "Birinci şekil")
    _seq_basligi_ekle(doc, "Şekil", 2, "İkinci şekil")
    _seq_basligi_ekle(doc, "Tablo", 1, "Birinci tablo")

    assert docx_baslik_numaralandirma_hatalari(doc) == []


def test_sabit_ve_yinelenen_baslik_numarasi_reddedilir():
    doc = Document()
    _seq_basligi_ekle(doc, "Şekil", 1, "Birinci şekil")
    doc.add_paragraph("Şekil 1 Sabit numaralı şekil", style="Caption")

    hatalar = docx_baslik_numaralandirma_hatalari(doc)

    assert any("SEQ alanı" in hata for hata in hatalar)


def test_rapor_sablonundaki_k1_basliklari_seq_alanidir():
    sablon = (
        Path(__file__).resolve().parents[1]
        / "ornek_sablonlar"
        / "rapor"
        / "TASLAK.docx"
    )
    doc = Document(sablon)
    beklenen = {
        "K1_Sekil_Numune_Lokasyon": "Şekil",
        "K1_Tablo_Numune_Koordinatlari": "Tablo",
        "K1_Sekil_Muhendislik_Jeolojisi": "Şekil",
        "K1_Tablo_Yerel_Zemin_Sinifi": "Tablo",
    }
    bulunan = {}
    for paragraf in doc.paragraphs:
        for bookmark in paragraf._p.xpath(".//w:bookmarkStart"):
            ad = bookmark.get(qn("w:name"))
            if ad in beklenen:
                bulunan[ad] = _docx_paragraf_alan_kodlari(paragraf)

    assert set(bulunan) == set(beklenen)
    for ad, etiket in beklenen.items():
        assert any(f"SEQ {etiket}" in kod for kod in bulunan[ad])
