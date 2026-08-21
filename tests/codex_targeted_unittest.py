import io
import sys
import tempfile
import unittest
import urllib.error
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

import jeoloji_yapay_zeka as yapay_zeka
from cizimler import CizimUretici
from ekler import ek_kategori_durumunu_hazirla, ek_taahhutname_mi, ekleri_denetle
from jeoloji_bolum_paketi import (
    STRATIGRAFIK_KESIT_CAPTION,
    stratigrafik_kesit_bolumunu_ayir,
)
from laboratuvar import laboratuvar_numune_anahtari, laboratuvar_satirlarini_birlestir
from rapor import RaporUretici
from word_jeoloji_birlestirme import _eski_rapor_muhendislik_jeolojisi_bolumunu_cikar


def _png(path, color, size=(640, 360)):
    Image.new("RGB", size, color).save(path)


class CodexTargetedTests(unittest.TestCase):
    def test_stratigrafik_kesit_normal_stil(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            image = root / "kesit.png"
            _png(image, "green")
            doc = Document()
            doc.add_paragraph("2.1 Bölgesel Jeoloji", style="Heading 2")
            doc.add_paragraph("Stratigrafik kesit Şekil 5'te verilmiştir.")
            doc.add_paragraph("Stratigrafik Kesit", style="Heading 3")
            doc.add_picture(str(image))
            doc.add_paragraph("Şekil 5 Çalışma alanı stratigrafik kesiti")
            doc.add_paragraph("2.1.1 Yapısal Jeoloji", style="Heading 3")
            source = root / "source.docx"
            output = root / "output.docx"
            doc.save(source)
            stratigrafik_kesit_bolumunu_ayir(source, output)
            result = Document(output)
            self.assertEqual(result.paragraphs[0].text, STRATIGRAFIK_KESIT_CAPTION)
            self.assertEqual(result.paragraphs[0].style.name, "Caption")
            self.assertNotIn("Stratigrafik Kesit", "\n".join(p.text for p in result.paragraphs))
            self.assertEqual(len(result.inline_shapes), 1)

    def test_eski_1_3_2_bolumu_cikarilir(self):
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "old.docx"
            doc = Document()
            doc.add_paragraph("1.3.1. Önceki Bölüm", style="Heading 3")
            doc.add_paragraph("Korunacak önceki metin.")
            doc.add_paragraph("1.3.2.İnceleme Alanı Mühendislik Jeolojisi", style="Heading 3")
            doc.add_paragraph("Seröiler ve çevresinde gözlenen eski metin.")
            doc.add_paragraph("1.3.3. Sonraki Bölüm", style="Heading 3")
            doc.add_paragraph("Korunacak sonraki metin.")
            doc.save(path)
            self.assertTrue(_eski_rapor_muhendislik_jeolojisi_bolumunu_cikar(path))
            text = "\n".join(p.text for p in Document(path).paragraphs)
            self.assertNotIn("Seröiler", text)
            self.assertIn("1.3.3. Sonraki Bölüm", text)

    def test_pga_sabit_gorseli_degistirir(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            old, pga = root / "old.png", root / "pga.png"
            _png(old, "red")
            _png(pga, "blue")
            doc = Document()
            caption = doc.add_paragraph("Şekil 3: Çanakkale bölgesi deprem tehlike haritası")
            doc.add_paragraph().add_run().add_picture(str(old))
            doc.add_paragraph("(Kaynak)")
            RaporUretici(SimpleNamespace()).rapor_pga_haritasini_yerlestir(doc, caption, str(pga))
            counts = [len(list(p._p.iter(qn("w:drawing")))) for p in doc.paragraphs]
            self.assertEqual(counts, [0, 1, 0])

    def test_laboratuvar_esleme(self):
        keys = {laboratuvar_numune_anahtari(v) for v in ("AÇ1", "AÇ-1", "AÇ 1", "AC1")}
        self.assertEqual(keys, {"AC1"})
        result = laboratuvar_satirlarini_birlestir(
            [("AÇ1", "0.00-1.00", "eski")],
            [("AÇ-1", "0,00 - 1,00", "yeni")],
        )
        self.assertEqual(result["eklenen"], 0)
        self.assertEqual(result["satirlar"][0][0], "AÇ1")

    @staticmethod
    def _http_error(code):
        return urllib.error.HTTPError(
            "https://example.invalid", code, "test", None,
            io.BytesIO(b'{"error":{"message":"test"}}'),
        )

    def test_gemini_retry(self):
        calls, waits = [], []

        def fail(*_args, **_kwargs):
            calls.append(True)
            raise self._http_error(503)

        with patch.object(yapay_zeka.urllib.request, "urlopen", fail), patch.object(
            yapay_zeka.time, "sleep", waits.append
        ):
            with self.assertRaises(yapay_zeka.JeolojiYapayZekaHatasi):
                yapay_zeka._post_json(
                    "https://example.invalid", {}, {}, 1,
                    retry_attempts=3, retry_http_statuses={503}, retry_label="Gemini",
                )
        self.assertEqual(len(calls), 3)
        self.assertEqual(waits, [1, 2])

    def test_ekler_istege_bagli_ve_taahhut_haric(self):
        text, status = ek_kategori_durumunu_hazirla({"EVRAKLAR": []}, "EVRAKLAR")
        self.assertIn("İsteğe bağlı", text)
        self.assertEqual(status, "secondary")
        taahhut = {"baslik": "Jeoloji Taahhütnamesi", "yol": "C:/taahhutname.pdf"}
        self.assertTrue(ek_taahhutname_mi(taahhut))
        self.assertEqual(ekleri_denetle({"EVRAKLAR": [taahhut]}, ["EVRAKLAR"])["toplam"], 0)

    def test_harita_ciktilari(self):
        app = SimpleNamespace(
            yuklu_kml_yolu="C:/proje/TKGM_0_656.kml",
            parsel_haritasi_kaynak_url="https://cbsapi.tkgm.gov.tr/api/parsel/123/0/656",
        )
        text = CizimUretici(app).parsel_haritasi_kaynak_metni()
        self.assertIn("Altlık: Google Uydu", text)
        self.assertNotIn("Taslak parsel gösterimidir", text)
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "yerbulduru.jpg"
            producer = CizimUretici(SimpleNamespace())
            producer.yerbulduru_iki_panel_resim_olustur(
                Image.new("RGB", (1000, 600), "green"),
                Image.new("RGB", (1000, 600), "blue"),
                path,
                (500, 300),
                (500, 300),
            )
            with Image.open(path) as image:
                ratio = image.height / image.width
            self.assertGreaterEqual(ratio, 1.35)
            self.assertLessEqual(ratio, 1.55)


if __name__ == "__main__":
    unittest.main(verbosity=2)
