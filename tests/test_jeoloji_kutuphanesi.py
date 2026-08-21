from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

from formasyon_metin_kutuphanesi import eski_metin_birimlere_dagit

from jeoloji_kutuphanesi import (
    AyniJeolojiKaydiHatasi,
    JeolojiKutuphanesi,
    JeolojiKutuphanesiHatasi,
)
from jeoloji_word_aktarimi import word_raporunu_oku
from jeoloji_klasor_aktarimi import proje_klasorunu_incele
from jeoloji_toplu_aktarim import ilce_klasorunu_tara, toplu_kayitlari_aktar


def _kayit(**overrides):
    record = {
        "il": "Çanakkale",
        "ilce": "Ayvacık",
        "yerlesim": "Küçükkuyu",
        "ada": "12",
        "parsel": "3",
        "formasyon": "Ayvacık Volkaniti (Tmay)",
        "genel_jeoloji_metni": "Genel jeoloji açıklaması.",
        "inceleme_alani_jeolojisi": "İnceleme alanı jeolojisi açıklaması.",
        "onay_durumu": "onayli",
    }
    record.update(overrides)
    return record


def test_oligo_miyosen_granitoyid_basligi_tg_birimine_ayrilir():
    metin = (
        "Oligosen-Miyosen Granitoyidleri (Tg)\n"
        "Biga Yarımadası'ndaki granodiyoritik bileşimli sığ sokulumlar "
        "Oligosen-Geç Miyosen aralığında bölgeye yerleşmiştir. "
        "Jeokronolojik yaşlandırmalar birimin bölgesel konumunu doğrulamaktadır.\n"
        "2.1.1 Yapısal Jeoloji ve Aktif Tektonik"
    )

    sonuc = eski_metin_birimlere_dagit(
        metin,
        [{"kod": "Tg", "ad": "Oligo-Miyosen Granitoyidleri"}],
    )

    assert len(next(iter(sonuc.values()))) > 80


def test_granitoyid_basligi_kod_olmadan_yas_ve_yazim_esdegeriyle_eslesir():
    metin = (
        "Oligosen-Miyosen Granitoyidleri\n"
        "Biga Yarımadası'ndaki granodiyoritik bileşimli sığ sokulumlar "
        "Oligosen-Geç Miyosen aralığında bölgeye yerleşmiştir. "
        "Jeokronolojik yaşlandırmalar birimin bölgesel konumunu doğrulamaktadır.\n"
        "2.1.1 Yapısal Jeoloji ve Aktif Tektonik"
    )

    sonuc = eski_metin_birimlere_dagit(
        metin,
        [{"kod": "", "ad": "Üst Oligosen-Alt Miyosen Granitoyitleri"}],
    )

    assert len(next(iter(sonuc.values()))) > 80


def _proje_raporu_yaz(path, *, ilce, yerlesim, ada, parsel):
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    table = document.add_table(rows=4, cols=2)
    for row, label, value in (
        (0, "İl", "Çanakkale"),
        (1, "İlçe", ilce),
        (2, "Köy", yerlesim),
        (3, "Ada / Parsel", f"{ada} / {parsel}"),
    ):
        table.cell(row, 0).text = label
        table.cell(row, 1).text = value
    document.add_heading("JEOLOJİ", level=1)
    document.add_paragraph("Bölgesel jeoloji açıklaması.")
    document.add_heading("İnceleme Alanı Jeolojisi", level=2)
    document.add_paragraph("Ayvacık Volkaniti birimleri inceleme alanında gözlenmektedir.")
    document.save(path)


def test_kayit_oneri_revizyon_ve_arsiv_akisi(tmp_path):
    library = JeolojiKutuphanesi(tmp_path / "jeoloji" / "canakkale.db")

    record_id = library.kaydet(_kayit())
    saved = library.getir(record_id)
    assert saved["ilce"] == "Ayvacık"
    assert saved["revizyon_no"] == 1

    listed = library.listele(ilce="ayvacik", formasyon="Ayvacık Volkaniti (Tmay)")
    assert [item["id"] for item in listed] == [record_id]

    suggestion = library.uygun_icerigi_bul(
        il="Çanakkale",
        ilce="Ayvacık",
        yerlesim="Küçükkuyu",
        ada="12",
        parsel="3",
        formasyon="Ayvacık Volkaniti (Tmay)",
    )
    assert suggestion["genel_jeoloji_metni"] == "Genel jeoloji açıklaması."
    assert suggestion["inceleme_alani_jeolojisi"] == "İnceleme alanı jeolojisi açıklaması."
    assert suggestion["kayit_idleri"] == [record_id]

    library.kaydet(
        _kayit(genel_jeoloji_metni="İkinci revizyon."),
        kayit_id=record_id,
    )
    assert library.getir(record_id)["revizyon_no"] == 2
    assert len(library.revizyonlar(record_id)) == 2

    assert library.arsivle(record_id) is True
    assert library.getir(record_id) is None
    assert library.uygun_icerigi_bul(ilce="Ayvacık") is None
    assert library.getir(record_id, aktif_olmayan=True)["aktif"] is False


def test_ayni_kunye_etkin_kaydi_engeller(tmp_path):
    library = JeolojiKutuphanesi(tmp_path / "canakkale.db")
    record_id = library.kaydet(_kayit())

    with pytest.raises(AyniJeolojiKaydiHatasi) as error:
        library.kaydet(_kayit(genel_jeoloji_metni="Başka bir metin."))

    assert error.value.kayit_id == record_id


def test_gecersiz_kayit_acik_hata_verir(tmp_path):
    library = JeolojiKutuphanesi(tmp_path / "canakkale.db")

    with pytest.raises(JeolojiKutuphanesiHatasi, match="İlçe"):
        library.kaydet({"il": "Çanakkale", "onay_durumu": "taslak"})


def test_word_raporu_okunur_ve_jeoloji_bolumu_kutuphaneye_paketlenir(tmp_path):
    source = tmp_path / "Ayvacık_Küçükkuyu.docx"
    document = Document()
    table = document.add_table(rows=4, cols=2)
    for row, label, value in (
        (0, "İl", "Çanakkale"),
        (1, "İlçe", "Ayvacık"),
        (2, "Köy", "Küçükkuyu"),
        (3, "Ada / Parsel", "12 / 3"),
    ):
        table.cell(row, 0).text = label
        table.cell(row, 1).text = value
    document.add_heading("JEOLOJİ", level=1)
    document.add_paragraph("Bölgesel jeoloji açıklaması.")
    document.add_heading("İnceleme Alanı Jeolojisi", level=2)
    document.add_paragraph("Ayvacık Volkaniti birimleri inceleme alanında gözlenmektedir.")
    document.save(source)

    result = word_raporunu_oku(source)
    assert result.hata == ""
    assert result.ilce == "Ayvacık"
    assert result.yerlesim == "Küçükkuyu"
    assert result.ada == "12"
    assert result.parsel == "3"
    assert result.genel_jeoloji_metni
    assert result.inceleme_alani_jeolojisi

    library = JeolojiKutuphanesi(tmp_path / "library" / "canakkale.db")
    record_id = library.kaydet(result.kutuphane_kaydi())
    saved = library.getir(record_id)
    assert Path(saved["bolum_docx_path"]).is_file()
    assert saved["bolum_hash"]

    backup = tmp_path / "backup" / "canakkale.zip"
    library.yedek_paketi_olustur(backup)
    with ZipFile(backup) as archive:
        names = set(archive.namelist())
    assert "canakkale.db" in names
    assert "manifest.json" in names
    assert any(name.startswith("jeoloji_bolumleri/") for name in names)


def test_klasor_word_ve_kml_adaylarini_bulur(tmp_path):
    source = tmp_path / "raporlar" / "Ayvacik_12_3_rapor.docx"
    source.parent.mkdir()
    document = Document()
    table = document.add_table(rows=4, cols=2)
    for row, label, value in (
        (0, "İl", "Çanakkale"),
        (1, "İlçe", "Ayvacık"),
        (2, "Köy", "Küçükkuyu"),
        (3, "Ada / Parsel", "12 / 3"),
    ):
        table.cell(row, 0).text = label
        table.cell(row, 1).text = value
    document.add_heading("JEOLOJİ", level=1)
    document.add_paragraph("Genel jeoloji açıklaması.")
    document.add_heading("İnceleme Alanı Jeolojisi", level=2)
    document.add_paragraph("Ayvacık Volkaniti inceleme alanında gözlenir.")
    document.save(source)
    kml = tmp_path / "veri" / "tkgm-12-ada-3-parsel.kml"
    kml.parent.mkdir()
    kml.write_text(
        "<kml><Placemark><Polygon><outerBoundaryIs><LinearRing><coordinates>"
        "26,40 26.01,40 26.01,40.01 26,40"
        "</coordinates></LinearRing></outerBoundaryIs></Polygon></Placemark></kml>",
        encoding="utf-8",
    )

    sonuc = proje_klasorunu_incele(tmp_path)

    assert Path(sonuc["word_adaylari"][0]["sonuc"].dosya_yolu) == source
    assert Path(sonuc["kml_adaylari"][0]["dosya_yolu"]) == kml


def test_geometri_saklanir_ve_harita_alaniyla_sorgulanir(tmp_path):
    library = JeolojiKutuphanesi(tmp_path / "library" / "canakkale.db")
    approved_id = library.kaydet(_kayit())
    draft_id = library.kaydet(
        _kayit(parsel="4", onay_durumu="taslak", genel_jeoloji_metni="Taslak rapor.")
    )
    kml = tmp_path / "parsel.kml"
    kml.write_text("<kml/>", encoding="utf-8")
    approved_polygon = {"ad": "12/3", "noktalar": [(40, 26), (40, 26.01), (40.01, 26.01), (40, 26)]}
    draft_polygon = {"ad": "12/4", "noktalar": [(40.02, 26), (40.02, 26.01), (40.03, 26.01), (40.02, 26)]}
    library.geometrileri_degistir(approved_id, [approved_polygon], kml)
    library.geometrileri_degistir(draft_id, [draft_polygon], kml)

    records = library.harita_kayitlari(
        min_enlem=39.9, max_enlem=40.1, min_boylam=25.9, max_boylam=26.1
    )
    assert [record["id"] for record in records] == [approved_id]
    records_with_drafts = library.harita_kayitlari(
        min_enlem=39.9,
        max_enlem=40.1,
        min_boylam=25.9,
        max_boylam=26.1,
        taslaklari_goster=True,
    )
    assert {record["id"] for record in records_with_drafts} == {approved_id, draft_id}
    assert library.getir(approved_id)["geometri_sayisi"] == 1

    backup = tmp_path / "backup.zip"
    library.yedek_paketi_olustur(backup)
    with ZipFile(backup) as archive:
        assert any(name.startswith("parsel_kml/") for name in archive.namelist())


def test_ilce_klasoru_projeleri_ayirir_toplu_aktarir_ve_tekrari_tanir(tmp_path):
    root = tmp_path / "AYVACIK"
    kozlu = root / "KOZLU" / "Ayvacık Kozlu 151-10 Ercan Şahin"
    tuzla = root / "TUZLA" / "Ayvacık Tuzla 20-4 Örnek"
    _proje_raporu_yaz(
        kozlu / "Ayvacık Kozlu 151-10 Rapor.docx",
        ilce="Ayvacık",
        yerlesim="Kozlu",
        ada="151",
        parsel="10",
    )
    _proje_raporu_yaz(
        tuzla / "Ayvacık Tuzla 20-4 Rapor.docx",
        ilce="Ayvacık",
        yerlesim="Tuzla",
        ada="20",
        parsel="4",
    )
    yardimci = kozlu / "JEOFİZİK" / "MASW değerlendirme.docx"
    yardimci.parent.mkdir()
    Document().save(yardimci)
    (kozlu / "tkgm-parsel-sorgu-sonuc-151-ada-10-parsel.kml").write_text(
        "<kml><Placemark><name>Kozlu 151 ada 10 parsel</name><Polygon>"
        "<outerBoundaryIs><LinearRing><coordinates>"
        "26,40 26.01,40 26.01,40.01 26,40"
        "</coordinates></LinearRing></outerBoundaryIs></Polygon></Placemark></kml>",
        encoding="utf-8",
    )
    library = JeolojiKutuphanesi(tmp_path / "library" / "canakkale.db")

    tarama = ilce_klasorunu_tara(root, kutuphane=library)
    by_folder = {Path(item["klasor"]).name: item for item in tarama["projeler"]}
    assert by_folder[kozlu.name]["durum_kodu"] == "hazir"
    assert by_folder[kozlu.name]["word_adayi"]["sonuc"].ada == "151"
    assert by_folder[tuzla.name]["durum_kodu"] == "kml_yok"
    assert all("JEOFİZİK" not in item["goreli_klasor"] for item in tarama["projeler"])

    sonuc = toplu_kayitlari_aktar(tarama["projeler"], library, onayli=True)
    assert len(sonuc["basarili"]) == 1
    saved = library.getir(sonuc["basarili"][0]["kayit_id"])
    assert saved["onay_durumu"] == "onayli"
    assert Path(saved["kaynak_klasor_path"]) == kozlu
    assert saved["geometri_sayisi"] == 1

    ikinci_tarama = ilce_klasorunu_tara(root, kutuphane=library)
    ikinci = {Path(item["klasor"]).name: item for item in ikinci_tarama["projeler"]}
    assert ikinci[kozlu.name]["durum_kodu"] == "mevcut"
