"""Gerçek rapor şablonunda eski ve toplu etiket değiştirmeyi karşılaştırır."""

from __future__ import annotations

import argparse
import re
import statistics
import sys
import time
from pathlib import Path

from docx import Document


KOK = Path(__file__).resolve().parents[1]
if str(KOK) not in sys.path:
    sys.path.insert(0, str(KOK))

from rapor import RaporUretici  # noqa: E402


class _App:
    def hata_kaydet(self, *_args, **_kwargs):
        pass


def _etiketleri_oku(doc):
    metin = []
    for part in doc.part.package.parts:
        partname = str(getattr(part, "partname", ""))
        if not partname.startswith("/word/") or not hasattr(part, "element"):
            continue
        try:
            metin.extend(node.text or "" for node in part.element.xpath(".//w:t"))
        except Exception:
            continue
    return sorted(set(re.findall(r"\[[A-Z0-9_]+\]", "".join(metin))))


def _eski_metin_degistir(uretici, doc, tag, value):
    val_str = str(value)
    for paragraph in uretici.docx_paragraflarini_dolas(doc):
        if tag not in paragraph.text:
            continue
        replaced = False
        for run in paragraph.runs:
            if tag in run.text:
                run.text = run.text.replace(tag, val_str)
                replaced = True
        if not replaced:
            full_text = "".join(run.text for run in paragraph.runs)
            if tag in full_text and paragraph.runs:
                paragraph.runs[0].text = full_text.replace(tag, val_str)
                for run in paragraph.runs[1:]:
                    run.text = ""
    uretici.rapor_xml_metin_degistir(doc, tag, val_str)


def _olc(sablon, degisimler, toplu):
    doc = Document(sablon)
    uretici = RaporUretici(_App())
    baslangic = time.perf_counter()
    if toplu:
        uretici.rapor_metinleri_toplu_degistir(doc, degisimler)
    else:
        for etiket, deger in degisimler.items():
            _eski_metin_degistir(uretici, doc, etiket, deger)
    return time.perf_counter() - baslangic


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--template",
        type=Path,
        default=KOK / "ornek_sablonlar" / "rapor" / "TASLAK.docx",
    )
    parser.add_argument("--repeat", type=int, default=3)
    args = parser.parse_args()

    etiketler = [
        etiket
        for etiket in _etiketleri_oku(Document(args.template))
        if etiket != "[FORMASYON]"
    ]
    degisimler = {
        etiket: f"Örnek değer {index}"
        for index, etiket in enumerate(etiketler, start=1)
    }

    # Disk önbelleğini ısıt; ölçüme yalnız değiştirme süresi dahildir.
    Document(args.template)
    eski = [_olc(args.template, degisimler, toplu=False) for _ in range(args.repeat)]
    yeni = [_olc(args.template, degisimler, toplu=True) for _ in range(args.repeat)]
    eski_medyan = statistics.median(eski)
    yeni_medyan = statistics.median(yeni)
    hizlanma = eski_medyan / yeni_medyan if yeni_medyan else float("inf")
    azalma = (1 - yeni_medyan / eski_medyan) * 100 if eski_medyan else 0

    print(f"Şablon: {args.template}")
    print(f"Etiket: {len(degisimler)}; tekrar: {args.repeat}")
    print("Eski (s): " + ", ".join(f"{sure:.4f}" for sure in eski))
    print("Yeni (s): " + ", ".join(f"{sure:.4f}" for sure in yeni))
    print(f"Medyan: {eski_medyan:.4f}s -> {yeni_medyan:.4f}s")
    print(f"Hızlanma: {hizlanma:.1f}x; süre azalması: %{azalma:.1f}")


if __name__ == "__main__":
    main()
