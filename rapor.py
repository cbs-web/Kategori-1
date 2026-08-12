import csv
import math
import os
import re
import unicodedata
from bisect import bisect_right
from copy import deepcopy
import tempfile
import tkinter as tk

import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.parts.numbering import NumberingPart
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from PIL import Image, ImageOps

from laboratuvar import (
    rapor_lab_tablolarini_ekle as laboratuvar_rapor_lab_tablolarini_ekle,
    rapor_lab_tablosu_olustur as laboratuvar_rapor_lab_tablosu_olustur,
)
from tasima import BOWLES_ZEMIN_KATSAYISI, TBDY2018TasimaGucu
from word_jeoloji_birlestirme import JEOLOJI_WORD_EKLEME_ISARETI
from word_numaralandirma import docx_baslik_numaralandirma_hatalari


RAPOR_TAAHHUT_BILGI_ALANLARI = (
    "JEOFIZIK_MUH_AD",
    "JEOFIZIK_MUH_SICIL",
    "JEOFIZIK_MUH_ADRES",
    "JEOFIZIK_MUH_TELEFON",
    "JEOLOJI_MUH_AD",
    "JEOLOJI_MUH_SICIL",
    "JEOLOJI_MUH_ADRES",
    "JEOLOJI_MUH_TELEFON",
)


class RaporUretici:
    def __init__(self, app):
        object.__setattr__(self, "app", app)

    def __getattr__(self, name):
        return getattr(self.app, name)

    def __setattr__(self, name, value):
        if name == "app":
            object.__setattr__(self, name, value)
        else:
            setattr(self.app, name, value)

    def docx_paragraflarini_dolas(self, doc):
        for p in doc.paragraphs:
            yield p
        for t in doc.tables:
            for row in t.rows:
                for c in row.cells:
                    for p in c.paragraphs:
                        yield p
        for section in doc.sections:
            headers_footers = [
                section.header, section.first_page_header, section.even_page_header,
                section.footer, section.first_page_footer, section.even_page_footer
            ]
            for hf in headers_footers:
                if hf:
                    for p in hf.paragraphs:
                        yield p
                    for table in hf.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    yield p

    def docx_xml_metin_dugumlerini_dolas(self, doc):
        for part in doc.part.package.parts:
            partname = str(getattr(part, "partname", ""))
            if not partname.startswith("/word/") or not hasattr(part, "element"):
                continue
            try:
                for node in part.element.xpath(".//w:t"):
                    yield node
            except Exception:
                continue

    def desteklenen_sablon_etiketleri(self):
        statik_etiketler = {
            "[BINA]", "[CALISMA]", "[ARAZI_CALISMA_ACIKLAMA]", "[ARAZI_CALISMA_OZETI]", "[LITOLOJI]", "[QK]", "[QT]", "[KS]", "[DF]", "[FORMASYON]",
            "[FORMASYON_ADI]", "[FORMASYON_KISA]", "[BIRIM_TANIMI]", "[MUHENDISLIK_JEOLOJISI_METNI]", "[JEOLOJI_BOLUMU]",
            "[VS30_CUMLESI]", "[YASS_SONUC_METNI]",
            "[MJH]", "[JEOFIZIK_LOKASYON]", "[JEOLOJI_LOKASYON]", "[YERBULDURU]", "[PARSEL_HARITASI]",
            "[JEOFIZIK_KOORDINAT]", "[JEOLOJI_KOORDINAT]", "[JEOFIZIK_PARAMETRE]", "[JEOFON_DIZILIM]",
            "[MASW]", "[VP]", "[TASIMA_GUCU]", "[KESIT]", "[KESIT_ACIKLAMA]", "[AC_LOGLARI]", "[LAB]"
        }
        taahhut_etiketleri = {f"[{kod}]" for kod in RAPOR_TAAHHUT_BILGI_ALANLARI}
        return {f"[{kod}]" for kod in self.veri_alanlari.keys()} | statik_etiketler | taahhut_etiketleri

    def rapor_kayit_tipi(self, isim):
        temiz = str(isim or "").strip().upper().replace(" ", "")
        if temiz.startswith(("AÇ", "AC")):
            return "AC"
        if temiz.startswith("YN"):
            return "YN"
        return ""

    def rapor_sayi_int(self, deger, varsayilan=0, minimum=0, maksimum=99):
        try:
            sayi = int(float(str(deger).replace(",", ".")))
        except Exception:
            sayi = varsayilan
        return max(minimum, min(maksimum, sayi))

    def rapor_jeofizik_profil_adedi(self):
        bilgiler = {}
        if hasattr(self.app, "jeofon_dizilim_bilgileri_al"):
            bilgiler = self.jeofon_dizilim_bilgileri_al() or {}
        deger = bilgiler.get("jeofizik_profil_adedi", "0")
        return self.rapor_sayi_int(deger, varsayilan=0, minimum=0, maksimum=99)

    def rapor_arazi_calisma_ozeti_olustur(self, ac_sayisi, yn_sayisi):
        jeofizik_adedi = self.rapor_jeofizik_profil_adedi()

        cumleler = []
        if ac_sayisi > 0:
            cumleler.append(f"İnceleme alanında {ac_sayisi} adet araştırma çukuru açılmıştır.")
        if yn_sayisi > 0:
            cumleler.append(f"İnceleme alanından {yn_sayisi} adet yüzey numunesi alınmıştır.")

        if jeofizik_adedi > 0:
            cumleler.append(
                f"Ayrıca jeofizik çalışmalar kapsamında {jeofizik_adedi} profil Sismik Kırılma "
                "(P dalgası) ve Yüzey Dalgalarının Çok Kanallı Analizi (MASW) ölçümleri yapılmıştır."
            )

        return " ".join(cumleler) if cumleler else "Arazi çalışması verisi bulunmamaktadır."

    def rapor_sayi_kisa_formatla(self, deger):
        sayi = float(str(deger).strip().replace(",", "."))
        if math.isclose(sayi, round(sayi), rel_tol=0.0, abs_tol=1e-9):
            return str(int(round(sayi)))
        return f"{sayi:.2f}".rstrip("0").rstrip(".").replace(".", ",")

    def rapor_arastirma_cukuru_aciklamasi_olustur(self, ac_yn_kayitlari):
        ac_kayitlari = [
            kayit for kayit in ac_yn_kayitlari
            if self.rapor_kayit_tipi(kayit.get("isim", "")) == "AC"
        ]
        if not ac_kayitlari:
            return "Çalışma alanında araştırma çukuru açılmamıştır."

        derinlikler = []
        for kayit in ac_kayitlari:
            entry = kayit.get("derinlik_entry")
            ham = str(entry.get()).strip() if entry is not None else ""
            if not ham:
                continue
            try:
                deger = self.rapor_sayi_kisa_formatla(ham)
            except ValueError:
                deger = ham
            if deger not in derinlikler:
                derinlikler.append(deger)

        adet = len(ac_kayitlari)
        if len(derinlikler) == 1:
            return f"Çalışma alanında {derinlikler[0]} m derinliğinde {adet} adet araştırma çukuru kazılmıştır."
        if derinlikler:
            derinlik_metni = ", ".join(derinlikler[:-1])
            if len(derinlikler) > 1:
                derinlik_metni += f" ve {derinlikler[-1]}"
            return f"Çalışma alanında derinlikleri {derinlik_metni} m olan {adet} adet araştırma çukuru kazılmıştır."
        return f"Çalışma alanında {adet} adet araştırma çukuru kazılmıştır."

    def rapor_muhendislik_jeolojisi_metnini_hazirla(self, metin, lito_parcalari):
        # Kütüphane Word'ündeki eski/sabit kayaç cümlesini taşımak yerine her
        # projede arazi sekmelerine girilmiş güncel litoloji açıklamasını kullan.
        del metin
        benzersiz = []
        for parca in lito_parcalari or []:
            temiz = " ".join(str(parca or "").split()).strip(" .;,:")
            if temiz and temiz.casefold() not in {deger.casefold() for deger in benzersiz}:
                benzersiz.append(temiz)
        if benzersiz:
            if len(benzersiz) == 1:
                litoloji = benzersiz[0]
            else:
                litoloji = ", ".join(benzersiz[:-1]) + f" ve {benzersiz[-1]}"
            return f"Çalışma alanında birimler, {litoloji} olarak gözlenmiştir."
        return "Çalışma alanında gözlenen birimlere ilişkin litoloji açıklaması girilmemiştir."

    def rapor_vs30_cumlesi_olustur(self):
        if not hasattr(self.app, "jeofizik_excel_yolu_al") or not self.jeofizik_excel_yolu_al():
            return "Vs30 değeri jeofizik çalışma sonuçlarına göre değerlendirilmiştir."
        param_ss_list = self.rapor_jeofizik_parametrelerini_oku()
        self._rapor_param_ss_list = param_ss_list
        degerler = []
        for serim in param_ss_list:
            for deger in serim.get("raw_vs30", [])[2:]:
                if pd.isna(deger) or not str(deger).strip():
                    continue
                metin = self.rapor_deger_formatla(str(deger).strip())
                if metin not in degerler:
                    degerler.append(metin)
                break
        if not degerler:
            return "Vs30 değeri jeofizik çalışma sonuçlarına göre değerlendirilmiştir."
        if len(degerler) == 1:
            return f"Jeofizik çalışmalarla belirlenen Vs30 değeri {degerler[0]} m/s'dir."
        return "Jeofizik çalışmalarda belirlenen Vs30 değerleri " + ", ".join(degerler) + " m/s'dir."

    def rapor_yass_sonuc_metni_olustur(self, ac_yn_kayitlari):
        degerler = []
        for kayit in ac_yn_kayitlari:
            if self.rapor_kayit_tipi(kayit.get("isim", "")) != "AC":
                continue
            for satir in self.ac_yn_satirlari(kayit):
                ham = str(satir[2]).strip() if len(satir) > 2 else ""
                if not ham or ham == "-":
                    continue
                try:
                    deger = self.rapor_sayi_kisa_formatla(ham)
                except ValueError:
                    deger = ham
                if deger not in degerler:
                    degerler.append(deger)
        if degerler:
            return "Araştırma çukurlarında yeraltı suyu " + ", ".join(degerler) + " m derinlikte gözlenmiştir."
        return "Araştırma çukurlarında yeraltı suyuna rastlanmamıştır."

    def rapor_derinlik_araligi_coz(self, derinlik):
        sayilar = re.findall(r"\d+(?:[.,]\d+)?", str(derinlik or ""))
        if len(sayilar) < 2:
            return None
        try:
            bas = float(sayilar[0].replace(",", "."))
            bit = float(sayilar[1].replace(",", "."))
        except ValueError:
            return None
        if bit < bas:
            raise ValueError(
                f"Derinlik aralığı ters girilmiş: {derinlik}. Başlangıç derinliği bitişten büyük olamaz."
            )
        return bas, bit

    def rapor_metre_formatla(self, deger):
        return f"{float(deger):.2f}"

    def rapor_kesit_katmanlarini_topla(self, ac_yn_kayitlari):
        katmanlar = []
        gorulen = set()
        for kayit in ac_yn_kayitlari:
            tip = self.rapor_kayit_tipi(kayit.get("isim", ""))
            for satir in self.ac_yn_satirlari(kayit):
                if len(satir) < 4:
                    continue
                derinlik = str(satir[0]).strip()
                tanim = str(satir[3]).strip()
                if not tanim or tanim == "-":
                    continue
                aralik = self.rapor_derinlik_araligi_coz(derinlik)
                bas, bit = aralik if aralik else (None, None)
                anahtar = (
                    tip,
                    round(bas, 3) if bas is not None else derinlik,
                    round(bit, 3) if bit is not None else derinlik,
                    tanim.lower(),
                )
                if anahtar in gorulen:
                    continue
                gorulen.add(anahtar)
                katmanlar.append({
                    "tip": tip,
                    "isim": kayit.get("isim", ""),
                    "derinlik": derinlik,
                    "bas": bas,
                    "bit": bit,
                    "tanim": tanim,
                })

        return sorted(
            katmanlar,
            key=lambda k: (
                9999 if k["bas"] is None else k["bas"],
                9999 if k["bit"] is None else k["bit"],
                k["tanim"],
            )
        )

    def rapor_katman_cumlesi(self, katman, sira):
        tanim = katman["tanim"]
        bas = katman["bas"]
        bit = katman["bit"]
        if bas is not None and bit is not None:
            aralik = f"{self.rapor_metre_formatla(bas)}-{self.rapor_metre_formatla(bit)} m aralığında"
        else:
            aralik = f"{katman['derinlik']} m aralığında" if katman["derinlik"] else "devamında"

        tanim_kucuk = tanim.lower()
        if sira == 0 and ("bitkisel" in tanim_kucuk or "toprak" in tanim_kucuk):
            return f"En üstte yaklaşık {aralik} {tanim} gözlenmiştir."
        if sira == 0:
            return f"Yüzeyde yaklaşık {aralik} {tanim} gözlenmiştir."
        return f"Bunun altında {aralik} {tanim} birimi devam etmektedir."

    def rapor_kesit_aciklama_olustur(self, ac_yn_kayitlari):
        ac_sayisi = sum(1 for kayit in ac_yn_kayitlari if self.rapor_kayit_tipi(kayit.get("isim", "")) == "AC")
        yn_sayisi = sum(1 for kayit in ac_yn_kayitlari if self.rapor_kayit_tipi(kayit.get("isim", "")) == "YN")
        katmanlar = self.rapor_kesit_katmanlarini_topla(ac_yn_kayitlari)

        if not katmanlar:
            if ac_sayisi == 0 and yn_sayisi > 0:
                return (
                    f"Araştırma çukuru kaydı bulunmamakta, {yn_sayisi} adet yüzey numunesi kaydı bulunmaktadır; "
                    "kesit açıklaması yalnız programa girilmiş verilerle sınırlıdır."
                )
            if ac_sayisi > 0:
                return "Araştırma çukuru verilerinde kesit açıklaması oluşturacak zemin tanımı bulunmamaktadır."
            return "Kesit açıklaması için araştırma çukuru veya yüzey numunesi verisi girilmemiştir."

        ac_katmanlari = [katman for katman in katmanlar if katman["tip"] == "AC"]
        kaynak_katmanlar = ac_katmanlari or katmanlar
        cumleler = []

        if ac_sayisi > 0 and yn_sayisi > 0:
            cumleler.append(
                f"Açılan {ac_sayisi} adet araştırma çukuru ve alınan {yn_sayisi} adet yüzey numunesi verilerine göre kesit değerlendirmesi yapılmıştır."
            )
        elif ac_sayisi > 0:
            cumleler.append(f"Açılan {ac_sayisi} adet araştırma çukuru verilerine göre kesit değerlendirmesi yapılmıştır.")
        else:
            cumleler.append(
                f"Araştırma çukuru kaydı bulunmamakta, {yn_sayisi} adet yüzey numunesi kaydı bulunmaktadır."
            )

        for index, katman in enumerate(kaynak_katmanlar[:4]):
            cumleler.append(self.rapor_katman_cumlesi(katman, index))

        return " ".join(cumleler)

    def sablon_etiketlerini_oku(self, dosya_yolu):
        import docx
        doc = docx.Document(dosya_yolu)
        return self.rapor_docx_etiketlerini_oku(doc)

    def rapor_docx_etiketlerini_oku(self, doc):
        metinler = [p.text for p in self.docx_paragraflarini_dolas(doc)]
        xml_dugumleri = list(self.docx_xml_metin_dugumlerini_dolas(doc))
        metinler.extend(node.text or "" for node in xml_dugumleri)
        metinler.append("".join(node.text or "" for node in xml_dugumleri))
        metin = "\n".join(metinler)
        etiketler = set(re.findall(r"\[[A-Z0-9_]+\]", metin))
        return {e for e in etiketler if not e.strip("[]").isdigit()}

    def rapor_update_fields_ayarla(self, doc):
        settings = doc.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    def rapor_belge_akisini_duzenle(self, doc):
        for paragraph in self.docx_paragraflarini_dolas(doc):
            style = getattr(paragraph, "style", None)
            style_text = " ".join(
                str(value or "")
                for value in (
                    getattr(style, "style_id", ""),
                    getattr(style, "name", ""),
                )
            ).casefold()
            if "heading" in style_text or "başlık" in style_text or "baslik" in style_text:
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.keep_together = True
            if "caption" in style_text or "resim yaz" in style_text:
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.keep_together = True

        for table in doc.tables:
            if not table.rows:
                continue
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:tblHeader")) is None:
                tr_pr.append(OxmlElement("w:tblHeader"))

    def rapor_word_alan_hedeflerini_denetle(self, doc):
        bookmark_adlari = set()
        alanlar = []
        for part in doc.part.package.parts:
            if not hasattr(part, "element"):
                continue
            for bookmark in part.element.xpath(".//w:bookmarkStart"):
                ad = bookmark.get(qn("w:name"))
                if ad:
                    bookmark_adlari.add(ad)
            for node in part.element.xpath(".//w:instrText"):
                kod = str(node.text or "").strip()
                if kod:
                    alanlar.append(kod)

        eksik = []
        for kod in alanlar:
            eslesme = re.search(r"\b(?:REF|PAGEREF)\s+([^\s\\]+)", kod, re.IGNORECASE)
            if eslesme and eslesme.group(1) not in bookmark_adlari:
                eksik.append(eslesme.group(1))
        return sorted(set(eksik))

    def rapor_cikti_belgesini_dogrula(self, doc):
        self.rapor_cozulemeyen_etiketleri_dogrula(doc)
        metin = "\n".join(node.text or "" for node in self.docx_xml_metin_dugumlerini_dolas(doc))
        hata_ifadeleri = (
            "Hata! Başvuru kaynağı bulunamadı",
            "Error! Reference source not found",
        )
        bulunan_hatalar = [ifade for ifade in hata_ifadeleri if ifade.casefold() in metin.casefold()]
        if bulunan_hatalar:
            raise ValueError("Rapor Word alanlarında bozuk çapraz başvuru bulundu.")

        eksik_hedefler = self.rapor_word_alan_hedeflerini_denetle(doc)
        if eksik_hedefler:
            raise ValueError(
                "Rapor Word alanlarında hedefi bulunmayan çapraz başvurular var: "
                + ", ".join(eksik_hedefler[:10])
            )

        numaralandirma_hatalari = docx_baslik_numaralandirma_hatalari(doc)
        if numaralandirma_hatalari:
            raise ValueError(
                "Rapor şekil/tablo numaralandırması tutarsız: "
                + " ".join(numaralandirma_hatalari[:5])
            )

        if re.search(r"q\s*(?:_|ₒ|o)\s*(?:≤|<=)\s*q\s*(?:_|ₜ|t)", metin, re.IGNORECASE):
            raise ValueError("Raporda kullanılmaması gereken q_o ≤ q_t karşılaştırması bulundu.")

        jeoloji_basliklari = [
            p for p in doc.paragraphs
            if self.rapor_ana_jeoloji_basligi_mi(p.text)
        ]
        if len(jeoloji_basliklari) != 1:
            raise ValueError(
                f"Raporda 2. JEOLOJİ ana başlığı bir kez bulunmalıdır; bulunan: {len(jeoloji_basliklari)}."
            )

        if JEOLOJI_WORD_EKLEME_ISARETI in metin:
            raise ValueError("Jeoloji Word içeriği rapora eklenememiş; ekleme işareti hâlâ mevcut.")

    def rapor_cozulemeyen_etiketleri_dogrula(self, doc):
        kalan = sorted(self.rapor_docx_etiketlerini_oku(doc))
        if kalan:
            raise ValueError(
                "Rapor üretimi tamamlanamadı; çözülemeyen şablon etiketleri: "
                + ", ".join(kalan)
            )

    def rapor_etiket_verilerini_hazirla(self):
        ac_sekmeleri_lito = []
        ac_yn_kayitlari = self.ac_yn_sekme_kayitlari()
        for kayit in ac_yn_kayitlari:
            ac_sekmeleri_lito.append(self.ac_yn_satirlari(kayit))

        lito_text_parts = []
        for satirlar in ac_sekmeleri_lito:
            for satir in satirlar:
                zemin_tanimi = str(satir[3]).strip() if len(satir) > 3 else ""
                if zemin_tanimi and zemin_tanimi != "-" and zemin_tanimi not in lito_text_parts:
                    lito_text_parts.append(zemin_tanimi)

        final_lito_text = ", ".join(lito_text_parts) if lito_text_parts else "Araştırma Çukuru / Yüzey Numunesi verisi bulunmamaktadır."

        ac_sayisi = sum(1 for kayit in ac_yn_kayitlari if self.rapor_kayit_tipi(kayit.get("isim", "")) == "AC")
        yn_sayisi = sum(1 for kayit in ac_yn_kayitlari if self.rapor_kayit_tipi(kayit.get("isim", "")) == "YN")
        arazi_calisma_ozeti = self.rapor_arazi_calisma_ozeti_olustur(ac_sayisi, yn_sayisi)

        if ac_sayisi > 0 and yn_sayisi > 0:
            final_calisma_text = f"{ac_sayisi} adet noktada araştırma çukuru çalışması yapılmış ve {yn_sayisi} adet noktadan yüzey numunesi alınmıştır."
        elif ac_sayisi > 0:
            final_calisma_text = f"{ac_sayisi} adet noktada araştırma çukuru çalışması yapılmıştır."
        elif yn_sayisi > 0:
            final_calisma_text = f"{yn_sayisi} adet noktadan yüzey numunesi alınmıştır."
        else:
            final_calisma_text = "Arazi çalışması verisi bulunmamaktadır."

        arazi_calisma_aciklama = self.rapor_arastirma_cukuru_aciklamasi_olustur(ac_yn_kayitlari)

        nihai_qt = self.entry_qt_nihai.get().strip() if hasattr(self, "entry_qt_nihai") else ""
        nihai_ks = self.entry_ks_nihai.get().strip() if hasattr(self, "entry_ks_nihai") else ""
        df_metni = ""
        df_val = self.tg_girdiler.get("Df") if hasattr(self, "tg_girdiler") else None
        if df_val:
            ham_df = str(df_val.get()).strip()
            if ham_df:
                try:
                    df_sayi = float(ham_df.replace(",", "."))
                except ValueError as exc:
                    raise ValueError("Temel derinliği [DF] sayısal olmalıdır.") from exc
                if not math.isfinite(df_sayi) or df_sayi < 0:
                    raise ValueError("Temel derinliği [DF] sonlu ve negatif olmayan bir sayı olmalıdır.")
                df_metni = f"{df_sayi:.2f}"
        formasyon_bilgileri = self.formasyon_bilgilerini_hazirla()
        birim_tanimi = str(
            formasyon_bilgileri.get("adi")
            or formasyon_bilgileri.get("birim_tanimi")
            or ""
        ).strip()
        fm_metin = str(
            formasyon_bilgileri.get("formasyon_metni")
            or birim_tanimi
        ).strip()
        muhendislik_metni = self.rapor_muhendislik_jeolojisi_metnini_hazirla(
            formasyon_bilgileri.get("muhendislik_metni", ""),
            lito_text_parts,
        )
        vs30_cumlesi = self.rapor_vs30_cumlesi_olustur()

        degisim_sozlugu = {f"[{kod}]": entry.get() for kod, entry in self.veri_alanlari.items()}
        taahhut_bilgileri = getattr(self, "taahhut_bilgileri", {})
        if not isinstance(taahhut_bilgileri, dict):
            taahhut_bilgileri = {}
        for kod in RAPOR_TAAHHUT_BILGI_ALANLARI:
            degisim_sozlugu[f"[{kod}]"] = str(taahhut_bilgileri.get(kod, "")).strip()
        degisim_sozlugu["[FORMASYON]"] = fm_metin
        degisim_sozlugu["[FORMASYON_ADI]"] = birim_tanimi
        degisim_sozlugu["[FORMASYON_KISA]"] = str(formasyon_bilgileri.get("kisa") or "").strip()
        degisim_sozlugu["[BIRIM_TANIMI]"] = birim_tanimi
        degisim_sozlugu["[MUHENDISLIK_JEOLOJISI_METNI]"] = muhendislik_metni
        degisim_sozlugu["[VS30_CUMLESI]"] = vs30_cumlesi
        degisim_sozlugu["[YASS_SONUC_METNI]"] = self.rapor_yass_sonuc_metni_olustur(ac_yn_kayitlari)
        son_qk = getattr(self, "son_qk", "")
        degisim_sozlugu["[QK]"] = "" if son_qk in (None, "", "-") else str(son_qk)
        degisim_sozlugu["[QT]"] = nihai_qt
        degisim_sozlugu["[KS]"] = nihai_ks
        degisim_sozlugu["[DF]"] = df_metni
        degisim_sozlugu["[LITOLOJI]"] = final_lito_text
        degisim_sozlugu["[CALISMA]"] = final_calisma_text
        degisim_sozlugu["[ARAZI_CALISMA_ACIKLAMA]"] = arazi_calisma_aciklama
        degisim_sozlugu["[ARAZI_CALISMA_OZETI]"] = arazi_calisma_ozeti
        degisim_sozlugu["[KESIT_ACIKLAMA]"] = self.rapor_kesit_aciklama_olustur(ac_yn_kayitlari)

        return degisim_sozlugu, fm_metin, nihai_qt, nihai_ks, df_metni

    def rapor_paragraf_bul(self, doc, tag):
        for p in self.docx_paragraflarini_dolas(doc):
            if tag in p.text:
                return p
        return None

    def rapor_metin_normalize(self, metin):
        metin = str(metin or "").strip().lower()
        ceviri = str.maketrans({
            "ç": "c", "ğ": "g", "ı": "i", "i": "i", "ö": "o", "ş": "s", "ü": "u",
            "Ç": "c", "Ğ": "g", "İ": "i", "I": "i", "Ö": "o", "Ş": "s", "Ü": "u",
        })
        metin = metin.translate(ceviri)
        metin = unicodedata.normalize("NFKD", metin)
        metin = "".join(ch for ch in metin if not unicodedata.combining(ch))
        metin = re.sub(r"[^a-z0-9]+", "_", metin).strip("_")
        return metin

    def rapor_jeoloji_sablon_adaylari(self):
        ilce = self.veri_alanlari.get("ILCE").get().strip() if self.veri_alanlari.get("ILCE") else ""
        koy = self.veri_alanlari.get("KOY").get().strip() if self.veri_alanlari.get("KOY") else ""
        ilce_n = self.rapor_metin_normalize(ilce)
        koy_n = self.rapor_metin_normalize(koy)
        adaylar = []
        if ilce_n and koy_n:
            adaylar.extend([
                f"{ilce_n}_{koy_n}.docx",
                f"{ilce_n}-{koy_n}.docx",
                f"{ilce_n} {koy_n}.docx",
            ])
        if koy_n:
            adaylar.append(f"{koy_n}.docx")
        if ilce and koy:
            adaylar.extend([
                f"{ilce}_{koy}.docx",
                f"{ilce} {koy}.docx",
            ])
        if koy:
            adaylar.append(f"{koy}.docx")
        return adaylar

    def rapor_jeoloji_sablon_yolu_bul(self):
        if getattr(self, "jeoloji_kutuphanesi_bolumu_aktif", False):
            resolver = getattr(self.app, "jeoloji_kutuphanesi_islemleri", None)
            if callable(resolver):
                try:
                    path = resolver().jeoloji_kutuphanesi_bolum_yolunu_coz()
                except (OSError, ValueError, TypeError) as exc:
                    self.hata_kaydet("Kütüphane JEOLOJİ Word paketi çözülemedi", exc)
                    path = ""
            else:
                path = getattr(self, "jeoloji_kutuphanesi_bolum_yolu", "")
            return path if path and os.path.isfile(path) else ""

        secili = getattr(self, "jeoloji_sablon_yolu", "")
        if secili and os.path.exists(secili):
            return secili

        klasorler = []
        if hasattr(self.app, "sablon_alt_klasoru"):
            klasorler.append(self.sablon_alt_klasoru("jeoloji"))
        if hasattr(self.app, "sablon_kok_klasoru"):
            klasorler.append(os.path.join(self.sablon_kok_klasoru(), "jeoloji"))

        aday_adlari = self.rapor_jeoloji_sablon_adaylari()
        for klasor in klasorler:
            if not os.path.isdir(klasor):
                continue
            for ad in aday_adlari:
                yol = os.path.join(klasor, ad)
                if os.path.exists(yol):
                    return yol

        hedefler = [self.rapor_metin_normalize(os.path.splitext(ad)[0]) for ad in aday_adlari]
        hedefler = [h for h in hedefler if h]
        for klasor in klasorler:
            if not os.path.isdir(klasor):
                continue
            try:
                for dosya_adi in os.listdir(klasor):
                    if dosya_adi.startswith("~$") or not dosya_adi.lower().endswith(".docx"):
                        continue
                    kok = self.rapor_metin_normalize(os.path.splitext(dosya_adi)[0])
                    if any(hedef == kok or hedef in kok or kok in hedef for hedef in hedefler):
                        return os.path.join(klasor, dosya_adi)
            except OSError:
                continue

        return ""

    def rapor_paragrafi_sil(self, paragraph):
        p = paragraph._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)

    def rapor_docx_govde_elemanlari(self, doc):
        for child in doc.element.body.iterchildren():
            if child.tag == qn("w:sectPr"):
                continue
            yield child

    def rapor_xml_paragraf_metni(self, eleman):
        return "".join(node.text or "" for node in eleman.iter(qn("w:t"))).strip()

    def rapor_ana_jeoloji_basligi_mi(self, metin):
        anahtar = self.rapor_metin_normalize(metin)
        return bool(re.fullmatch(r"(?:\d+_)*jeoloji", anahtar))

    def rapor_jeoloji_govde_elemanlarini_hazirla(self, kaynak_doc):
        elemanlar = list(self.rapor_docx_govde_elemanlari(kaynak_doc))
        while elemanlar and elemanlar[0].tag == qn("w:p") and not self.rapor_xml_paragraf_metni(elemanlar[0]):
            elemanlar.pop(0)
        if (
            elemanlar
            and elemanlar[0].tag == qn("w:p")
            and self.rapor_ana_jeoloji_basligi_mi(self.rapor_xml_paragraf_metni(elemanlar[0]))
        ):
            elemanlar.pop(0)
        while elemanlar and elemanlar[0].tag == qn("w:p") and not self.rapor_xml_paragraf_metni(elemanlar[0]):
            elemanlar.pop(0)
        return elemanlar

    def rapor_paragraf_xml_stili_ayarla(self, paragraf, stil_id, keep_next=False):
        p_pr = paragraf.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            paragraf.insert(0, p_pr)
        p_style = p_pr.find(qn("w:pStyle"))
        if p_style is None:
            p_style = OxmlElement("w:pStyle")
            p_pr.insert(0, p_style)
        p_style.set(qn("w:val"), stil_id)
        if keep_next and p_pr.find(qn("w:keepNext")) is None:
            p_pr.append(OxmlElement("w:keepNext"))
        if keep_next and p_pr.find(qn("w:keepLines")) is None:
            p_pr.append(OxmlElement("w:keepLines"))

    def rapor_sabit_gorselleri_satir_icine_al(self, eleman):
        """Sayfaya sabitlenmiş görselleri rapor akışında güvenli satır içi nesnelere çevir."""
        for anchor in list(eleman.iter(qn("wp:anchor"))):
            graphic = anchor.find(qn("a:graphic"))
            extent = anchor.find(qn("wp:extent"))
            doc_pr = anchor.find(qn("wp:docPr"))
            if graphic is None or extent is None or doc_pr is None:
                continue
            inline = OxmlElement("wp:inline")
            for ad in ("distT", "distB", "distL", "distR"):
                inline.set(ad, "0")
            for etiket in ("wp:extent", "wp:effectExtent", "wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
                alt = anchor.find(qn(etiket))
                if alt is not None:
                    inline.append(deepcopy(alt))
            parent = anchor.getparent()
            if parent is not None:
                parent.replace(anchor, inline)

    def rapor_jeoloji_eleman_bicimini_duzenle(self, eleman, ilk_eleman=False):
        for sect_pr in list(eleman.iter(qn("w:sectPr"))):
            parent = sect_pr.getparent()
            if parent is not None:
                parent.remove(sect_pr)
        self.rapor_sabit_gorselleri_satir_icine_al(eleman)

        if ilk_eleman and eleman.tag == qn("w:p"):
            for page_break in list(eleman.iter(qn("w:br"))):
                if page_break.get(qn("w:type")) == "page":
                    parent = page_break.getparent()
                    if parent is not None:
                        parent.remove(page_break)
            p_pr = eleman.find(qn("w:pPr"))
            if p_pr is not None:
                page_break_before = p_pr.find(qn("w:pageBreakBefore"))
                if page_break_before is not None:
                    p_pr.remove(page_break_before)

        if eleman.tag == qn("w:p"):
            metin = self.rapor_xml_paragraf_metni(eleman)
            numara = re.match(r"^\s*(\d+(?:\.\d+)+)[.)]?\s+", metin)
            if numara:
                seviye = min(3, numara.group(1).count(".") + 1)
                self.rapor_paragraf_xml_stili_ayarla(eleman, f"Heading{seviye}", keep_next=True)
            elif re.match(r"^\s*(Şekil|Tablo|Çizelge)\b", metin, re.IGNORECASE):
                self.rapor_paragraf_xml_stili_ayarla(eleman, "Caption", keep_next=True)
        elif eleman.tag == qn("w:tbl"):
            ilk_satir = eleman.find(qn("w:tr"))
            if ilk_satir is not None:
                tr_pr = ilk_satir.find(qn("w:trPr"))
                if tr_pr is None:
                    tr_pr = OxmlElement("w:trPr")
                    ilk_satir.insert(0, tr_pr)
                if tr_pr.find(qn("w:tblHeader")) is None:
                    tr_pr.append(OxmlElement("w:tblHeader"))

    def rapor_jeoloji_kopya_kimliklerini_duzenle(self, hedef_doc, elemanlar):
        kullanilan_adlar = {
            node.get(qn("w:name"))
            for node in hedef_doc.element.body.iter(qn("w:bookmarkStart"))
            if node.get(qn("w:name"))
        }
        kullanilan_idler = {
            int(node.get(qn("w:id")))
            for node in hedef_doc.element.body.iter(qn("w:bookmarkStart"))
            if str(node.get(qn("w:id"), "")).isdigit()
        }
        sonraki_id = max(kullanilan_idler, default=0) + 1
        id_esleme = {}
        ad_esleme = {}
        for eleman in elemanlar:
            for baslangic in eleman.iter(qn("w:bookmarkStart")):
                eski_id = baslangic.get(qn("w:id"))
                if eski_id not in id_esleme:
                    id_esleme[eski_id] = str(sonraki_id)
                    sonraki_id += 1
                eski_ad = baslangic.get(qn("w:name")) or "Jeoloji"
                if eski_ad not in ad_esleme:
                    temel = "K1J_" + self.rapor_metin_normalize(eski_ad)[:28]
                    aday = temel
                    sira = 2
                    while aday in kullanilan_adlar:
                        aday = f"{temel[:34]}_{sira}"
                        sira += 1
                    kullanilan_adlar.add(aday)
                    ad_esleme[eski_ad] = aday

        kullanilan_docpr = {
            int(node.get("id"))
            for node in hedef_doc.element.body.iter(qn("wp:docPr"))
            if str(node.get("id", "")).isdigit()
        }
        sonraki_docpr = max(kullanilan_docpr, default=0) + 1
        for eleman in elemanlar:
            for baslangic in eleman.iter(qn("w:bookmarkStart")):
                eski_id = baslangic.get(qn("w:id"))
                baslangic.set(qn("w:id"), id_esleme[eski_id])
                eski_ad = baslangic.get(qn("w:name")) or "Jeoloji"
                baslangic.set(qn("w:name"), ad_esleme[eski_ad])
            for bitis in eleman.iter(qn("w:bookmarkEnd")):
                eski_id = bitis.get(qn("w:id"))
                if eski_id in id_esleme:
                    bitis.set(qn("w:id"), id_esleme[eski_id])
            for alan in eleman.iter(qn("w:instrText")):
                kod = alan.text or ""
                for eski_ad, yeni_ad in ad_esleme.items():
                    kod = re.sub(
                        rf"(?i)(\b(?:REF|PAGEREF)\s+){re.escape(eski_ad)}(?=\s|\\|$)",
                        rf"\g<1>{yeni_ad}",
                        kod,
                    )
                alan.text = kod
            for doc_pr in eleman.iter(qn("wp:docPr")):
                doc_pr.set("id", str(sonraki_docpr))
                sonraki_docpr += 1

    def rapor_benzersiz_parca_adi(self, kaynak_parca, kullanilan_adlar):
        """Kaynak OPC parçası için hedef pakette çakışmayan bir ad üret."""
        kaynak_ad = str(kaynak_parca.partname)
        klasor, dosya = os.path.split(kaynak_ad)
        kok, uzanti = os.path.splitext(dosya)
        kok = re.sub(r"\d+$", "", kok) or "part"
        sira = 1
        while True:
            aday = f"{klasor}/k1_{kok}{sira}{uzanti}"
            if aday not in kullanilan_adlar:
                kullanilan_adlar.add(aday)
                return PackURI(aday)
            sira += 1

    def rapor_xml_iliski_kimliklerini_degistir(self, xml_koku, iliski_esleme):
        iliski_uzayi = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        for eleman in xml_koku.iter():
            for nitelik, deger in list(eleman.attrib.items()):
                if nitelik.startswith(f"{{{iliski_uzayi}}}") and deger in iliski_esleme:
                    eleman.set(nitelik, iliski_esleme[deger])

    def rapor_iliskili_parcayi_kopyala(
        self,
        kaynak_parca,
        hedef_paket,
        parca_esleme,
        kullanilan_adlar,
    ):
        """Bir OPC parçasını alt ilişkileriyle hedef pakete güvenli biçimde kopyala."""
        if kaynak_parca in parca_esleme:
            return parca_esleme[kaynak_parca]

        yeni_ad = self.rapor_benzersiz_parca_adi(kaynak_parca, kullanilan_adlar)
        yeni_parca = Part(yeni_ad, kaynak_parca.content_type, kaynak_parca.blob, hedef_paket)
        parca_esleme[kaynak_parca] = yeni_parca

        iliski_esleme = {}
        for eski_rid, iliski in kaynak_parca.rels.items():
            if iliski.is_external:
                yeni_rid = yeni_parca.relate_to(iliski.target_ref, iliski.reltype, True)
            else:
                yeni_hedef = self.rapor_iliskili_parcayi_kopyala(
                    iliski.target_part,
                    hedef_paket,
                    parca_esleme,
                    kullanilan_adlar,
                )
                yeni_rid = yeni_parca.relate_to(yeni_hedef, iliski.reltype)
            iliski_esleme[eski_rid] = yeni_rid

        if iliski_esleme:
            try:
                xml_koku = parse_xml(kaynak_parca.blob)
            except Exception:
                xml_koku = None
            if xml_koku is not None:
                self.rapor_xml_iliski_kimliklerini_degistir(xml_koku, iliski_esleme)
                from lxml import etree

                yeni_parca._blob = etree.tostring(
                    xml_koku,
                    encoding="UTF-8",
                    xml_declaration=True,
                    standalone=True,
                )
        return yeni_parca

    def rapor_eleman_iliskilerini_kopyala(
        self,
        eleman,
        kaynak_parca,
        hedef_parca,
        parca_esleme,
        kullanilan_adlar,
    ):
        iliski_uzayi = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        kullanilan_ridler = {
            deger
            for dugum in eleman.iter()
            for nitelik, deger in dugum.attrib.items()
            if nitelik.startswith(f"{{{iliski_uzayi}}}")
        }
        iliski_esleme = {}
        for eski_rid in kullanilan_ridler:
            if eski_rid not in kaynak_parca.rels:
                raise ValueError(f"Kaynak Word içeriğinde çözülemeyen ilişki var: {eski_rid}")
            iliski = kaynak_parca.rels[eski_rid]
            if iliski.is_external:
                yeni_rid = hedef_parca.relate_to(iliski.target_ref, iliski.reltype, True)
            else:
                yeni_hedef = self.rapor_iliskili_parcayi_kopyala(
                    iliski.target_part,
                    hedef_parca.package,
                    parca_esleme,
                    kullanilan_adlar,
                )
                yeni_rid = hedef_parca.relate_to(yeni_hedef, iliski.reltype)
            iliski_esleme[eski_rid] = yeni_rid
        self.rapor_xml_iliski_kimliklerini_degistir(eleman, iliski_esleme)

    def rapor_stil_eslemesini_hazirla(self, hedef_doc, kaynak_doc):
        hedef_kok = hedef_doc.styles.element
        kaynak_stiller = list(kaynak_doc.styles.element.findall(qn("w:style")))
        hedefler = {
            stil.get(qn("w:styleId")): stil
            for stil in hedef_kok.findall(qn("w:style"))
            if stil.get(qn("w:styleId"))
        }
        kullanilan = set(hedefler)
        esleme = {}
        for kaynak_stil in kaynak_stiller:
            stil_id = kaynak_stil.get(qn("w:styleId"))
            if not stil_id:
                continue
            hedef_stil = hedefler.get(stil_id)
            if hedef_stil is None:
                esleme[stil_id] = stil_id
                kullanilan.add(stil_id)
                continue
            from lxml import etree

            if etree.tostring(kaynak_stil) == etree.tostring(hedef_stil):
                esleme[stil_id] = stil_id
                continue
            temel = f"K1_{re.sub(r'[^A-Za-z0-9_]+', '_', stil_id)}"
            aday = temel
            sira = 2
            while aday in kullanilan:
                aday = f"{temel}_{sira}"
                sira += 1
            kullanilan.add(aday)
            esleme[stil_id] = aday
        return esleme

    def rapor_stilleri_kopyala(self, hedef_doc, kaynak_doc, stil_esleme):
        hedef_kok = hedef_doc.styles.element
        hedef_idler = {
            stil.get(qn("w:styleId"))
            for stil in hedef_kok.findall(qn("w:style"))
        }
        kopyalananlar = []
        for kaynak_stil in kaynak_doc.styles.element.findall(qn("w:style")):
            eski_id = kaynak_stil.get(qn("w:styleId"))
            yeni_id = stil_esleme.get(eski_id, eski_id)
            if not yeni_id or (yeni_id == eski_id and yeni_id in hedef_idler):
                continue
            yeni_stil = deepcopy(kaynak_stil)
            yeni_stil.set(qn("w:styleId"), yeni_id)
            for bag_etiketi in ("w:basedOn", "w:next", "w:link"):
                for bag in yeni_stil.iter(qn(bag_etiketi)):
                    eski = bag.get(qn("w:val"))
                    if eski in stil_esleme:
                        bag.set(qn("w:val"), stil_esleme[eski])
            hedef_kok.append(yeni_stil)
            hedef_idler.add(yeni_id)
            kopyalananlar.append(yeni_stil)
        return kopyalananlar

    def rapor_numaralandirma_parcasi(self, doc):
        try:
            return doc.part.numbering_part
        except (KeyError, NotImplementedError):
            element = parse_xml(f"<w:numbering {nsdecls('w')}/>")
            parca = NumberingPart(
                PackURI("/word/numbering.xml"),
                CT.WML_NUMBERING,
                element,
                doc.part.package,
            )
            doc.part.relate_to(parca, RT.NUMBERING)
            return parca

    def rapor_numaralandirmayi_kopyala(self, hedef_doc, kaynak_doc, stil_esleme):
        try:
            kaynak_parca = kaynak_doc.part.numbering_part
        except (KeyError, NotImplementedError):
            return {}
        hedef_parca = self.rapor_numaralandirma_parcasi(hedef_doc)
        kaynak_kok = kaynak_parca.element
        hedef_kok = hedef_parca.element

        hedef_abstract_idler = [
            int(x.get(qn("w:abstractNumId")))
            for x in hedef_kok.findall(qn("w:abstractNum"))
            if str(x.get(qn("w:abstractNumId"), "")).isdigit()
        ]
        hedef_num_idler = [
            int(x.get(qn("w:numId")))
            for x in hedef_kok.findall(qn("w:num"))
            if str(x.get(qn("w:numId"), "")).isdigit()
        ]
        sonraki_abstract = max(hedef_abstract_idler, default=-1) + 1
        sonraki_num = max(hedef_num_idler, default=0) + 1

        abstract_esleme = {}
        for kaynak in kaynak_kok.findall(qn("w:abstractNum")):
            eski = kaynak.get(qn("w:abstractNumId"))
            abstract_esleme[eski] = str(sonraki_abstract)
            sonraki_abstract += 1
        num_esleme = {}
        for kaynak in kaynak_kok.findall(qn("w:num")):
            eski = kaynak.get(qn("w:numId"))
            num_esleme[eski] = str(sonraki_num)
            sonraki_num += 1

        parca_esleme = {}
        kullanilan_adlar = {str(p.partname) for p in hedef_doc.part.package.iter_parts()}
        for kaynak in kaynak_kok.findall(qn("w:abstractNum")):
            yeni = deepcopy(kaynak)
            yeni.set(qn("w:abstractNumId"), abstract_esleme[kaynak.get(qn("w:abstractNumId"))])
            for stil_etiketi in ("w:pStyle", "w:rStyle", "w:numStyleLink", "w:styleLink"):
                for stil in yeni.iter(qn(stil_etiketi)):
                    eski = stil.get(qn("w:val"))
                    if eski in stil_esleme:
                        stil.set(qn("w:val"), stil_esleme[eski])
            self.rapor_eleman_iliskilerini_kopyala(
                yeni,
                kaynak_parca,
                hedef_parca,
                parca_esleme,
                kullanilan_adlar,
            )
            hedef_kok.append(yeni)

        for kaynak in kaynak_kok.findall(qn("w:num")):
            yeni = deepcopy(kaynak)
            yeni.set(qn("w:numId"), num_esleme[kaynak.get(qn("w:numId"))])
            abstract = yeni.find(qn("w:abstractNumId"))
            if abstract is not None:
                eski = abstract.get(qn("w:val"))
                if eski in abstract_esleme:
                    abstract.set(qn("w:val"), abstract_esleme[eski])
            hedef_kok.append(yeni)
        return num_esleme

    def rapor_stil_numara_referanslarini_degistir(self, eleman, stil_esleme, num_esleme):
        for etiket in ("w:pStyle", "w:rStyle", "w:tblStyle"):
            for dugum in eleman.iter(qn(etiket)):
                eski = dugum.get(qn("w:val"))
                if eski in stil_esleme:
                    dugum.set(qn("w:val"), stil_esleme[eski])
        for dugum in eleman.iter(qn("w:numId")):
            eski = dugum.get(qn("w:val"))
            if eski in num_esleme:
                dugum.set(qn("w:val"), num_esleme[eski])

    def rapor_docx_icerigi_ekle(self, hedef_p, kaynak_doc):
        anchor = hedef_p._p
        hedef_doc = hedef_p.part.document
        stil_esleme = self.rapor_stil_eslemesini_hazirla(hedef_doc, kaynak_doc)
        kopyalanan_stiller = self.rapor_stilleri_kopyala(hedef_doc, kaynak_doc, stil_esleme)
        num_esleme = self.rapor_numaralandirmayi_kopyala(hedef_doc, kaynak_doc, stil_esleme)
        for stil in kopyalanan_stiller:
            self.rapor_stil_numara_referanslarini_degistir(stil, stil_esleme, num_esleme)

        parca_esleme = {}
        kullanilan_adlar = {str(p.partname) for p in hedef_doc.part.package.iter_parts()}
        kaynak_elemanlar = self.rapor_jeoloji_govde_elemanlarini_hazirla(kaynak_doc)
        yeni_elemanlar = [deepcopy(child) for child in kaynak_elemanlar]
        self.rapor_jeoloji_kopya_kimliklerini_duzenle(hedef_doc, yeni_elemanlar)
        for index, yeni in enumerate(yeni_elemanlar):
            self.rapor_jeoloji_eleman_bicimini_duzenle(yeni, ilk_eleman=index == 0)
            self.rapor_stil_numara_referanslarini_degistir(yeni, stil_esleme, num_esleme)
            self.rapor_eleman_iliskilerini_kopyala(
                yeni,
                kaynak_doc.part,
                hedef_doc.part,
                parca_esleme,
                kullanilan_adlar,
            )
        for yeni in reversed(yeni_elemanlar):
            anchor.addnext(yeni)
        return len(yeni_elemanlar)

    def rapor_jeoloji_bolumu_ekle(self, doc):
        tag = "[JEOLOJI_BOLUMU]"
        hedef_p = self.rapor_paragraf_bul(doc, tag)
        if not hedef_p:
            return

        sablon_yolu = self.rapor_jeoloji_sablon_yolu_bul()
        if not sablon_yolu:
            if getattr(self, "jeoloji_kutuphanesi_bolumu_aktif", False):
                record_id = getattr(self, "jeoloji_kutuphanesi_kayit_id", None)
                ek = f" #{record_id}" if record_id else ""
                mesaj = (
                    f"Seçili kütüphane kaydının{ek} JEOLOJİ Word paketi bulunamadı. "
                    "Kütüphaneden kaydı yeniden uygulayın veya başka bir jeoloji şablonu seçin."
                )
            else:
                mesaj = "Jeoloji bölümü şablonu bulunamadı."
            self.hata_kaydet(mesaj)
            raise ValueError(mesaj)

        self.rapor_metin_degistir(doc, tag, JEOLOJI_WORD_EKLEME_ISARETI)
        return sablon_yolu

    def rapor_tabloyu_ortala(self, tbl):
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in tbl.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)

    def rapor_tablo_label_kolonlari(self, label_columns):
        if not label_columns:
            return set()
        if isinstance(label_columns, int):
            return set(range(max(0, label_columns)))
        return set(label_columns)

    def rapor_hucre_kenarlik_ayarla(self, cell, renk="1F1F1F", boyut="6"):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)

        for kenar in ("top", "left", "bottom", "right"):
            edge = borders.find(qn(f"w:{kenar}"))
            if edge is None:
                edge = OxmlElement(f"w:{kenar}")
                borders.append(edge)
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), boyut)
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), renk)

    def rapor_hucre_yazi_stili_uygula(self, cell, bold=False, renk=None, punto=8):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            if not paragraph.runs:
                paragraph.add_run("")
            for run in paragraph.runs:
                run.font.bold = bool(bold)
                run.font.size = Pt(punto)
                if renk:
                    run.font.color.rgb = RGBColor.from_string(renk)

    def rapor_tablo_stili_uygula(self, tablo, header_rows=1, label_columns=0, font_size=None, banded_rows=True):
        tablo.style = "Table Grid"
        tablo.alignment = WD_TABLE_ALIGNMENT.CENTER
        tablo.autofit = True
        label_cols = self.rapor_tablo_label_kolonlari(label_columns)
        if not tablo.rows:
            return

        kolon_sayisi = len(tablo.rows[0].cells)
        punto = font_size if font_size is not None else (8 if kolon_sayisi >= 10 else 9)
        baslik_rengi = "2F5597"
        alternatif_rengi = "D9EAF7"

        for r_idx, row in enumerate(tablo.rows):
            govde_satir_no = max(0, r_idx - header_rows)
            for c_idx, cell in enumerate(row.cells):
                baslik_hucresi = r_idx < header_rows or c_idx in label_cols
                if baslik_hucresi:
                    self.rapor_hucre_arkaplan_rengi(cell, baslik_rengi)
                    self.rapor_hucre_yazi_stili_uygula(cell, bold=True, renk="FFFFFF", punto=punto)
                else:
                    dolgu = alternatif_rengi if banded_rows and govde_satir_no % 2 == 1 else "FFFFFF"
                    self.rapor_hucre_arkaplan_rengi(cell, dolgu)
                    self.rapor_hucre_yazi_stili_uygula(cell, bold=False, renk="000000", punto=punto)
                self.rapor_hucre_kenarlik_ayarla(cell)

    def _rapor_etiket_deseni_hazirla(self, degisim_sozlugu):
        degisimler = {
            str(etiket): str(deger)
            for etiket, deger in dict(degisim_sozlugu or {}).items()
            if str(etiket)
        }
        if not degisimler:
            return None, {}
        etiketler = sorted(degisimler, key=len, reverse=True)
        return re.compile("|".join(re.escape(etiket) for etiket in etiketler)), degisimler

    def _rapor_xml_metin_ata(self, dugum, metin):
        metin = str(metin)
        dugum.text = metin
        bosluk_ozelligi = "{http://www.w3.org/XML/1998/namespace}space"
        if metin[:1].isspace() or metin[-1:].isspace():
            dugum.set(bosluk_ozelligi, "preserve")
        else:
            dugum.attrib.pop(bosluk_ozelligi, None)

    def _rapor_xml_metinleri_toplu_degistir(self, doc, desen, degisimler):
        degisim_sayisi = 0
        for part in doc.part.package.parts:
            partname = str(getattr(part, "partname", ""))
            if not partname.startswith("/word/") or not hasattr(part, "element"):
                continue
            try:
                dugumler = list(part.element.xpath(".//w:t"))
            except Exception:
                continue
            if not dugumler:
                continue

            parcalar = [dugum.text or "" for dugum in dugumler]
            birlesik = "".join(parcalar)
            eslesmeler = list(desen.finditer(birlesik))
            if not eslesmeler:
                continue

            bitisler = []
            toplam = 0
            for parca in parcalar:
                toplam += len(parca)
                bitisler.append(toplam)

            # Sağdan sola çalışmak, bir değişimin kendisinden önceki özgün
            # karakter ofsetlerini kaydırmasını engeller.
            for eslesme in reversed(eslesmeler):
                baslangic, bitis = eslesme.span()
                ilk = bisect_right(bitisler, baslangic)
                son = bisect_right(bitisler, bitis - 1)
                if ilk >= len(dugumler) or son >= len(dugumler):
                    continue
                ilk_baslangici = bitisler[ilk] - len(parcalar[ilk])
                son_baslangici = bitisler[son] - len(parcalar[son])
                ilk_ofset = baslangic - ilk_baslangici
                son_ofset = bitis - son_baslangici
                val_str = degisimler[eslesme.group(0)]
                if ilk == son:
                    mevcut = dugumler[ilk].text or ""
                    self._rapor_xml_metin_ata(
                        dugumler[ilk],
                        mevcut[:ilk_ofset] + val_str + mevcut[son_ofset:],
                    )
                else:
                    ilk_metin = dugumler[ilk].text or ""
                    son_metin = dugumler[son].text or ""
                    self._rapor_xml_metin_ata(
                        dugumler[ilk],
                        ilk_metin[:ilk_ofset] + val_str,
                    )
                    for index in range(ilk + 1, son):
                        self._rapor_xml_metin_ata(dugumler[index], "")
                    self._rapor_xml_metin_ata(dugumler[son], son_metin[son_ofset:])
                degisim_sayisi += 1
        return degisim_sayisi

    def rapor_metinleri_toplu_degistir(self, doc, degisim_sozlugu):
        """Word etiketlerini belgeyi etiket başına yeniden taramadan değiştirir.

        Normal paragraflar önce tek geçişte ele alınır; böylece tek bir run
        içindeki etiketin mevcut karakter biçimi korunur. Run'lara bölünmüş
        etiketler ile metin kutusu gibi python-docx'in paragraf API'sinde
        görünmeyen parçalar, ardından Word XML'i üzerinde tek geçişte çözülür.
        """
        desen, degisimler = self._rapor_etiket_deseni_hazirla(degisim_sozlugu)
        if desen is None:
            return 0

        degisim_sayisi = 0
        gorulen_paragraflar = set()
        for paragraf in self.docx_paragraflarini_dolas(doc):
            paragraf_kimligi = id(paragraf._p)
            if paragraf_kimligi in gorulen_paragraflar:
                continue
            gorulen_paragraflar.add(paragraf_kimligi)

            runlar = list(paragraf.runs)
            if not runlar:
                continue
            tam_metin = "".join(run.text for run in runlar)
            eslesmeler = list(desen.finditer(tam_metin))
            if not eslesmeler:
                continue

            parcalar = [run.text for run in runlar]
            bitisler = []
            toplam = 0
            for parca in parcalar:
                toplam += len(parca)
                bitisler.append(toplam)

            for eslesme in reversed(eslesmeler):
                baslangic, bitis = eslesme.span()
                ilk = bisect_right(bitisler, baslangic)
                son = bisect_right(bitisler, bitis - 1)
                if ilk >= len(runlar) or son >= len(runlar):
                    continue
                ilk_baslangici = bitisler[ilk] - len(parcalar[ilk])
                son_baslangici = bitisler[son] - len(parcalar[son])
                ilk_ofset = baslangic - ilk_baslangici
                son_ofset = bitis - son_baslangici
                val_str = degisimler[eslesme.group(0)]
                if ilk == son:
                    mevcut = runlar[ilk].text
                    runlar[ilk].text = mevcut[:ilk_ofset] + val_str + mevcut[son_ofset:]
                else:
                    ilk_metin = runlar[ilk].text
                    son_metin = runlar[son].text
                    runlar[ilk].text = ilk_metin[:ilk_ofset] + val_str
                    for index in range(ilk + 1, son):
                        runlar[index].text = ""
                    runlar[son].text = son_metin[son_ofset:]
                degisim_sayisi += 1

        # Hiperbağ, metin kutusu ve diğer ham Word parçalarında kalmış ya da
        # farklı XML metin düğümlerine bölünmüş etiketleri tamamlar.
        degisim_sayisi += self._rapor_xml_metinleri_toplu_degistir(
            doc,
            desen,
            degisimler,
        )
        return degisim_sayisi

    def rapor_metin_degistir(self, doc, tag, value):
        return self.rapor_metinleri_toplu_degistir(doc, {tag: value})

    def rapor_xml_metin_degistir(self, doc, tag, value):
        desen, degisimler = self._rapor_etiket_deseni_hazirla({tag: value})
        if desen is None:
            return 0
        return self._rapor_xml_metinleri_toplu_degistir(doc, desen, degisimler)

    def rapor_deger_formatla(self, v):
        if pd.isna(v):
            return "-"
        v_str = str(v).strip(" \t\n\r\xa0")
        if v_str == "" or v_str == "-":
            return "-"
        try:
            f = float(v_str.replace(",", "."))
            if f == int(f):
                return str(int(f))
            return "{:.2f}".format(f).replace(".", ",")
        except:
            return v_str.replace(".", ",")

    def rapor_formatli_metin_ekle(self, p_obj, metin_satiri):
        i = 0
        uzunluk = len(metin_satiri)
        while i < uzunluk:
            if metin_satiri[i] == '^':
                i += 1
                if i < uzunluk and metin_satiri[i] == '(':
                    basla = i
                    while i < uzunluk and metin_satiri[i] != ')':
                        i += 1
                    if i < uzunluk:
                        p_obj.add_run(metin_satiri[basla:i + 1]).font.superscript = True
                        i += 1
                    else:
                        p_obj.add_run(metin_satiri[basla:i]).font.superscript = True
                else:
                    basla = i
                    while i < uzunluk and (metin_satiri[i].isalnum() or metin_satiri[i] in ["'"]):
                        i += 1
                    p_obj.add_run(metin_satiri[basla:i]).font.superscript = True
            elif metin_satiri[i] == '_':
                i += 1
                if i < uzunluk and metin_satiri[i] == '(':
                    basla = i
                    while i < uzunluk and metin_satiri[i] != ')':
                        i += 1
                    if i < uzunluk:
                        p_obj.add_run(metin_satiri[basla:i + 1]).font.subscript = True
                        i += 1
                    else:
                        p_obj.add_run(metin_satiri[basla:i]).font.subscript = True
                else:
                    basla = i
                    while i < uzunluk and (metin_satiri[i].isalnum() or metin_satiri[i] in ["'"]):
                        i += 1
                    p_obj.add_run(metin_satiri[basla:i]).font.subscript = True
            else:
                basla = i
                while i < uzunluk and metin_satiri[i] not in ['^', '_']:
                    i += 1
                p_obj.add_run(metin_satiri[basla:i])

    def rapor_run_stilini_kopyala(self, kaynak, hedef):
        hedef.bold = kaynak[0]
        hedef.italic = kaynak[1]
        hedef.underline = kaynak[2]
        hedef.font.size = kaynak[3]
        hedef.font.name = kaynak[4]
        if kaynak[5]:
            hedef.font.color.rgb = kaynak[5]

    def rapor_cok_satirli_etiket_degistir(self, doc, tag, value):
        val_str = str(value or "")
        for p in self.docx_paragraflarini_dolas(doc):
            if tag not in p.text:
                continue

            if not any(tag in run.text for run in p.runs) and p.runs:
                full_text = "".join(run.text for run in p.runs)
                p.runs[0].text = full_text
                for run in p.runs[1:]:
                    run.text = ""

            original_runs = []
            for run in p.runs:
                try:
                    color = run.font.color.rgb if run.font.color else None
                except:
                    color = None
                original_runs.append((run.text, (run.bold, run.italic, run.underline, run.font.size, run.font.name, color)))

            p.clear()
            for text, stil in original_runs:
                if not text:
                    continue
                if tag not in text:
                    new_run = p.add_run(text)
                    self.rapor_run_stilini_kopyala(stil, new_run)
                    continue

                parts = text.split(tag)
                for idx, part in enumerate(parts):
                    if part:
                        new_run = p.add_run(part)
                        self.rapor_run_stilini_kopyala(stil, new_run)
                    if idx < len(parts) - 1 and val_str:
                        satirlar = val_str.split('\n')
                        for s_idx, satir in enumerate(satirlar):
                            onceki_run_sayisi = len(p.runs)
                            if satir.strip():
                                self.rapor_formatli_metin_ekle(p, satir)
                                for new_run in p.runs[onceki_run_sayisi:]:
                                    self.rapor_run_stilini_kopyala(stil, new_run)
                            if s_idx < len(satirlar) - 1:
                                br = p.add_run('\n')
                                self.rapor_run_stilini_kopyala(stil, br)

    def rapor_xml_sonrasina_ekle(self, anchor, yeni_eleman):
        anchor.addnext(yeni_eleman)
        return yeni_eleman

    def rapor_kenarlikli_gorsel_yolu(self, img_path):
        if not img_path or not os.path.exists(img_path):
            return img_path
        try:
            with Image.open(img_path) as img:
                img.load()
                if img.mode in ("RGBA", "LA"):
                    arka_plan = Image.new("RGB", img.size, "white")
                    arka_plan.paste(img, mask=img.getchannel("A"))
                    img = arka_plan
                else:
                    img = img.convert("RGB")

                kenarlik_px = max(4, min(24, round(img.width * 0.0025)))
                kenarlikli = ImageOps.expand(img, border=kenarlik_px, fill="#2b2b2b")
                with tempfile.NamedTemporaryFile(
                    prefix="k1_kenarlikli_",
                    suffix=".jpg",
                    delete=False,
                ) as f:
                    temp_yol = f.name
                kenarlikli.save(temp_yol, quality=95)
                return temp_yol
        except Exception as e:
            if "temp_yol" in locals():
                try:
                    os.remove(temp_yol)
                except OSError:
                    pass
            self.hata_kaydet(f"Görsel kenarlığı oluşturulamadı: {img_path}", e)
            return img_path

    def rapor_gorsel_ekle(self, paragraph, img_path, width):
        gorsel_yolu = self.rapor_kenarlikli_gorsel_yolu(img_path)
        try:
            return paragraph.add_run().add_picture(gorsel_yolu, width=width)
        finally:
            if gorsel_yolu and gorsel_yolu != img_path:
                try:
                    os.remove(gorsel_yolu)
                except OSError:
                    pass

    def rapor_gorsel_kullanilabilir_genisligi(self, doc, maksimum_cm=16.5):
        """Görseli A4 sayfanın kenar boşlukları içinde mümkün olduğunca büyüt."""
        try:
            section = doc.sections[0]
            kullanilabilir = section.page_width - section.left_margin - section.right_margin
            return min(Cm(maksimum_cm), kullanilabilir)
        except Exception:
            return Cm(15)

    def rapor_pga_haritasi_yolu(self):
        yol = str(getattr(self, "img_pga_haritasi", "") or "").strip()
        if yol and os.path.isfile(yol):
            return yol
        try:
            from harita_dosyalari import proje_klasorundeki_rapor_haritalarini_bul

            bulunan = proje_klasorundeki_rapor_haritalarini_bul(
                getattr(self, "guncel_dosya_yolu", "")
            )
            yol = bulunan.get("img_pga_haritasi", "")
            if yol:
                self.img_pga_haritasi = yol
        except Exception:
            yol = ""
        return yol if yol and os.path.isfile(yol) else ""

    def rapor_paragraf_gorsel_iceriyor(self, paragraph):
        return any(
            next(paragraph._p.iter(qn(etiket)), None) is not None
            for etiket in ("w:drawing", "w:pict", "w:object")
        )

    def rapor_paragraf_icerigini_temizle(self, paragraph):
        """Paragraf biçimini koruyup içindeki metin/görsel/alanları temizle."""
        p_pr = paragraph._p.find(qn("w:pPr"))
        for child in list(paragraph._p):
            if child is not p_pr:
                paragraph._p.remove(child)

    def rapor_pga_haritasini_yerlestir(self, doc, caption, pga_yolu):
        """Şablondaki sabit deprem haritasını seçilen PGA görseliyle değiştir.

        Şekil başlığı aynı yerde kalır. Başlıktan sonraki ilk görsel paragrafı
        kaynak satırından önce aranır ve yerinde değiştirilir; böylece şablon
        görseli ile PGA görseli raporda üst üste/peş peşe kalmaz.
        """
        hedef = None
        aday = caption._p.getnext()
        while aday is not None and aday.tag != qn("w:sectPr"):
            if aday.tag != qn("w:p"):
                break
            paragraph = Paragraph(aday, caption._parent)
            metin = self.rapor_metin_normalize(paragraph.text)
            if self.rapor_paragraf_gorsel_iceriyor(paragraph):
                hedef = paragraph
                break
            if metin:
                break
            aday = aday.getnext()

        if hedef is None:
            image_xml = OxmlElement("w:p")
            caption._p.addnext(image_xml)
            hedef = Paragraph(image_xml, caption._parent)
        else:
            self.rapor_paragraf_icerigini_temizle(hedef)

        hedef.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.rapor_gorsel_ekle(
            hedef,
            pga_yolu,
            width=self.rapor_gorsel_kullanilabilir_genisligi(doc),
        )
        return hedef

    def rapor_tasima_etiket_degerlerini_dogrula(self, doc):
        etiketler = self.rapor_docx_etiketlerini_oku(doc)
        kritik = etiketler & {"[QK]", "[QT]", "[KS]", "[DF]", "[TASIMA_GUCU]"}
        if not kritik:
            return
        if hasattr(self, "tasima_raporu_guncel_mi") and not self.tasima_raporu_guncel_mi():
            raise ValueError("Taşıma gücü girdileri değişmiş; taşıma rapor metnini yeniden oluşturun.")

        kontroller = []
        if "[QK]" in kritik:
            kontroller.append(("QK", getattr(self, "son_qk", None), False))
        if "[QT]" in kritik:
            entry = getattr(self, "entry_qt_nihai", None)
            kontroller.append(("QT", entry.get() if entry is not None else None, False))
        if "[KS]" in kritik:
            entry = getattr(self, "entry_ks_nihai", None)
            kontroller.append(("KS", entry.get() if entry is not None else None, False))
        if "[DF]" in kritik:
            entry = getattr(self, "tg_girdiler", {}).get("Df")
            kontroller.append(("DF", entry.get() if entry is not None else None, True))

        for ad, ham, sifir_kabul in kontroller:
            if ham is None or str(ham).strip() in ("", "-"):
                raise ValueError(f"Şablondaki [{ad}] etiketi için doğrulanmış değer bulunamadı.")
            try:
                sayi = float(str(ham).strip().replace(",", "."))
            except ValueError as exc:
                raise ValueError(f"Şablondaki [{ad}] değeri sayısal olmalıdır.") from exc
            if not math.isfinite(sayi) or sayi < 0 or (not sifir_kabul and sayi <= 0):
                kosul = "negatif olmayan" if sifir_kabul else "sıfırdan büyük"
                raise ValueError(f"Şablondaki [{ad}] değeri sonlu ve {kosul} olmalıdır.")

    def rapor_statik_etiketleri_degistir(self, doc):
        self.rapor_tasima_etiket_degerlerini_dogrula(doc)
        degisim_sozlugu, fm_metin, _, _, _ = self.rapor_etiket_verilerini_hazirla()
        self._rapor_muhendislik_jeolojisi_cumlesi = str(
            degisim_sozlugu.get("[MUHENDISLIK_JEOLOJISI_METNI]", "") or ""
        ).strip()
        cok_satirli_etiketler = {
            "[FORMASYON]": fm_metin,
        }
        tek_satirli_etiketler = {
            eski: yeni
            for eski, yeni in degisim_sozlugu.items()
            if eski not in cok_satirli_etiketler
        }
        self.rapor_metinleri_toplu_degistir(doc, tek_satirli_etiketler)
        for eski, yeni in cok_satirli_etiketler.items():
            self.rapor_cok_satirli_etiket_degistir(doc, eski, yeni)

    def rapor_bina_tablosu_ekle(self, doc):
        bina_p = self.rapor_paragraf_bul(doc, "[BINA]")
        if not bina_p:
            return
        self.rapor_metin_degistir(doc, "[BINA]", "")
        if not self.bina_alanlari:
            return
        tablo_bina = doc.add_table(rows=len(self.bina_alanlari), cols=2)
        tablo_bina.style = 'Table Grid'
        for i, (etiket, entry) in enumerate(self.bina_alanlari.items()):
            tablo_bina.rows[i].cells[0].text = etiket
            tablo_bina.rows[i].cells[1].text = entry.get()
        self.rapor_tablo_stili_uygula(tablo_bina, header_rows=0, label_columns=1)
        bina_p._p.addnext(tablo_bina._tbl)

    def rapor_resimleri_ekle(self, doc):
        resim_etiketleri = [
            ("[MJH]", getattr(self, "img_mjh", None)),
            ("[JEOFIZIK_LOKASYON]", getattr(self, "img_jeofizik_lok", None)),
            ("[JEOLOJI_LOKASYON]", getattr(self, "img_jeoloji_lok", None)),
            ("[YERBULDURU]", getattr(self, "img_yerbulduru", None)),
            ("[PARSEL_HARITASI]", getattr(self, "img_parsel_haritasi", None)),
        ]
        for tag, img_path in resim_etiketleri:
            p = self.rapor_paragraf_bul(doc, tag)
            if not p:
                continue
            self.rapor_metin_degistir(doc, tag, "")
            if img_path and os.path.exists(img_path):
                width = self.rapor_gorsel_kullanilabilir_genisligi(doc)
                self.rapor_gorsel_ekle(p, img_path, width=width)
            else:
                p.add_run("(Harita resmi yok. Haritalar sekmesinden Hazırla butonuna basın)")

        pga_yolu = self.rapor_pga_haritasi_yolu()
        pga_caption_key = "canakkale_bolgesi_deprem_tehlike_haritasi"
        for caption in self.docx_paragraflarini_dolas(doc):
            if pga_caption_key not in self.rapor_metin_normalize(caption.text):
                continue
            if pga_yolu:
                self.rapor_pga_haritasini_yerlestir(doc, caption, pga_yolu)
            else:
                uyari = (
                    "PGA haritası görseli bulunamadı; bu konum haritasız bırakıldı. "
                    "Ön Değerler sekmesinden isteğe bağlı bir PGA görseli seçin."
                )
                caption.add_run(" (" + uyari + ")")
                rapor_uyarilari = list(getattr(self, "_rapor_uyarilari", []))
                if uyari not in rapor_uyarilari:
                    rapor_uyarilari.append(uyari)
                self._rapor_uyarilari = rapor_uyarilari

    def rapor_jeofizik_koordinat_tablosu_ekle(self, doc):
        jeo_koor_p = self.rapor_paragraf_bul(doc, "[JEOFIZIK_KOORDINAT]")
        if not jeo_koor_p:
            return
        self.rapor_metin_degistir(doc, "[JEOFIZIK_KOORDINAT]", "")
        koordinatlar = self.jeofizik_koordinatlari_al()
        if not koordinatlar:
            return

        tablo = doc.add_table(rows=len(koordinatlar) + 2, cols=3)
        tablo.style = 'Table Grid'
        h1 = tablo.rows[0].cells
        h1[0].text = "Çalışma No"
        h1[1].text = "Koordinatlar (WGS84)"
        h1[1].merge(h1[2])
        h2 = tablo.rows[1].cells
        h2[1].text = "Enlem"
        h2[2].text = "Boylam"

        satir_idx = 2
        for degerler in koordinatlar:
            tablo.rows[satir_idx].cells[0].text = str(degerler[0])
            tablo.rows[satir_idx].cells[1].text = str(degerler[1]) + "°"
            tablo.rows[satir_idx].cells[2].text = str(degerler[2]) + "°"
            satir_idx += 1
        self.rapor_tablo_stili_uygula(tablo, header_rows=2)
        jeo_koor_p._p.addnext(tablo._tbl)

    def rapor_jeoloji_koordinat_tablosu_ekle(self, doc):
        jeoloji_koor_p = self.rapor_paragraf_bul(doc, "[JEOLOJI_KOORDINAT]")
        if not jeoloji_koor_p:
            return
        self.rapor_metin_degistir(doc, "[JEOLOJI_KOORDINAT]", "")
        jeo_points = [(isim, d["lat"], d["lon"]) for isim, d in self.harita_isaretleri.items() if d["tip"] in ["AÇ", "YN"]]
        if len(jeo_points) == 0:
            return

        tablo = doc.add_table(rows=len(jeo_points) + 2, cols=3)
        tablo.style = 'Table Grid'
        h1 = tablo.rows[0].cells
        h1[0].text = "Çalışma No"
        h1[1].text = "Koordinatlar (WGS84)"
        h1[1].merge(h1[2])
        h2 = tablo.rows[1].cells
        h2[1].text = "Enlem"
        h2[2].text = "Boylam"

        satir_idx = 2
        for isim, lat, lon in jeo_points:
            tablo.rows[satir_idx].cells[0].text = isim
            tablo.rows[satir_idx].cells[1].text = f"{lat:.6f}°"
            tablo.rows[satir_idx].cells[2].text = f"{lon:.6f}°"
            satir_idx += 1
        self.rapor_tablo_stili_uygula(tablo, header_rows=2)
        jeoloji_koor_p._p.addnext(tablo._tbl)

    def rapor_koordinat_tablolarini_ekle(self, doc):
        self.rapor_jeofizik_koordinat_tablosu_ekle(doc)
        self.rapor_jeoloji_koordinat_tablosu_ekle(doc)

    def rapor_jeofizik_excel_yolu(self):
        return self.jeofizik_excel_yolu_al()

    def rapor_jeofizik_tablo_dosyasi_oku(self, dosya_yolu):
        if not dosya_yolu or not os.path.isfile(dosya_yolu):
            raise ValueError("Jeofizik parametre dosyası bulunamadı.")
        uzanti = os.path.splitext(dosya_yolu)[1].lower()
        if uzanti == ".csv":
            son_hata = None
            for kodlama in ("utf-8-sig", "utf-8", "utf-16", "cp1254", "latin-1"):
                try:
                    with open(dosya_yolu, "r", encoding=kodlama, newline="") as f:
                        ornek = f.read(65536)
                        if not ornek.strip():
                            raise ValueError("Jeofizik CSV dosyası boş.")
                        try:
                            lehce = csv.Sniffer().sniff(ornek, delimiters=",;\t|")
                        except csv.Error:
                            ayirac = max(",;\t|", key=ornek.count)
                            lehce = None
                        f.seek(0)
                        okuyucu = csv.reader(f, lehce) if lehce is not None else csv.reader(f, delimiter=ayirac)
                        satirlar = list(okuyucu)
                    return pd.DataFrame(satirlar)
                except UnicodeError as exc:
                    son_hata = exc
                    continue
                except csv.Error as exc:
                    raise ValueError(f"Jeofizik CSV dosyası ayrıştırılamadı: {exc}") from exc
            raise ValueError(f"Jeofizik CSV dosyasının karakter kodlaması okunamadı: {son_hata}")
        if uzanti not in {".xlsx", ".xls", ".xlsm"}:
            raise ValueError(f"Desteklenmeyen jeofizik parametre dosyası: {uzanti or 'uzantısız'}")
        sayfalar = pd.read_excel(dosya_yolu, header=None, sheet_name=None)
        dolu_sayfalar = [df for df in sayfalar.values() if not df.dropna(how="all").empty]
        if not dolu_sayfalar:
            raise ValueError("Jeofizik Excel dosyasında okunabilir veri bulunamadı.")
        return pd.concat(dolu_sayfalar, ignore_index=True)

    def rapor_jeofizik_parametrelerini_oku(self):
        param_ss_list = []
        excel_yolu = self.rapor_jeofizik_excel_yolu()
        if not excel_yolu:
            return param_ss_list
        try:
            df_jeo = self.rapor_jeofizik_tablo_dosyasi_oku(excel_yolu)
            current_serim = None
            for _, row in df_jeo.iterrows():
                row_str = [str(x).strip(" \t\n\r\xa0") for x in row if pd.notna(x)]
                if not row_str:
                    continue

                cell_0_str = str(row.iloc[0]).strip(" \t\n\r\xa0") if pd.notna(row.iloc[0]) else ""
                cell_0_normal = self.rapor_metin_normalize(cell_0_str)
                if "sismik_olcu_ve_hesaplarinin_sahibi" in cell_0_normal:
                    s_name = "SS"
                    for cell in row_str:
                        cell_normal = self.rapor_metin_normalize(cell)
                        if "serim" in cell_normal or re.search(r"(^|_)ss\d*($|_)", cell_normal):
                            s_name = cell
                            break
                    current_serim = {"ad": s_name, "layers": []}
                    param_ss_list.append(current_serim)
                elif current_serim is not None:
                    r0 = cell_0_normal
                    if r0.startswith("vs30"):
                        current_serim["raw_vs30"] = row.tolist()
                    elif r0.startswith("vp") or "boyuna_dalga" in r0:
                        current_serim["raw_vp"] = row.tolist()
                    elif r0.startswith("vs") or "enine_dalga" in r0:
                        current_serim["raw_vs"] = row.tolist()
                    elif "tabaka_kalinligi" in r0:
                        current_serim["raw_h"] = row.tolist()
                    elif "tabaka_yogunlugu" in r0:
                        current_serim["raw_rho"] = row.tolist()
                    elif "poisson_orani" in r0:
                        current_serim["raw_nu"] = row.tolist()
                    elif "elastisite" in r0 or "young" in r0:
                        current_serim["raw_E"] = row.tolist()
                    elif "kayma_modulu" in r0 or "gmax" in r0:
                        current_serim["raw_G"] = row.tolist()
                    elif "bulk" in r0 or "sikismazlik" in r0:
                        current_serim["raw_K"] = row.tolist()

            for s in param_ss_list:
                s["ad"] = str(s.get("ad", "SS")).strip(" \t\n\r\xa0")
                if "raw_vp" not in s:
                    continue
                vs30_val = "-"
                if "raw_vs30" in s:
                    for val in s["raw_vs30"][2:]:
                        if pd.notna(val):
                            v_str = str(val).strip(" \t\n\r\xa0")
                            if v_str != "":
                                vs30_val = v_str
                                break

                def ham_deger(key, col_idx):
                    raw = s.get(key, [])
                    return raw[col_idx] if col_idx < len(raw) else "-"

                for col_idx in range(2, len(s["raw_vp"])):
                    vp_val = s["raw_vp"][col_idx]
                    if pd.isna(vp_val) or str(vp_val).strip(" \t\n\r\xa0") == "":
                        continue
                    layer = {
                        "vp": vp_val,
                        "vs": ham_deger("raw_vs", col_idx),
                        "h": ham_deger("raw_h", col_idx),
                        "rho": ham_deger("raw_rho", col_idx),
                        "nu": ham_deger("raw_nu", col_idx),
                        "E": ham_deger("raw_E", col_idx),
                        "G": ham_deger("raw_G", col_idx),
                        "K": ham_deger("raw_K", col_idx),
                        "vs30": vs30_val if col_idx == 2 else "-",
                    }
                    s["layers"].append(layer)
        except Exception as e:
            self.hata_kaydet("Rapor için jeofizik parametreleri okunamadı", e)
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Jeofizik parametre dosyası okunamadı: {e}") from e
        if not param_ss_list:
            raise ValueError(
                "Jeofizik parametre dosyasında 'Sismik Ölçü ve Hesaplarının Sahibi' bölümü bulunamadı."
            )
        if not any(s.get("layers") for s in param_ss_list):
            raise ValueError("Jeofizik parametre dosyasında Vp tabaka verisi bulunamadı.")
        return param_ss_list

    def rapor_tablo_basligini_kalin_yap(self, tablo):
        for cell in tablo.rows[0].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True

    def rapor_jeofon_sayisi_parse(self, deger, varsayilan=12):
        try:
            sayi = int(float(str(deger).replace(",", ".")))
        except Exception:
            sayi = varsayilan
        return max(1, min(48, sayi))

    def rapor_float_parse(self, deger, varsayilan=0.0):
        try:
            return float(str(deger).replace(",", "."))
        except Exception:
            return varsayilan

    def rapor_jeofon_dizilim_bilgileri(self):
        if hasattr(self.app, "jeofon_dizilim_bilgileri_al"):
            bilgiler = self.jeofon_dizilim_bilgileri_al()
        else:
            bilgiler = {}
        return {
            "jeofon_sayisi": self.rapor_jeofon_sayisi_parse(bilgiler.get("jeofon_sayisi", "12"), 12),
            "jeofon_araligi": self.rapor_float_parse(bilgiler.get("jeofon_araligi", "2"), 2.0),
            "duz_offset": self.rapor_float_parse(bilgiler.get("duz_offset", "0"), 0.0),
            "ters_offset": self.rapor_float_parse(bilgiler.get("ters_offset", "2"), 2.0),
        }

    def rapor_hucre_arkaplan_rengi(self, cell, renk):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), renk)

    def rapor_hucre_metni_ayarla(self, cell, metin, bold=False, renk=None, punto=8):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(metin))
        run.font.bold = bold
        run.font.size = Pt(punto)
        if renk:
            run.font.color.rgb = RGBColor.from_string(renk)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def rapor_jeofon_dizilim_satirlari(self):
        bilgiler = self.rapor_jeofon_dizilim_bilgileri()
        jeofon_sayisi = bilgiler["jeofon_sayisi"]
        aralik = bilgiler["jeofon_araligi"]
        duz_offset = bilgiler["duz_offset"]
        ters_offset = bilgiler["ters_offset"]

        headers = ["", "Düz\nVuruş"]
        headers.extend([f"{i}\nJeof" for i in range(1, jeofon_sayisi + 1)])
        headers.append("Ters\nVuruş")

        kaynak_mesafeleri = ["Kaynak ve Jeofon mesafeleri (m.)", self.rapor_deger_formatla(duz_offset)]
        kaynak_mesafeleri.extend([self.rapor_deger_formatla(aralik)] * jeofon_sayisi)
        kaynak_mesafeleri.append(self.rapor_deger_formatla(ters_offset))

        duzden_mesafeler = ["Düz Vuruştan Mesafeler (m.)", self.rapor_deger_formatla(duz_offset)]
        duzden_mesafeler.extend(
            [self.rapor_deger_formatla(duz_offset + (i * aralik)) for i in range(1, jeofon_sayisi + 1)]
        )
        duzden_mesafeler.append(self.rapor_deger_formatla(duz_offset + (jeofon_sayisi * aralik) + ters_offset))

        return headers, kaynak_mesafeleri, duzden_mesafeler

    def rapor_jeofon_dizilim_tablosu_ekle(self, doc):
        tag = "[JEOFON_DIZILIM]"
        hedef_p = self.rapor_paragraf_bul(doc, tag)
        if not hedef_p:
            return
        self.rapor_metin_degistir(doc, tag, "")

        satirlar = self.rapor_jeofon_dizilim_satirlari()
        kolon_sayisi = len(satirlar[0])
        tablo = doc.add_table(rows=3, cols=kolon_sayisi)
        tablo.style = "Table Grid"
        tablo.alignment = WD_TABLE_ALIGNMENT.CENTER
        tablo.autofit = True

        for r_idx, satir in enumerate(satirlar):
            for c_idx, metin in enumerate(satir):
                cell = tablo.rows[r_idx].cells[c_idx]
                cell.text = str(metin)

        self.rapor_tablo_stili_uygula(tablo, header_rows=1, label_columns=1)
        hedef_p._p.addnext(tablo._tbl)

    def rapor_jeofizik_parametre_tablosu_ekle(self, doc, param_ss_list):
        jeo_param_p = self.rapor_paragraf_bul(doc, "[JEOFIZIK_PARAMETRE]")
        if not (jeo_param_p and param_ss_list):
            return
        self.rapor_metin_degistir(doc, "[JEOFIZIK_PARAMETRE]", "")
        headers = ["Serim", "Tabaka", "Kalınlık (m)", "Vp (m/s)", "Vs (m/s)", "Yoğ. (g/cm3)", "Poisson", "Elast. (kg/cm2)", "Kayma (kg/cm2)", "Bulk (kg/cm2)"]
        total_rows = sum(len(ss.get("layers", [])) for ss in param_ss_list)
        if total_rows <= 0:
            return

        tablo = doc.add_table(rows=total_rows + 1, cols=len(headers))
        tablo.style = 'Table Grid'
        for j, h in enumerate(headers):
            tablo.rows[0].cells[j].text = h
        self.rapor_tablo_basligini_kalin_yap(tablo)

        row_idx = 1
        for ss in param_ss_list:
            layers = ss.get("layers", [])
            start_merge_idx = row_idx
            for i, layer in enumerate(layers):
                r = tablo.rows[row_idx].cells
                r[0].text = str(ss.get("ad", "SS")) if i == 0 else ""
                r[1].text = str(i + 1)
                r[2].text = "-" if i == len(layers) - 1 else self.rapor_deger_formatla(layer.get("h", "-"))
                r[3].text = self.rapor_deger_formatla(layer.get("vp", "-"))
                r[4].text = self.rapor_deger_formatla(layer.get("vs", "-"))
                r[5].text = self.rapor_deger_formatla(layer.get("rho", "-"))
                r[6].text = self.rapor_deger_formatla(layer.get("nu", "-"))
                r[7].text = self.rapor_deger_formatla(layer.get("E", "-"))
                r[8].text = self.rapor_deger_formatla(layer.get("G", "-"))
                r[9].text = self.rapor_deger_formatla(layer.get("K", "-"))
                row_idx += 1

            if (row_idx - 1) > start_merge_idx:
                mc = tablo.rows[start_merge_idx].cells[0].merge(tablo.rows[row_idx - 1].cells[0])
                mc.text = str(ss.get("ad", "SS")).strip()
                for p in mc.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        self.rapor_tablo_stili_uygula(tablo, header_rows=1, font_size=8)
        jeo_param_p._p.addnext(tablo._tbl)

    def rapor_masw_liste_hazirla(self, param_ss_list):
        masw_ozel_liste = []
        total_masw_rows = 0
        for ss in param_ss_list:
            f_layers = []
            last_vs = None
            ortam_no = 1
            current_h_sum = 0.0
            has_inf = False

            vs30_val = "-"
            if "raw_vs30" in ss:
                for val in ss["raw_vs30"][2:]:
                    if pd.notna(val) and str(val).strip() != "":
                        vs30_val = self.rapor_deger_formatla(str(val).strip())
                        break

            layers = ss.get("layers", [])
            for i, layer in enumerate(layers):
                val_vs = self.rapor_deger_formatla(layer.get("vs", "-"))
                val_h_raw = layer.get("h", "-")

                is_inf = False
                h_val = 0.0
                if val_h_raw == "-" or pd.isna(val_h_raw) or str(val_h_raw).strip() == "":
                    is_inf = True
                else:
                    try:
                        h_val = float(str(val_h_raw).replace(",", "."))
                    except:
                        is_inf = True

                if i == len(layers) - 1:
                    is_inf = True

                if val_vs != last_vs and val_vs != "-":
                    if last_vs is not None:
                        f_layers[-1]["h"] = "-" if has_inf else self.rapor_deger_formatla(current_h_sum)
                    f_layers.append({"ortam_no": ortam_no, "vs": val_vs, "h": "-"})
                    ortam_no += 1
                    last_vs = val_vs
                    current_h_sum = h_val
                    has_inf = is_inf
                elif val_vs == last_vs and val_vs != "-":
                    current_h_sum += h_val
                    if is_inf:
                        has_inf = True

            if f_layers:
                f_layers[-1]["h"] = "-" if has_inf else self.rapor_deger_formatla(current_h_sum)
                masw_ozel_liste.append({"ad": ss.get("ad", "SS"), "vs30": vs30_val, "layers": f_layers})
                total_masw_rows += len(f_layers)
        return masw_ozel_liste, total_masw_rows

    def rapor_masw_tablosu_ekle(self, doc, param_ss_list):
        masw_p = self.rapor_paragraf_bul(doc, "[MASW]")
        if not (masw_p and param_ss_list):
            return
        self.rapor_metin_degistir(doc, "[MASW]", "")

        masw_ozel_liste, total_masw_rows = self.rapor_masw_liste_hazirla(param_ss_list)
        headers_masw = ["Serim No", "Ortam No", "Vs(m/sn)", "Kalınlık h (m)", "Vs30(m/sn)"]
        if total_masw_rows <= 0:
            return

        tablo_masw = doc.add_table(rows=total_masw_rows + 1, cols=len(headers_masw))
        tablo_masw.style = 'Table Grid'
        for j, h in enumerate(headers_masw):
            tablo_masw.rows[0].cells[j].text = h
        self.rapor_tablo_basligini_kalin_yap(tablo_masw)

        row_idx = 1
        for ss in masw_ozel_liste:
            layers = ss.get("layers", [])
            start_merge_idx = row_idx
            for i, layer in enumerate(layers):
                r = tablo_masw.rows[row_idx].cells
                r[0].text = str(ss.get("ad", "SS")) if i == 0 else ""
                r[1].text = str(layer.get("ortam_no", i + 1))
                r[2].text = str(layer.get("vs", "-"))
                r[3].text = str(layer.get("h", "-"))
                r[4].text = str(ss.get("vs30", "-")) if i == 0 else ""
                row_idx += 1

            if (row_idx - 1) > start_merge_idx:
                mc1 = tablo_masw.rows[start_merge_idx].cells[0].merge(tablo_masw.rows[row_idx - 1].cells[0])
                mc1.text = str(ss.get("ad", "SS")).strip()
                for p in mc1.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mc1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                mc2 = tablo_masw.rows[start_merge_idx].cells[4].merge(tablo_masw.rows[row_idx - 1].cells[4])
                mc2.text = str(ss.get("vs30", "-")).strip()
                for p in mc2.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mc2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        self.rapor_tablo_stili_uygula(tablo_masw, header_rows=1)
        masw_p._p.addnext(tablo_masw._tbl)

    def rapor_vp_liste_hazirla(self, param_ss_list):
        vp_ozel_liste = []
        total_vp_rows = 0
        for ss in param_ss_list:
            f_layers = []
            last_vp = None
            ortam_no = 1
            for layer in ss.get("layers", []):
                val_fmt = self.rapor_deger_formatla(layer.get("vp", "-"))
                if val_fmt != last_vp and val_fmt != "-":
                    f_layers.append({"ortam_no": ortam_no, "vp": val_fmt})
                    ortam_no += 1
                    last_vp = val_fmt
            if f_layers:
                vp_ozel_liste.append({"ad": ss.get("ad", "SS"), "layers": f_layers})
                total_vp_rows += len(f_layers)
        return vp_ozel_liste, total_vp_rows

    def rapor_vp_tablosu_ekle(self, doc, param_ss_list):
        vp_p = self.rapor_paragraf_bul(doc, "[VP]")
        if not (vp_p and param_ss_list):
            return
        self.rapor_metin_degistir(doc, "[VP]", "")

        vp_ozel_liste, total_vp_rows = self.rapor_vp_liste_hazirla(param_ss_list)
        headers_vp = ["Ölçü No", "Ortam No", "Vp (m/sn)"]
        if total_vp_rows <= 0:
            return

        tablo_vp = doc.add_table(rows=total_vp_rows + 1, cols=len(headers_vp))
        tablo_vp.style = 'Table Grid'
        for j, h in enumerate(headers_vp):
            tablo_vp.rows[0].cells[j].text = h
        self.rapor_tablo_basligini_kalin_yap(tablo_vp)

        row_idx = 1
        for ss in vp_ozel_liste:
            layers = ss.get("layers", [])
            start_merge_idx = row_idx
            for i, layer in enumerate(layers):
                r = tablo_vp.rows[row_idx].cells
                r[0].text = str(ss.get("ad", "SS")) if i == 0 else ""
                r[1].text = str(layer.get("ortam_no", i + 1))
                r[2].text = layer.get("vp", "-")
                row_idx += 1

            if (row_idx - 1) > start_merge_idx:
                mc = tablo_vp.rows[start_merge_idx].cells[0].merge(tablo_vp.rows[row_idx - 1].cells[0])
                mc.text = str(ss.get("ad", "SS")).strip()
                for p in mc.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        self.rapor_tablo_stili_uygula(tablo_vp, header_rows=1)
        vp_p._p.addnext(tablo_vp._tbl)

    def rapor_jeofizik_tablolarini_ekle(self, doc):
        param_ss_list = getattr(self, "_rapor_param_ss_list", None)
        if param_ss_list is None:
            param_ss_list = self.rapor_jeofizik_parametrelerini_oku()
            self._rapor_param_ss_list = param_ss_list
        self.rapor_jeofon_dizilim_tablosu_ekle(doc)
        self.rapor_jeofizik_parametre_tablosu_ekle(doc, param_ss_list)
        self.rapor_masw_tablosu_ekle(doc, param_ss_list)
        self.rapor_vp_tablosu_ekle(doc, param_ss_list)

    def rapor_tasima_float_al(self, anahtar, varsayilan=None):
        entry = getattr(self, "tg_girdiler", {}).get(anahtar)
        if entry is None:
            raise ValueError(f"Taşıma gücü girdisi bulunamadı: {anahtar}")
        ham = str(entry.get()).strip()
        if not ham:
            raise ValueError(f"Taşıma gücü girdisi boş: {anahtar}")
        try:
            deger = float(ham.replace(",", "."))
        except ValueError as exc:
            raise ValueError(f"Taşıma gücü girdisi sayısal değil: {anahtar}") from exc
        if not math.isfinite(deger):
            raise ValueError(f"Taşıma gücü girdisi sonlu değil: {anahtar}")
        return deger

    def rapor_tasima_sonuc_float_al(self, alan_adi, etiket):
        entry = getattr(self, alan_adi, None)
        ham = str(entry.get()).strip() if entry is not None else ""
        if not ham:
            raise ValueError(f"Taşıma gücü sonucu boş: {etiket}")
        try:
            deger = float(ham.replace(",", "."))
        except ValueError as exc:
            raise ValueError(f"Taşıma gücü sonucu sayısal değil: {etiket}") from exc
        if not math.isfinite(deger) or deger <= 0:
            raise ValueError(f"Taşıma gücü sonucu sonlu ve sıfırdan büyük olmalıdır: {etiket}")
        return deger

    def rapor_tasima_gucu_tablosu_olustur(self, doc):
        c_v = self.rapor_tasima_float_al("c")
        phi_v = self.rapor_tasima_float_al("phi")
        gn_v = self.rapor_tasima_float_al("gn")
        gsat_v = self.rapor_tasima_float_al("gsat")
        gn_kNm3 = gn_v * 9.81
        gsat_kNm3 = gsat_v * 9.81
        B_v = self.rapor_tasima_float_al("B")
        L_v = self.rapor_tasima_float_al("L")
        Df_v = self.rapor_tasima_float_al("Df")
        yass_variable = getattr(self, "tasima_yass_var", None)
        yass_var = bool(yass_variable.get()) if yass_variable is not None else True
        yass_v = (
            self.rapor_tasima_float_al("yass")
            if yass_var
            else Df_v + B_v
        )
        girilen_gamma_rv = self.rapor_tasima_float_al("RvGk")
        gamma_rv = 1.4
        if not math.isclose(girilen_gamma_rv, gamma_rv, rel_tol=0.0, abs_tol=1e-9):
            raise ValueError("Zemin taşıma gücü için γ_Rv değeri 1.40 olmalıdır.")

        a = TBDY2018TasimaGucu(c_v, phi_v, gn_kNm3, gsat_kNm3, yass_v)
        qk_kn, qt_kn = a.analiz_yap(B_v, L_v, Df_v, gamma_Rv=gamma_rv)
        qk_tm2 = qk_kn / 9.81
        hesaplanan_qt_tm2 = qt_kn / 9.81
        nihai_qt_tm2 = self.rapor_tasima_sonuc_float_al("entry_qt_nihai", "nihai q_t")
        tolerans = max(1e-9, abs(hesaplanan_qt_tm2) * 1e-12)
        if nihai_qt_tm2 > hesaplanan_qt_tm2 + tolerans:
            raise ValueError(
                f"Raporda kullanılacak q_t ({nihai_qt_tm2:.2f} t/m²), hesaplanan tasarım "
                f"dayanımından ({hesaplanan_qt_tm2:.2f} t/m²) büyük olamaz."
            )

        ks_tm3 = nihai_qt_tm2 * BOWLES_ZEMIN_KATSAYISI * gamma_rv
        girilen_ks_tm3 = self.rapor_tasima_sonuc_float_al("entry_ks_nihai", "nihai k_s")
        if not math.isclose(girilen_ks_tm3, round(ks_tm3, 2), rel_tol=0.0, abs_tol=0.005000001):
            raise ValueError(
                f"Nihai k_s değeri, {BOWLES_ZEMIN_KATSAYISI:.0f} × q_t × G_k bağıntısından "
                f"hesaplanan {ks_tm3:.2f} t/m³ değeriyle uyuşmuyor; taşıma rapor metnini yeniden oluşturun."
            )

        tablo_tbdy = doc.add_table(rows=9, cols=6)
        tablo_tbdy.style = 'Table Grid'
        tablo_tbdy.allow_autofit = True
        tablo_tbdy.autofit = True

        for tblW in tablo_tbdy._tbl.tblPr.xpath("./w:tblW"):
            tblW.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'auto')
            tblW.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '0')

        for row in tablo_tbdy.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.get_or_add_tcW()
                tcW.type = 'auto'
                tcW.w = 0

        ts_data = [
            ["Kohezyon (c) [t/m^2]", f"{c_v / 9.81:.2f}", "N_q", f"{a.Nq:.2f}", "i_q", f"{a.iq:.2f}"],
            ["İçsel Sürtünme Açısı ϕ [derece]", f"{phi_v:.1f}", "N_c", f"{a.Nc:.2f}", "i_c", f"{a.ic:.2f}"],
            ["Birim Hacim Ağırlık γ [t/m^3]", f"{gn_v:.2f}", "N_γ", f"{a.Ngamma:.2f}", "i_γ", f"{a.igamma:.2f}"],
            ["Temel Genişliği (B) [metre]", f"{B_v:.2f}", "s_q", f"{a.sq:.2f}", "g_q", f"{a.gq:.2f}"],
            ["Temel Uzunluğu (L) [metre]", f"{L_v:.2f}", "s_c", f"{a.sc:.2f}", "g_c", f"{a.gc:.2f}"],
            ["Temel Derinliği (D_f) [metre]", f"{Df_v:.2f}", "s_γ", f"{a.sgamma:.2f}", "g_γ", f"{a.ggamma:.2f}"],
            ["Yükleme Eğikliği Açısı [derece]", "0", "d_q", f"{a.dq:.2f}", "b_q", f"{a.bq:.2f}"],
            ["Temel Zemini Eğim Açısı [derece]", "0", "d_c", f"{a.dc:.2f}", "b_c", f"{a.bc:.2f}"],
            ["Temel Taban Eğim Açısı [derece]", "0", "d_γ", f"{a.dgamma:.2f}", "b_γ", f"{a.bgamma:.2f}"],
        ]

        for r_idx, satir_veri in enumerate(ts_data):
            for c_idx, hucre_veri in enumerate(satir_veri):
                cell = tablo_tbdy.rows[r_idx].cells[c_idx]
                cell.text = ""
                p_cell = cell.paragraphs[0]
                self.rapor_formatli_metin_ekle(p_cell, hucre_veri)
                if c_idx in [0, 2, 4]:
                    for run in p_cell.runs:
                        run.bold = True

        self.rapor_tablo_stili_uygula(tablo_tbdy, header_rows=0, label_columns={0, 2, 4}, font_size=9)
        return tablo_tbdy

    def rapor_tasima_tekrar_satiri_mi(self, satir):
        normalize = self.rapor_metin_normalize(satir)
        if normalize.startswith("hesapta_c_"):
            return True
        return (
            normalize.startswith(("n_c_", "nc_"))
            and "n_q_" in normalize
            and "s_c_" in normalize
            and "s_q_" in normalize
            and "d_c_" in normalize
            and "d_q_" in normalize
        )

    def rapor_tasima_metin_satirlarini_temizle(self, tam_metin):
        temiz_satirlar = []
        tablo_sonrasi = False
        for satir in str(tam_metin or "").splitlines():
            if self.rapor_tasima_tekrar_satiri_mi(satir):
                continue
            if "[TABLO_BURADA]" in satir:
                while temiz_satirlar and not temiz_satirlar[-1].strip():
                    temiz_satirlar.pop()
                temiz_satirlar.append("[TABLO_BURADA]")
                tablo_sonrasi = True
                continue
            if not satir.strip():
                if tablo_sonrasi or not temiz_satirlar or not temiz_satirlar[-1].strip():
                    continue
                temiz_satirlar.append("")
                continue
            tablo_sonrasi = False
            temiz_satirlar.append(satir)

        while temiz_satirlar and not temiz_satirlar[-1].strip():
            temiz_satirlar.pop()
        return temiz_satirlar

    def rapor_tasima_gucu_ekle(self, doc):
        tg_p = self.rapor_paragraf_bul(doc, "[TASIMA_GUCU]")
        if not tg_p:
            return
        self.rapor_metin_degistir(doc, "[TASIMA_GUCU]", "")
        if hasattr(self, "tasima_raporu_guncel_mi") and not self.tasima_raporu_guncel_mi():
            raise ValueError("Taşıma gücü girdileri değişmiş; taşıma rapor metnini yeniden oluşturun.")
        tam_metin = ""
        if hasattr(self, "txt_tasima_rapor"):
            tam_metin = self.txt_tasima_rapor.get("1.0", tk.END).strip()
        if not tam_metin:
            raise ValueError("Taşıma gücü rapor metni oluşturulmamış.")

        # Eski proje dosyalarında saklanan girdi/katsayı tekrarlarını ve bunlardan
        # kalan yığılmış boş paragrafları çıktı sırasında da ayıkla.
        temiz_satirlar = self.rapor_tasima_metin_satirlarini_temizle(tam_metin)

        tablo_eklendi = False
        onceki_paragraf = None
        for satir in temiz_satirlar:
            if "[TABLO_BURADA]" in satir:
                if onceki_paragraf is not None:
                    onceki_paragraf.paragraph_format.space_after = Pt(8)
                tablo_tbdy = self.rapor_tasima_gucu_tablosu_olustur(doc)
                tg_p._p.addprevious(tablo_tbdy._tbl)
                tablo_eklendi = True
                onceki_paragraf = None
            elif satir.strip() == "":
                onceki_paragraf = tg_p.insert_paragraph_before("")
            else:
                paragraf = tg_p.insert_paragraph_before("")
                self.rapor_formatli_metin_ekle(paragraf, satir)
                onceki_paragraf = paragraf
        tasima_turu = ""
        if getattr(self, "zemin_kaya_var", None) is not None:
            tasima_turu = self.zemin_kaya_var.get()
        if tasima_turu != "kaya" and not tablo_eklendi:
            raise ValueError("Taşıma gücü rapor metninde [TABLO_BURADA] etiketi bulunamadı.")
        if not tg_p.text.strip():
            ebeveyn = tg_p._p.getparent()
            if ebeveyn is not None:
                ebeveyn.remove(tg_p._p)

    def rapor_kesit_ekle(self, doc):
        kesit_p = self.rapor_paragraf_bul(doc, "[KESIT]")
        if not kesit_p:
            return
        self.rapor_metin_degistir(doc, "[KESIT]", "")
        ac_sekmeleri = []
        for kayit in self.ac_yn_sekme_kayitlari():
            ac_sekmeleri.append({"isim": kayit["isim"], "satirlar": self.ac_yn_satirlari(kayit)})

        with tempfile.NamedTemporaryFile(prefix="rapor_kesit_", suffix=".jpg", delete=False) as f:
            k_resim_yolu = f.name
        try:
            self.kesit_ciz_olustur(ac_sekmeleri, k_resim_yolu)
            self.rapor_gorsel_ekle(kesit_p, k_resim_yolu, width=Inches(6.0))
        finally:
            try:
                os.remove(k_resim_yolu)
            except OSError:
                pass
        kesit_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def rapor_ac_loglarini_ekle(self, doc):
        log_p = self.rapor_paragraf_bul(doc, "[AC_LOGLARI]")
        if not log_p:
            return
        self.rapor_metin_degistir(doc, "[AC_LOGLARI]", "")
        anchor = log_p._p
        for kayit in self.ac_yn_sekme_kayitlari():
            veri = self.ac_yn_kaydi_verisini_oku(kayit)
            isim = veri["isim"]
            derinlik = veri["derinlik"]
            enlem = veri["enlem"]
            boylam = veri["boylam"]
            aciklama = veri["aciklama"]
            tree = kayit["tree"]

            baslik_p = doc.add_paragraph(f"{isim} LOGU", style='Heading 2')
            anchor = self.rapor_xml_sonrasina_ekle(anchor, baslik_p._p)

            bilgi_tablo = doc.add_table(rows=2, cols=4)
            bilgi_tablo.style = 'Table Grid'
            bilgi_tablo.rows[0].cells[0].text = "Çukur/Numune No:"
            bilgi_tablo.rows[0].cells[1].text = isim
            bilgi_tablo.rows[0].cells[2].text = "Koordinatlar:"
            bilgi_tablo.rows[0].cells[3].text = f"Y: {enlem} / X: {boylam}"
            bilgi_tablo.rows[1].cells[0].text = "Derinlik:"
            bilgi_tablo.rows[1].cells[1].text = f"{derinlik} m"
            bilgi_tablo.rows[1].cells[2].text = "Açıklama:"
            bilgi_tablo.rows[1].cells[3].text = aciklama
            self.rapor_tablo_stili_uygula(bilgi_tablo, header_rows=0, label_columns={0, 2})
            anchor = self.rapor_xml_sonrasina_ekle(anchor, bilgi_tablo._tbl)

            veri_tablo = doc.add_table(rows=1, cols=len(tree["columns"]))
            veri_tablo.style = 'Table Grid'
            for c_idx, col in enumerate(tree["columns"]):
                veri_tablo.rows[0].cells[c_idx].text = col
            for row_id in tree.get_children():
                satir = tree.item(row_id)["values"]
                r = veri_tablo.add_row()
                for c_idx, val in enumerate(satir):
                    r.cells[c_idx].text = str(val) if val is not None and str(val) != "" else "-"
            self.rapor_tablo_stili_uygula(veri_tablo, header_rows=1)
            anchor = self.rapor_xml_sonrasina_ekle(anchor, veri_tablo._tbl)

            bosluk_p = doc.add_paragraph("")
            anchor = self.rapor_xml_sonrasina_ekle(anchor, bosluk_p._p)

    def rapor_lab_tablosu_olustur(self, doc, tree, kolon_map):
        return laboratuvar_rapor_lab_tablosu_olustur(
            doc,
            tree,
            kolon_map,
            self.rapor_tabloyu_ortala,
            self.rapor_tablo_stili_uygula,
        )

    def rapor_lab_tablolarini_ekle(self, doc):
        laboratuvar_rapor_lab_tablolarini_ekle(self, doc)

    def rapor_icerigini_doldur(self, doc):
        self.rapor_statik_etiketleri_degistir(doc)
        jeoloji_word_yolu = self.rapor_jeoloji_bolumu_ekle(doc)
        self.rapor_bina_tablosu_ekle(doc)
        self.rapor_resimleri_ekle(doc)
        self.rapor_koordinat_tablolarini_ekle(doc)
        self.rapor_jeofizik_tablolarini_ekle(doc)
        self.rapor_tasima_gucu_ekle(doc)
        self.rapor_kesit_ekle(doc)
        self.rapor_ac_loglarini_ekle(doc)
        self.rapor_lab_tablolarini_ekle(doc)
        self.rapor_belge_akisini_duzenle(doc)
        self.rapor_cozulemeyen_etiketleri_dogrula(doc)
        self.rapor_update_fields_ayarla(doc)
        bolgesel_jeoloji_verisi = deepcopy(getattr(self, "genel_jeoloji_verisi", {}))
        if isinstance(bolgesel_jeoloji_verisi, dict) and bolgesel_jeoloji_verisi.get("birimler"):
            bolgesel_jeoloji_verisi["gorsel_yolu"] = getattr(self, "img_genel_jeoloji", "") or ""
        return {
            "jeoloji_word_yolu": jeoloji_word_yolu or "",
            "bolgesel_jeoloji_verisi": bolgesel_jeoloji_verisi,
            "muhendislik_jeolojisi_cumlesi": getattr(
                self,
                "_rapor_muhendislik_jeolojisi_cumlesi",
                "",
            ),
        }

